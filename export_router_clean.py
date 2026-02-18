
from fastapi import FastAPI

# === Default FastAPI app (auto-generated) ===
app = FastAPI(title="Export Service", version="1.0")

from fastapi import APIRouter
from fastapi.responses import JSONResponse
from pathlib import Path
from datetime import datetime
import json

from export_formatter import format_export_to_word

router = APIRouter(tags=["export"])

EXPORT_DIR = Path("exports")
EXPORT_DIR.mkdir(exist_ok=True)

@router.get("/export")
def export_latest(response_type: str = "json"):
    """
    导出评分结果与追溯日志的综合文件。
    response_type: json / docx / pdf
    """
    score_file = Path("score_result.json")
    if not score_file.exists():
        return JSONResponse({"error": "尚未生成评分结果文件，请先运行 /score 接口。"}, status_code=404)

    with score_file.open("r", encoding="utf-8") as f:
        data = json.load(f)

    export_time = datetime.now().strftime("%Y%m%d_%H%M%S")
    export_name = f"export_{export_time}"
    meta = {
        "export_name": export_name,
        "generated_at": export_time,
        "source": str(score_file),
        "trace_chain": ["评分模块 /score", "导出模块 /export"],
    }

    output_path = EXPORT_DIR / f"{export_name}.json"
    with output_path.open("w", encoding="utf-8") as wf:
        json.dump({"meta": meta, "data": data}, wf, ensure_ascii=False, indent=2)

    if response_type == "json":
        result = {"message": "导出成功", "file": str(output_path)}
        try:
            format_export_to_word(str(output_path))
            result["formatted_docx"] = str(output_path.with_suffix(".docx"))
        except Exception as e:
            result["warning"] = f"Word 导出失败：{e}"
        return JSONResponse(result)

    elif response_type == "docx":
        try:
            format_export_to_word(str(output_path))
            return JSONResponse({
                "message": "已生成 Word 报告",
                "file": str(output_path),
                "formatted_docx": str(output_path.with_suffix(".docx"))
            })
        except Exception as e:
            return JSONResponse({"error": f"Word 导出失败：{e}"}, status_code=500)

    elif response_type == "pdf":
        return JSONResponse({"message": "PDF 导出功能占位，稍后集成排版模块。"})

    return JSONResponse({"error": "不支持的导出类型"}, status_code=400)


# === Codex assist integration (optional) ===
try:
    from backend.routers.assist_codex import router as _codex_router
    app.include_router(_codex_router, prefix="/assist", tags=["assist"])
    print("[codex] router mounted at /assist")
except Exception as _e:
    print(f"[codex] integration skipped: {_e}")


# === 自动附加 AI 校核说明到导出报告 ===
import os
from pathlib import Path
from docx import Document

def append_ai_note_to_report(report_path: str):
    note_path = Path("exports/AI校核说明.txt")
    if not os.path.exists(report_path) or not note_path.exists():
        print(f"[AI附页] 缺少文件：{report_path} 或 {note_path}")
        return

    doc = Document(report_path)
    doc.add_page_break()
    doc.add_heading("附录C  AI校核说明", level=1)
    doc.add_paragraph("以下内容由系统自动生成，用于辅助人工复核：")

    for line in note_path.read_text(encoding="utf-8").splitlines():
        doc.add_paragraph(line)

    new_path = report_path.replace(".docx", "+AI说明.docx")
    doc.save(new_path)
    print(f"[AI附页] 已自动生成报告：{new_path}")


# === M5: audit log demo route ===
try:
    from fastapi import Body
    from backend.audit.audit_log import write_export_log

    @app.post("/assist/audit/log_demo")
    def _assist_audit_log_demo(payload: dict = Body(default={})):
        # 这里只写一份演示日志，返回日志路径
        inputs = payload.get("inputs", {})
        outputs = payload.get("outputs", {})
        meta = payload.get("meta", {})
        log_path = write_export_log(
            event="export_demo",
            inputs=inputs,
            outputs=outputs,
            meta=meta
        )
        return {"ok": True, "log_path": log_path}
except Exception as _e:
    print(f"[audit] route init skipped: {_e}")


# === M5: audit log replay demo ===
import json
from fastapi import Query
from backend.audit.audit_log import _sha256
from pathlib import Path

@app.get("/assist/audit/replay_demo")
def _assist_audit_replay_demo(log_file: str = Query(..., description="日志文件名, 位于 logs/exports/<date>/ 下")):
    path = Path("logs/exports")/log_file
    if not path.exists():
        return {"ok": False, "error": f"日志文件不存在: {path}"}
    data = json.loads(path.read_text(encoding="utf-8"))
    raw = json.dumps(data.get("payload"), ensure_ascii=False, sort_keys=True)
    sha_local = _sha256(raw)
    return {
        "ok": True,
        "rid": data.get("rid"),
        "event": data.get("event"),
        "env": data.get("env"),
        "sha256_logged": data.get("payload_sha256"),
        "sha256_recomputed": sha_local,
        "verified": sha_local == data.get("payload_sha256"),
        "payload": data.get("payload")
    }


# === M5: 产物指纹入库（文件路径 + SHA256） ===
import hashlib, os
from backend.audit.audit_log import write_export_log

def _sha256_file(path:str)->str:
    h=hashlib.sha256()
    with open(path,'rb') as f:
        for chunk in iter(lambda: f.read(1<<16), b''):
            h.update(chunk)
    return h.hexdigest()

@app.post("/assist/audit/log_bind")
def _assist_audit_log_bind():
    # 按你的现有导出目录，收集典型产物（你也可以换成真实导出文件名）
    report = "exports/缺口分析附录.with-indexes+AI说明.docx"
    ai_note = "exports/AI校核说明.txt"

    inputs  = {"citations": {"file_path": "citations.sample.json"},
               "scores":    {"file_path": "scores.sample.json"}}
    outputs = {}
    for fp in [report, ai_note]:
        if os.path.exists(fp):
            outputs[os.path.basename(fp)] = {"file_path": fp, "sha256": _sha256_file(fp)}

    meta = {"purpose":"export-fingerprint","who":"api","route":"/assist/audit/log_bind"}
    log_path = write_export_log(event="export_fingerprint", inputs=inputs, outputs=outputs, meta=meta)
    return {"ok": True, "log_path": log_path, "outputs": outputs}
