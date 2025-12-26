
from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

from docx.oxml.ns import qn
def _apply_run_font(run, font_name: str, east_asia_name: str, size_pt: float):
    from docx.shared import Pt
    try:
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        # 同步设置东亚字体，防止中文回退为系统默认
        if run._element.rPr is None:
            run._element.get_or_add_rPr()
        if run._element.rPr.rFonts is None:
            run._element.rPr.get_or_add_rFonts()
        run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_name)
        run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    except Exception:
        pass


# 可选 PDF 转换
try:
    from docx2pdf import convert as docx2pdf_convert
except Exception:
    docx2pdf_convert = None

router = APIRouter()

class PublishRequest(BaseModel):
    response_file: Optional[str] = None
    title_prefix: Optional[str] = "专业排版导出"
    format: Optional[str] = "docx"              # 支持: docx / pdf
    style: Optional[Dict[str, Any]] = None      # 样式：font_size_pt/line_spacing/lines_per_page/max_chars_per_line/font_name/monospace

def _read_response(base: Path, pr: PublishRequest):
    p = Path(pr.response_file) if pr.response_file else (base/'last_score_response.json')
    if not p.exists():
        raise HTTPException(status_code=400, detail=f"未找到响应文件：{p}")
    return json.loads(p.read_text(encoding='utf-8')), p

def _build_docx(base: Path, data: Dict[str, Any], src_path: Path, title_prefix: str, style: Dict[str, Any] | None = None) -> Dict[str, str]:
    style = style or {}
    east_asia_font = style.get("east_asia_font", style.get("font_name", "PingFang SC"))
    font_name = style.get("font_name", "Times New Roman")
    font_size_pt = float(style.get("font_size_pt", 11))
    line_spacing = float(style.get("line_spacing", 1.2))
    lines_per_page = style.get("lines_per_page")                 # int or None
    max_chars = style.get("max_chars_per_line")                  # int or None
    use_mono = bool(style.get("monospace", False))

    details = data.get('details', [])
    total_score = data.get('total_score')
    exports = data.get('exports', {})

    outdir = base/'exports'
    outdir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    docx_path = outdir/f'{title_prefix}_{stamp}.docx'
    meta_path = outdir/f'{title_prefix}_{stamp}.meta.json'

    # --- 文档排版 ---
    doc = Document()

    # 设置默认 Normal 样式
    try:
        st = doc.styles["Normal"]
        st.font.name = font_name
        st.font.size = Pt(font_size_pt)
        st.paragraph_format.line_spacing = line_spacing
    except Exception:
        pass

    for s in doc.sections:
        # 基础边距
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
        # 页眉/页脚
        header = s.header.paragraphs[0]
        header.text = "专业级可追溯文档自动化生成系统 | 审查报告"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.runs[0].font.size = Pt(9)
        footer = s.footer.paragraphs[0]
        footer.text = "Confidential · 自动生成"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(9)

        # 近似控制每页行数：按 行高≈字号×行距 计算可写高度并微调底边距
        try:
            if isinstance(lines_per_page, int) and lines_per_page > 5:
                line_pitch_pt = font_size_pt * line_spacing
                content_pt = s.page_height.pt - s.top_margin.pt - s.bottom_margin.pt
                need_pt = lines_per_page * line_pitch_pt
                if need_pt < s.page_height.pt * 0.95:
                    new_bottom = max(36.0, s.page_height.pt - s.top_margin.pt - need_pt)  # ≥0.5英寸≈36pt
                    s.bottom_margin = Pt(new_bottom)
        except Exception:
            pass

        # 近似控制每行字符数：按 字宽≈(0.6或0.5)*字号 调整右边距（建议配合等宽字体）
        try:
            if isinstance(max_chars, int) and max_chars > 10:
                char_w = (0.5 if use_mono else 0.6) * font_size_pt
                text_w_pt = max_chars * char_w
                page_w_pt = s.page_width.pt - s.left_margin.pt - s.right_margin.pt
                if text_w_pt < page_w_pt:
                    shrink = page_w_pt - text_w_pt
                    s.right_margin = Pt(max(36.0, s.right_margin.pt + shrink))
        except Exception:
            pass

    # 封面
    t = doc.add_paragraph(); r = t.add_run("审查与缺口分析报告")
    r.bold = True; r.font.size = Pt(24); r.font.name = font_name; _apply_run_font(r, font_name, east_asia_font, 24)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2 = doc.add_paragraph()
    rr = t2.add_run(datetime.now().strftime("导出时间：%Y-%m-%d %H:%M:%S"))
    rr.font.size = Pt(11); rr.font.name = font_name; _apply_run_font(rr, font_name, east_asia_font, 11)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # 总览
    h = doc.add_paragraph(); h1 = h.add_run("一、评分总览")
    h1.bold = True; h1.font.size = Pt(14); h1.font.name = font_name; _apply_run_font(h1, font_name, east_asia_font, 14)
    p = doc.add_paragraph(f"总分（total_score）：{total_score if total_score is not None else 'N/A'}")

    # 明细
    doc.add_paragraph("")
    h2 = doc.add_paragraph(); h2r = h2.add_run("二、评分点明细")
    h2r.bold = True; h2r.font.size = Pt(14); h2r.font.name = font_name; _apply_run_font(h2r, font_name, east_asia_font, 14)

    def as_rows(items):
        if not items: return [], []
        if isinstance(items, dict): items = [items]
        cols = set()
        for it in items:
            if isinstance(it, dict): cols.update(it.keys())
            else: cols.update(["item"])
        cols = list(cols)
        rows = []
        for it in items:
            if isinstance(it, dict): rows.append([str(it.get(c, "")) for c in cols])
            else: rows.append([str(it)])
        return rows, cols

    rows, cols = as_rows(details)
    if rows and cols:
        table = doc.add_table(rows=1, cols=len(cols))
        hdr = table.rows[0].cells
        for i, c in enumerate(cols): hdr[i].text = str(c)
        for rrow in rows:
            cells = table.add_row().cells
            for i, v in enumerate(rrow): cells[i].text = str(v)
    else:
        doc.add_paragraph("（无明细可展示）")

    # 追溯链
    doc.add_paragraph("")
    h3 = doc.add_paragraph(); h3r = h3.add_run("三、追溯链信息（Trace Chain）")
    h3r.bold = True; h3r.font.size = Pt(14); h3r.font.name = font_name; _apply_run_font(h3r, font_name, east_asia_font, 14)
    trace = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "engine_ruleset": "rules_sample.json",
        "response_has_exports": bool(exports),
        "exports_excel": exports.get("excel"),
        "exports_word": exports.get("word"),
        "source_response_file": str(src_path),
        "publisher_version": "M4-style-v1"
    }
    for k, v in trace.items():
        doc.add_paragraph(f"{k}: {v}")

    _normalize_all(doc, font_name, east_asia_font, font_size_pt, line_spacing)
    doc.save(docx_path)

    meta = {
        "publisher_version": trace["publisher_version"],
        "generated_at": trace["generated_at"],
        "inputs": {"response_file": str(src_path)},
        "outputs": {"docx": str(docx_path)},
        "links": {"excel": exports.get("excel"), "word": exports.get("word")},
        "style": {
            "font_name": font_name, "font_size_pt": font_size_pt, "line_spacing": line_spacing,
            "lines_per_page": lines_per_page, "max_chars_per_line": max_chars, "monospace": use_mono
        }
    }
    Path(meta_path).write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding='utf-8')
    return {"docx": str(docx_path), "meta": str(meta_path)}


def _normalize_all(doc, font_name: str, east_asia_name: str, font_size_pt: float, line_spacing: float):
    # 段落
    for p in doc.paragraphs:
        try:
            p.paragraph_format.line_spacing = line_spacing
        except Exception:
            pass
        for r in p.runs:
            _apply_run_font(r, font_name, east_asia_name, font_size_pt)
    # 表格
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    try:
                        p.paragraph_format.line_spacing = line_spacing
                    except Exception:
                        pass
                    for r in p.runs:
                        _apply_run_font(r, font_name, east_asia_name, font_size_pt)


def _to_pdf(docx_path: Path) -> Path:
    pdf_path = docx_path.with_suffix(".pdf")
    if docx2pdf_convert is None:
        raise RuntimeError("PDF 转换不可用：未检测到 docx2pdf / 无法调用 Word")
    docx2pdf_convert(str(docx_path), str(pdf_path))
    if not pdf_path.exists():
        raise RuntimeError("docx2pdf 执行后未生成 PDF")
    return pdf_path

@router.post("/publish")
def publish(req: PublishRequest):
    base = Path.home()/'Desktop/文档生成系统'
    data, src = _read_response(base, req)
    fmt = (req.format or "docx").lower()
    outs = _build_docx(base, data, src, req.title_prefix or "专业排版导出", req.style or {})
    if fmt == "docx":
        return {"ok": True, "format": "docx", "outputs": outs}
    if fmt == "pdf":
        try:
            pdf_path = _to_pdf(Path(outs["docx"]))
            outs["pdf"] = str(pdf_path)
            return {"ok": True, "format": "pdf", "outputs": outs}
        except RuntimeError as e:
            return {"ok": False, "format": "pdf", "error": str(e), "outputs": outs}
    raise HTTPException(status_code=400, detail="当前仅支持 format=docx 或 pdf")
