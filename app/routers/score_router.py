
try:
    import pandas as pd
    from docx import Document
    DOC_EXPORT=True
except Exception:
    DOC_EXPORT=False

# -*- coding: utf-8 -*-
from fastapi import APIRouter, HTTPException
from pydantic import BaseModel, Field
from ..core.rule_engine import RuleEngine

router = APIRouter(prefix="/score", tags=["score"])

# 规则引擎单例（加载 rules_sample.json）
_engine = RuleEngine("rules_sample.json")

class ScoreRequest(BaseModel):
    text: str = Field(..., min_length=1, description="待核验与评分的文本")

@router.get("/ping")
def ping():
    return {"ok": True, "service": "score"}

@router.post("", summary="根据规则引擎对文本进行评分核验")
def score(req: ScoreRequest):
    text = req.text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="text 不能为空")
    result = _engine.evaluate(text)
    resp = {"ok": True, "total_score": result["total_score"], "details": result["details"]}
    if DOC_EXPORT:
        from pathlib import Path
        from datetime import datetime
        outdir = Path.home()/"Desktop/文档生成系统/exports"
        outdir.mkdir(parents=True, exist_ok=True)
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        xlsx_path = outdir / f"评分点覆盖清单_{stamp}.xlsx"
        docx_path = outdir / f"缺口分析附录_{stamp}.docx"
        try:
            import pandas as pd
            df = pd.DataFrame(result.get("details", []))
            with pd.ExcelWriter(xlsx_path) as w:
                df.to_excel(w, index=False, sheet_name="coverage")
        except Exception:
            pass
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("缺口分析附录", level=1)
            for item in result.get("details", []):
                doc.add_paragraph(str(item))
            doc.save(docx_path)
        except Exception:
            pass
        resp["exports"] = {"excel": str(xlsx_path), "word": str(docx_path)}
    return resp
