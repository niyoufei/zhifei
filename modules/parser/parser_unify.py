import os
from typing import Dict, Any, Tuple, List

class UnifiedParser:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_type = self._detect_type()

    def _detect_type(self) -> str:
        ext = os.path.splitext(self.file_path)[1].lower()
        if ext in ['.docx']:
            return 'word'
        elif ext in ['.pdf']:
            return 'pdf'
        elif ext in ['.xlsx', '.xls']:
            return 'excel'
        elif ext in ['.png', '.jpg', '.jpeg']:
            return 'image'
        elif ext in ['.dxf']:
            return 'cad'
        elif ext in ['.dwg']:
            return 'dwg'
        return 'unknown'

    def _read_docx_text(self) -> Tuple[str, Dict[str, Any]]:
        from docx import Document
        doc = Document(self.file_path)
        lines: List[str] = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                lines.append(p.text.strip())
        return "\n".join(lines), {"paragraphs": len(doc.paragraphs)}

    def _read_pdf_text(self) -> Tuple[str, Dict[str, Any]]:
        from pypdf import PdfReader
        reader = PdfReader(self.file_path)
        texts = []
        for page in reader.pages:
            texts.append(page.extract_text() or "")
        text = "\n\n".join(texts)
        return text, {"pages": len(reader.pages)}

    def _read_excel_text(self) -> Tuple[str, Dict[str, Any]]:
        ext = os.path.splitext(self.file_path)[1].lower()
        if ext == ".xlsx":
            from openpyxl import load_workbook

            wb = load_workbook(self.file_path, data_only=True, read_only=True)
            lines: List[str] = []
            for ws in wb.worksheets:
                lines.append(f"[Sheet] {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    row_vals = [str(v) for v in row if v not in (None, "")]
                    if row_vals:
                        lines.append(" | ".join(row_vals))
            return "\n".join(lines), {"sheets": len(wb.worksheets), "engine": "openpyxl"}

        # xls: 优先用 pandas 读取，若环境无 xlrd 则抛错并由上层记录 meta.error
        import pandas as pd

        xls = pd.read_excel(self.file_path, sheet_name=None, header=None)
        lines: List[str] = []
        for sheet_name, df in xls.items():
            lines.append(f"[Sheet] {sheet_name}")
            for row in df.fillna("").values.tolist():
                row_vals = [str(v) for v in row if str(v).strip()]
                if row_vals:
                    lines.append(" | ".join(row_vals))
        return "\n".join(lines), {"sheets": len(xls), "engine": "pandas"}

    def _read_image_meta(self) -> Dict[str, Any]:
        from PIL import Image
        with Image.open(self.file_path) as im:
            return {"width": im.width, "height": im.height, "mode": im.mode}

    def _cad_meta_to_text(self, meta: Dict[str, Any]) -> str:
        if not isinstance(meta, dict):
            return ""
        blocks = meta.get("insert_blocks") or {}
        top_blocks = []
        if isinstance(blocks, dict):
            for k, v in sorted(blocks.items(), key=lambda x: x[1], reverse=True)[:8]:
                top_blocks.append(f"{k}:{v}")
        topo_text = ""
        try:
            from modules.parser.drawing_topology import topology_summary_text

            topo_text = topology_summary_text(meta.get("topology") if isinstance(meta.get("topology"), dict) else {})
        except Exception:
            topo_text = ""
        topo_line = f"\n{topo_text}" if topo_text else ""
        return (
            f"CAD图纸信息\n"
            f"图层数量: {meta.get('layers_count')}\n"
            f"实体数量: {meta.get('entities_count')}\n"
            f"块引用: {'; '.join(top_blocks)}"
            f"{topo_line}"
        )

    def parse(self) -> Dict[str, Any]:
        ext = os.path.splitext(self.file_path)[1].lower()[1:]

        if ext == "dxf":
            from modules.parser.parse_cad import parse_cad_from_dxf
            meta = parse_cad_from_dxf(self.file_path)
            return {
                "type": "cad",
                "text": self._cad_meta_to_text(meta),
                "meta": meta,
            }

        if ext == "docx":
            text, meta = self._read_docx_text()
            return {"type": "word", "text": text, "meta": meta}

        if ext == "pdf":
            text, meta = self._read_pdf_text()
            return {"type": "pdf", "text": text, "meta": meta}

        if ext in {"xlsx", "xls"}:
            text, meta = self._read_excel_text()
            return {"type": "excel", "text": text, "meta": meta}

        if ext in {"png", "jpg", "jpeg"}:
            meta = self._read_image_meta()
            text = f"图片信息\n宽: {meta.get('width')}\n高: {meta.get('height')}\n模式: {meta.get('mode')}"
            return {"type": "image", "text": text, "meta": meta}

        if ext == "dwg":
            note = "dwg 暂不支持解析"
            return {"type": "dwg", "text": f"图纸信息\n类型: DWG\n说明: {note}", "meta": {"note": note}}

        return {"type": self.file_type, "text": "", "meta": {"note": "该类型尚未实现解析"}}
