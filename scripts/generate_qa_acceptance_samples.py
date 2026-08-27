#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path
from typing import Any

from docx import Document

from backend.zhifei_autoplan.engineering_graphics import (
    GraphicEdge,
    GraphicNode,
    GraphicSpec,
    render_engineering_graphic,
)
from backend.zhifei_autoplan.exporter import export_autoplan_docx


REPO_ROOT = Path(__file__).resolve().parents[1]
QA_ROOT = REPO_ROOT / "artifacts" / "qa"
DOCX_ROOT = QA_ROOT / "docx"
INPUT_ROOT = QA_ROOT / "inputs"
GRAPHICS_ROOT = QA_ROOT / "graphics"

STYLE = {
    "paper": "A4",
    "body_font": "宋体",
    "body_latin_font": "Times New Roman",
    "body_size": 14,
    "title_font": "宋体",
    "title_latin_font": "Times New Roman",
    "title_size": 16,
    "doc_title_size": 16,
    "line_spacing_rule": "fixed",
    "line_spacing_pt": 22,
    "margin_top_cm": 2.5,
    "margin_right_cm": 2.0,
    "margin_bottom_cm": 2.0,
    "margin_left_cm": 2.0,
    "chapter_start_new_page": True,
    "front_matter": {"toc_pages": 2, "full_index_enabled": True},
}

CHAPTER_TOPICS = (
    "编制说明与项目响应",
    "施工总体部署",
    "施工准备与场地组织",
    "施工进度计划与资源配置",
    "测量复核与技术管理",
    "道路与基层施工方法",
    "排水与管线施工方法",
    "交通组织与接口协调",
    "质量保证体系",
    "安全生产与应急管理",
    "文明施工与环境保护",
    "成品保护与资料移交",
)

D_GRAPHIC_PROFILES = (
    ("组织架构图", 7, "tree"),
    ("工艺流程图", 6, "auto"),
    ("安全管理体系图", 8, "tree"),
    ("进度时间轴", 4, "three_row"),
    ("两排流程图", 10, "two_row"),
    ("左右树状图", 9, "tree"),
    ("照片组合图", 6, "three_row"),
    ("参数标注图", 5, "auto"),
    ("大型横向图", 12, "three_row"),
)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _paragraph(chapter_index: int, paragraph_index: int) -> str:
    return (
        f"第{chapter_index}章第{paragraph_index}项为合成验收数据，仅用于检验施组编制系统的分页、样式、目录、表格和渲染稳定性。"
        "实施过程按照施工准备、技术交底、样板确认、过程检查、整改复验和资料归档形成闭环；"
        "具体工程量、工期、材料型号及现场条件必须在真实项目中以招标文件、设计图纸和经批准的专项方案为准。"
    )


def _markdown_table(chapter_index: int) -> str:
    rows = [
        "| 序号 | 控制事项 | 实施要求 | 验收记录 |",
        "| --- | --- | --- | --- |",
    ]
    for index in range(1, 7):
        rows.append(
            f"| {index} | 第{chapter_index}章控制点{index} | 交底、检查、复核并形成闭环记录 | 合成验收记录{chapter_index}-{index} |"
        )
    return "\n".join(rows)


def _make_sections(
    *,
    section_count: int,
    paragraphs_per_section: int,
    include_markdown_tables: bool = False,
    extra_paragraph_sections: int = 0,
) -> list[dict[str, str]]:
    sections: list[dict[str, str]] = []
    for section_index in range(1, section_count + 1):
        topic = CHAPTER_TOPICS[(section_index - 1) % len(CHAPTER_TOPICS)]
        title = f"第{section_index}章 {topic}"
        content = ["## 实施原则与验收边界"]
        if section_index == 1:
            content.extend(
                (
                    "## 信息化管理",
                    "信息化管理台账明确参数、检查频次、责任岗位和验收记录，所有字段均为合成验收数据。",
                    "## 绿色工地",
                    "绿色工地措施覆盖扬尘、噪声、污水、固体废物和节材管理，检查问题整改后复验。",
                    "## 劳保用品",
                    "劳保用品按作业风险配置、领用、检查和更换，形成劳保用品配置矩阵。",
                    "## 关键工序控制点表",
                    "风险：合成接口条件不一致导致排版与内容脱节；控制：逐章复核参数、频次、责任和记录；验证：执行自动化结构检查与真实渲染验收。",
                )
            )
        paragraph_count = paragraphs_per_section + int(section_index <= extra_paragraph_sections)
        content.extend(
            _paragraph(section_index, paragraph_index)
            for paragraph_index in range(1, paragraph_count + 1)
        )
        if include_markdown_tables:
            content.extend(("## 本章控制矩阵", _markdown_table(section_index)))
        sections.append({"title": title, "content": "\n".join(content)})
    return sections


def _graphic_media(sample: str, count: int) -> list[dict[str, Any]]:
    target = GRAPHICS_ROOT / sample
    target.mkdir(parents=True, exist_ok=True)
    media: list[dict[str, Any]] = []
    for index in range(1, count + 1):
        # Vary node cardinality and layout independently.  The previous
        # arithmetic tied both cycles together and produced only twelve visual
        # structures, so a 30-image acceptance sample was correctly rejected by
        # the perceptual duplicate gate.  This matrix yields up to 36 distinct
        # deterministic structures before any title/detail variation is used.
        node_count = 3 + ((index - 1) % 10)
        layout = ("tree", "two_row", "three_row")[((index - 1) // 10) % 3]
        graphic_kind = "施工控制与验收闭环"
        if sample == "sample_D" and index <= len(D_GRAPHIC_PROFILES):
            graphic_kind, node_count, layout = D_GRAPHIC_PROFILES[index - 1]
        if sample != "sample_D" and layout == "three_row" and node_count == 4:
            # (2, 1, 1) has an almost identical 64-bit dHash to the three-node
            # two-row shape despite different content.  A 13-node three-band
            # matrix is both more demanding and unambiguously distinct.
            node_count = 13
        nodes = tuple(
            GraphicNode(
                node_id=f"N{node_index + 1}",
                title=f"{index}-{node_index + 1}控制节点",
                detail=(
                    "施工准备与复核" if (node_index + index) % 3 == 0
                    else "过程检查与整改" if (node_index + index) % 3 == 1
                    else "验收签认与归档"
                ),
            )
            for node_index in range(node_count)
        )
        if layout == "tree" and len(nodes) > 2:
            edges = tuple(
                GraphicEdge(source=nodes[(node_index - 1) // 2].node_id, target=nodes[node_index].node_id)
                for node_index in range(1, len(nodes))
            )
        elif layout in {"two_row", "three_row"} and len(nodes) > 3:
            midpoint = (len(nodes) + 1) // 2
            routed: list[GraphicEdge] = []
            routed.extend(
                GraphicEdge(source=nodes[node_index].node_id, target=nodes[node_index + 1].node_id)
                for node_index in range(max(0, midpoint - 1))
            )
            routed.extend(
                GraphicEdge(source=nodes[node_index].node_id, target=nodes[node_index + 1].node_id)
                for node_index in range(midpoint, len(nodes) - 1)
            )
            routed.append(GraphicEdge(source=nodes[midpoint - 1].node_id, target=nodes[midpoint].node_id))
            edges = tuple(routed)
        else:
            edges = tuple(
                GraphicEdge(source=nodes[node_index].node_id, target=nodes[node_index + 1].node_id)
                for node_index in range(max(0, len(nodes) - 1))
            )
        spec = GraphicSpec(
            title=f"工程图示{index:02d}：{graphic_kind}",
            subtitle=f"{sample} 合成验收样本 · {node_count}节点 · 非真实项目",
            nodes=nodes,
            edges=edges,
            layout=layout,
            caption=f"{graphic_kind}仅用于程序化排版与渲染验收",
        )
        png_path = target / f"graphic_{index:02d}.png"
        svg_path = target / f"graphic_{index:02d}.svg"
        receipt = render_engineering_graphic(spec, png_path=png_path, svg_path=svg_path)
        media.append(
            {
                "path": str(png_path),
                "svg_path": str(svg_path),
                "caption": f"{sample}—工程图示{index:02d}{graphic_kind}",
                "source_kind": "deterministic_project_diagram",
                "source_ref": f"{sample}/graphic_{index:02d}.svg",
                "semantic_terms": ["施工控制", "验收闭环", graphic_kind, f"图示{index:02d}"],
                "text_verified": True,
                "required": True,
                "render_receipt": receipt,
            }
        )
    return media


def _standard_tables(count: int, *, landscape_from: int | None = None) -> list[dict[str, Any]]:
    result: list[dict[str, Any]] = []
    for table_index in range(1, count + 1):
        landscape = bool(landscape_from is not None and table_index >= landscape_from)
        if landscape:
            result.append(
                {
                    "title": f"附表{table_index} 合成横向工序控制台账",
                    "orientation": "landscape",
                    "headers": [
                        "序号", "分部", "工序", "资源", "数量", "参数", "控制要求", "检查频次", "责任岗位", "验收资料",
                    ],
                    "merge_header_groups": [
                        {"start": 0, "end": 2, "label": "工作分解"},
                        {"start": 3, "end": 5, "label": "资源参数"},
                        {"start": 6, "end": 9, "label": "执行与验收"},
                    ],
                    "rows": [
                        [
                            str(row_index),
                            f"分部{table_index}",
                            f"工序{table_index}-{row_index}",
                            "施工班组",
                            f"{row_index}组",
                            "合成参数",
                            "交底、检查、复核、整改",
                            "每道工序",
                            "施工员与质量员",
                            f"验收记录{table_index}-{row_index}",
                        ]
                        for row_index in range(1, 17)
                    ],
                }
            )
            continue
        result.append(
            {
                "title": f"附表{table_index} 合成工序控制台账",
                "headers": ["序号", "工序", "责任岗位", "控制要求", "验收资料"],
                "rows": [
                    [
                        str(row_index),
                        f"合成工序{table_index}-{row_index}",
                        "施工员与质量员",
                        "技术交底、过程检查、整改复验",
                        f"验收记录{table_index}-{row_index}",
                    ]
                    for row_index in range(1, 9)
                ],
            }
        )
    return result


def _complex_tables() -> list[dict[str, Any]]:
    tables: list[dict[str, Any]] = []
    for table_index in range(1, 9):
        landscape = table_index >= 5
        headers = ["序号", "工序", "资源", "数量", "控制措施", "验收记录", "责任岗位", "状态"]
        rows: list[list[Any]] = []
        for row_index in range(1, 15):
            resource: Any = "测量与施工班组"
            if table_index == 1 and row_index == 1:
                resource = {
                    "text": "测量组",
                    "nested": {
                        "headers": ["设备", "校验状态"],
                        "rows": [["全站仪", "有效期内"], ["水准仪", "有效期内"]],
                    },
                }
            rows.append(
                [
                    str(row_index),
                    f"复杂表格工序{table_index}-{row_index}",
                    resource,
                    f"{row_index}组",
                    "轴线、标高、材料和工序接口逐项复核，异常整改后复验",
                    f"合成检查记录{table_index}-{row_index}",
                    "专业负责人",
                    "待验收" if row_index % 2 else "已复核",
                ]
            )
        tables.append(
            {
                "title": f"复杂表格{table_index}—{'横向' if landscape else '纵向'}资源与验收矩阵",
                "orientation": "landscape" if landscape else "portrait",
                "headers": headers,
                "merge_header_groups": [
                    {"start": 0, "end": 1, "label": "工作分解"},
                    {"start": 2, "end": 3, "label": "资源配置"},
                    {"start": 4, "end": 7, "label": "执行闭环"},
                ],
                "rows": rows,
            }
        )
    return tables


def _payload(sample: str) -> dict[str, Any]:
    base: dict[str, Any] = {
        "topic": f"{sample} 合成市政道路工程施工组织设计（非真实项目）",
        "project_name": f"{sample} 合成市政道路工程（非真实项目）",
        "project_code": f"QA-{sample}",
        "bidder_company": "合成验收单位（非真实主体）",
        "style": dict(STYLE),
    }
    if sample == "sample_A":
        base.update(
            sections=_make_sections(section_count=10, paragraphs_per_section=9),
            tables=_standard_tables(4),
            media=_graphic_media(sample, 6),
        )
    elif sample == "sample_B":
        base.update(
            sections=_make_sections(
                section_count=32,
                paragraphs_per_section=32,
                include_markdown_tables=True,
            ),
            tables=_standard_tables(8, landscape_from=7),
            media=_graphic_media(sample, 30),
        )
    elif sample == "sample_C":
        base.update(
            sections=_make_sections(
                section_count=34,
                paragraphs_per_section=32,
                extra_paragraph_sections=23,
            ),
            tables=_standard_tables(6),
        )
    elif sample == "sample_D":
        base.update(
            sections=_make_sections(section_count=10, paragraphs_per_section=15),
            tables=_standard_tables(3),
            media=_graphic_media(sample, 9),
        )
    elif sample == "sample_E":
        base.update(
            sections=_make_sections(section_count=8, paragraphs_per_section=8),
            tables=_complex_tables(),
        )
    else:
        raise ValueError(f"unknown sample: {sample}")
    return base


def _artifact_receipt(sample: str, output: Path) -> dict[str, Any]:
    document = Document(str(output))
    structural = json.loads(output.with_suffix(".structural_quality.json").read_text(encoding="utf-8"))
    figures = json.loads(output.with_suffix(".figure_manifest.json").read_text(encoding="utf-8"))
    return {
        "sample": sample,
        "docx": str(output),
        "sha256": _sha256(output),
        "size_bytes": output.stat().st_size,
        "sections": len(document.sections),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "structural_status": structural.get("status"),
        "structural_digest": structural.get("decision_digest"),
        "figure_count": figures.get("figure_count"),
        "figure_delivery_allowed": figures.get("delivery_allowed"),
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--sample",
        action="append",
        choices=["sample_A", "sample_B", "sample_C", "sample_D", "sample_E"],
        help="Generate only the selected sample; may be repeated. Defaults to all.",
    )
    args = parser.parse_args()
    selected = args.sample or ["sample_A", "sample_B", "sample_C", "sample_D", "sample_E"]
    for path in (DOCX_ROOT, INPUT_ROOT, GRAPHICS_ROOT):
        path.mkdir(parents=True, exist_ok=True)

    receipts: list[dict[str, Any]] = []
    for sample in selected:
        payload = _payload(sample)
        input_path = INPUT_ROOT / f"{sample}.json"
        input_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        output_path = DOCX_ROOT / f"{sample}.docx"
        export_autoplan_docx(payload, str(output_path))
        receipts.append(_artifact_receipt(sample, output_path))

    receipt_path = QA_ROOT / "docx_generation_receipt.json"
    prior: dict[str, Any] = {}
    if receipt_path.exists():
        try:
            prior = json.loads(receipt_path.read_text(encoding="utf-8"))
        except Exception:
            prior = {}
    merged = {item["sample"]: item for item in prior.get("samples") or [] if isinstance(item, dict)}
    merged.update({item["sample"]: item for item in receipts})
    receipt_path.write_text(
        json.dumps(
            {
                "schema": "zhifei.qa.docx_generation.v1",
                "synthetic_data_only": True,
                "samples": [merged[key] for key in sorted(merged)],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    print(json.dumps(receipts, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
