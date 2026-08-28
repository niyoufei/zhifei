from __future__ import annotations

import hashlib
import warnings
import zipfile
from collections.abc import Callable
from pathlib import Path
from xml.etree import ElementTree

import pytest
from docx import Document
from openpyxl import Workbook

from backend.zhifei_autoplan.formal_artifact_integrity import (
    FormalArtifactIntegrityError,
    validate_formal_ooxml_artifact,
)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _validate(path: Path, kind: str) -> dict[str, object]:
    return validate_formal_ooxml_artifact(
        path,
        artifact_kind=kind,  # type: ignore[arg-type]
        expected_size=path.stat().st_size,
        expected_sha256=_sha256(path),
    )


def _write_docx(path: Path) -> None:
    document = Document()
    document.add_heading("施工部署", level=1)
    document.add_paragraph("现场复核、过程检查和验收记录形成闭环。")
    document.save(path)


def _write_xlsx(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "summary"
    sheet["A1"] = "project_id"
    sheet["B1"] = "P-001"
    workbook.save(path)


def _rewrite_zip(
    source: Path,
    destination: Path,
    transform: Callable[[str, bytes], bytes | None],
) -> None:
    with zipfile.ZipFile(source, "r") as reader, zipfile.ZipFile(
        destination,
        "w",
        compression=zipfile.ZIP_DEFLATED,
    ) as writer:
        for info in reader.infolist():
            payload = transform(info.filename, reader.read(info))
            if payload is not None:
                writer.writestr(info, payload)


def test_validates_real_minimal_docx_and_returns_stable_projection(
    tmp_path: Path,
) -> None:
    path = tmp_path / "formal.docx"
    _write_docx(path)

    first = _validate(path, "docx")
    second = _validate(path, "docx")

    assert first == second
    assert first["schema_version"] == "formal-ooxml-integrity-v1"
    assert first["artifact_kind"] == "docx"
    assert first["main_part"] == "word/document.xml"
    assert first["body_present"] is True
    assert len(str(first["projection_digest"])) == 64


def test_validates_real_minimal_xlsx_and_resolves_worksheets(
    tmp_path: Path,
) -> None:
    path = tmp_path / "formal.xlsx"
    _write_xlsx(path)

    projection = _validate(path, "xlsx")

    assert projection["artifact_kind"] == "xlsx"
    assert projection["main_part"] == "xl/workbook.xml"
    assert projection["sheet_count"] == 1
    assert len(str(projection["worksheet_parts_digest"])) == 64


def test_rejects_arbitrary_bytes_even_when_size_and_hash_match(
    tmp_path: Path,
) -> None:
    path = tmp_path / "fake.docx"
    path.write_bytes(b"not-an-ooxml-package")

    with pytest.raises(FormalArtifactIntegrityError) as error:
        _validate(path, "docx")

    assert error.value.code == "FORMAL_ARTIFACT_OOXML_INVALID"


@pytest.mark.parametrize(
    ("source_kind", "declared_kind"),
    (("docx", "xlsx"), ("xlsx", "docx")),
)
def test_rejects_docx_xlsx_masquerading(
    tmp_path: Path,
    source_kind: str,
    declared_kind: str,
) -> None:
    path = tmp_path / f"masquerade.{declared_kind}"
    if source_kind == "docx":
        _write_docx(path)
    else:
        _write_xlsx(path)

    with pytest.raises(FormalArtifactIntegrityError) as error:
        _validate(path, declared_kind)

    assert error.value.code == "FORMAL_ARTIFACT_OOXML_KIND_MISMATCH"


def test_rejects_package_with_missing_required_main_part(tmp_path: Path) -> None:
    original = tmp_path / "original.docx"
    malformed = tmp_path / "missing-main.docx"
    _write_docx(original)
    _rewrite_zip(
        original,
        malformed,
        lambda name, payload: None if name == "word/document.xml" else payload,
    )

    with pytest.raises(FormalArtifactIntegrityError) as error:
        _validate(malformed, "docx")

    assert error.value.code == "FORMAL_ARTIFACT_OOXML_INVALID"


def test_rejects_dangling_internal_relationship(tmp_path: Path) -> None:
    original = tmp_path / "original.docx"
    malformed = tmp_path / "dangling.docx"
    _write_docx(original)

    def add_dangling_relationship(name: str, payload: bytes) -> bytes:
        if name != "word/_rels/document.xml.rels":
            return payload
        root = ElementTree.fromstring(payload)
        ElementTree.SubElement(
            root,
            "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship",
            {
                "Id": "rIdDangling",
                "Type": (
                    "http://schemas.openxmlformats.org/officeDocument/2006/"
                    "relationships/image"
                ),
                "Target": "media/missing.png",
            },
        )
        return ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)

    _rewrite_zip(original, malformed, add_dangling_relationship)

    with pytest.raises(FormalArtifactIntegrityError) as error:
        _validate(malformed, "docx")

    assert error.value.code == "FORMAL_ARTIFACT_OOXML_INVALID"


def test_rejects_duplicate_zip_member(tmp_path: Path) -> None:
    path = tmp_path / "duplicate.docx"
    _write_docx(path)
    with zipfile.ZipFile(path, "r") as archive:
        document_xml = archive.read("word/document.xml")
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        with zipfile.ZipFile(path, "a") as archive:
            archive.writestr("word/document.xml", document_xml)

    with pytest.raises(FormalArtifactIntegrityError) as error:
        _validate(path, "docx")

    assert error.value.code == "FORMAL_ARTIFACT_OOXML_INVALID"


def test_rejects_zip_member_path_traversal(tmp_path: Path) -> None:
    path = tmp_path / "traversal.xlsx"
    _write_xlsx(path)
    with zipfile.ZipFile(path, "a") as archive:
        archive.writestr("../escape.xml", b"<escape/>")

    with pytest.raises(FormalArtifactIntegrityError) as error:
        _validate(path, "xlsx")

    assert error.value.code == "FORMAL_ARTIFACT_OOXML_INVALID"


def test_rejects_size_or_sha_witness_mismatch(tmp_path: Path) -> None:
    path = tmp_path / "formal.docx"
    _write_docx(path)

    with pytest.raises(FormalArtifactIntegrityError) as size_error:
        validate_formal_ooxml_artifact(
            path,
            artifact_kind="docx",
            expected_size=path.stat().st_size + 1,
            expected_sha256=_sha256(path),
        )
    assert size_error.value.code == "FORMAL_ARTIFACT_UNTRUSTED"

    with pytest.raises(FormalArtifactIntegrityError) as hash_error:
        validate_formal_ooxml_artifact(
            path,
            artifact_kind="docx",
            expected_size=path.stat().st_size,
            expected_sha256="0" * 64,
        )
    assert hash_error.value.code == "FORMAL_ARTIFACT_HASH_MISMATCH"
