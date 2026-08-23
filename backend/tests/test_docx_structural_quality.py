from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from backend.zhifei_autoplan.docx_structural_quality import audit_docx_structural_quality
from backend.zhifei_autoplan.exporter import export_autoplan_docx


pytestmark = pytest.mark.usefixtures("allow_legacy_export_contract")


def _contract() -> dict:
    path = Path(__file__).parent / "golden" / "technical_bid_quality_contract.json"
    return json.loads(path.read_text(encoding="utf-8"))


def test_representative_technical_bid_matches_structural_golden_contract(tmp_path: Path) -> None:
    contract = _contract()
    output = tmp_path / "technical_bid.docx"
    export_autoplan_docx(
        {
            "topic": "综合楼改造工程施工组织设计",
            "style": {
                "paper": contract["paper"],
                "body_font": contract["body"]["font_family"],
                "body_size_pt": contract["body"]["size_pt"],
                "line_spacing_rule": "fixed",
                "line_spacing_pt": contract["body"]["line_spacing_pt"],
                "margin_top_cm": contract["margins_cm"]["top"],
                "margin_right_cm": contract["margins_cm"]["right"],
                "margin_bottom_cm": contract["margins_cm"]["bottom"],
                "margin_left_cm": contract["margins_cm"]["left"],
            },
            "sections": [
                {
                    "title": "第一章 施工总体部署",
                    "content": "施工准备、工序衔接、过程检查和验收记录形成闭环。" * 20,
                },
                {
                    "title": "第二章 质量保证措施",
                    "content": "质量责任、实测实量、复核验收和资料归档全过程可追溯。" * 20,
                },
            ],
            "tables": [
                {
                    "title": "关键工序控制表",
                    "headers": ["工序", "控制点", "验收记录"],
                    "rows": [["测量放线", "轴线与标高", "复核记录"]],
                }
            ],
        },
        str(output),
    )

    receipt = json.loads(output.with_suffix(".structural_quality.json").read_text(encoding="utf-8"))
    assert receipt["status"] == "pass"
    assert receipt["section_metrics"][0]["width_cm"] == pytest.approx(contract["page_width_cm"], abs=0.12)
    assert receipt["section_metrics"][0]["height_cm"] == pytest.approx(contract["page_height_cm"], abs=0.12)
    assert receipt["body_style"]["size_pt"] == contract["body"]["size_pt"]
    assert receipt["body_style"]["line_spacing_pt"] == contract["body"]["line_spacing_pt"]
    assert all(receipt["word_fields"][name.lower()] for name in contract["required_fields"])
    assert receipt["word_fields"]["update_on_open"] is contract["update_fields_on_open"]


def test_structural_gate_blocks_uncontrolled_word_package(tmp_path: Path) -> None:
    output = tmp_path / "uncontrolled.docx"
    document = Document()
    document.add_paragraph("没有目录、页码和标题结构的正文")
    document.save(output)

    result = audit_docx_structural_quality(output, strict=False)
    codes = {item["code"] for item in result["hard_failures"]}

    assert result["status"] == "blocked"
    assert {"TOC_FIELD_MISSING", "PAGE_FIELD_MISSING", "NUMPAGES_FIELD_MISSING"} <= codes
    assert "HEADING_STRUCTURE_MISSING" in codes


def test_structural_gate_blocks_internal_implementation_leak(tmp_path: Path) -> None:
    output = tmp_path / "leak.docx"
    document = Document()
    document.add_heading("施工部署", level=1)
    document.add_paragraph("main:anthropic -> fallback_1；job_id=hidden")
    document.save(output)

    result = audit_docx_structural_quality(output, strict=False)

    assert "INTERNAL_IMPLEMENTATION_LEAK" in {
        item["code"] for item in result["hard_failures"]
    }


def test_structural_decision_digest_is_stable_for_same_package(tmp_path: Path) -> None:
    output = tmp_path / "stable.docx"
    export_autoplan_docx(
        {
            "topic": "稳定性测试",
            "sections": [{"title": "第一章 施工准备", "content": "准备工作形成检查记录。"}],
        },
        str(output),
    )

    first = audit_docx_structural_quality(output, strict=True)
    second = audit_docx_structural_quality(output, strict=True)

    assert first["decision_digest"] == second["decision_digest"]
