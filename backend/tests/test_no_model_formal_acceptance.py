from __future__ import annotations

import argparse
import errno
import hashlib
import inspect
import json
import socket
import stat
import subprocess
import urllib.request
from pathlib import Path
from types import SimpleNamespace
from typing import Any

import pytest
from docx import Document
from openpyxl import Workbook

from backend.zhifei_autoplan import no_model_formal_acceptance as acceptance
from backend.zhifei_autoplan.delivery_receipt import (
    build_delivery_receipt,
    canonical_delivery_receipt_digest,
)
from backend.zhifei_autoplan.no_model_formal_acceptance import (
    AcceptanceError,
    _artifact_bundle,
    _assert_path_without_symlinks,
    _candidate_source,
    _capture_latest_state,
    _PreparedAcceptance,
    _validate_event_and_provider_chain,
    canonical_digest,
    publish_acceptance_receipt,
    read_regular_file_snapshot,
    validate_acceptance_receipt,
)
from backend.zhifei_autoplan.provider_admission import LAYER_NAMES
from backend.zhifei_autoplan.provider_admission import (
    canonical_digest as provider_canonical_digest,
)
from backend.zhifei_autoplan.provider_admission import (
    decide_required_roles as decide_provider_required_roles,
)
from backend.zhifei_autoplan.provider_admission import (
    public_snapshot as provider_public_snapshot,
)
from backend.zhifei_autoplan.provider_admission import (
    write_snapshot as write_provider_snapshot,
)
from backend.zhifei_autoplan.sealed_compliance import (
    SEALED_COMPLIANCE_ROOT_RELATIVE_PATH,
    SEALED_OFFICIAL_REGISTRY_RELATIVE_PATH,
    sealed_official_registry_path,
)
from scripts import refresh_no_model_formal_acceptance as refresh_cli
from scripts.build_local_release import ReleaseBuildError
from scripts.launch_latest_release import LaunchError
from scripts.runtime_supervisor import SupervisorError

_TEST_FIXED_WRITE_CONTEXT: dict[str, Any] = {}


@pytest.fixture(autouse=True)
def _fixed_write_context_seam(monkeypatch: pytest.MonkeyPatch):
    """Keep production write authority private while permitting atomic writer tests."""

    _TEST_FIXED_WRITE_CONTEXT.clear()

    def fixed_context() -> dict[str, Any]:
        return dict(_TEST_FIXED_WRITE_CONTEXT)

    monkeypatch.setattr(acceptance, "_fixed_current_write_context", fixed_context)
    yield
    _TEST_FIXED_WRITE_CONTEXT.clear()


def _authorize_for_publish(prepared: _PreparedAcceptance) -> _PreparedAcceptance:
    release = prepared.release_projection
    current = next(
        witness
        for witness in prepared.witnesses
        if isinstance(witness, acceptance.FileSnapshot)
        and witness.sha256 == release["current_json_sha256"]
    )
    authority_digest = canonical_digest(
        {
            "kind": "test-fixed-current-write-context",
            "current_sha256": current.sha256,
            "release": release,
            "data_root": str(prepared.data_root),
            "registry_path": str(prepared.registry_path),
        }
    )
    _TEST_FIXED_WRITE_CONTEXT.update(
        {
            "authority_digest": authority_digest,
            "release_identity": release,
            "current_witness": current,
            "data_root": prepared.data_root,
            "registry_path": prepared.registry_path,
            "release_validator": prepared.release_validator,
        }
    )
    acceptance._CURRENT_WRITE_AUTHORITIES[prepared] = authority_digest
    acceptance._PREPARED_CAPABILITIES[prepared] = acceptance._prepared_signature(
        prepared
    )
    return prepared


def _release(tmp_path: Path, current_sha256: str) -> dict[str, Any]:
    release_id = "release-" + "1" * 24
    return {
        "system_id": "docgen-system",
        "release_id": release_id,
        "release_root": str(tmp_path / release_id),
        "health_status": "verified_healthy",
        "supervisor_instance_id": "instance-1",
        "manifest_digest": "2" * 64,
        "source_digest": "3" * 64,
        "runtime_digest": "4" * 64,
        "current_json_sha256": current_sha256,
        "supervisor_state_sha256": "5" * 64,
        "backend_health_sha256": "6" * 64,
        "supervisor_pid": 101,
        "backend_pid": 102,
        "ui_pid": 103,
    }


def _stages(tmp_path: Path, registry_sha256: str) -> dict[str, Any]:
    digest = "a" * 64
    return {
        "drawing_index": {
            "digest": digest,
            "ok": False,
            "processed": 0,
            "indexed": 0,
            "graphics_only_pages": 0,
            "integrity_rejections": 0,
            "identity_errors": 0,
            "text_status": "missing",
            "page_coverage_status": "missing",
            "chapter_binding_status": "missing",
        },
        "standard_index": {
            "digest": digest,
            "ok": False,
            "official_registry_path": str(tmp_path / "registry.json"),
            "official_registry_sha256": registry_sha256,
            "indexed": 0,
            "official_verified": 0,
            "integrity_rejections": 0,
            "identity_errors": 0,
            "missing_text_or_ocr": 0,
            "locator_unavailable": 0,
            "chapter_binding_count": 0,
            "chapter_binding_status": "missing",
            "standards": [],
        },
        "project_parameter_evidence": {
            "digest": digest,
            "status": "HOLD",
            "ready": False,
            "coverage_complete": False,
            "conflict_count": 0,
            "evidence_set_receipt_digest": None,
            "validation": {"ok": False},
        },
        "project_fact_ledger": {
            "digest": digest,
            "validation": {"ok": False},
            "formal_parameter_readiness": "hold",
            "unresolved_fields": ["resource_peak"],
        },
        "confirmation_checklist": {
            "digest": digest,
            "formal_ready": False,
            "resolved_fields": [],
            "blocked_fields": ["resource_peak"],
        },
        "cross_index": {
            "digest": digest,
            "validation": {"ok": False},
            "focus_count": 0,
            "mentioned_count": 0,
            "closed_ok_count": 0,
            "missing_drawing_locator_count": 0,
            "missing_standard_locator_count": 0,
        },
        "formal_delivery_gate": {
            "digest": digest,
            "delivery_allowed": False,
            "blocker_codes": ["NO_CURRENT_FORMAL_SOURCE"],
            "blocker_count": 1,
            "formal_contract_version": "formal-delivery-v1",
        },
    }


def _receipt(
    tmp_path: Path,
    *,
    run_id: str,
    release: dict[str, Any],
    supersedes: str | None = None,
) -> dict[str, Any]:
    registry_sha256 = "7" * 64
    inputs = {
        "tender": {
            "label": "tender_matrix.json",
            "status": "present",
            "sha256": "8" * 64,
            "size": 10,
            "absence_digest": None,
        },
        "boq": {
            "label": "boq_data.json",
            "status": "present",
            "sha256": "9" * 64,
            "size": 10,
            "absence_digest": None,
        },
        "plan": {
            "label": "plan.json",
            "status": "missing",
            "sha256": None,
            "size": 0,
            "absence_digest": "c" * 64,
        },
        "ingest_audit": {
            "label": "audit/ingest.jsonl",
            "status": "missing",
            "sha256": None,
            "size": 0,
            "absence_digest": "d" * 64,
        },
        "approval_audit": {
            "label": "audit/project_fact_approvals.jsonl",
            "status": "missing",
            "sha256": None,
            "size": 0,
            "absence_digest": "e" * 64,
            "requested_count": 0,
            "verified_count": 0,
            "rejected_count": 0,
            "rejections_digest": canonical_digest([]),
        },
        "ingest_evidence_set": {
            "status": "unavailable",
            "digest": "b" * 64,
            "record_count": 0,
        },
        "jobs": {"status": "present", "digest": canonical_digest([]), "file_count": 0},
        "events_directory": {
            "status": "missing",
            "members_digest": canonical_digest([]),
            "member_count": 0,
            "absence_digest": "f" * 64,
        },
        "provider_admission_state": {
            "label": "autoplan/provider_admission/provider-admission-v1.latest.json",
            "status": "missing",
            "sha256": None,
            "size": 0,
            "absence_digest": "1" * 64,
        },
        "official_registry": {
            "label": SEALED_OFFICIAL_REGISTRY_RELATIVE_PATH.as_posix(),
            "status": "present",
            "sha256": registry_sha256,
            "size": 10,
            "entry_count": 3,
            "realpath": str(tmp_path / "registry.json"),
            "source_kind": "current_sealed_registry_bytes",
            "standard_index_sha256": registry_sha256,
            "absence_digest": None,
        },
    }
    source_fields = {
        "eligible": False,
        "machine_code": "HOLD_NO_CURRENT_FORMAL_SOURCE",
        "job_id": None,
        "variant_index": None,
        "variant_id": None,
        "result_sha256": None,
        "source_input_receipt_digest": None,
        "checkpoint_digest": None,
        "event_evidence": None,
        "artifact_evidence": None,
    }
    core = {
        "schema_version": "autoplan-no-model-acceptance-v2",
        "run_id": run_id,
        "generated_at": "2026-08-28T08:00:00Z",
        "project_id": "P-1",
        "mode": "no_model_formal_acceptance",
        "decision": "HOLD",
        "model_calls": 0,
        "provider_probes": 0,
        "runtime_state": "verified_healthy_sealed_release",
        "provider_admission": "not_probed",
        "provenance_trust": "local_owner_controlled",
        "cryptographic_attestation": False,
        "release": release,
        "inputs": inputs,
        "formal_source_eligibility": source_fields,
        "stages": _stages(tmp_path, registry_sha256),
        "machine_codes": ["HOLD_NO_CURRENT_FORMAL_SOURCE"],
        "supersedes_receipt_digest": supersedes,
    }
    return {**core, "receipt_digest": canonical_digest(core)}


def _prepared(tmp_path: Path, *, run_id: str) -> _PreparedAcceptance:
    data_root, registry_path, release, current = _collect_fixture(tmp_path)
    prepared = acceptance.collect_acceptance_snapshot(
        project_id="P-1",
        data_root=data_root,
        registry_path=registry_path,
        release_identity=release,
        release_witnesses=[current],
        release_validator=lambda: dict(release),
        run_id=run_id,
        generated_at="2026-08-28T08:00:00Z",
    )
    return _authorize_for_publish(prepared)


def _collect_fixture(
    tmp_path: Path,
    *,
    release_root: Path | None = None,
) -> tuple[Path, Path, dict[str, Any], Any]:
    workspace = tmp_path / "workspace"
    data_root = workspace / "backend" / "data"
    project_root = data_root / "autoplan" / "projects" / "P-1"
    project_root.mkdir(parents=True, exist_ok=True)
    tender_path = project_root / "tender_matrix.json"
    if not tender_path.exists():
        tender_path.write_text(
            json.dumps(
                {
                    "project_name": "离线验收项目",
                    "outline": [],
                    "items": [],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
    boq_path = project_root / "boq_data.json"
    if not boq_path.exists():
        boq_path.write_text(
            json.dumps({"items": [], "stats": {}}, ensure_ascii=False),
            encoding="utf-8",
        )
    audit_path = data_root / "audit" / "ingest.jsonl"
    audit_path.parent.mkdir(parents=True, exist_ok=True)
    if not audit_path.exists():
        audit_path.write_bytes(b"")
    current_path = tmp_path / "current.json"
    if not current_path.exists():
        current_path.write_text('{"release":"current"}', encoding="utf-8")
    current = read_regular_file_snapshot(current_path)
    assert current is not None
    release = _release(tmp_path, current.sha256)
    if release_root is not None:
        release["release_root"] = str(release_root)
        release["release_id"] = release_root.name
    actual_release_root = Path(release["release_root"])
    actual_release_root.mkdir(parents=True, exist_ok=True)
    mutable_kg = tmp_path / "mutable-knowledge-graph"
    mutable_kg.mkdir(parents=True, exist_ok=True)
    mutable_link = actual_release_root / "知识图谱"
    if not (mutable_link.exists() or mutable_link.is_symlink()):
        mutable_link.symlink_to(mutable_kg.resolve(), target_is_directory=True)
    registry_path = sealed_official_registry_path(actual_release_root)
    registry_path.parent.mkdir(parents=True, exist_ok=True)
    if not registry_path.exists():
        registry_path.write_text(
            json.dumps(
                {
                    "standards": [
                        {
                            "standard_code": "GB 55037-2022",
                            "source_name": "建筑防火通用规范",
                            "official_source": "https://ha.119.gov.cn/example",
                            "official_document_url": "https://oss.example/gb55037.pdf",
                            "official_content_sha256": "a" * 64,
                            "effective_status": "active",
                            "current_version": "GB 55037-2022",
                            "latest": True,
                        }
                    ]
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
    registry_path.chmod(0o444)
    registry_sha256 = hashlib.sha256(registry_path.read_bytes()).hexdigest()
    manifest = {
        "schema_version": 1,
        "release_id": release["release_id"],
        "source_digest": release["source_digest"],
        "runtime_digest": release["runtime_digest"],
        "files": [
            {
                "path": SEALED_OFFICIAL_REGISTRY_RELATIVE_PATH.as_posix(),
                "size": registry_path.stat().st_size,
                "mode": 0o444,
                "sha256": registry_sha256,
            }
        ],
        "directories": [
            {
                "path": SEALED_COMPLIANCE_ROOT_RELATIVE_PATH.as_posix(),
                "mode": 0o555,
            }
        ],
        "mutable_links": [
            {
                "path": "知识图谱",
                "target": str(mutable_kg.resolve()),
            }
        ],
    }
    manifest_path = actual_release_root / "release-manifest.json"
    manifest_bytes = json.dumps(
        manifest,
        ensure_ascii=False,
        sort_keys=True,
    ).encode("utf-8")
    if not manifest_path.exists():
        manifest_path.write_bytes(manifest_bytes)
        manifest_path.chmod(0o444)
    else:
        assert not manifest_path.is_symlink()
        assert manifest_path.read_bytes() == manifest_bytes
    release["manifest_digest"] = hashlib.sha256(manifest_bytes).hexdigest()
    return data_root, registry_path, release, current


def test_strict_validator_rejects_minimal_self_digested_mapping(tmp_path: Path) -> None:
    forged_core = {"run_id": "forged", "project_id": "P-1"}
    forged = {**forged_core, "receipt_digest": canonical_digest(forged_core)}
    result = validate_acceptance_receipt(forged)
    assert result["ok"] is False
    assert "receipt_fields_invalid" in result["errors"]


def test_strict_validator_rejects_non_json_and_empty_source_semantics(
    tmp_path: Path,
) -> None:
    prepared = _prepared(tmp_path, run_id="strict-json")
    non_json = json.loads(json.dumps(prepared.receipt))
    non_json["inputs"]["tender"]["size"] = float("nan")
    result = validate_acceptance_receipt(non_json)
    assert result["ok"] is False
    assert "receipt_not_strict_json" in result["errors"]

    empty_source = json.loads(json.dumps(prepared.receipt))
    empty_source["formal_source_eligibility"] = {
        "eligible": True,
        "machine_code": "CURRENT_FORMAL_SOURCE_ELIGIBLE",
        "job_id": "a" * 32,
        "variant_index": 0,
        "variant_id": 1,
        "result_sha256": "b" * 64,
        "source_input_receipt_digest": "c" * 64,
        "checkpoint_digest": "d" * 64,
        "event_evidence": {
            "event_count": 0,
            "provider_attempt_count": 0,
            "successful_chapter_count": 0,
            "event_file_count": 0,
            "event_bundle_digest": "e" * 64,
        },
        "artifact_evidence": {
            "artifact_count": 0,
            "delivery_receipt_digest": "f" * 64,
            "artifact_set_digest": "1" * 64,
        },
    }
    empty_source["machine_codes"] = [
        "CURRENT_FORMAL_SOURCE_ELIGIBLE",
        "NO_CURRENT_FORMAL_SOURCE",
    ]
    core = {
        key: value
        for key, value in empty_source.items()
        if key != "receipt_digest"
    }
    empty_source["receipt_digest"] = canonical_digest(core)
    result = validate_acceptance_receipt(empty_source)
    assert result["ok"] is False
    assert "formal_source_invalid" in result["errors"]


def test_strict_validator_rejects_verified_approval_without_audit_bytes(
    tmp_path: Path,
) -> None:
    prepared = _prepared(tmp_path, run_id="approval-audit-binding")
    forged = prepared.receipt
    forged["inputs"]["approval_audit"].update(
        {
            "status": "missing",
            "sha256": None,
            "size": 0,
            "requested_count": 1,
            "verified_count": 1,
            "rejected_count": 0,
        }
    )
    core = {key: value for key, value in forged.items() if key != "receipt_digest"}
    forged["receipt_digest"] = canonical_digest(core)

    result = validate_acceptance_receipt(forged)

    assert result["ok"] is False
    assert "approval_audit_input_invalid" in result["errors"]


def test_strict_validator_binds_first_machine_code_to_source(
    tmp_path: Path,
) -> None:
    prepared = _prepared(tmp_path, run_id="machine-code-binding")
    forged = json.loads(json.dumps(prepared.receipt))
    forged["machine_codes"] = ["HOLD_DIFFERENT_SOURCE"]
    core = {key: value for key, value in forged.items() if key != "receipt_digest"}
    forged["receipt_digest"] = canonical_digest(core)

    result = validate_acceptance_receipt(forged)

    assert result["ok"] is False
    assert "machine_codes_invalid" in result["errors"]


def test_prepared_capability_is_opaque_and_returns_receipt_copies(
    tmp_path: Path,
) -> None:
    prepared = _prepared(tmp_path, run_id="opaque-capability")
    exposed = prepared.receipt
    exposed["decision"] = "PASS"
    exposed["stages"]["drawing_index"]["ok"] = True

    assert prepared.receipt["decision"] == "HOLD"
    assert prepared.receipt["stages"]["drawing_index"]["ok"] is False
    with pytest.raises(AttributeError):
        prepared.jobs_digest = "0" * 64  # type: ignore[misc]
    with pytest.raises(TypeError):
        _PreparedAcceptance()
    forged = object.__new__(_PreparedAcceptance)
    with pytest.raises(AcceptanceError) as error:
        publish_acceptance_receipt(forged)
    assert error.value.code == "ACCEPTANCE_SNAPSHOT_INVALID"


def test_synthetic_snapshot_is_permanently_unpublishable(
    tmp_path: Path,
) -> None:
    data_root, registry_path, release, current = _collect_fixture(tmp_path)
    prepared = acceptance.collect_acceptance_snapshot(
        project_id="P-1",
        data_root=data_root,
        registry_path=registry_path,
        release_identity=release,
        release_witnesses=[current],
        release_validator=lambda: dict(release),
        run_id="synthetic-never-publish",
        generated_at="2026-08-28T08:00:00Z",
    )
    output_parent = (
        data_root / "autoplan" / "acceptance_receipts" / "no_model_formal"
    )

    with pytest.raises(AcceptanceError) as error:
        publish_acceptance_receipt(prepared)

    assert error.value.code == "ACCEPTANCE_WRITE_ATTESTATION_REQUIRED"
    assert not output_parent.exists()


def test_public_write_path_rejects_before_collection(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    def forbidden_collect(**_kwargs: Any) -> Any:
        raise AssertionError("write rejection must precede collection")

    monkeypatch.setattr(
        acceptance,
        "collect_acceptance_snapshot",
        forbidden_collect,
    )
    with pytest.raises(AcceptanceError) as error:
        acceptance.run_acceptance(
            project_id="P-1",
            data_root=Path("/untrusted/data"),
            registry_path=Path("/untrusted/registry.json"),
            release_identity={},
            release_witnesses=[],
            release_validator=dict,
            write=True,
        )
    assert error.value.code == "ACCEPTANCE_WRITE_REQUIRES_CURRENT_CLI"
    assert set(inspect.signature(acceptance.run_current_runtime_acceptance_write).parameters) == {
        "project_id",
        "run_id",
        "source_job_id",
    }


def test_malformed_fixed_write_context_is_machine_coded(
    tmp_path: Path,
) -> None:
    prepared = _prepared(tmp_path, run_id="malformed-fixed-context")
    _TEST_FIXED_WRITE_CONTEXT["data_root"] = None

    with pytest.raises(AcceptanceError) as error:
        publish_acceptance_receipt(prepared)

    assert error.value.code == "ACCEPTANCE_WRITE_ATTESTATION_CHANGED"


def test_strict_validator_rejects_self_digested_pass_with_failed_stages(
    tmp_path: Path,
) -> None:
    prepared = _prepared(tmp_path, run_id="forged-pass-stages")
    forged = prepared.receipt
    forged["decision"] = "PASS"
    forged["inputs"]["ingest_evidence_set"]["status"] = "verified"
    forged["formal_source_eligibility"] = {
        "eligible": True,
        "machine_code": "CURRENT_FORMAL_SOURCE_ELIGIBLE",
        "job_id": "a" * 32,
        "variant_index": 0,
        "variant_id": 1,
        "result_sha256": "b" * 64,
        "source_input_receipt_digest": "c" * 64,
        "checkpoint_digest": "d" * 64,
        "event_evidence": {
            "event_count": 5,
            "provider_attempt_count": 1,
            "successful_chapter_count": 1,
            "event_file_count": 1,
            "event_bundle_digest": "e" * 64,
        },
        "artifact_evidence": {
            "artifact_count": 13,
            "delivery_receipt_digest": "f" * 64,
            "artifact_set_digest": "1" * 64,
        },
    }
    forged["machine_codes"] = ["CURRENT_FORMAL_SOURCE_ELIGIBLE"]
    forged["stages"]["standard_index"]["ok"] = True
    forged["stages"]["formal_delivery_gate"].update(
        {
            "delivery_allowed": True,
            "blocker_codes": [],
            "blocker_count": 0,
        }
    )
    core = {key: value for key, value in forged.items() if key != "receipt_digest"}
    forged["receipt_digest"] = canonical_digest(core)

    result = validate_acceptance_receipt(forged)

    assert result["ok"] is False
    assert "pass_preconditions_invalid" in result["errors"]


_PROVIDER_ROUTES = [
        {
            "slot": "text_primary",
            "role": "text_draft",
            "provider": "provider-a",
            "model": "model-a",
        },
        {
            "slot": "text_review",
            "role": "text_review",
            "provider": "provider-a",
            "model": "model-review",
        },
        {
            "slot": "document_render",
            "role": "document_render",
            "provider": "provider-render",
            "model": "model-render",
        },
]


def _provider_admission_internal() -> dict[str, Any]:
    slots: list[dict[str, Any]] = []
    admitted_chain: list[dict[str, str]] = []
    for index, route in enumerate(_PROVIDER_ROUTES, start=1):
        credential_fingerprint = hashlib.sha256(
            f"fixture-credential-{index}".encode()
        ).hexdigest()
        identity = {
            "slot": route["slot"],
            "role": route["role"],
            "provider": route["provider"],
            "model": route["model"],
            "credential_fingerprint": credential_fingerprint,
        }
        identity_digest = provider_canonical_digest(
            {
                field: identity[field]
                for field in (
                    "slot",
                    "provider",
                    "model",
                    "credential_fingerprint",
                )
            }
        )
        chain_entry = {**identity, "identity_digest": identity_digest}
        admitted_chain.append(chain_entry)
        layers = {
            name: {"status": "pass", "code": "probe_passed"}
            for name in LAYER_NAMES
        }
        layers["stream"] = {
            "status": "skipped",
            "code": "stream_not_required",
        }
        slots.append(
            {
                **chain_entry,
                "admitted": True,
                "layers": layers,
                "reason_codes": [],
                "checked_at": 1.0,
                "expires_at": 3601.0,
                "stream_required": False,
                "cache_hit": False,
                "probe_duration_ms": 1,
            }
        )
    required_roles = ["text_draft", "text_review", "document_render"]
    decision = decide_provider_required_roles(slots, required_roles)
    core = {
        "schema_version": "provider-admission-v1",
        "generated_at": 1.0,
        "ttl_seconds": 3600.0,
        "required_roles": required_roles,
        "slots": slots,
        "admitted_chain": admitted_chain,
        "role_decision": decision["roles"],
        "missing_roles": decision["missing_roles"],
        "generation_allowed": decision["generation_allowed"],
        "fallback_configured": decision["fallback_configured"],
        "fallback_ready": decision["fallback_ready"],
        "resilience_degraded": decision["resilience_degraded"],
        "degraded": decision["degraded"],
    }
    return {**core, "admission_digest": provider_canonical_digest(core)}


def _reseal_provider_admission(value: dict[str, Any]) -> dict[str, Any]:
    core = {key: item for key, item in value.items() if key != "admission_digest"}
    value["admission_digest"] = provider_canonical_digest(core)
    return value


def _provider_admission_public() -> dict[str, Any]:
    return provider_public_snapshot(_provider_admission_internal())


def _provider_binding_digest() -> str:
    internal = _provider_admission_internal()
    return provider_canonical_digest(
        {
            "schema_version": "provider-admission-binding-v1",
            "required_roles": internal["required_roles"],
            "admitted_route_identities": [
                {
                    field: row[field]
                    for field in (
                        "slot",
                        "role",
                        "provider",
                        "model",
                        "identity_digest",
                    )
                }
                for row in internal["admitted_chain"]
            ],
        }
    )


_PROVIDER_ADMISSION_BINDING_DIGEST = _provider_binding_digest()


def _provider_admission_projection() -> dict[str, Any]:
    admission = _provider_admission_public()
    document_render = next(
        row
        for row in admission["admitted_chain"]
        if row["role"] == "document_render"
    )
    return {
        "public_digest": admission["public_digest"],
        "binding_digest": _PROVIDER_ADMISSION_BINDING_DIGEST,
        "required_roles": admission["required_roles"],
        "admitted_chain": admission["admitted_chain"],
        "document_render": document_render,
    }


def _passing_receipt(tmp_path: Path) -> dict[str, Any]:
    receipt = _receipt(
        tmp_path,
        run_id="strict-pass",
        release=_release(tmp_path, "0" * 64),
    )
    internal_admission = _provider_admission_internal()
    admission = _provider_admission_projection()
    provider_state_sha = "2" * 64
    event_members_digest = "3" * 64
    admitted_identities = [
        {
            field: row[field]
            for field in ("slot", "role", "provider", "model", "identity_digest")
        }
        for row in internal_admission["admitted_chain"]
    ]
    admission_evidence = {
        **admission,
        "durable_snapshot_digest": internal_admission["admission_digest"],
        "durable_file_sha256": provider_state_sha,
        "admitted_route_identities": admitted_identities,
    }
    event_evidence = {
        "event_count": 8,
        "provider_attempt_count": 1,
        "provider_attempts": {"provider-a": 1},
        "successful_chapter_count": 1,
        "event_file_count": 1,
        "event_bundle_digest": "4" * 64,
        "event_directory_members_digest": event_members_digest,
        "provider_admission": admission_evidence,
        "attempt_id": _ATTEMPT_ID,
        "owner_instance_id": _OWNER_INSTANCE_ID,
        "job_revision": _JOB_REVISION,
        "chapter_routes": [
            {
                "chapter_index": 1,
                "slot": "text_primary",
                "provider": "provider-a",
                "model": "model-a",
            }
        ],
    }
    artifact_evidence = {
        "artifact_count": 13,
        "delivery_receipt_digest": "5" * 64,
        "artifact_set_digest": "6" * 64,
        "render_attempt_count": 1,
        "render_provider_attempts": {"provider-render": 1},
    }
    receipt.update(
        {
            "decision": "PASS",
            "formal_source_eligibility": {
                "eligible": True,
                "machine_code": "CURRENT_FORMAL_SOURCE_ELIGIBLE",
                "job_id": "a" * 32,
                "variant_index": 0,
                "variant_id": 1,
                "result_sha256": "7" * 64,
                "source_input_receipt_digest": "8" * 64,
                "checkpoint_digest": "9" * 64,
                "event_evidence": event_evidence,
                "artifact_evidence": artifact_evidence,
            },
            "machine_codes": ["CURRENT_FORMAL_SOURCE_ELIGIBLE"],
        }
    )
    receipt["inputs"]["jobs"] = {
        "status": "present",
        "digest": "a" * 64,
        "file_count": 1,
    }
    receipt["inputs"]["events_directory"] = {
        "status": "present",
        "members_digest": event_members_digest,
        "member_count": 1,
        "absence_digest": None,
    }
    receipt["inputs"]["provider_admission_state"] = {
        "label": "autoplan/provider_admission/provider-admission-v1.latest.json",
        "status": "present",
        "sha256": provider_state_sha,
        "size": 100,
        "absence_digest": None,
    }
    receipt["inputs"]["ingest_evidence_set"].update(
        {"status": "verified", "record_count": 1}
    )
    stages = receipt["stages"]
    stages["drawing_index"].update(
        {
            "ok": True,
            "processed": 1,
            "indexed": 1,
            "text_status": "complete",
            "page_coverage_status": "complete",
            "chapter_binding_status": "complete",
        }
    )
    stages["standard_index"].update(
        {
            "ok": True,
            "indexed": 1,
            "official_verified": 1,
            "chapter_binding_count": 1,
            "chapter_binding_status": "complete",
            "standards": [
                {
                    "identity_status": "identified",
                    "registry_status": "verified_filename_and_cover",
                    "page_anchor_count": 1,
                }
            ],
        }
    )
    stages["project_parameter_evidence"].update(
        {
            "status": "READY",
            "ready": True,
            "coverage_complete": True,
            "evidence_set_receipt_digest": "b" * 64,
            "validation": {"ok": True},
        }
    )
    stages["project_fact_ledger"].update(
        {
            "validation": {"ok": True},
            "formal_parameter_readiness": {
                "ready": True,
                "missing_fields": [],
                "provisional_fields": [],
            },
            "unresolved_fields": [],
        }
    )
    stages["confirmation_checklist"].update(
        {"formal_ready": True, "resolved_fields": ["all"], "blocked_fields": []}
    )
    stages["cross_index"].update(
        {
            "validation": {"ok": True},
            "focus_count": 1,
            "mentioned_count": 1,
            "closed_ok_count": 1,
        }
    )
    stages["formal_delivery_gate"].update(
        {"delivery_allowed": True, "blocker_codes": [], "blocker_count": 0}
    )
    core = {key: value for key, value in receipt.items() if key != "receipt_digest"}
    receipt["receipt_digest"] = canonical_digest(core)
    return receipt


def test_strict_pass_receipt_cross_binds_present_inputs_and_attempt_counts(
    tmp_path: Path,
) -> None:
    valid = _passing_receipt(tmp_path)
    assert validate_acceptance_receipt(valid) == {"ok": True, "errors": []}

    mutations = {
        "jobs_missing": lambda value: value["inputs"]["jobs"].update(
            {"status": "missing", "file_count": 0}
        ),
        "events_missing": lambda value: value["inputs"]["events_directory"].update(
            {
                "status": "missing",
                "members_digest": canonical_digest([]),
                "member_count": 0,
                "absence_digest": "c" * 64,
            }
        ),
        "event_member_count": lambda value: value["inputs"][
            "events_directory"
        ].update({"member_count": 2}),
        "event_file_count": lambda value: value[
            "formal_source_eligibility"
        ]["event_evidence"].update({"event_file_count": 2}),
        "provider_attempt_count": lambda value: value[
            "formal_source_eligibility"
        ]["event_evidence"].update({"provider_attempt_count": 2}),
        "provider_attempts": lambda value: value["formal_source_eligibility"][
            "event_evidence"
        ]["provider_attempts"].update({"provider-a": 2}),
        "render_attempt_count": lambda value: value[
            "formal_source_eligibility"
        ]["artifact_evidence"].update({"render_attempt_count": 2}),
        "render_provider_attempts": lambda value: value[
            "formal_source_eligibility"
        ]["artifact_evidence"]["render_provider_attempts"].update(
            {"provider-render": 2}
        ),
    }
    for name, mutate in mutations.items():
        forged = json.loads(json.dumps(valid))
        mutate(forged)
        core = {
            key: value for key, value in forged.items() if key != "receipt_digest"
        }
        forged["receipt_digest"] = canonical_digest(core)

        result = validate_acceptance_receipt(forged)

        assert result["ok"] is False, name
        assert "pass_preconditions_invalid" in result["errors"], name


_ATTEMPT_ID = "b" * 32
_OWNER_INSTANCE_ID = "c" * 32
_JOB_REVISION = 2


def _job_execution_identity(job_id: str = "a" * 32) -> dict[str, Any]:
    return {
        "job_id": job_id,
        "attempt_id": _ATTEMPT_ID,
        "owner_instance_id": _OWNER_INSTANCE_ID,
        "job_revision": _JOB_REVISION,
    }


def _install_offline_fail_on_call_guards(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import httpx
    import requests

    from backend.zhifei_autoplan.provider_admission import ProviderAdmissionManager
    from backend.zhifei_autoplan.providers.base import BaseProvider
    from backend.zhifei_autoplan.utils.llm_client import LLMClient

    def forbidden(*_args: Any, **_kwargs: Any) -> Any:
        raise AssertionError("offline acceptance attempted an external call")

    provider_types = set(BaseProvider.__subclasses__())
    pending = list(provider_types)
    while pending:
        provider_type = pending.pop()
        children = set(provider_type.__subclasses__())
        provider_types.update(children)
        pending.extend(children)
    monkeypatch.setattr(BaseProvider, "complete", forbidden)
    for provider_type in provider_types:
        monkeypatch.setattr(provider_type, "complete", forbidden)
    monkeypatch.setattr(LLMClient, "_init_provider", forbidden)
    monkeypatch.setattr(ProviderAdmissionManager, "admit", forbidden)
    monkeypatch.setattr(ProviderAdmissionManager, "admit_chain", forbidden)
    monkeypatch.setattr(subprocess, "run", forbidden)
    monkeypatch.setattr(subprocess, "Popen", forbidden)
    monkeypatch.setattr(subprocess, "check_output", forbidden)
    monkeypatch.setattr(socket, "socket", forbidden)
    monkeypatch.setattr(socket, "create_connection", forbidden)
    monkeypatch.setattr(urllib.request, "urlopen", forbidden)
    monkeypatch.setattr(requests.sessions.Session, "request", forbidden)
    monkeypatch.setattr(httpx.Client, "request", forbidden)
    monkeypatch.setattr(httpx.AsyncClient, "request", forbidden)


def _formal_event_chain() -> list[dict[str, Any]]:
    common = {
        **_job_execution_identity(),
        "ts": 1.0,
    }
    admission = _provider_admission_public()
    return [
        {**common, "event": "job_started"},
        {
            **common,
            "event": "provider_admission_started",
            "required_roles": admission["required_roles"],
            "candidate_count": len(admission["admitted_chain"]),
        },
        {
            **common,
            "event": "provider_admission_completed",
            "schema_version": admission["schema_version"],
            "status": admission["status"],
            "required_roles": admission["required_roles"],
            "generation_allowed": admission["generation_allowed"],
            "degraded": admission["degraded"],
            "admitted_chain": admission["admitted_chain"],
            "missing_roles": admission["missing_roles"],
            "public_digest": admission["public_digest"],
            "binding_digest": _PROVIDER_ADMISSION_BINDING_DIGEST,
        },
        {
            **common,
            "event": "provider_attempt_started",
            "variant_id": 1,
            "chapter_index": 1,
            "provider": "provider-a",
            "model": "model-a",
            "slot": "text_primary",
        },
        {
            **common,
            "event": "provider_attempt_finished",
            "variant_id": 1,
            "chapter_index": 1,
            "provider": "provider-a",
            "model": "model-a",
            "slot": "text_primary",
            "ok": True,
        },
        {
            **common,
            "event": "chapter_checkpoint_saved",
            "variant_id": 1,
            "chapter_index": 1,
        },
        {
            **common,
            "event": "chapter_completed",
            "variant_id": 1,
            "chapter_index": 1,
            "ok": True,
        },
        {
            **common,
            "event": "job_succeeded",
            "dry_run": False,
            "delivery_scope": "document",
        },
    ]


def test_formal_event_chain_requires_ordered_attempt_checkpoint_and_completion() -> None:
    sections = [
        {
            "title": "第一章",
            "content": "正文",
            "provider": "provider-a",
            "model": "model-a",
            "model_slot": "text_primary",
        }
    ]
    admitted_routes = _provider_admission_projection()["admitted_chain"]
    evidence = _validate_event_and_provider_chain(
        events=_formal_event_chain(),
        variant_id=1,
        sections=sections,
        admitted_routes=admitted_routes,
        attempt_id=_ATTEMPT_ID,
        owner_instance_id=_OWNER_INSTANCE_ID,
        job_revision=_JOB_REVISION,
    )
    assert evidence["provider_attempt_count"] == 1

    reordered = _formal_event_chain()
    attempt_started = next(
        index
        for index, row in enumerate(reordered)
        if row["event"] == "provider_attempt_started"
    )
    attempt_finished = next(
        index
        for index, row in enumerate(reordered)
        if row["event"] == "provider_attempt_finished"
    )
    reordered[attempt_started], reordered[attempt_finished] = (
        reordered[attempt_finished],
        reordered[attempt_started],
    )
    with pytest.raises(AcceptanceError) as attempt_error:
        _validate_event_and_provider_chain(
            events=reordered,
            variant_id=1,
            sections=sections,
            admitted_routes=admitted_routes,
            attempt_id=_ATTEMPT_ID,
            owner_instance_id=_OWNER_INSTANCE_ID,
            job_revision=_JOB_REVISION,
        )
    assert attempt_error.value.code == "HOLD_SOURCE_PROVIDER_ATTEMPTS_INCOMPLETE"

    checkpoint_after_completion = _formal_event_chain()
    checkpoint_saved = next(
        index
        for index, row in enumerate(checkpoint_after_completion)
        if row["event"] == "chapter_checkpoint_saved"
    )
    chapter_completed = next(
        index
        for index, row in enumerate(checkpoint_after_completion)
        if row["event"] == "chapter_completed"
    )
    checkpoint_after_completion[checkpoint_saved], checkpoint_after_completion[
        chapter_completed
    ] = (
        checkpoint_after_completion[chapter_completed],
        checkpoint_after_completion[checkpoint_saved],
    )
    with pytest.raises(AcceptanceError) as checkpoint_error:
        _validate_event_and_provider_chain(
            events=checkpoint_after_completion,
            variant_id=1,
            sections=sections,
            admitted_routes=admitted_routes,
            attempt_id=_ATTEMPT_ID,
            owner_instance_id=_OWNER_INSTANCE_ID,
            job_revision=_JOB_REVISION,
        )
    assert checkpoint_error.value.code == "HOLD_SOURCE_PROVIDER_ATTEMPTS_INCOMPLETE"

    unknown_variant = _formal_event_chain()
    for row in unknown_variant:
        if row["event"] in {"provider_attempt_started", "provider_attempt_finished"}:
            row["variant_id"] = 2
    with pytest.raises(AcceptanceError) as variant_error:
        _validate_event_and_provider_chain(
            events=unknown_variant,
            variant_id=1,
            allowed_variant_ids={1},
            sections=sections,
            admitted_routes=admitted_routes,
            attempt_id=_ATTEMPT_ID,
            owner_instance_id=_OWNER_INSTANCE_ID,
            job_revision=_JOB_REVISION,
        )
    assert variant_error.value.code == "HOLD_SOURCE_PROVIDER_ATTEMPTS_INCOMPLETE"

    attempt_after_success = _formal_event_chain()
    attempt_started_row = next(
        dict(row)
        for row in attempt_after_success
        if row["event"] == "provider_attempt_started"
    )
    attempt_finished_row = next(
        {**row, "ok": False}
        for row in attempt_after_success
        if row["event"] == "provider_attempt_finished"
    )
    terminal_index = next(
        index
        for index, row in enumerate(attempt_after_success)
        if row["event"] == "job_succeeded"
    )
    attempt_after_success[terminal_index:terminal_index] = [
        attempt_started_row,
        attempt_finished_row,
    ]
    with pytest.raises(AcceptanceError) as late_attempt_error:
        _validate_event_and_provider_chain(
            events=attempt_after_success,
            variant_id=1,
            sections=sections,
            admitted_routes=admitted_routes,
            attempt_id=_ATTEMPT_ID,
            owner_instance_id=_OWNER_INSTANCE_ID,
            job_revision=_JOB_REVISION,
        )
    assert late_attempt_error.value.code == (
        "HOLD_SOURCE_PROVIDER_ATTEMPTS_INCOMPLETE"
    )

    unadmitted_route = _formal_event_chain()
    for row in unadmitted_route:
        if row["event"] in {"provider_attempt_started", "provider_attempt_finished"}:
            row["provider"] = "provider-unadmitted"
            row["model"] = "model-unadmitted"
            row["slot"] = "slot-unadmitted"
    with pytest.raises(AcceptanceError) as route_error:
        _validate_event_and_provider_chain(
            events=unadmitted_route,
            variant_id=1,
            sections=sections,
            admitted_routes=admitted_routes,
            attempt_id=_ATTEMPT_ID,
            owner_instance_id=_OWNER_INSTANCE_ID,
            job_revision=_JOB_REVISION,
        )
    assert route_error.value.code == "HOLD_SOURCE_PROVIDER_ATTEMPTS_INCOMPLETE"


@pytest.mark.parametrize(
    ("field", "value"),
    [
        ("attempt_id", "d" * 32),
        ("owner_instance_id", "e" * 32),
        ("job_revision", _JOB_REVISION + 1),
    ],
)
def test_event_chain_rejects_mixed_terminal_lineage(
    field: str,
    value: Any,
) -> None:
    events = _formal_event_chain()
    events[3][field] = value
    sections = [
        {
            "title": "第一章",
            "content": "正文",
            "provider": "provider-a",
            "model": "model-a",
            "model_slot": "text_primary",
        }
    ]

    with pytest.raises(AcceptanceError) as error:
        _validate_event_and_provider_chain(
            events=events,
            variant_id=1,
            sections=sections,
            admitted_routes=_provider_admission_projection()["admitted_chain"],
            attempt_id=_ATTEMPT_ID,
            owner_instance_id=_OWNER_INSTANCE_ID,
            job_revision=_JOB_REVISION,
        )

    assert error.value.code == "HOLD_SOURCE_EVENTS_INCOMPLETE"


def test_provider_admission_binding_is_recomputed_from_durable_identities() -> None:
    durable = acceptance._validated_internal_provider_admission(
        _provider_admission_internal()
    )
    events = _formal_event_chain()
    validated = acceptance._validated_provider_admission(
        _provider_admission_public(),
        events=events,
        durable=durable,
    )
    assert validated["binding_digest"] == _PROVIDER_ADMISSION_BINDING_DIGEST
    assert validated["admitted_route_identities"]

    forged_events = json.loads(json.dumps(events))
    completed = next(
        row
        for row in forged_events
        if row["event"] == "provider_admission_completed"
    )
    completed["binding_digest"] = "f" * 64
    with pytest.raises(AcceptanceError) as binding_error:
        acceptance._validated_provider_admission(
            _provider_admission_public(),
            events=forged_events,
            durable=durable,
        )
    assert (
        binding_error.value.code
        == "HOLD_SOURCE_PROVIDER_ADMISSION_INCOMPLETE"
    )

    non_boolean = _provider_admission_internal()
    non_boolean["slots"][0]["admitted"] = 1
    core = {
        key: value
        for key, value in non_boolean.items()
        if key != "admission_digest"
    }
    non_boolean["admission_digest"] = provider_canonical_digest(core)
    with pytest.raises(AcceptanceError) as identity_error:
        acceptance._validated_internal_provider_admission(non_boolean)
    assert (
        identity_error.value.code
        == "HOLD_SOURCE_PROVIDER_ADMISSION_INCOMPLETE"
    )


@pytest.mark.parametrize(
    "failed_layer",
    ["configuration", "credentials", "model", "quota", "circuit", "stream"],
)
def test_durable_provider_admission_recomputes_admitted_from_layers(
    failed_layer: str,
) -> None:
    forged = _provider_admission_internal()
    slot = forged["slots"][0]
    if failed_layer == "stream":
        slot["stream_required"] = True
    slot["layers"][failed_layer] = {
        "status": "fail",
        "code": f"{failed_layer}_failed",
    }
    slot["reason_codes"] = [f"{failed_layer}_failed"]
    _reseal_provider_admission(forged)

    with pytest.raises(AcceptanceError) as error:
        acceptance._validated_internal_provider_admission(forged)

    assert error.value.code == "HOLD_SOURCE_PROVIDER_ADMISSION_INCOMPLETE"


@pytest.mark.parametrize(
    ("scope", "field", "value"),
    [
        ("snapshot", "generated_at", "not-a-time"),
        ("snapshot", "generated_at", float("nan")),
        ("snapshot", "generated_at", float("inf")),
        ("snapshot", "ttl_seconds", -1.0),
        ("slot", "checked_at", 4000.0),
        ("slot", "expires_at", 3600.0),
        ("slot", "expires_at", float("-inf")),
    ],
)
def test_durable_provider_admission_rejects_invalid_time_contract(
    scope: str,
    field: str,
    value: Any,
) -> None:
    forged = _provider_admission_internal()
    target = forged if scope == "snapshot" else forged["slots"][0]
    target[field] = value
    try:
        _reseal_provider_admission(forged)
    except ValueError:
        # Non-finite JSON is rejected before a recomputed digest can exist.
        pass

    with pytest.raises(AcceptanceError) as error:
        acceptance._validated_internal_provider_admission(forged)

    assert error.value.code == "HOLD_SOURCE_PROVIDER_ADMISSION_INCOMPLETE"


def test_durable_provider_admission_rejects_reason_code_drift() -> None:
    forged = _provider_admission_internal()
    forged["slots"][0]["reason_codes"] = ["probe_failed"]
    _reseal_provider_admission(forged)

    with pytest.raises(AcceptanceError) as error:
        acceptance._validated_internal_provider_admission(forged)

    assert error.value.code == "HOLD_SOURCE_PROVIDER_ADMISSION_INCOMPLETE"


def _write_minimal_docx(path: Path, *, heading: str, body: str) -> None:
    document = Document()
    document.add_heading(heading, level=1)
    document.add_paragraph(body)
    document.save(path)


def _write_minimal_xlsx(path: Path, *, label: str) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "summary"
    sheet.append(["kind", "value"])
    sheet.append([label, "ok"])
    workbook.save(path)


def _formal_artifact_fixture(
    tmp_path: Path,
    *,
    job_id: str = "a" * 32,
    formal_variant: dict[str, Any] | None = None,
    source_variant: dict[str, Any] | None = None,
) -> tuple[Path, Path, dict[str, Any], Path, dict[str, Any]]:
    release_root = tmp_path / "release"
    workspace_root = tmp_path / "workspace"
    build_root = workspace_root / "build"
    release_root.mkdir()
    build_root.mkdir(parents=True)

    def json_artifact(name: str, payload: dict[str, Any]) -> Path:
        path = build_root / name
        path.write_text(
            json.dumps(payload, ensure_ascii=False, sort_keys=True),
            encoding="utf-8",
        )
        return path

    provider_admission = _provider_admission_public()
    document_route = _provider_admission_projection()["document_render"]
    if formal_variant is None:
        formal_variant = {
            "variant_id": 1,
            "sections": [
                {
                    "title": "第一章",
                    "content": "正式正文",
                    "original_content": "正式正文",
                    "provider": "provider-a",
                    "model": "model-a",
                    "model_slot": "text_primary",
                    "professional_render": {"status": "refined"},
                }
            ],
            "source_input_receipt": {},
            "generation_release_identity": {},
            "model_routing": {
                "provider_admission": provider_admission,
                "document_render": document_route,
            },
            "professional_render": {
                "provider": document_route["provider"],
                "model_id": document_route["model"],
            },
        }
    if source_variant is None:
        source_variant = json.loads(json.dumps(formal_variant))

    source_json = json_artifact(
        f"actions_{job_id}.json",
        {"variants": [source_variant]},
    )
    source = build_root / "source.docx"
    professional = build_root / "professional.docx"
    compare = build_root / "compare.docx"
    expert = build_root / "expert.docx"
    focus = build_root / "focus.xlsx"
    score = build_root / "score.xlsx"
    _write_minimal_docx(source, heading="源文档", body="生成正文")
    _write_minimal_docx(professional, heading="专业文档", body="专业终稿正文")
    _write_minimal_docx(compare, heading="对比文档", body="修订对比")
    _write_minimal_docx(expert, heading="专家复核", body="复核通过")
    _write_minimal_xlsx(focus, label="重点清单")
    _write_minimal_xlsx(score, label="评分证据")

    professional_json = json_artifact(
        "professional.json",
        {
            "variants": [formal_variant],
            "professional_render_source_variant": 1,
        },
    )
    professional_sha256 = hashlib.sha256(professional.read_bytes()).hexdigest()
    source_sha256 = hashlib.sha256(source.read_bytes()).hexdigest()
    source_json_sha256 = hashlib.sha256(source_json.read_bytes()).hexdigest()
    professional_json_sha256 = hashlib.sha256(
        professional_json.read_bytes()
    ).hexdigest()

    figure_core = {
        "schema_version": "docx_figure_delivery.v2",
        "figure_count": 0,
        "source_media": {},
        "insertions": [],
        "insertion_failures": [],
        "embedded_media_verification": {"ok": True},
        "issues": [],
        "warnings": [],
        "delivery_allowed": True,
        "status": "pass",
    }
    figure_payload = {
        **figure_core,
        "decision_digest": canonical_digest(figure_core),
    }
    (build_root / "professional.figure_manifest.json").write_text(
        json.dumps(figure_payload),
        encoding="utf-8",
    )

    structural = build_root / "professional.structural_quality.json"
    structural_core = {
        "schema": "zhifei.docx_structural_quality.v1",
        "created_at": "2026-08-28T08:00:00Z",
        "status": "pass",
        "docx": str(professional),
        "docx_sha256": professional_sha256,
        "visible_chars": 6,
        "paragraph_count": 2,
        "heading_count": 1,
        "table_count": 0,
        "section_metrics": [
            {
                "section": 1,
                "width_cm": 21.0,
                "height_cm": 29.7,
                "margins_cm": {
                    "top": 2.5,
                    "right": 2.0,
                    "bottom": 2.0,
                    "left": 2.0,
                },
                "orientation": "portrait",
            }
        ],
        "section_story_references": [
            {
                "section": 1,
                "header_types": ["default"],
                "footer_types": ["default"],
                "default_header": True,
                "default_footer": True,
            }
        ],
        "body_style": {
            "font": "宋体",
            "size_pt": 14.0,
            "line_spacing_pt": 28.0,
            "first_line_chars": "200",
            "space_before_twips": "0",
            "space_after_twips": "0",
        },
        "word_fields": {
            "toc": True,
            "page": True,
            "numpages": True,
            "update_on_open": True,
        },
        "figure_delivery": {
            "delivery_allowed": True,
            "figure_count": 0,
            "decision_digest": figure_payload["decision_digest"],
        },
        "package_integrity": {
            "invalid_xml": [],
            "duplicate_relationship_ids": [],
            "dangling_relationships": [],
            "duplicate_bookmark_ids": [],
            "duplicate_bookmark_names": [],
            "custom_parts": [],
        },
        "hard_failures": [],
        "warnings": [],
    }
    structural_material = {
        key: value
        for key, value in structural_core.items()
        if key not in {"created_at", "docx"}
    }
    structural.write_text(
        json.dumps(
            {
                **structural_core,
                "decision_digest": canonical_digest(structural_material),
            }
        ),
        encoding="utf-8",
    )

    visual = build_root / "professional.visual_quality.json"
    visual_pdf = build_root / "professional.pdf"
    visual_pdf.write_bytes(b"%PDF-1.4\n1 0 obj<</Type/Catalog>>endobj\n%%EOF\n")
    visual_pdf_sha256 = hashlib.sha256(visual_pdf.read_bytes()).hexdigest()
    preview_dir = build_root / "professional-preview"
    preview_dir.mkdir()
    visual_core = {
        "schema": "zhifei.docx_visual_quality.v1",
        "created_at": "2026-08-28T08:00:00Z",
        "status": "pass",
        "docx": str(professional),
        "docx_sha256": professional_sha256,
        "pdf": str(visual_pdf),
        "pdf_sha256": visual_pdf_sha256,
        "preview_dir": str(preview_dir),
        "receipt": str(visual),
        "page_count": 1,
        "page_metrics": [
            {
                "page": 1,
                "text_chars": 120,
                "ink_ratio": 0.08,
                "edge_ink_ratio": 0.001,
                "pixel_width": 1240,
                "pixel_height": 1754,
                "blank": False,
                "sparse": False,
                "orphan_heading": False,
                "edge_clipping_risk": False,
            }
        ],
        "blank_pages": [],
        "sparse_pages": [],
        "sparse_page_budget": 1,
        "orphan_heading_pages": [],
        "edge_clipping_risk_pages": [],
        "sparse_page_streaks": [],
        "page_geometry_outliers": [],
        "cjk_glyph_integrity": {
            "status": "pass",
            "inspected_glyphs": 2,
            "empty_glyphs": 0,
            "empty_glyph_ratio": 0.0,
            "unique_cjk_characters": 2,
            "unique_glyph_shapes": 2,
            "shape_retention": 1.0,
            "largest_shape_collision": 1,
            "hard_failures": [],
        },
        "hard_failures": [],
        "warnings": [],
    }
    visual_material = {
        key: value
        for key, value in visual_core.items()
        if key not in {"created_at", "docx", "pdf", "preview_dir", "receipt"}
    }
    visual.write_text(
        json.dumps(
            {
                **visual_core,
                "decision_digest": canonical_digest(visual_material),
            }
        ),
        encoding="utf-8",
    )
    structural_payload = json.loads(structural.read_text(encoding="utf-8"))
    visual_payload = json.loads(visual.read_text(encoding="utf-8"))
    structural_sha256 = hashlib.sha256(structural.read_bytes()).hexdigest()
    visual_sha256 = hashlib.sha256(visual.read_bytes()).hexdigest()

    render_attempt_core = {
        "schema_version": "document-render-attempt-evidence-v1",
        "execution_control_schema_version": "execution-control-v1",
        "role": "document_render",
        "slot": document_route["slot"],
        "provider": document_route["provider"],
        "model": document_route["model"],
        "job_id": job_id,
        "variant": 1,
        "model_attempts_before": 0,
        "model_attempts_after": 1,
        "attempt_count": 1,
        "provider_attempts_before": 0,
        "provider_attempts_after": 1,
        "provider_attempt_count": 1,
    }
    render_attempt = {
        **render_attempt_core,
        "evidence_digest": canonical_digest(render_attempt_core),
    }
    render_receipt = build_root / "professional.receipt.json"
    source_sections = source_variant.get("sections") or []
    formal_sections = formal_variant.get("sections") or []
    render_core = {
        "schema": "zhifei.professional_document_render.v1",
        "job_id": job_id,
        "variant": 1,
        "role": "document_render",
        "slot": document_route["slot"],
        "provider": document_route["provider"],
        "model_id": document_route["model"],
        "source_json": str(source_json),
        "source_json_sha256": source_json_sha256,
        "source_docx": str(source),
        "source_docx_sha256": source_sha256,
        "professional_docx": str(professional),
        "professional_json": str(professional_json),
        "professional_json_sha256": professional_json_sha256,
        "section_count": len(formal_sections),
        "source_char_count": sum(
            len(str(section.get("content") or ""))
            for section in source_sections
            if isinstance(section, dict)
        ),
        "professional_char_count": sum(
            len(str(section.get("content") or ""))
            for section in formal_sections
            if isinstance(section, dict)
        ),
        "professional_docx_sha256": professional_sha256,
        "render_attempt_evidence": render_attempt,
        "quality_gate": {
            "original_preserved": True,
            "titles_preserved": True,
            "evidence_not_reduced": True,
            "tender_style_fields_preserved": True,
            "export_succeeded": True,
            "structural_quality_passed": True,
            "visual_page_quality_passed": True,
            "no_blank_pages": True,
            "no_orphan_headings": True,
        },
        "structural_quality": {
            "receipt": str(structural),
            "receipt_sha256": structural_sha256,
            "status": "pass",
            "hard_failures": [],
            "docx_sha256": professional_sha256,
            "decision_digest": structural_payload["decision_digest"],
            "heading_count": structural_payload["heading_count"],
            "table_count": structural_payload["table_count"],
            "word_fields": structural_payload["word_fields"],
            "body_style": structural_payload["body_style"],
            "section_metrics": structural_payload["section_metrics"],
            "figure_delivery": structural_payload["figure_delivery"],
        },
        "visual_quality": {
            "receipt": str(visual),
            "receipt_sha256": visual_sha256,
            "status": "pass",
            "hard_failures": [],
            "docx_sha256": professional_sha256,
            "pdf_sha256": visual_pdf_sha256,
            "decision_digest": visual_payload["decision_digest"],
            "page_count": visual_payload["page_count"],
            "blank_pages": visual_payload["blank_pages"],
            "sparse_pages": visual_payload["sparse_pages"],
            "orphan_heading_pages": visual_payload["orphan_heading_pages"],
            "edge_clipping_risk_pages": visual_payload[
                "edge_clipping_risk_pages"
            ],
            "pdf": str(visual_pdf),
            "preview_dir": str(preview_dir),
        },
    }
    render_receipt.write_text(
        json.dumps(
            {
                **render_core,
                "receipt_digest": canonical_digest(render_core),
            }
        ),
        encoding="utf-8",
    )
    sealed = build_delivery_receipt(
        job_id=job_id,
        job_execution_identity=_job_execution_identity(job_id),
        source_docx=[source],
        professional_docx=[professional],
        professional_json=[professional_json],
        professional_receipts=[render_receipt],
        compare_docx=[compare],
        focus_xlsx=[focus],
        score_overview_xlsx=[score],
        expert_review_docx=[expert],
        receipt_path=build_root / "delivery-receipt.json",
    )
    result = {
        "json": str(source_json),
        "source_docx": [str(source)],
        "docx": [str(professional)],
        "professional_docx": [str(professional)],
        "professional_json": [str(professional_json)],
        "professional_render_receipt": [str(render_receipt)],
        "compare_docx": [str(compare)],
        "focus_xlsx": [str(focus)],
        "score_overview_xlsx": [str(score)],
        "expert_review_docx": [str(expert)],
        "delivery_receipt": str(sealed["receipt"]),
        "delivery_decision_digest": sealed["decision_digest"],
    }
    return release_root, workspace_root, result, professional, formal_variant


def test_artifact_bundle_rejects_bytes_changed_after_delivery_receipt(
    tmp_path: Path,
) -> None:
    job_id = "a" * 32
    (
        release_root,
        workspace_root,
        result,
        professional,
        formal_variant,
    ) = _formal_artifact_fixture(
        tmp_path,
        job_id=job_id,
    )
    witnesses, evidence, formal_variants = _artifact_bundle(
        result=result,
        release_root=release_root,
        workspace_root=workspace_root,
        expected_variants=[formal_variant],
        job_id=job_id,
        provider_admission=_provider_admission_projection(),
        job_execution_identity=_job_execution_identity(job_id),
    )
    assert witnesses
    assert evidence["delivery_receipt_digest"] == result[
        "delivery_decision_digest"
    ]
    assert formal_variants == [formal_variant]

    professional.write_bytes(b"tampered-professional")
    with pytest.raises(AcceptanceError) as error:
        _artifact_bundle(
            result=result,
            release_root=release_root,
            workspace_root=workspace_root,
            expected_variants=[formal_variant],
            job_id=job_id,
            provider_admission=_provider_admission_projection(),
            job_execution_identity=_job_execution_identity(job_id),
        )
    assert error.value.code == "HOLD_SOURCE_OUTPUT_UNTRUSTED"


def test_artifact_bundle_rejects_visual_pdf_changed_after_quality_receipt(
    tmp_path: Path,
) -> None:
    job_id = "a" * 32
    (
        release_root,
        workspace_root,
        result,
        _professional,
        formal_variant,
    ) = _formal_artifact_fixture(tmp_path, job_id=job_id)
    _artifact_bundle(
        result=result,
        release_root=release_root,
        workspace_root=workspace_root,
        expected_variants=[formal_variant],
        job_id=job_id,
        provider_admission=_provider_admission_projection(),
        job_execution_identity=_job_execution_identity(job_id),
    )

    (workspace_root / "build" / "professional.pdf").write_bytes(
        b"%PDF-1.4\ntampered\n%%EOF\n"
    )

    with pytest.raises(AcceptanceError) as error:
        _artifact_bundle(
            result=result,
            release_root=release_root,
            workspace_root=workspace_root,
            expected_variants=[formal_variant],
            job_id=job_id,
            provider_admission=_provider_admission_projection(),
            job_execution_identity=_job_execution_identity(job_id),
        )
    assert error.value.code == "HOLD_SOURCE_OUTPUT_UNTRUSTED"


def test_artifact_bundle_rejects_delivery_execution_identity_tamper(
    tmp_path: Path,
) -> None:
    job_id = "a" * 32
    release_root, workspace_root, result, _professional, formal_variant = (
        _formal_artifact_fixture(tmp_path, job_id=job_id)
    )
    receipt_path = Path(result["delivery_receipt"])
    receipt = json.loads(receipt_path.read_text(encoding="utf-8"))
    receipt["job_execution_identity"]["attempt_id"] = "d" * 32
    receipt["decision_digest"] = canonical_delivery_receipt_digest(receipt)
    receipt_path.write_text(json.dumps(receipt), encoding="utf-8")
    result["delivery_decision_digest"] = receipt["decision_digest"]

    with pytest.raises(AcceptanceError) as error:
        _artifact_bundle(
            result=result,
            release_root=release_root,
            workspace_root=workspace_root,
            expected_variants=[formal_variant],
            job_id=job_id,
            provider_admission=_provider_admission_projection(),
            job_execution_identity=_job_execution_identity(job_id),
        )

    assert error.value.code == "HOLD_SOURCE_OUTPUT_UNTRUSTED"


def test_artifact_bundle_rejects_final_section_route_drift(
    tmp_path: Path,
) -> None:
    source_variant = {
        "variant_id": 1,
        "sections": [
            {
                "title": "第一章",
                "content": "源正文",
                "provider": "provider-a",
                "model": "model-a",
                "model_slot": "text_primary",
            }
        ],
        "source_input_receipt": {},
        "generation_release_identity": {},
        "model_routing": {
            "provider_admission": _provider_admission_public(),
            "document_render": _provider_admission_projection()[
                "document_render"
            ],
        },
    }
    formal_variant = json.loads(json.dumps(source_variant))
    formal_variant["sections"] = [
        {
            **source_variant["sections"][0],
            "content": "专业正文",
            "original_content": "源正文",
            "provider": "provider-drift",
            "professional_render": {"status": "refined"},
        }
    ]
    formal_variant["professional_render"] = {
        "provider": "provider-render",
        "model_id": "model-render",
    }
    release_root, workspace_root, result, _professional, _variant = (
        _formal_artifact_fixture(
            tmp_path,
            formal_variant=formal_variant,
            source_variant=source_variant,
        )
    )

    with pytest.raises(AcceptanceError) as error:
        _artifact_bundle(
            result=result,
            release_root=release_root,
            workspace_root=workspace_root,
            expected_variants=[source_variant],
            job_id="a" * 32,
            provider_admission=_provider_admission_projection(),
            job_execution_identity=_job_execution_identity(),
        )

    assert error.value.code == "HOLD_SOURCE_OUTPUT_UNTRUSTED"


@pytest.mark.parametrize("symlink_level", ["build", "nested"])
def test_trusted_artifact_path_rejects_each_symlink_component(
    tmp_path: Path,
    symlink_level: str,
) -> None:
    release_root = tmp_path / "release"
    workspace_root = tmp_path / "workspace"
    release_root.mkdir()
    workspace_root.mkdir()
    external = tmp_path / "external"
    external.mkdir()
    (external / "artifact.json").write_text("{}", encoding="utf-8")
    if symlink_level == "build":
        (workspace_root / "build").symlink_to(
            external,
            target_is_directory=True,
        )
        artifact = workspace_root / "build" / "artifact.json"
    else:
        build_root = workspace_root / "build"
        build_root.mkdir()
        (build_root / "nested").symlink_to(
            external,
            target_is_directory=True,
        )
        artifact = build_root / "nested" / "artifact.json"

    with pytest.raises(AcceptanceError) as error:
        acceptance._trusted_build_snapshot(
            artifact,
            release_root=release_root,
            workspace_root=workspace_root,
        )

    assert error.value.code == "HOLD_SOURCE_OUTPUT_UNTRUSTED"


def test_artifact_bundle_rejects_symlinked_preview_directory(
    tmp_path: Path,
) -> None:
    release_root, workspace_root, result, _professional, formal_variant = (
        _formal_artifact_fixture(tmp_path)
    )
    preview = workspace_root / "build" / "professional-preview"
    preview.rmdir()
    external = tmp_path / "external-preview"
    external.mkdir()
    preview.symlink_to(external, target_is_directory=True)

    with pytest.raises(AcceptanceError) as error:
        _artifact_bundle(
            result=result,
            release_root=release_root,
            workspace_root=workspace_root,
            expected_variants=[formal_variant],
            job_id="a" * 32,
            provider_admission=_provider_admission_projection(),
            job_execution_identity=_job_execution_identity(),
        )

    assert error.value.code == "HOLD_SOURCE_OUTPUT_UNTRUSTED"


def test_artifact_bundle_revalidates_quality_receipt_semantics(
    tmp_path: Path,
) -> None:
    job_id = "a" * 32
    (
        release_root,
        workspace_root,
        result,
        _professional,
        formal_variant,
    ) = _formal_artifact_fixture(
        tmp_path,
        job_id=job_id,
    )
    render_path = Path(result["professional_render_receipt"][0])
    render = json.loads(render_path.read_text(encoding="utf-8"))
    render["quality_gate"]["no_blank_pages"] = False
    render_path.write_text(json.dumps(render), encoding="utf-8")

    receipt_path = Path(result["delivery_receipt"])
    receipt = json.loads(receipt_path.read_text(encoding="utf-8"))
    render_artifact = receipt["variants"][0]["professional_render_receipt"]
    render_bytes = render_path.read_bytes()
    render_artifact["size"] = len(render_bytes)
    render_artifact["sha256"] = hashlib.sha256(render_bytes).hexdigest()
    receipt["decision_digest"] = canonical_delivery_receipt_digest(receipt)
    receipt_path.write_text(json.dumps(receipt), encoding="utf-8")
    result["delivery_decision_digest"] = receipt["decision_digest"]

    with pytest.raises(AcceptanceError) as error:
        _artifact_bundle(
            result=result,
            release_root=release_root,
            workspace_root=workspace_root,
            expected_variants=[formal_variant],
            job_id=job_id,
            provider_admission=_provider_admission_projection(),
            job_execution_identity=_job_execution_identity(job_id),
        )
    assert error.value.code == "HOLD_SOURCE_OUTPUT_UNTRUSTED"


def test_quality_receipts_require_full_production_nested_schema(
    tmp_path: Path,
) -> None:
    (
        _release_root,
        workspace_root,
        _result,
        professional,
        _formal_variant,
    ) = _formal_artifact_fixture(tmp_path)
    build_root = workspace_root / "build"
    structural_path = build_root / "professional.structural_quality.json"
    visual_path = build_root / "professional.visual_quality.json"
    pdf_path = build_root / "professional.pdf"
    figure_path = build_root / "professional.figure_manifest.json"
    structural = json.loads(structural_path.read_text(encoding="utf-8"))
    visual = json.loads(visual_path.read_text(encoding="utf-8"))
    figure = json.loads(figure_path.read_text(encoding="utf-8"))
    docx_sha256 = hashlib.sha256(professional.read_bytes()).hexdigest()

    assert acceptance._structural_quality_receipt_valid(
        structural,
        docx_sha256=docx_sha256,
        docx_path=professional.resolve(),
        figure_manifest=figure,
    )
    assert acceptance._visual_quality_receipt_valid(
        visual,
        docx_sha256=docx_sha256,
        docx_path=professional.resolve(),
        pdf_path=pdf_path.resolve(),
        receipt_path=visual_path.resolve(),
        preview_dir_path=(build_root / "professional-preview").resolve(),
    )

    structural["section_metrics"] = []
    structural["section_story_references"] = []
    structural_material = {
        key: value
        for key, value in structural.items()
        if key not in {"created_at", "decision_digest", "docx", "receipt"}
    }
    structural["decision_digest"] = canonical_digest(structural_material)
    assert not acceptance._structural_quality_receipt_valid(
        structural,
        docx_sha256=docx_sha256,
        docx_path=professional.resolve(),
        figure_manifest=figure,
    )

    visual["page_metrics"] = [{"page": 1}]
    visual_material = {
        key: value
        for key, value in visual.items()
        if key
        not in {
            "created_at",
            "decision_digest",
            "docx",
            "pdf",
            "preview_dir",
            "receipt",
        }
    }
    visual["decision_digest"] = canonical_digest(visual_material)
    assert not acceptance._visual_quality_receipt_valid(
        visual,
        docx_sha256=docx_sha256,
        docx_path=professional.resolve(),
        pdf_path=pdf_path.resolve(),
        receipt_path=visual_path.resolve(),
        preview_dir_path=(build_root / "professional-preview").resolve(),
    )


def test_artifact_bundle_rejects_visual_receipt_without_explicit_docx_hash(
    tmp_path: Path,
) -> None:
    job_id = "a" * 32
    (
        release_root,
        workspace_root,
        result,
        _professional,
        formal_variant,
    ) = _formal_artifact_fixture(tmp_path, job_id=job_id)
    receipt_path = Path(result["delivery_receipt"])
    receipt = json.loads(receipt_path.read_text(encoding="utf-8"))
    visual_artifact = receipt["variants"][0]["visual_quality_receipt"]
    visual_path = Path(visual_artifact["path"])
    visual = json.loads(visual_path.read_text(encoding="utf-8"))
    visual.pop("docx_sha256")
    visual_path.write_text(json.dumps(visual), encoding="utf-8")
    visual_bytes = visual_path.read_bytes()
    visual_artifact["size"] = len(visual_bytes)
    visual_artifact["sha256"] = hashlib.sha256(visual_bytes).hexdigest()
    receipt["decision_digest"] = canonical_delivery_receipt_digest(receipt)
    receipt_path.write_text(json.dumps(receipt), encoding="utf-8")
    result["delivery_decision_digest"] = receipt["decision_digest"]

    with pytest.raises(AcceptanceError) as error:
        _artifact_bundle(
            result=result,
            release_root=release_root,
            workspace_root=workspace_root,
            expected_variants=[formal_variant],
            job_id=job_id,
            provider_admission=_provider_admission_projection(),
            job_execution_identity=_job_execution_identity(job_id),
        )

    assert error.value.code == "HOLD_SOURCE_OUTPUT_UNTRUSTED"


@pytest.mark.parametrize("malformed_count", ["1", 1.0, True, [], {}])
def test_artifact_bundle_rejects_malformed_variant_count_without_escaping(
    tmp_path: Path,
    malformed_count: Any,
) -> None:
    job_id = "a" * 32
    (
        release_root,
        workspace_root,
        result,
        _professional,
        formal_variant,
    ) = _formal_artifact_fixture(tmp_path, job_id=job_id)
    receipt_path = Path(result["delivery_receipt"])
    receipt = json.loads(receipt_path.read_text(encoding="utf-8"))
    receipt["variant_count"] = malformed_count
    receipt["decision_digest"] = canonical_delivery_receipt_digest(receipt)
    receipt_path.write_text(json.dumps(receipt), encoding="utf-8")
    result["delivery_decision_digest"] = receipt["decision_digest"]

    with pytest.raises(AcceptanceError) as error:
        _artifact_bundle(
            result=result,
            release_root=release_root,
            workspace_root=workspace_root,
            expected_variants=[formal_variant],
            job_id=job_id,
            provider_admission=_provider_admission_projection(),
            job_execution_identity=_job_execution_identity(job_id),
        )

    assert error.value.code == "HOLD_SOURCE_OUTPUT_UNTRUSTED"


@pytest.mark.parametrize("timestamp", [float("nan"), float("inf"), float("-inf")])
def test_event_bundle_rejects_non_finite_timestamp(
    tmp_path: Path,
    timestamp: float,
) -> None:
    job_id = "a" * 32
    events_dir = tmp_path / "events"
    events_dir.mkdir()
    (events_dir / f"{job_id}.jsonl").write_text(
        json.dumps({"job_id": job_id, "event": "job_started", "ts": timestamp})
        + "\n",
        encoding="utf-8",
    )

    with pytest.raises(AcceptanceError) as error:
        acceptance._event_bundle(events_dir=events_dir, job_id=job_id)

    assert error.value.code == "HOLD_SOURCE_EVENTS_INCOMPLETE"


def test_event_bundle_rejects_non_utf8_and_non_monotonic_fragments(
    tmp_path: Path,
) -> None:
    job_id = "a" * 32
    events_dir = tmp_path / "events"
    events_dir.mkdir()
    first = events_dir / f"{job_id}.1.jsonl"
    second = events_dir / f"{job_id}.2.jsonl"
    first.write_bytes(b"\xff")
    second.write_text(
        json.dumps({"job_id": job_id, "event": "job_succeeded", "ts": 2.0})
        + "\n",
        encoding="utf-8",
    )
    with pytest.raises(AcceptanceError) as encoding_error:
        acceptance._event_bundle(events_dir=events_dir, job_id=job_id)
    assert encoding_error.value.code == "HOLD_SOURCE_EVENTS_INCOMPLETE"

    first.write_text(
        json.dumps({"job_id": job_id, "event": "job_started", "ts": 2.0})
        + "\n",
        encoding="utf-8",
    )
    second.write_text(
        json.dumps({"job_id": job_id, "event": "job_succeeded", "ts": 1.0})
        + "\n",
        encoding="utf-8",
    )
    with pytest.raises(AcceptanceError) as chronology_error:
        acceptance._event_bundle(events_dir=events_dir, job_id=job_id)
    assert chronology_error.value.code == "HOLD_SOURCE_EVENTS_INCOMPLETE"


def test_candidate_source_accepts_complete_job_event_checkpoint_artifact_chain(
    tmp_path: Path,
) -> None:
    job_id = "a" * 32
    project_id = "P-FORMAL-SOURCE"
    tender = {"project_name": "项目"}
    boq = {"items": []}
    sections = [
        {
            "title": "第一章",
            "content": "生成正文",
            "provider": "provider-a",
            "model": "model-a",
            "model_slot": "text_primary",
        }
    ]
    professional_sections = [
        {
            "title": "第一章",
            "content": "专业终稿正文",
            "original_content": "生成正文",
            "provider": "provider-a",
            "model": "model-a",
            "model_slot": "text_primary",
            "professional_render": {"status": "refined"},
        }
    ]
    outline = ["第一章"]
    provider_admission = _provider_admission_public()
    provider_projection = _provider_admission_projection()
    document_route = provider_projection["document_render"]
    release_root = tmp_path / "release"
    release_identity = {
        "system_id": "docgen-system",
        "release_id": release_root.name,
        "manifest_digest": "1" * 64,
        "source_digest": "2" * 64,
        "runtime_digest": "3" * 64,
        "release_root": str(release_root),
    }
    generation_release = {
        "schema_version": "autoplan-generation-release-v1",
        **release_identity,
        "runtime_mode": "sealed_release",
        "release_managed": True,
    }
    input_core = {
        "schema_version": "autoplan-source-input-v1",
        "project_id": project_id,
        "tender_digest": canonical_digest(tender),
        "boq_digest": canonical_digest(boq),
    }
    gate_core = {"delivery_allowed": True, "blockers": []}
    variant = {
        "variant_id": 1,
        "delivery_scope": "document",
        "delivery_ready": True,
        "dry_run": False,
        "sections": sections,
        "outline": outline,
        "model_routing": {
            "provider_admission": provider_admission,
            "document_render": document_route,
        },
        "delivery_quality_gate": {
            **gate_core,
            "decision_digest": canonical_digest(gate_core),
        },
        "source_input_receipt": {
            **input_core,
            "receipt_digest": canonical_digest(input_core),
        },
        "generation_release_identity": generation_release,
    }
    formal_variant = json.loads(json.dumps(variant))
    formal_variant["sections"] = professional_sections
    formal_variant["professional_render"] = {
        "provider": document_route["provider"],
        "model_id": document_route["model"],
    }
    (
        release_root,
        workspace_root,
        result,
        _professional,
        _formal_variant,
    ) = _formal_artifact_fixture(
        tmp_path,
        job_id=job_id,
        formal_variant=formal_variant,
        source_variant=variant,
    )
    data_root = workspace_root / "backend" / "data" / "autoplan"
    jobs_dir = data_root / "jobs"
    events_dir = data_root / "events"
    checkpoint_dir = data_root / "checkpoints" / job_id
    for directory in (jobs_dir, events_dir, checkpoint_dir):
        directory.mkdir(parents=True, exist_ok=True)
    provider_admission_dir = data_root / "provider_admission"
    provider_admission_dir.mkdir(parents=True)
    provider_admission_path = write_provider_snapshot(
        _provider_admission_internal(),
        root=provider_admission_dir,
    )
    provider_admission_snapshot = read_regular_file_snapshot(
        provider_admission_path
    )
    assert provider_admission_snapshot is not None

    def current_events_state() -> acceptance.DirectoryStateWitness:
        return acceptance._capture_directory_state(
            events_dir,
            code="HOLD_SOURCE_EVENTS_INCOMPLETE",
        )
    output_path = Path(result["json"])
    result.update(
        {
            "delivery_ready": True,
            "validation_scope": "document",
            "delivery_profile": "sonnet5_professional_word",
            "generation_release_identity": generation_release,
            "job_execution_identity": _job_execution_identity(job_id),
        }
    )

    event_path = events_dir / f"{job_id}.jsonl"
    event_path.write_text(
        "".join(
            json.dumps(row, ensure_ascii=False) + "\n"
            for row in [
                {
                    **event,
                    **(
                        {"generation_release_identity": generation_release}
                        if event.get("event") == "job_started"
                        else {}
                    ),
                }
                for event in _formal_event_chain()
            ]
        ),
        encoding="utf-8",
    )
    binding_core = {
        "schema_version": "generation-checkpoint-v3",
        **_job_execution_identity(job_id),
        "topic": "项目施工组织设计",
        "project_id": project_id,
        "project_type": "房屋建筑工程",
        "style": {"tone": "professional"},
        "chapter_pages": {"第一章": 1},
        "delivery_scope": "document",
        "variant_id": "1",
        "outline": outline,
        "project_fact_digest": "4" * 64,
        "requirement_plan_digest": "5" * 64,
        "provider_admission_digest": _PROVIDER_ADMISSION_BINDING_DIGEST,
        "provider_routes": [
            {
                "slot": row["slot"],
                "provider": row["provider"],
                "model": row["model"],
            }
            for row in provider_admission["admitted_chain"]
            if row["role"] != "document_render"
        ],
        "prompt_contract_digest": "6" * 64,
    }
    binding = {
        **binding_core,
        "binding_digest": canonical_digest(binding_core),
    }
    section_core = {
        "chapter_index": 0,
        "chapter_title": "第一章",
        "chapter_context_digest": "7" * 64,
        "saved_at": 1.0,
        "result": sections[0],
    }
    checkpoint_core = {
        "schema_version": "generation-checkpoint-v3",
        "status": "complete",
        "binding": binding,
        "binding_digest": binding["binding_digest"],
        "sections": {
            "0": {
                **section_core,
                "section_digest": canonical_digest(section_core),
            }
        },
    }
    checkpoint_path = checkpoint_dir / "variant-1.json"
    checkpoint_path.write_text(
        json.dumps(
            {
                **checkpoint_core,
                "integrity_digest": canonical_digest(checkpoint_core),
            }
        ),
        encoding="utf-8",
    )
    job = {
        "job_id": job_id,
        "status": "succeeded",
        "attempt_id": None,
        "owner_instance_id": None,
        "error": None,
        "last_attempt_id": _ATTEMPT_ID,
        "last_owner_instance_id": _OWNER_INSTANCE_ID,
        "last_job_revision": _JOB_REVISION,
        "lease_revoke_reason": "transition:succeeded",
        "revision": 3,
        "created_at": 0.5,
        "lease_acquired_at": 0.6,
        "lease_revoked_at": 0.9,
        "updated_at": 1.0,
        "payload": {
            "project_id": project_id,
            "delivery_scope": "document",
            "dry_run": False,
            "resume_from_job_id": None,
        },
        "result": result,
        "agent_runtime": {
            "generation_release_identity": generation_release,
            "execution_control": {
                "schema_version": "execution-control-v1",
                "cancelled": False,
                "limits": {
                    "max_concurrency": 1,
                    "max_model_attempts": 4,
                    "max_input_chars": 10_000,
                    "max_requested_output_tokens": 1_000,
                },
                "usage": {
                    "model_attempts": 2,
                    "input_chars": 100,
                    "requested_output_tokens": 20,
                    "actual_input_tokens": 10,
                    "actual_uncached_input_tokens": 10,
                    "cache_creation_input_tokens": 0,
                    "cache_read_input_tokens": 0,
                    "cache_hit_ratio": 0.0,
                    "actual_output_tokens": 5,
                    "actual_output_chars": 20,
                    "active": 0,
                    "peak_active": 1,
                    "provider_attempts": {
                        "provider-a": 1,
                        "provider-render": 1,
                    },
                },
                "elapsed_seconds": 1.0,
            }
        },
    }
    job_path = jobs_dir / f"{job_id}.json"
    job_path.write_text(json.dumps(job), encoding="utf-8")
    job_snapshot = read_regular_file_snapshot(job_path)
    assert job_snapshot is not None

    selected, machine_code = _candidate_source(
        jobs=[(job_snapshot, job)],
        project_id=project_id,
        tender=tender,
        boq=boq,
        release_identity=release_identity,
        release_root=release_root,
        workspace_root=workspace_root,
        requested_job_id=job_id,
        provider_admission_snapshot=provider_admission_snapshot,
        events_state=current_events_state(),
    )

    assert machine_code == "CURRENT_FORMAL_SOURCE_ELIGIBLE"
    assert selected is not None
    assert selected["job_id"] == job_id
    assert selected["event_projection"]["successful_chapter_count"] == 1
    assert selected["artifact_projection"]["artifact_count"] >= 13
    assert selected["variant"]["sections"] == professional_sections
    assert {witness.path for witness in selected["witnesses"]} >= {
        output_path.resolve(),
        event_path.resolve(),
        (checkpoint_dir / "variant-1.json").resolve(),
    }

    provider_admission_path.chmod(0o644)
    broad_mode_snapshot = read_regular_file_snapshot(provider_admission_path)
    assert broad_mode_snapshot is not None
    broad_mode_selected, broad_mode_code = _candidate_source(
        jobs=[(job_snapshot, job)],
        project_id=project_id,
        tender=tender,
        boq=boq,
        release_identity=release_identity,
        release_root=release_root,
        workspace_root=workspace_root,
        requested_job_id=job_id,
        provider_admission_snapshot=broad_mode_snapshot,
        events_state=current_events_state(),
    )
    assert broad_mode_selected is None
    assert broad_mode_code == "HOLD_SOURCE_PROVIDER_ADMISSION_INCOMPLETE"
    provider_admission_path.chmod(0o600)
    provider_admission_snapshot = read_regular_file_snapshot(
        provider_admission_path
    )
    assert provider_admission_snapshot is not None

    original_checkpoint_bytes = checkpoint_path.read_bytes()
    incomplete_checkpoint = json.loads(original_checkpoint_bytes)
    incomplete_checkpoint["binding"]["provider_routes"] = incomplete_checkpoint[
        "binding"
    ]["provider_routes"][:-1]
    incomplete_binding_core = {
        key: value
        for key, value in incomplete_checkpoint["binding"].items()
        if key != "binding_digest"
    }
    incomplete_checkpoint["binding"]["binding_digest"] = canonical_digest(
        incomplete_binding_core
    )
    incomplete_checkpoint["binding_digest"] = incomplete_checkpoint["binding"][
        "binding_digest"
    ]
    incomplete_checkpoint_core = {
        key: value
        for key, value in incomplete_checkpoint.items()
        if key != "integrity_digest"
    }
    incomplete_checkpoint["integrity_digest"] = canonical_digest(
        incomplete_checkpoint_core
    )
    checkpoint_path.write_text(json.dumps(incomplete_checkpoint), encoding="utf-8")
    missing_route_selected, missing_route_code = _candidate_source(
        jobs=[(job_snapshot, job)],
        project_id=project_id,
        tender=tender,
        boq=boq,
        release_identity=release_identity,
        release_root=release_root,
        workspace_root=workspace_root,
        requested_job_id=job_id,
        provider_admission_snapshot=provider_admission_snapshot,
        events_state=current_events_state(),
    )
    assert missing_route_selected is None
    assert missing_route_code == "HOLD_SOURCE_CHECKPOINT_INCOMPLETE"
    checkpoint_path.write_bytes(original_checkpoint_bytes)

    wrong_lineage_checkpoint = json.loads(original_checkpoint_bytes)
    wrong_lineage_checkpoint["binding"]["attempt_id"] = "d" * 32
    wrong_binding_core = {
        key: value
        for key, value in wrong_lineage_checkpoint["binding"].items()
        if key != "binding_digest"
    }
    wrong_lineage_checkpoint["binding"]["binding_digest"] = canonical_digest(
        wrong_binding_core
    )
    wrong_lineage_checkpoint["binding_digest"] = wrong_lineage_checkpoint[
        "binding"
    ]["binding_digest"]
    wrong_checkpoint_core = {
        key: value
        for key, value in wrong_lineage_checkpoint.items()
        if key != "integrity_digest"
    }
    wrong_lineage_checkpoint["integrity_digest"] = canonical_digest(
        wrong_checkpoint_core
    )
    checkpoint_path.write_text(
        json.dumps(wrong_lineage_checkpoint),
        encoding="utf-8",
    )
    wrong_lineage_selected, wrong_lineage_code = _candidate_source(
        jobs=[(job_snapshot, job)],
        project_id=project_id,
        tender=tender,
        boq=boq,
        release_identity=release_identity,
        release_root=release_root,
        workspace_root=workspace_root,
        requested_job_id=job_id,
        provider_admission_snapshot=provider_admission_snapshot,
        events_state=current_events_state(),
    )
    assert wrong_lineage_selected is None
    assert wrong_lineage_code == "HOLD_SOURCE_CHECKPOINT_INCOMPLETE"
    checkpoint_path.write_bytes(original_checkpoint_bytes)

    wrong_result_identity = json.loads(json.dumps(job))
    wrong_result_identity["result"]["job_execution_identity"][
        "owner_instance_id"
    ] = "d" * 32
    job_path.write_text(json.dumps(wrong_result_identity), encoding="utf-8")
    wrong_result_snapshot = read_regular_file_snapshot(job_path)
    assert wrong_result_snapshot is not None
    wrong_result_selected, wrong_result_code = _candidate_source(
        jobs=[(wrong_result_snapshot, wrong_result_identity)],
        project_id=project_id,
        tender=tender,
        boq=boq,
        release_identity=release_identity,
        release_root=release_root,
        workspace_root=workspace_root,
        requested_job_id=job_id,
        provider_admission_snapshot=provider_admission_snapshot,
        events_state=current_events_state(),
    )
    assert wrong_result_selected is None
    assert wrong_result_code == "HOLD_SOURCE_JOB_UNTRUSTED"
    job_path.write_text(json.dumps(job), encoding="utf-8")
    job_snapshot = read_regular_file_snapshot(job_path)
    assert job_snapshot is not None

    malformed_job_id = "b" * 32
    malformed_job = json.loads(json.dumps(job))
    malformed_job["job_id"] = malformed_job_id
    malformed_job["updated_at"] = float("nan")
    malformed_path = jobs_dir / f"{malformed_job_id}.json"
    malformed_path.write_text(json.dumps(malformed_job), encoding="utf-8")
    malformed_snapshot = read_regular_file_snapshot(malformed_path)
    assert malformed_snapshot is not None
    coexist_selected, coexist_code = _candidate_source(
        jobs=[(malformed_snapshot, malformed_job), (job_snapshot, job)],
        project_id=project_id,
        tender=tender,
        boq=boq,
        release_identity=release_identity,
        release_root=release_root,
        workspace_root=workspace_root,
        requested_job_id=None,
        provider_admission_snapshot=provider_admission_snapshot,
        events_state=current_events_state(),
    )
    assert coexist_code == "CURRENT_FORMAL_SOURCE_ELIGIBLE"
    assert coexist_selected is not None
    assert coexist_selected["job_id"] == job_id
    malformed_selected, malformed_code = _candidate_source(
        jobs=[(malformed_snapshot, malformed_job)],
        project_id=project_id,
        tender=tender,
        boq=boq,
        release_identity=release_identity,
        release_root=release_root,
        workspace_root=workspace_root,
        requested_job_id=malformed_job_id,
        provider_admission_snapshot=provider_admission_snapshot,
        events_state=current_events_state(),
    )
    assert malformed_selected is None
    assert malformed_code == "HOLD_SOURCE_JOB_UNTRUSTED"

    invalid_event_job_id = "c" * 32
    invalid_event_job = json.loads(json.dumps(job))
    invalid_event_job["job_id"] = invalid_event_job_id
    invalid_event_job["updated_at"] = 2.0
    invalid_event_job["result"]["job_execution_identity"]["job_id"] = (
        invalid_event_job_id
    )
    invalid_output_path = (
        workspace_root / "build" / f"actions_{invalid_event_job_id}.json"
    )
    invalid_output_path.write_text(
        json.dumps({"variants": [variant]}),
        encoding="utf-8",
    )
    invalid_event_job["result"]["json"] = str(invalid_output_path)
    invalid_event_path = events_dir / f"{invalid_event_job_id}.jsonl"
    invalid_event_path.write_bytes(b"\xff")
    invalid_job_path = jobs_dir / f"{invalid_event_job_id}.json"
    invalid_job_path.write_text(json.dumps(invalid_event_job), encoding="utf-8")
    invalid_job_snapshot = read_regular_file_snapshot(invalid_job_path)
    assert invalid_job_snapshot is not None
    event_coexist_selected, event_coexist_code = _candidate_source(
        jobs=[(invalid_job_snapshot, invalid_event_job), (job_snapshot, job)],
        project_id=project_id,
        tender=tender,
        boq=boq,
        release_identity=release_identity,
        release_root=release_root,
        workspace_root=workspace_root,
        requested_job_id=None,
        provider_admission_snapshot=provider_admission_snapshot,
        events_state=current_events_state(),
    )
    assert event_coexist_code == "CURRENT_FORMAL_SOURCE_ELIGIBLE"
    assert event_coexist_selected is not None
    assert event_coexist_selected["job_id"] == job_id
    invalid_event_selected, invalid_event_code = _candidate_source(
        jobs=[(invalid_job_snapshot, invalid_event_job)],
        project_id=project_id,
        tender=tender,
        boq=boq,
        release_identity=release_identity,
        release_root=release_root,
        workspace_root=workspace_root,
        requested_job_id=invalid_event_job_id,
        provider_admission_snapshot=provider_admission_snapshot,
        events_state=current_events_state(),
    )
    assert invalid_event_selected is None
    assert invalid_event_code == "HOLD_SOURCE_EVENTS_INCOMPLETE"

    stale_job = json.loads(json.dumps(job))
    stale_job["result"]["generation_release_identity"]["manifest_digest"] = (
        "9" * 64
    )
    job_path.write_text(json.dumps(stale_job), encoding="utf-8")
    stale_snapshot = read_regular_file_snapshot(job_path)
    assert stale_snapshot is not None
    stale_selected, stale_code = _candidate_source(
        jobs=[(stale_snapshot, stale_job)],
        project_id=project_id,
        tender=tender,
        boq=boq,
        release_identity=release_identity,
        release_root=release_root,
        workspace_root=workspace_root,
        requested_job_id=job_id,
        provider_admission_snapshot=provider_admission_snapshot,
        events_state=current_events_state(),
    )
    assert stale_selected is None
    assert stale_code == "HOLD_SOURCE_JOB_UNTRUSTED"

    job_path.write_text(json.dumps(job), encoding="utf-8")
    ordinal_output = {"variants": [{**variant, "variant_id": 2}]}
    output_path.write_text(json.dumps(ordinal_output), encoding="utf-8")
    restored_snapshot = read_regular_file_snapshot(job_path)
    assert restored_snapshot is not None
    ordinal_selected, ordinal_code = _candidate_source(
        jobs=[(restored_snapshot, job)],
        project_id=project_id,
        tender=tender,
        boq=boq,
        release_identity=release_identity,
        release_root=release_root,
        workspace_root=workspace_root,
        requested_job_id=job_id,
        provider_admission_snapshot=provider_admission_snapshot,
        events_state=current_events_state(),
    )
    assert ordinal_selected is None
    assert ordinal_code == "HOLD_SOURCE_OUTPUT_UNTRUSTED"


def test_reserved_latest_run_id_and_bare_mapping_are_rejected(tmp_path: Path) -> None:
    prepared = _prepared(tmp_path, run_id="run-1")
    latest = dict(prepared.receipt)
    latest["run_id"] = "latest"
    core = {key: value for key, value in latest.items() if key != "receipt_digest"}
    latest["receipt_digest"] = canonical_digest(core)
    assert validate_acceptance_receipt(latest)["ok"] is False
    with pytest.raises(AcceptanceError, match="内部签发"):
        publish_acceptance_receipt(prepared.receipt)  # type: ignore[arg-type]


def test_collect_uses_one_audit_and_registry_snapshot_with_zero_external_calls(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    data_root, registry_path, release, current = _collect_fixture(tmp_path)
    audit_path = data_root / "audit" / "ingest.jsonl"
    seen_audit_objects: list[tuple[str, ...]] = []
    seen_registry_objects: list[bytes] = []
    original_drawing = acceptance.build_drawing_index
    original_standard = acceptance.build_standard_index
    original_parameter = acceptance.build_project_parameter_evidence
    original_ledger = acceptance.build_project_fact_ledger_from_inputs
    original_gate = acceptance.build_delivery_quality_gate

    def drawing_wrapper(*args: Any, **kwargs: Any) -> dict[str, Any]:
        seen_audit_objects.append(kwargs["audit_lines"])
        return original_drawing(*args, **kwargs)

    def standard_wrapper(*args: Any, **kwargs: Any) -> dict[str, Any]:
        seen_audit_objects.append(kwargs["audit_lines"])
        seen_registry_objects.append(kwargs["official_registry_bytes"])
        return original_standard(*args, **kwargs)

    def parameter_wrapper(*args: Any, **kwargs: Any) -> dict[str, Any]:
        seen_audit_objects.append(kwargs["audit_lines"])
        return original_parameter(*args, **kwargs)

    def ledger_wrapper(*args: Any, **kwargs: Any) -> dict[str, Any]:
        seen_audit_objects.append(kwargs["trusted_ingest_audit_lines"])
        return original_ledger(*args, **kwargs)

    def gate_wrapper(*args: Any, **kwargs: Any) -> dict[str, Any]:
        seen_audit_objects.append(kwargs["trusted_ingest_audit_lines"])
        seen_registry_objects.append(kwargs["trusted_standard_registry_bytes"])
        return original_gate(*args, **kwargs)

    original_read_text = Path.read_text
    original_read_bytes = Path.read_bytes

    def guarded_read_text(path: Path, *args: Any, **kwargs: Any) -> str:
        if path == audit_path:
            raise AssertionError("captured audit bytes were reread")
        return original_read_text(path, *args, **kwargs)

    def guarded_read_bytes(path: Path) -> bytes:
        if path == registry_path:
            raise AssertionError("captured registry bytes were reread")
        return original_read_bytes(path)

    monkeypatch.setattr(acceptance, "build_drawing_index", drawing_wrapper)
    monkeypatch.setattr(acceptance, "build_standard_index", standard_wrapper)
    monkeypatch.setattr(
        acceptance,
        "build_project_parameter_evidence",
        parameter_wrapper,
    )
    monkeypatch.setattr(
        acceptance,
        "build_project_fact_ledger_from_inputs",
        ledger_wrapper,
    )
    monkeypatch.setattr(acceptance, "build_delivery_quality_gate", gate_wrapper)
    monkeypatch.setattr(Path, "read_text", guarded_read_text)
    monkeypatch.setattr(Path, "read_bytes", guarded_read_bytes)
    _install_offline_fail_on_call_guards(monkeypatch)

    prepared = acceptance.collect_acceptance_snapshot(
        project_id="P-1",
        data_root=data_root,
        registry_path=registry_path,
        release_identity=release,
        release_witnesses=[current],
        release_validator=lambda: dict(release),
        run_id="offline-collect",
        generated_at="2026-08-28T08:00:00Z",
    )

    receipt = prepared.receipt
    assert receipt["decision"] == "HOLD"
    assert receipt["model_calls"] == 0
    assert receipt["provider_probes"] == 0
    assert len(seen_audit_objects) == 5
    assert all(value is seen_audit_objects[0] for value in seen_audit_objects)
    assert len(seen_registry_objects) == 2
    assert all(value is seen_registry_objects[0] for value in seen_registry_objects)
    assert receipt["inputs"]["ingest_audit"]["sha256"] == hashlib.sha256(
        b""
    ).hexdigest()


def test_collect_rejects_non_utf8_audit_before_any_builder(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    data_root, registry_path, release, current = _collect_fixture(tmp_path)
    (data_root / "audit" / "ingest.jsonl").write_bytes(b"\xff")

    def forbidden(*_args: Any, **_kwargs: Any) -> Any:
        raise AssertionError("builder must not run before audit UTF-8 validation")

    monkeypatch.setattr(acceptance, "build_drawing_index", forbidden)
    monkeypatch.setattr(acceptance, "build_standard_index", forbidden)
    monkeypatch.setattr(
        acceptance,
        "build_project_parameter_evidence",
        forbidden,
    )
    with pytest.raises(AcceptanceError) as error:
        acceptance.collect_acceptance_snapshot(
            project_id="P-1",
            data_root=data_root,
            registry_path=registry_path,
            release_identity=release,
            release_witnesses=[current],
            release_validator=lambda: dict(release),
            run_id="offline-invalid-audit",
            generated_at="2026-08-28T08:00:00Z",
        )

    assert error.value.code == "ACCEPTANCE_INGEST_AUDIT_INVALID"


@pytest.mark.parametrize("missing_input", ["plan", "ingest", "approval"])
def test_missing_input_witness_rejects_missing_to_present_transition(
    tmp_path: Path,
    missing_input: str,
) -> None:
    data_root, registry_path, release, current = _collect_fixture(tmp_path)
    project_root = data_root / "autoplan" / "projects" / "P-1"
    targets = {
        "plan": project_root / "plan.json",
        "ingest": data_root / "audit" / "ingest.jsonl",
        "approval": data_root / "audit" / "project_fact_approvals.jsonl",
    }
    target = targets[missing_input]
    if target.exists():
        target.unlink()
    prepared = acceptance.collect_acceptance_snapshot(
        project_id="P-1",
        data_root=data_root,
        registry_path=registry_path,
        release_identity=release,
        release_witnesses=[current],
        release_validator=lambda: dict(release),
        run_id=f"missing-{missing_input}",
        generated_at="2026-08-28T08:00:00Z",
    )
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text("{}\n" if missing_input == "plan" else "\n", encoding="utf-8")

    with pytest.raises(AcceptanceError) as error:
        acceptance.verify_snapshot_stability(prepared)

    assert error.value.code == "ACCEPTANCE_INPUT_CHANGED"


def test_event_directory_membership_witness_rejects_new_fragment(
    tmp_path: Path,
) -> None:
    data_root, registry_path, release, current = _collect_fixture(tmp_path)
    events_dir = data_root / "autoplan" / "events"
    events_dir.mkdir(parents=True)
    prepared = acceptance.collect_acceptance_snapshot(
        project_id="P-1",
        data_root=data_root,
        registry_path=registry_path,
        release_identity=release,
        release_witnesses=[current],
        release_validator=lambda: dict(release),
        run_id="event-membership-change",
        generated_at="2026-08-28T08:00:00Z",
    )
    assert prepared.receipt["inputs"]["events_directory"] == {
        "status": "present",
        "members_digest": canonical_digest([]),
        "member_count": 0,
        "absence_digest": None,
    }
    (events_dir / ("a" * 32 + ".1.jsonl")).write_text(
        "{}\n",
        encoding="utf-8",
    )

    with pytest.raises(AcceptanceError) as error:
        acceptance.verify_snapshot_stability(prepared)

    assert error.value.code == "ACCEPTANCE_INPUT_CHANGED"


@pytest.mark.parametrize("target", ["audit", "registry"])
def test_snapshot_stability_rejects_same_bytes_inode_replacement(
    tmp_path: Path,
    target: str,
) -> None:
    data_root, registry_path, release, current = _collect_fixture(tmp_path)
    prepared = acceptance.collect_acceptance_snapshot(
        project_id="P-1",
        data_root=data_root,
        registry_path=registry_path,
        release_identity=release,
        release_witnesses=[current],
        release_validator=lambda: dict(release),
        run_id=f"same-bytes-{target}",
        generated_at="2026-08-28T08:00:00Z",
    )
    target_path = (
        data_root / "audit" / "ingest.jsonl"
        if target == "audit"
        else registry_path
    )
    before = read_regular_file_snapshot(target_path)
    assert before is not None
    replacement = target_path.with_name(f".{target_path.name}.replacement")
    replacement.write_bytes(before.raw)
    acceptance.os.replace(replacement, target_path)
    after = read_regular_file_snapshot(target_path)
    assert after is not None
    assert after.sha256 == before.sha256
    assert after.inode != before.inode

    with pytest.raises(AcceptanceError) as error:
        acceptance.verify_snapshot_stability(prepared)

    assert error.value.code == "ACCEPTANCE_INPUT_CHANGED"


def test_release_registry_path_rejects_parent_symlink(tmp_path: Path) -> None:
    release_root = tmp_path / "release"
    outside = tmp_path / "outside"
    release_root.mkdir(parents=True)
    outside.mkdir()
    (outside / "_official_registry.json").write_text("{}", encoding="utf-8")
    (release_root / SEALED_COMPLIANCE_ROOT_RELATIVE_PATH).symlink_to(
        outside,
        target_is_directory=True,
    )

    with pytest.raises(AcceptanceError) as error:
        _assert_path_without_symlinks(
            sealed_official_registry_path(release_root),
            root=release_root,
        )

    assert error.value.code == "ACCEPTANCE_REGISTRY_UNTRUSTED"


def test_collect_rejects_mutable_registry_fallback_even_when_bytes_match(
    tmp_path: Path,
) -> None:
    data_root, sealed_registry, release, current = _collect_fixture(tmp_path)
    release_root = Path(release["release_root"])
    mutable_registry = release_root / "知识图谱" / "compliance" / sealed_registry.name
    mutable_registry.parent.mkdir(parents=True, exist_ok=True)
    mutable_registry.write_bytes(sealed_registry.read_bytes())

    with pytest.raises(AcceptanceError) as error:
        acceptance.collect_acceptance_snapshot(
            project_id="P-1",
            data_root=data_root,
            registry_path=mutable_registry,
            release_identity=release,
            release_witnesses=[current],
            release_validator=lambda: dict(release),
            run_id="mutable-registry-forbidden",
            generated_at="2026-08-28T08:00:00Z",
        )

    assert error.value.code == "ACCEPTANCE_REGISTRY_UNTRUSTED"


@pytest.mark.parametrize("mutation", ["tamper", "missing", "symlink"])
def test_collect_rejects_unsealed_registry_state(
    tmp_path: Path,
    mutation: str,
) -> None:
    data_root, registry_path, release, current = _collect_fixture(tmp_path)
    registry_path.chmod(0o644)
    if mutation == "tamper":
        registry_path.write_bytes(registry_path.read_bytes() + b"\n")
        registry_path.chmod(0o444)
    elif mutation == "missing":
        registry_path.unlink()
    else:
        outside = tmp_path / "outside-registry.json"
        outside.write_text("{}", encoding="utf-8")
        registry_path.unlink()
        registry_path.symlink_to(outside)

    with pytest.raises(AcceptanceError) as error:
        acceptance.collect_acceptance_snapshot(
            project_id="P-1",
            data_root=data_root,
            registry_path=registry_path,
            release_identity=release,
            release_witnesses=[current],
            release_validator=lambda: dict(release),
            run_id=f"sealed-registry-{mutation}",
            generated_at="2026-08-28T08:00:00Z",
        )

    assert error.value.code == "ACCEPTANCE_REGISTRY_UNTRUSTED"


def test_latest_must_match_same_project_immutable_receipt(tmp_path: Path) -> None:
    prepared = _prepared(tmp_path, run_id="run-history")
    published = publish_acceptance_receipt(prepared)
    latest = Path(published["latest_receipt"])
    forged = _receipt(
        tmp_path,
        run_id="run-forged",
        release=prepared.release_projection,
    )
    latest.write_text(json.dumps(forged), encoding="utf-8")

    with pytest.raises(AcceptanceError) as error:
        _capture_latest_state(
            data_root=prepared.data_root,
            project_id="P-1",
            output_root=None,
        )

    assert error.value.code == "ACCEPTANCE_LATEST_INVALID"


def test_atomic_publish_permissions_duplicate_and_latest_cas(tmp_path: Path) -> None:
    first = _prepared(tmp_path, run_id="run-1")
    second = _prepared(tmp_path, run_id="run-2")
    result = publish_acceptance_receipt(first)
    immutable = Path(result["immutable_receipt"])
    latest = Path(result["latest_receipt"])
    assert stat.S_IMODE(immutable.stat().st_mode) == 0o400
    assert stat.S_IMODE(latest.stat().st_mode) == 0o600
    assert validate_acceptance_receipt(json.loads(latest.read_text()))["ok"] is True

    with pytest.raises(AcceptanceError) as stale:
        publish_acceptance_receipt(second)
    assert stale.value.code == "ACCEPTANCE_LATEST_CONCURRENT_UPDATE"

    duplicate = _prepared(tmp_path, run_id="run-1")
    with pytest.raises(AcceptanceError) as duplicate_error:
        publish_acceptance_receipt(duplicate)
    assert duplicate_error.value.code == "ACCEPTANCE_RUN_ID_EXISTS"


def test_publish_revalidates_witness_before_any_receipt_write(tmp_path: Path) -> None:
    prepared = _prepared(tmp_path, run_id="run-tamper")
    evidence_path = prepared.witnesses[0].path
    evidence_path.write_bytes(b"tampered")

    with pytest.raises(AcceptanceError) as error:
        publish_acceptance_receipt(prepared)

    assert error.value.code == "ACCEPTANCE_INPUT_CHANGED"
    output_dir = prepared.data_root / "autoplan" / "acceptance_receipts"
    assert not output_dir.exists()


def test_publish_removes_orphan_if_latest_commit_fails(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    prepared = _prepared(tmp_path, run_id="run-no-orphan")
    original_verify = acceptance.verify_snapshot_stability
    calls = 0

    def fail_before_latest(snapshot: _PreparedAcceptance) -> None:
        nonlocal calls
        calls += 1
        original_verify(snapshot)
        if calls == 3:
            raise AcceptanceError(
                "ACCEPTANCE_INPUT_CHANGED",
                "simulated latest precommit failure",
            )

    monkeypatch.setattr(
        acceptance,
        "verify_snapshot_stability",
        fail_before_latest,
    )
    with pytest.raises(AcceptanceError) as error:
        publish_acceptance_receipt(prepared)
    assert error.value.code == "ACCEPTANCE_INPUT_CHANGED"

    output_dir = (
        prepared.data_root
        / "autoplan"
        / "acceptance_receipts"
        / "no_model_formal"
        / "P-1"
    )
    assert not (output_dir / "run-no-orphan.json").exists()
    assert not (output_dir / "latest.json").exists()

    monkeypatch.setattr(
        acceptance,
        "verify_snapshot_stability",
        original_verify,
    )
    assert publish_acceptance_receipt(prepared)["written"] is True


def test_publish_rejects_successor_directory_replacement(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    prepared = _prepared(tmp_path, run_id="run-dir-swap")
    output_dir = (
        prepared.data_root
        / "autoplan"
        / "acceptance_receipts"
        / "no_model_formal"
        / "P-1"
    )
    original_link = acceptance.os.link
    swapped = False

    def swap_then_link(*args: Any, **kwargs: Any) -> None:
        nonlocal swapped
        if not swapped:
            swapped = True
            successors = output_dir / "successors"
            successors.rename(output_dir / "successors-old")
            successors.mkdir(mode=0o700)
        original_link(*args, **kwargs)

    monkeypatch.setattr(acceptance.os, "link", swap_then_link)
    with pytest.raises(AcceptanceError) as error:
        publish_acceptance_receipt(prepared)

    assert error.value.code == "ACCEPTANCE_OUTPUT_DIRECTORY_CHANGED"
    assert not (output_dir / "latest.json").exists()
    assert not (output_dir / "run-dir-swap.json").exists()
    assert list((output_dir / "successors").iterdir()) == []


def test_publish_recovers_durable_successor_after_latest_rename_failure(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    prepared = _prepared(tmp_path, run_id="run-recover")
    original_rename = acceptance.os.rename
    failed = False

    def fail_latest_once(src: str, dst: str, **kwargs: Any) -> None:
        nonlocal failed
        if not failed and src.startswith(".latest.") and dst == "latest.json":
            failed = True
            raise OSError(errno.EIO, "simulated rename failure")
        original_rename(src, dst, **kwargs)

    monkeypatch.setattr(acceptance.os, "rename", fail_latest_once)
    with pytest.raises(AcceptanceError) as first_error:
        publish_acceptance_receipt(prepared)
    assert first_error.value.code == "ACCEPTANCE_LATEST_UPDATE_FAILED"

    monkeypatch.setattr(acceptance.os, "rename", original_rename)
    with pytest.raises(AcceptanceError) as recovery_signal:
        publish_acceptance_receipt(prepared)
    assert recovery_signal.value.code == "ACCEPTANCE_LATEST_CONCURRENT_UPDATE"

    output_dir = (
        prepared.data_root
        / "autoplan"
        / "acceptance_receipts"
        / "no_model_formal"
        / "P-1"
    )
    latest = read_regular_file_snapshot(output_dir / "latest.json")
    immutable = read_regular_file_snapshot(output_dir / "run-recover.json")
    assert latest is not None and immutable is not None
    assert latest.raw == immutable.raw
    recovered = json.loads(latest.raw)
    assert recovered["receipt_digest"] == prepared.receipt["receipt_digest"]
    followup = _prepared(tmp_path, run_id="run-after-recovery")
    assert publish_acceptance_receipt(followup)["written"] is True


def test_publish_partial_claim_write_leaves_no_authoritative_orphan(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    prepared = _prepared(tmp_path, run_id="run-partial-claim")
    original_write = acceptance._write_all
    calls = 0

    def fail_claim_write(descriptor: int, payload: bytes) -> None:
        nonlocal calls
        calls += 1
        if calls == 3:
            acceptance.os.write(descriptor, payload[:12])
            raise OSError(errno.EIO, "simulated partial claim")
        original_write(descriptor, payload)

    monkeypatch.setattr(acceptance, "_write_all", fail_claim_write)
    with pytest.raises(AcceptanceError) as error:
        publish_acceptance_receipt(prepared)
    assert error.value.code == "ACCEPTANCE_SUCCESSOR_WRITE_FAILED"

    output_dir = (
        prepared.data_root
        / "autoplan"
        / "acceptance_receipts"
        / "no_model_formal"
        / "P-1"
    )
    assert not (output_dir / "run-partial-claim.json").exists()
    assert not (output_dir / "latest.json").exists()
    assert list((output_dir / "successors").iterdir()) == []

    monkeypatch.setattr(acceptance, "_write_all", original_write)
    assert publish_acceptance_receipt(prepared)["written"] is True


def test_publish_rejects_symlinked_output_parent(
    tmp_path: Path,
) -> None:
    prepared = _prepared(tmp_path, run_id="run-symlink-parent")
    autoplan_root = prepared.data_root / "autoplan"
    outside = prepared.data_root / "outside-receipts"
    outside.mkdir()
    (autoplan_root / "acceptance_receipts").symlink_to(
        outside,
        target_is_directory=True,
    )

    with pytest.raises(AcceptanceError) as error:
        publish_acceptance_receipt(prepared)

    assert error.value.code == "ACCEPTANCE_OUTPUT_PATH_UNTRUSTED"
    assert list(outside.iterdir()) == []


def test_publish_rejects_lock_inode_replacement(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    prepared = _prepared(tmp_path, run_id="run-lock-replaced")
    output_dir = (
        prepared.data_root
        / "autoplan"
        / "acceptance_receipts"
        / "no_model_formal"
        / "P-1"
    )
    lock_path = output_dir / ".publish.lock"
    original_flock = acceptance.fcntl.flock
    replaced = False

    def replace_after_lock(descriptor: int, operation: int) -> None:
        nonlocal replaced
        original_flock(descriptor, operation)
        if operation & acceptance.fcntl.LOCK_EX and not replaced:
            replaced = True
            lock_path.rename(output_dir / ".publish.lock.replaced")
            lock_path.write_bytes(b"replacement")
            lock_path.chmod(0o600)

    monkeypatch.setattr(acceptance.fcntl, "flock", replace_after_lock)
    with pytest.raises(AcceptanceError) as error:
        publish_acceptance_receipt(prepared)

    assert error.value.code == "ACCEPTANCE_PUBLICATION_LOCK_CHANGED"
    assert not (output_dir / "run-lock-replaced.json").exists()
    assert not (output_dir / "latest.json").exists()


def test_publish_rejects_world_writable_existing_output_directory_without_chmod(
    tmp_path: Path,
) -> None:
    prepared = _prepared(tmp_path, run_id="run-output-mode")
    output_dir = (
        prepared.data_root
        / "autoplan"
        / "acceptance_receipts"
        / "no_model_formal"
        / "P-1"
    )
    output_dir.mkdir(parents=True)
    output_dir.chmod(0o777)

    with pytest.raises(AcceptanceError) as error:
        publish_acceptance_receipt(prepared)

    assert error.value.code == "ACCEPTANCE_OUTPUT_PATH_UNTRUSTED"
    assert stat.S_IMODE(output_dir.stat().st_mode) == 0o777
    assert list(output_dir.iterdir()) == []


@pytest.mark.parametrize("interrupt", [KeyboardInterrupt(), SystemExit(17)])
def test_publish_cleans_uncommitted_immutable_before_reraising_interrupt(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    interrupt: BaseException,
) -> None:
    prepared = _prepared(tmp_path, run_id="run-interrupted")
    original_verify = acceptance._verify_publication_precommit
    calls = 0

    def interrupt_after_immutable(*args: Any, **kwargs: Any) -> None:
        nonlocal calls
        calls += 1
        original_verify(*args, **kwargs)
        if calls == 2:
            raise interrupt

    monkeypatch.setattr(
        acceptance,
        "_verify_publication_precommit",
        interrupt_after_immutable,
    )
    with pytest.raises(type(interrupt)):
        publish_acceptance_receipt(prepared)

    output_dir = (
        prepared.data_root
        / "autoplan"
        / "acceptance_receipts"
        / "no_model_formal"
        / "P-1"
    )
    assert not (output_dir / "run-interrupted.json").exists()
    assert not (output_dir / "latest.json").exists()
    assert not list(output_dir.glob(".immutable.*"))


@pytest.mark.parametrize("link_number", [1, 2])
@pytest.mark.parametrize("interrupt_type", [KeyboardInterrupt, SystemExit])
def test_publish_reconciles_hardlink_created_before_interrupt(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    link_number: int,
    interrupt_type: type[BaseException],
) -> None:
    prepared = _prepared(tmp_path, run_id=f"run-link-{link_number}")
    original_link = acceptance.os.link
    calls = 0

    def link_then_interrupt(*args: Any, **kwargs: Any) -> None:
        nonlocal calls
        calls += 1
        original_link(*args, **kwargs)
        if calls == link_number:
            if interrupt_type is SystemExit:
                raise SystemExit(19)
            raise KeyboardInterrupt

    monkeypatch.setattr(acceptance.os, "link", link_then_interrupt)
    with pytest.raises(interrupt_type):
        publish_acceptance_receipt(prepared)

    output_dir = (
        prepared.data_root
        / "autoplan"
        / "acceptance_receipts"
        / "no_model_formal"
        / "P-1"
    )
    assert not (output_dir / f"run-link-{link_number}.json").exists()
    assert not (output_dir / "latest.json").exists()
    assert list((output_dir / "successors").iterdir()) == []
    assert not list(output_dir.glob(".immutable.*"))
    assert not list((output_dir / "successors").glob(".claim.*"))


def test_publish_surfaces_claim_temp_cleanup_failure(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    prepared = _prepared(tmp_path, run_id="run-claim-cleanup")
    original_unlink = acceptance.os.unlink

    def reject_claim_temp(name: str, *args: Any, **kwargs: Any) -> None:
        if str(name).startswith(".claim."):
            raise OSError(errno.EIO, "simulated claim temp cleanup failure")
        original_unlink(name, *args, **kwargs)

    monkeypatch.setattr(acceptance.os, "unlink", reject_claim_temp)
    with pytest.raises(AcceptanceError) as error:
        publish_acceptance_receipt(prepared)

    assert error.value.code == "ACCEPTANCE_SUCCESSOR_TEMP_CLEANUP_FAILED"


def test_directory_handle_open_closes_fd_on_interrupt(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    opened: list[int] = []
    closed: list[int] = []
    original_open = acceptance.os.open
    original_close = acceptance.os.close
    original_validate = acceptance._validate_owned_directory_fd

    def tracked_open(*args: Any, **kwargs: Any) -> int:
        descriptor = original_open(*args, **kwargs)
        opened.append(descriptor)
        return descriptor

    def tracked_close(descriptor: int) -> None:
        closed.append(descriptor)
        original_close(descriptor)

    def interrupt_validation(_descriptor: int) -> Any:
        raise KeyboardInterrupt

    monkeypatch.setattr(acceptance.os, "open", tracked_open)
    monkeypatch.setattr(acceptance.os, "close", tracked_close)
    monkeypatch.setattr(
        acceptance,
        "_validate_owned_directory_fd",
        interrupt_validation,
    )

    with pytest.raises(KeyboardInterrupt):
        acceptance._open_output_directory_handle(
            data_root=tmp_path,
            project_id="P-1",
            output_root=None,
            create=True,
        )

    assert opened
    assert set(opened).issubset(closed)
    monkeypatch.setattr(
        acceptance,
        "_validate_owned_directory_fd",
        original_validate,
    )


@pytest.mark.parametrize(
    "error",
    [
        LaunchError("LAUNCH_CURRENT_CHANGED", "current changed"),
        SupervisorError("SUPERVISOR_STATE_INVALID", "state invalid"),
        ReleaseBuildError("RELEASE_HOME_UNAVAILABLE", "home invalid"),
    ],
)
def test_cli_expected_failures_are_machine_coded(
    monkeypatch: pytest.MonkeyPatch,
    capsys: pytest.CaptureFixture[str],
    error: Exception,
) -> None:
    monkeypatch.setattr(refresh_cli, "_execute", lambda _args: (_ for _ in ()).throw(error))
    assert refresh_cli.main(["--project-id", "P-1"]) == 2
    payload = json.loads(capsys.readouterr().out)
    assert payload["machine_code"] == error.code  # type: ignore[attr-defined]
    assert payload["model_calls"] == 0
    assert payload["provider_probes"] == 0


def test_cli_unexpected_failure_is_redacted(
    monkeypatch: pytest.MonkeyPatch,
    capsys: pytest.CaptureFixture[str],
) -> None:
    monkeypatch.setattr(
        refresh_cli,
        "_execute",
        lambda _args: (_ for _ in ()).throw(RuntimeError("secret detail")),
    )
    assert refresh_cli.main(["--project-id", "P-1"]) == 2
    payload = json.loads(capsys.readouterr().out)
    assert payload["machine_code"] == "ACCEPTANCE_UNEXPECTED_FAILURE"
    assert "secret detail" not in payload["message"]


def test_cli_argument_failures_are_machine_coded(
    capsys: pytest.CaptureFixture[str],
) -> None:
    assert refresh_cli.main([]) == 2
    payload = json.loads(capsys.readouterr().out)
    assert payload["machine_code"] == "ACCEPTANCE_ARGUMENTS_INVALID"
    assert payload["model_calls"] == 0
    assert payload["provider_probes"] == 0

    assert refresh_cli.main(["--help"]) == 0


def test_cli_dry_run_uses_absolute_current_paths_from_arbitrary_cwd(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    base = tmp_path / "release-base"
    release_dir = base / "releases" / ("release-" + "1" * 24)
    unrelated = tmp_path / "unrelated"
    unrelated.mkdir()
    spec = SimpleNamespace(base=base, release_dir=release_dir)
    captured: dict[str, Any] = {}
    monkeypatch.setattr(refresh_cli, "default_release_base", lambda: base)
    monkeypatch.setattr(refresh_cli, "_sealed_release_context", lambda _base: (object(), spec))
    monkeypatch.setattr(
        refresh_cli,
        "_release_bundle",
        lambda *_args, **_kwargs: ({"release_id": "fixture"}, [], dict),
    )

    def fake_run_acceptance(**kwargs: Any) -> dict[str, Any]:
        captured.update(kwargs)
        return {"ok": True, "mode": "dry_run"}

    monkeypatch.setattr(refresh_cli, "run_acceptance", fake_run_acceptance)
    monkeypatch.setenv(
        "ZF_COMPLIANCE_ROOT",
        str(tmp_path / "mutable-registry-override-must-be-ignored"),
    )
    monkeypatch.chdir(unrelated)
    args = argparse.Namespace(
        project_id="P-1",
        release_base=None,
        data_root=None,
        output_root=None,
        run_id=None,
        source_job_id=None,
        write=False,
    )
    result = refresh_cli._execute(args)
    assert result["mode"] == "dry_run"
    assert captured["data_root"] == (
        base / "state" / "workspace" / "backend" / "data"
    )
    assert Path(captured["registry_path"]).is_absolute()
    assert captured["registry_path"] == sealed_official_registry_path(release_dir)
    assert captured["write"] is False


def test_fixed_current_write_authority_uses_shared_sealed_registry_locator(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    base = tmp_path / "release-base"
    release_dir = base / "releases" / ("release-" + "1" * 24)
    current_path = base / "current.json"
    current_path.parent.mkdir(parents=True)
    current_path.write_text("{}", encoding="utf-8")
    current = read_regular_file_snapshot(current_path)
    assert current is not None
    spec = SimpleNamespace(base=base, release_dir=release_dir)
    sealed_snapshot = object()
    release = _release(tmp_path, current.sha256)
    release["release_root"] = str(release_dir)
    monkeypatch.setattr(refresh_cli, "default_release_base", lambda: base)
    monkeypatch.setattr(
        refresh_cli,
        "_sealed_release_context",
        lambda _base: (sealed_snapshot, spec),
    )
    monkeypatch.setattr(
        refresh_cli,
        "_release_bundle",
        lambda *_args, **_kwargs: (release, [current], lambda: dict(release)),
    )

    context = refresh_cli._fixed_current_write_context_impl()

    assert context["registry_path"] == sealed_official_registry_path(release_dir)
    assert context["data_root"] == base / "state" / "workspace" / "backend" / "data"


def test_runtime_health_attestation_binds_backend_process_pid(
    tmp_path: Path,
) -> None:
    identity = {
        "system_id": "docgen-system",
        "release_id": "release-" + "1" * 24,
        "manifest_digest": "2" * 64,
        "source_digest": "3" * 64,
        "runtime_digest": "4" * 64,
    }
    spec = SimpleNamespace(
        identity=SimpleNamespace(as_dict=lambda: dict(identity)),
        release_dir=tmp_path / identity["release_id"],
    )
    supervisor = {
        "release_id": identity["release_id"],
        "backend_pid": 102,
        "ui_pid": 103,
    }
    payload = {
        "ok": True,
        **identity,
        "release_root": str(spec.release_dir),
        "release_managed": True,
        "runtime_mode": "sealed_release",
        "process_pid": 102,
        "jobs": {
            "active": 0,
            "queued": 0,
            "running": 0,
            "cancel_requested": 0,
            "stale": 0,
        },
        "queue": {
            "queue_depth": 0,
            "dispatched_jobs": 0,
            "active_process_alive": False,
        },
        "supervisor": {
            "available": True,
            "managed": True,
            "status": "healthy",
            "release_id": identity["release_id"],
            "backend_pid": 102,
            "ui_pid": 103,
            "circuit_open": False,
            "health_degraded": False,
            "restart_count_window": 0,
        },
    }
    projection = refresh_cli._backend_health_projection(
        payload,
        spec=spec,
        supervisor=supervisor,
    )
    assert projection["process_pid"] == 102

    payload["process_pid"] = 999
    with pytest.raises(AcceptanceError) as error:
        refresh_cli._backend_health_projection(
            payload,
            spec=spec,
            supervisor=supervisor,
        )
    assert error.value.code == "ACCEPTANCE_RUNTIME_HEALTH_INVALID"


def test_cli_main_dry_run_executes_full_offline_collection_without_side_effects(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
    capsys: pytest.CaptureFixture[str],
) -> None:
    _install_offline_fail_on_call_guards(monkeypatch)
    base = tmp_path / "release-base"
    release_root = base / "releases" / ("release-" + "1" * 24)
    data_root, _registry_path, release, current = _collect_fixture(
        base / "state",
        release_root=release_root,
    )
    spec = SimpleNamespace(base=base, release_dir=release_root)
    sealed_snapshot = object()
    monkeypatch.setattr(refresh_cli, "default_release_base", lambda: base)
    monkeypatch.setattr(
        refresh_cli,
        "_sealed_release_context",
        lambda _base: (sealed_snapshot, spec),
    )

    def release_bundle(
        snapshot: Any,
        actual_spec: Any,
        *,
        runner: Any,
        http_get: Any,
    ) -> tuple[dict[str, Any], list[Any], Any]:
        assert snapshot is sealed_snapshot
        assert actual_spec is spec
        assert runner is not None
        assert http_get is not None
        return release, [current], lambda: dict(release)

    monkeypatch.setattr(refresh_cli, "_release_bundle", release_bundle)
    assert (
        refresh_cli.main(
            [
                "--project-id",
                "P-1",
                "--run-id",
                "cli-offline-e2e",
            ]
        )
        == 0
    )
    payload = json.loads(capsys.readouterr().out)
    assert payload["ok"] is True
    assert payload["mode"] == "dry_run"
    assert payload["decision"] == "HOLD"
    assert payload["receipt"]["model_calls"] == 0
    assert payload["receipt"]["provider_probes"] == 0
    assert payload["write_result"] is None
    assert not (
        data_root
        / "autoplan"
        / "acceptance_receipts"
        / "no_model_formal"
    ).exists()


def test_fixed_current_writer_remains_zero_model_and_zero_provider(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    data_root, registry_path, release, current = _collect_fixture(tmp_path)
    _install_offline_fail_on_call_guards(monkeypatch)
    authority_digest = canonical_digest(
        {
            "kind": "test-fixed-production-writer",
            "release": release,
            "current_sha256": current.sha256,
            "data_root": str(data_root),
            "registry_path": str(registry_path),
        }
    )
    _TEST_FIXED_WRITE_CONTEXT.update(
        {
            "authority_digest": authority_digest,
            "release_identity": release,
            "current_witness": current,
            "data_root": data_root,
            "registry_path": registry_path,
            "release_validator": lambda: dict(release),
        }
    )

    result = acceptance.run_current_runtime_acceptance_write(
        project_id="P-1",
        run_id="fixed-current-zero-model",
    )

    assert result["ok"] is True
    assert result["mode"] == "write"
    assert result["receipt"]["decision"] == "HOLD"
    assert result["receipt"]["model_calls"] == 0
    assert result["receipt"]["provider_probes"] == 0
    assert result["write_result"]["written"] is True
    assert set(inspect.signature(refresh_cli.main).parameters) == {"argv"}
    assert set(inspect.signature(refresh_cli._execute).parameters) == {"args"}


def test_cli_write_rejects_all_path_and_registry_overrides(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    monkeypatch.setattr(refresh_cli, "default_release_base", lambda: tmp_path)
    args = argparse.Namespace(
        project_id="P-1",
        release_base=tmp_path,
        data_root=None,
        output_root=None,
        run_id=None,
        source_job_id=None,
        write=True,
    )
    with pytest.raises(AcceptanceError) as path_error:
        refresh_cli._execute(args)
    assert path_error.value.code == "ACCEPTANCE_WRITE_PATH_OVERRIDE_FORBIDDEN"

    args.release_base = None
    monkeypatch.setenv("ZF_COMPLIANCE_ROOT", str(tmp_path / "override"))
    with pytest.raises(AcceptanceError) as registry_error:
        refresh_cli._execute(args)
    assert registry_error.value.code == (
        "ACCEPTANCE_WRITE_REGISTRY_OVERRIDE_FORBIDDEN"
    )
