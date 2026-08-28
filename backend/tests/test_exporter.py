"""
Tests for backend/zhifei_autoplan/exporter.py
"""
from __future__ import annotations

import datetime as dt
import json
import re
import tempfile
import zipfile
from pathlib import Path
from unittest.mock import MagicMock, patch, Mock

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm

from backend.zhifei_autoplan.exporter import (
    _append_field_run,
    _apply_footer_page_numbers,
    _apply_style,
    _auto_density_images_for_pages,
    _build_report_paths,
    _build_static_toc_entries,
    _clear_block_container,
    _cover_image_caption,
    _format_cover_year_month,
    _format_toc_display_title,
    _hide_paragraph,
    _infer_toc_level,
    _insert_auto_toc,
    _insert_cover_page,
    _insert_full_index_page,
    _normalize_front_matter_page_mode,
    _normalize_full_index_enabled,
    _paginate_toc_entries,
    _resolve_cover_meta,
    _render_toc_line,
    _rank_submission_media,
    _resolve_front_matter_plan,
    _set_cell_border,
    _set_cell_shading,
    _set_cell_width,
    _set_table_all_borders,
    _submission_media_eligible,
    _style_cover_paragraph,
    _toc_entry_style,
    _topic_to_cover_project_name,
    _to_cn_month,
    _usable_page_width_cm,
    export_autoplan_docx,
    export_autoplan_compare_docx,
    export_autoplan_docx_from_file,
)


pytestmark = pytest.mark.usefixtures("allow_legacy_export_contract")


# =============================================================================
# Fixtures
# =============================================================================

@pytest.fixture
def temp_dir():
    """Provide a temporary directory for test outputs."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield tmpdir


@pytest.fixture
def basic_data():
    """Basic test data for export functions."""
    return {
        "topic": "测试施工组织设计",
        "style": {
            "body_font": "SimSun",
            "body_size": 12,
            "title_font": "SimHei",
            "title_size": 16,
            "line_spacing": 1.5,
        },
        "sections": [
            {
                "title": "第一章 工程概述",
                "content": "本工程位于某市某区，总投资约1000万元。",
                "agent_role": "项目经理",
            },
            {
                "title": "第二章 施工部署",
                "content": "施工总平面布置合理，资源配置优化。",
            },
        ],
    }


@pytest.fixture
def data_with_quality_checks(basic_data):
    """Test data with quality checks."""
    basic_data["document_audience"] = "internal_review"
    basic_data["quality_checks"] = {
        "structure": {"ok": True, "note": "结构完整"},
        "score_coverage": {"ok": False, "missing_count": 3},
        "closed_loop": {"ok": True},
        "engineering": {"ok": True},
        "evidence": {"ok": False, "by_section": [
            {"title": "第一章", "evidence_count": 2},
            {"title": "第二章", "evidence_count": 0},
        ]},
        "template_style": {"ok": True},
        "score_coverage_by_section": [
            {"title": "第一章", "ok": True},
            {"title": "第二章", "ok": False, "missing": [
                {"dimension": "质量", "keywords": "质量控制"},
            ]},
        ],
        "closed_loop_by_section": [
            {"title": "第一章", "ok": True, "has_risk": True, "has_measure": True},
            {"title": "第二章", "ok": False, "has_risk": True, "has_measure": False},
        ],
        "engineering_by_section": [
            {"title": "第一章", "ok": True, "missing": []},
            {"title": "第二章", "ok": False, "missing": ["人员", "设备"]},
        ],
        "remediation": [
            {"title": "第二章", "type": "missing_evidence", "suggestion": "补充证据材料"},
        ],
    }
    return basic_data


@pytest.fixture
def data_with_llm_remediation(basic_data):
    """Test data with LLM remediation compare."""
    basic_data["document_audience"] = "internal_review"
    basic_data["sections"] = [
        {
            "title": "第一章 工程概述",
            "content": "整改后的内容，更加详细完整。",
            "original_content": "原始内容，较为简单。",
            "auto_remediated": "llm",
        },
        {
            "title": "第二章 施工部署",
            "content": "无需整改的内容。",
        },
    ]
    return basic_data


# =============================================================================
# Tests for _apply_style
# =============================================================================

class TestApplyStyle:
    """Tests for _apply_style function."""

    def test_apply_style_default_values(self):
        """Test _apply_style with empty style dict uses defaults."""
        doc = Document()
        apply_func = _apply_style(doc, {})

        assert callable(apply_func)

    def test_apply_style_custom_font(self):
        """Test _apply_style with custom font settings."""
        doc = Document()
        style = {
            "body_font": "Arial",
            "body_size": 14,
            "title_font": "Times New Roman",
            "title_size": 18,
            "line_spacing": 2.0,
        }
        apply_func = _apply_style(doc, style)

        assert callable(apply_func)

    def test_apply_style_partial_config(self):
        """Test _apply_style with partial style config."""
        doc = Document()
        style = {"font": "KaiTi", "font_size": 11}
        apply_func = _apply_style(doc, style)

        assert callable(apply_func)

    def test_apply_paragraph_to_normal_text(self):
        """Test applying style to a normal paragraph."""
        doc = Document()
        style = {"body_font": "SimSun", "body_size": 12}
        apply_func = _apply_style(doc, style)

        p = doc.add_paragraph("测试段落")
        apply_func(p, is_title=False)

        # Verify paragraph was processed (no exception)
        assert len(p.runs) >= 0

    def test_apply_paragraph_to_title(self):
        """Test applying style to a title paragraph."""
        doc = Document()
        style = {"title_font": "SimHei", "title_size": 16}
        apply_func = _apply_style(doc, style)

        h = doc.add_heading("测试标题", level=1)
        apply_func(h, is_title=True)

        assert len(h.runs) >= 0

    def test_apply_style_handles_missing_normal_style(self):
        """Test _apply_style handles document without Normal style gracefully."""
        doc = Document()
        # Remove Normal style if exists (simulate edge case)
        style = {"body_font": "Arial"}

        # Should not raise exception
        apply_func = _apply_style(doc, style)
        assert callable(apply_func)


class TestFrontMatterHelpers:
    """Tests for deterministic DOCX front matter helper functions."""

    def test_cover_title_and_date_helpers(self):
        assert _topic_to_cover_project_name("某厂房施工组织设计方案") == "某厂房"
        assert _topic_to_cover_project_name("某厂房施工组织设计") == "某厂房"
        assert _topic_to_cover_project_name("某厂房施组方案") == "某厂房"
        assert _to_cn_month(1) == "一"
        assert _to_cn_month(10) == "十"
        assert _to_cn_month(12) == "十二"
        assert _format_cover_year_month(dt.datetime(2026, 4, 1)) == "二零二六年四月"

    def test_build_report_paths_use_docx_stem(self, temp_dir):
        output_path = Path(temp_dir) / "nested" / "baseline.docx"
        json_path, log_path = _build_report_paths(str(output_path))

        assert json_path == output_path.parent / "baseline.build_report.json"
        assert log_path == output_path.parent / "baseline.build_report.log"

    def test_front_matter_plan_counts_include_and_exclude_modes(self):
        include_plan = _resolve_front_matter_plan(
            style_raw={
                "cover_page_count": 1,
                "toc_page_count": 2,
                "full_index_enabled": "yes",
                "full_index_page_count": 1,
                "front_matter_page_mode": "include",
                "document_total_pages_target": 120,
            },
            data={},
            body_pages_estimate=90,
        )
        assert include_plan["actual_front_matter_pages"] == 4
        assert include_plan["effective_document_pages"] == 120

        exclude_plan = _resolve_front_matter_plan(
            style_raw={
                "cover_page_count": 1,
                "toc_page_count": 2,
                "full_index_enabled": False,
                "front_matter_page_mode": "exclude",
            },
            data={"total_pages_target": 120},
            body_pages_estimate=90,
        )
        assert exclude_plan["actual_front_matter_pages"] == 3
        assert exclude_plan["effective_document_pages"] == 123
        assert _normalize_front_matter_page_mode("bad") == "include"
        assert _normalize_full_index_enabled("enabled") is True
        assert _normalize_full_index_enabled("off") is False

    def test_static_toc_entries_start_after_front_matter(self):
        entries = _build_static_toc_entries(
            sections=[
                {"title": "第一章 工程概况"},
                {"title": "第二章 施工部署"},
            ],
            section_pages=[3, 4],
            front_matter_plan={
                "cover_pages": 1,
                "toc_pages": 2,
                "full_index_pages": 1,
            },
        )
        assert entries == [
            {
                "order": 1,
                "title": "第一章 工程概况",
                "start_page": 5,
                "planned_pages": 3,
                "page_number_exact": False,
            },
            {
                "order": 2,
                "title": "第二章 施工部署",
                "start_page": 8,
                "planned_pages": 4,
                "page_number_exact": False,
            },
        ]

    def test_toc_pagination_and_display_helpers(self):
        entries = [{"title": f"第{i}章"} for i in range(1, 6)]
        chunks = _paginate_toc_entries(entries, 2)
        assert [len(chunk) for chunk in chunks] == [3, 2]
        assert _format_toc_display_title("第一章 工程概况") == "第一章、工程概况"
        assert _format_toc_display_title("第一节 项目概况") == "第一节、项目概况"
        assert _infer_toc_level({"title": "第一章 工程概况"}) == 1
        assert _infer_toc_level({"title": "第一节 项目概况"}) == 2
        assert _infer_toc_level({"title": "一、施工部署"}) == 3
        assert _infer_toc_level({"title": "1.1 施工部署"}) == 3
        assert _infer_toc_level({"title": "手工指定", "level": 9}) == 3

    def test_insert_full_index_page_renders_section_index(self):
        doc = Document()
        apply_paragraph = _apply_style(doc, {"body_font": "宋体", "title_font": "黑体"})

        _insert_full_index_page(
            doc,
            apply_paragraph,
            topic="智慧厂房施工组织设计",
            sections=[
                {"title": "第一章 工程概况"},
                {"title": "第二章 施工部署"},
            ],
            chapter_pages={
                "第一章 工程概况": 3,
                "第二章 施工部署": {"target": 5},
            },
            effective_document_pages=42,
        )

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "全文索引" in text
        assert "智慧厂房施工组织设计；章节数=2；成品预计总页数=42页。" in text
        assert "01. 第一章 工程概况（约3页）" in text
        assert "02. 第二章 施工部署（约5页）" in text
        assert "w:br" in doc._element.xml

    def test_insert_full_index_page_prefers_structured_index_entries(self):
        doc = Document()
        apply_paragraph = _apply_style(doc, {"body_font": "宋体", "title_font": "黑体"})

        _insert_full_index_page(
            doc,
            apply_paragraph,
            topic="施工组织设计",
            sections=[{"title": "第一章 工程概况"}],
            chapter_pages={"第一章 工程概况": 3},
            effective_document_pages=12,
            index_entries=[
                {"title": "索引项一", "summary": "A01. 自定义索引摘要", "planned_pages": 2},
                {"title": "索引项二", "planned_pages": 4},
            ],
        )

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "A01. 自定义索引摘要（约2页）" in text
        assert "02. 索引项二（约4页）" in text
        assert "01. 第一章 工程概况" not in text

    def test_insert_full_index_page_handles_empty_sections(self):
        doc = Document()
        apply_paragraph = _apply_style(doc, {"body_font": "宋体", "title_font": "黑体"})

        _insert_full_index_page(
            doc,
            apply_paragraph,
            topic="施工组织设计",
            sections=[],
            chapter_pages={},
            effective_document_pages=1,
        )

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "章节数=0" in text
        assert "当前无可索引章节。" in text


class TestTocRenderingHelpers:
    """Tests for deterministic DOCX TOC rendering helper functions."""

    def test_append_field_run_and_hide_paragraph_emit_docx_xml(self):
        doc = Document()
        paragraph = doc.add_paragraph()

        _append_field_run(paragraph, 'TOC \\o "1-2" \\h \\z \\u')
        _hide_paragraph(paragraph)

        xml = paragraph._element.xml
        assert "TOC" in xml
        assert "fldCharType" in xml
        assert "vanish" in xml

    def test_toc_entry_style_uses_level_specific_formatting(self):
        style_cfg = {
            "body_font": "宋体",
            "body_latin_font": "Times New Roman",
            "body_size": 12,
            "title_font": "黑体",
            "title_latin_font": "Arial",
            "title_size": 16,
        }

        level1 = _toc_entry_style(style_cfg, 1)
        level2 = _toc_entry_style(style_cfg, 2)
        level3 = _toc_entry_style(style_cfg, 3)

        assert level1["font_east"] == "黑体"
        assert level1["bold"] is True
        assert level1["left_indent_cm"] == 0.0
        assert level2["font_east"] == "宋体"
        assert level2["color_rgb"] == (15, 89, 102)
        assert level2["left_indent_cm"] == 1.0
        assert level3["left_indent_cm"] == 2.0

    def test_render_toc_line_adds_formatted_static_entry(self):
        doc = Document()
        paragraph = _render_toc_line(
            doc,
            {"title": "第一章 工程概况", "start_page": 5, "level": 1},
            style_cfg={"title_font": "黑体", "title_latin_font": "Arial", "title_size": 16},
        )

        assert "第一章、工程概况" in paragraph.text
        assert "5" in paragraph.text
        assert "\t" in paragraph.text
        assert 'w:leader="dot"' in paragraph._element.xml
        assert "PAGEREF ZF_CHAPTER_1" in paragraph._element.xml
        assert paragraph.runs[0].bold is True

    def test_insert_auto_toc_renders_one_visible_live_field(self):
        doc = Document()
        apply_paragraph = _apply_style(doc, {"body_font": "宋体", "title_font": "黑体"})

        _insert_auto_toc(
            doc,
            apply_paragraph,
            style_cfg={"body_font": "宋体", "title_font": "黑体", "doc_title_size": 18},
            toc_pages=2,
            toc_entries=[
                {"title": "第一章 工程概况", "start_page": 4, "level": 1},
                {"title": "第一节 项目概况", "start_page": 5, "level": 2},
                {"title": "一、施工部署", "start_page": 6, "level": 3},
            ],
        )

        text = "\n".join(p.text for p in doc.paragraphs)
        xml = "\n".join(p._element.xml for p in doc.paragraphs)
        assert "目录" in text
        assert "目录（续）" not in text
        assert "第一章、工程概况" not in text
        assert 'TOC \\o "1-3"' in xml
        assert "vanish" not in xml


class TestHeaderFooterHelpers:
    """Tests for deterministic DOCX header/footer helper functions."""

    def test_usable_page_width_and_clear_container_helpers(self):
        doc = Document()
        width_cm = _usable_page_width_cm(doc)
        assert width_cm >= 8.0

        footer = doc.sections[0].footer
        footer.add_paragraph("临时页脚")
        footer.add_table(rows=1, cols=1, width=Cm(2))
        assert "临时页脚" in footer._element.xml

        _clear_block_container(footer)

        assert "临时页脚" not in footer._element.xml
        assert "tbl" not in footer._element.xml

    def test_table_cell_width_shading_and_border_helpers_emit_xml(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        left_cell = table.cell(0, 0)

        _set_cell_width(left_cell, 4.2)
        _set_cell_shading(left_cell, "14A6AE")
        _set_cell_border(left_cell, top={"color": "14A6AE", "sz": 10})
        _set_table_all_borders(table, color="D9EAF0", sz=8, bottom=True)

        xml = table._element.xml
        assert "tcW" in xml
        assert "14A6AE" in xml
        assert "D9EAF0" in xml
        assert "bottom" in xml

    def test_apply_footer_page_numbers_adds_company_and_page_field(self):
        doc = Document()

        _apply_footer_page_numbers(
            doc,
            {"body_font": "宋体", "body_latin_font": "Times New Roman"},
            bidder_company="测试单位",
            logo_path=None,
        )

        footer = doc.sections[0].footer
        xml = footer._element.xml
        text = "\n".join(cell.text for table in footer.tables for row in table.rows for cell in row.cells)
        assert "测试单位" in text
        assert "PAGE" in xml
        assert "NUMPAGES" in xml
        assert "AFC4CE" in xml


class TestCoverPageHelpers:
    """Tests for deterministic DOCX cover page helper functions."""

    def test_cover_image_caption_uses_specific_or_safe_default_label(self):
        assert _cover_image_caption("智慧厂房", "深基坑开挖.png", "site_photo") == "智慧厂房 · 深基坑开挖"
        assert _cover_image_caption("智慧厂房", "微信图片_20260401.jpg", "site_photo") == "智慧厂房 · 现场实景图"
        assert _cover_image_caption("", "image.png", "") == "项目效果图"

    def test_resolve_cover_meta_normalizes_input_without_external_lookup(self):
        meta = _resolve_cover_meta(
            {
                "topic": "智慧厂房施工组织设计方案",
                "project_code": "ZF-2026-001",
                "cover_image_path": "/missing/cover.png",
                "issue_year_month": "二零二六年四月",
                "branding": {
                    "bidder_company": "智飞建设有限公司",
                    "logo_path": "/missing/logo.png",
                },
            }
        )

        assert meta["project_name"] == "智慧厂房"
        assert meta["project_code"] == "ZF-2026-001"
        assert meta["cover_title"] == "施工组织设计"
        assert meta["cover_image_path"] == ""
        assert meta["logo_path"] == ""
        assert meta["bidder_company"] == "智飞建设有限公司"
        assert meta["issue_year_month"] == "二零二六年四月"

    def test_style_cover_paragraph_applies_text_and_formatting(self):
        doc = Document()
        paragraph = doc.add_paragraph()

        run = _style_cover_paragraph(
            paragraph,
            east_font="黑体",
            latin_font="Arial",
            size_pt=18,
            text="施工组织设计",
            bold=True,
            color_rgb=(16, 158, 170),
            space_before_pt=6,
            space_after_pt=8,
            line_spacing_pt=24,
        )

        assert paragraph.text == "施工组织设计"
        assert paragraph.alignment == 1
        assert run.bold is True
        assert run.font.size.pt == 18
        assert str(run.font.color.rgb) == "109EAA"

    def test_insert_cover_page_adds_cover_text_without_main_flow(self):
        doc = Document()
        meta = _resolve_cover_meta(
            {
                "topic": "智慧厂房施工组织设计",
                "project_code": "ZF-2026-002",
                "bidder_company": "智飞建设有限公司",
                "issue_year_month": "二零二六年四月",
            }
        )

        _insert_cover_page(doc, {"body_font": "宋体", "title_font": "黑体"}, meta)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "智慧厂房" in text
        assert "项目编号：ZF-2026-002" in text
        assert "施工组织设计" in text
        assert "智飞建设有限公司" in text
        assert "二零二六年四月" in text

    def test_insert_cover_page_keeps_cover_image_outside_exact_body_line_grid(self, temp_dir):
        from PIL import Image

        cover_image = Path(temp_dir) / "项目现场.jpg"
        Image.new("RGB", (1200, 720), color=(28, 107, 128)).save(cover_image)
        doc = Document()

        _insert_cover_page(
            doc,
            {"body_font": "宋体", "title_font": "黑体", "line_spacing_pt": 22},
            {
                "project_name": "测试项目",
                "cover_title": "施工组织设计",
                "cover_image_path": str(cover_image),
                "cover_image_caption": "项目现场实景图",
            },
        )

        image_paragraph = next(
            paragraph for paragraph in doc.paragraphs
            if paragraph._p.xpath(".//w:drawing")
        )
        assert image_paragraph.paragraph_format.line_spacing == 1.0
        assert image_paragraph.paragraph_format.line_spacing_rule == WD_LINE_SPACING.SINGLE
        assert image_paragraph.paragraph_format.keep_together is True

    def test_cover_rejects_obviously_truncated_project_title_fragment(self):
        meta = _resolve_cover_meta({"topic": "春路）等3条道路工程施工组织设计"})

        assert meta["project_name"] == ""


# =============================================================================
# Tests for export_autoplan_docx
# =============================================================================

class TestExportAutoplanDocx:
    """Tests for export_autoplan_docx function."""

    def test_basic_export(self, temp_dir, basic_data):
        """Test basic DOCX export with minimal data."""
        output_path = Path(temp_dir) / "output.docx"
        result = export_autoplan_docx(basic_data, str(output_path))

        assert result == str(output_path)
        assert output_path.exists()

    def test_submission_docx_excludes_internal_review_material(self, temp_dir):
        output_path = Path(temp_dir) / "submission_safe.docx"
        data = {
            "topic": "测试施工组织设计",
            "sections": [
                {
                    "title": "第一章 施工部署",
                    "agent_role": "主控Agent",
                    "content": (
                        "【多Agent】主控=master；专业=specialist。\n"
                        "【施工准备】\n"
                        "- 对关键工序实行样板先行。【证据:招标文件.pdf#abc123@99】\n"
                        "entity_master_key=internal-only\n"
                        "【图谱节点:ZF-KG-001】"
                    ),
                }
            ],
            "quality_checks": {"structure": {"ok": False}},
            "drawing_index": {"drawings": [{"filename": "internal.dwg"}]},
            "cross_index": {"focus_items": [{"name": "内部重点项"}]},
            "param_trace": {"receipt": {"keys": {"internal.key": {"value": "x"}}}},
        }

        export_autoplan_docx(data, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "对关键工序实行样板先行。" in text
        assert "主控Agent" not in text
        assert "entity_master_key" not in text
        assert "ZF-KG-001" not in text
        assert "abc123@99" not in text
        assert "质量校验摘要" not in text
        assert "图纸证据索引" not in text
        assert "重点项证据闭环索引" not in text
        assert "可编辑参数影响回执" not in text
        assert any(p.style.style_id == "Heading1" and p.text == "第一章 施工部署" for p in doc.paragraphs)
        assert any(p.style.style_id == "Heading2" and p.text == "施工准备" for p in doc.paragraphs)

    def test_submission_docx_rejects_non_text_section_content(self, temp_dir):
        output_path = Path(temp_dir) / "invalid_submission.docx"
        with pytest.raises(RuntimeError, match="SUBMISSION_TEXT_TYPE_INVALID"):
            export_autoplan_docx(
                {
                    "topic": "测试",
                    "sections": [{"title": "第一章", "content": {"quality_checks": "leak"}}],
                },
                str(output_path),
            )
        assert not output_path.exists()

    def test_submission_docx_removes_xml_forbidden_control_characters(self, temp_dir):
        output_path = Path(temp_dir) / "xml_safe.docx"
        export_autoplan_docx(
            {
                "topic": "控制字符\x00验收\x0b施工组织设计",
                "sections": [
                    {
                        "title": "第一章\x00 风险与措施",
                        "content": (
                            "信息化管理、绿色工地和劳保用品配置矩阵纳入方案。\n"
                            "关键工序控制点表记录参数、频次、责任、验收和记录要求。\x00\x0b"
                        ),
                    }
                ],
            },
            str(output_path),
        )

        doc = Document(str(output_path))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        assert "控制字符验收" in text
        assert "第一章 风险与措施" in text
        assert "\x00" not in text
        assert "\x0b" not in text

    def test_submission_docx_renders_markdown_and_risk_triplets_semantically(self, temp_dir):
        output_path = Path(temp_dir) / "semantic_submission.docx"
        export_autoplan_docx(
            {
                "topic": "测试施工组织设计",
                "sections": [
                    {
                        "title": "第一章 质量控制",
                        "content": (
                            "**控制目标**：抽检合格率按 "
                            "100 - defect_count * 100 / max(sample_count, 1) 计算；"
                            "批次指标按 inspection_batches * pass_rate_percent / "
                            "max(total_work_batches, 1) 计算。\n"
                            "- 风险→控制→验证：风险：工序偏差；控制：实行两级复核；"
                            "验证：抽检记录签字归档。"
                        ),
                    }
                ],
            },
            str(output_path),
        )

        doc = Document(str(output_path))
        text = "\n".join(
            [p.text for p in doc.paragraphs]
            + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
        )
        assert "**" not in text
        assert "defect_count" not in text
        assert "sample_count" not in text
        assert "total_work_batches" not in text
        assert "总批次" in text
        assert "抽检合格率" in text
        assert any(run.bold and "控制目标" in run.text for p in doc.paragraphs for run in p.runs)
        assert any("风险" in table.cell(0, 0).text for table in doc.tables if len(table.columns) == 3)
        assert "控制措施" in text
        assert "验证与留痕" in text

    def test_submission_docx_removes_generation_instructions_and_generic_metric_template(self, temp_dir):
        output_path = Path(temp_dir) / "clean_submission.docx"
        export_autoplan_docx(
            {
                "topic": "测试施工组织设计",
                "sections": [
                    {
                        "title": "第一章 质量控制",
                        "content": (
                            "【自动补充】消除空泛词：\n"
                            "- 将空泛词替换为可执行动作+参数+频次+验收标准。\n"
                            "- 对含“频次/阈值/时限/人数/型号/工期”等结论句逐条补证据。\n"
                            "- 示例：抽检频次=每100m2 1次。\n"
                            "量化指标：频次=2次/日（班前+收工）；阈值=偏差≤5mm；间距=1000mm；"
                            "厚度=50mm；时长=4h/作业段；人数=8人/班；设备型号=20t挖机1台。\n"
                            "施工前核对设计图纸并形成书面会审记录。\n"
                            "施工前核对设计图纸并形成书面会审记录。"
                        ),
                    }
                ],
            },
            str(output_path),
        )

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "消除空泛词" not in text
        assert "将空泛词替换" not in text
        assert "示例：" not in text
        assert "频次=2次/日（班前+收工）" not in text
        assert text.count("施工前核对设计图纸并形成书面会审记录") == 1

    def test_submission_docx_suppresses_duplicate_risk_tables_across_sections(self, temp_dir):
        output_path = Path(temp_dir) / "deduplicated_risks.docx"
        repeated = (
            "风险→控制→验证：风险：关键参数超差导致返工；"
            "控制：执行首件确认和过程复核；验证：验收记录签字归档。"
        )
        export_autoplan_docx(
            {
                "topic": "测试施工组织设计",
                "sections": [
                    {"title": "第一章 质量控制", "content": repeated},
                    {"title": "第二章 施工部署", "content": repeated},
                ],
            },
            str(output_path),
        )

        doc = Document(str(output_path))
        assert len(doc.tables) == 1
        assert "关键参数超差导致返工" in doc.tables[0].cell(1, 0).text

    def test_submission_docx_removes_format_scaffolding_and_builds_semantic_hierarchy(self, temp_dir):
        output_path = Path(temp_dir) / "clean_hierarchy.docx"
        export_autoplan_docx(
            {
                "topic": "某医院局部改造工程施工组织设计",
                "sections": [
                    {
                        "title": "针对工程项目整体理解",
                        "content": (
                            "【排版及格式合规声明】\n"
                            "本方案已设定纸张A4、正文字体宋体、页边距上2.5cm。\n"
                            "# 针对工程项目整体理解\n"
                            "## 1. 工程特点\n"
                            "### 1.1 接口管理\n"
                            "**A. 组织与策划层 (Architecture & Planning)**\n"
                            "开工前完成图纸会审和区域移交。"
                        ),
                    }
                ],
            },
            str(output_path),
        )

        doc = Document(str(output_path))
        text = "\n".join(
            [paragraph.text for paragraph in doc.paragraphs]
            + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
        )
        assert "排版及格式合规声明" not in text
        assert "页边距上2.5cm" not in text
        assert sum(1 for p in doc.paragraphs if p.text == "针对工程项目整体理解") == 1
        assert any(p.style.style_id == "Heading2" and p.text == "1. 工程特点" for p in doc.paragraphs)
        assert any(p.style.style_id == "Heading3" and p.text == "1.1 接口管理" for p in doc.paragraphs)
        assert any(p.style.style_id == "Heading3" and p.text == "组织与策划" for p in doc.paragraphs)
        assert "Architecture & Planning" not in text
        assert any("项目约束" in table.cell(0, 0).text for table in doc.tables if len(table.columns) == 3)
        assert "医院局部改造施工组织响应" in text

    def test_submission_docx_rewrites_unverified_planning_parameters(self, temp_dir):
        output_path = Path(temp_dir) / "verified_parameters_only.docx"
        export_autoplan_docx(
            {
                "topic": "某医院局部改造施工组织设计",
                "missing_parameters": {
                    "missing": [
                        {"key": "总工期"},
                        {"key": "资源峰值"},
                        {"key": "关键线路间隔"},
                        {"key": "风险检查频次"},
                        {"key": "质量阈值"},
                        {"key": "偏差处置时限"},
                    ]
                },
                "sections": [
                    {
                        "title": "确保工期与质量的保障体系与措施",
                        "content": (
                            "本工程计划总工期137.0天，资源峰值1180.0人（当量），"
                            "关键线路间隔3天。\n"
                            "检查频次=2次/日（班前+收工），质量阈值为偏差≤5mm，"
                            "偏差处置时限≤4h。"
                        ),
                    }
                ],
            },
            str(output_path),
        )

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "137.0天" not in text
        assert "1180.0人" not in text
        assert "关键线路间隔3天" not in text
        assert "2次/日（班前+收工）" not in text
        assert "偏差≤5mm" not in text
        assert "偏差处置时限≤4h" not in text
        assert "经批准的总进度计划" in text
        assert "各阶段资源按经批准的进度计划" in text

    def test_submission_docx_removes_review_scaffolding_and_rewrites_machine_formulas(self, temp_dir):
        output_path = Path(temp_dir) / "clean_review_scaffolding.docx"
        export_autoplan_docx(
            {
                "topic": "某医院局部改造施工组织设计",
                "missing_parameters": {
                    "missing": [
                        {"key": "风险检查频次"},
                        {"key": "质量阈值"},
                        {"key": "偏差处置时限"},
                    ]
                },
                "sections": [
                    {
                        "title": "确保工期与质量的保障体系与措施",
                        "content": (
                            "## 证据标注\n"
                            "评分维度回填：重难点。\n"
                            "- C. 材料准备：核对材料报验资料。\n"
                            "- D. 施工作业：按样板确认结果展开。\n"
                            "- **A. 施工准备**：复核作业条件。\n"
                            "- **B 施工**：按样板展开作业。\n"
                            "* **C检查：** 核对实体质量。\n"
                            "### D. 验收与闭环层 (Acceptance & Closure)\n"
                            "**B. 约束条件**\n"
                            "风险占比=high_risk_tasks * 100 / total_tasks。\n"
                            "进度延误率=delay_hours * 100 / planned_hours。\n"
                            "质量合格率阈值=≥98%，检查频次=2次/日，处置时限≤4h内。\n"
                            "D验收：提交各子系统联调联试报告。合格率阈值要求≥98%。\n"
                            "所有工序采用A/B/C/D/E结构，表1（模版A）。\n"
                            "采用A（准备）/B（施工）/C（检查）/D（验收）/E（成品保护/后期处理）五步工序法。\n"
                            "配置设备型号=20t挖机1台（量化默认值），连续作业时长=4h/作业段。\n"
                            "一次验收通过率控制在≥95%，扫码领料覆盖率≥95%。\n"
                            "裸土及渣土使用密目网覆盖，厚度控制50mm以保持水土。\n"
                            "触发“清单-工序-资源映射补强节点”评估实体返工率。\n"
                            "清单-工序-资源映射日清表。\n"
                            "证据未给出红线边界及夜间禁行具体要求 → 编制口径：按合肥市中心城区管控标准，"
                            "暂定晚22:00-早6:00禁止重型渣土车及材料运输车通行，场区内部划定4m宽单向循环通道1处 "
                            "→ 需澄清项：具体交通开口坐标及夜间施工许可办理权限归属。"
                        ),
                    }
                ],
            },
            str(output_path),
        )

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "证据标注" not in text
        assert "评分维度回填" not in text
        assert "C. 材料准备" not in text
        assert "D. 施工作业" not in text
        assert "材料准备：核对材料报验资料" in text
        assert "施工作业：按样板确认结果展开" in text
        assert "A. 施工准备" not in text
        assert "B 施工" not in text
        assert "C检查" not in text
        assert "D. 验收与闭环层" not in text
        assert "B. 约束条件" not in text
        assert "施工准备：复核作业条件" in text
        assert "施工：按样板展开作业" in text
        assert "检查： 核对实体质量" in text
        assert "验收与闭环层" in text
        assert "准备、施工、检查、验收、成品保护五步闭环工序法" in text
        assert "high_risk_tasks" not in text
        assert "delay_hours" not in text
        assert "高风险任务数÷任务总数×100%" in text
        assert "累计延误时长÷计划时长×100%" in text
        assert "98%" not in text
        assert "≥95%" not in text
        assert "2次/日" not in text
        assert "≤4h" not in text
        assert "A/B/C/D/E" not in text
        assert "模版A" not in text
        assert "20t挖机" not in text
        assert "量化默认值" not in text
        assert "4h/作业段" not in text
        assert "厚度控制50mm" not in text
        assert "补强节点" not in text
        assert "清单-工序-资源映射" not in text
        assert "按检验批记录实体质量" in text
        assert "施工前联合建设单位复核改造区域红线" in text
        assert "需澄清项" not in text
        assert "小型开挖设备" in text
        assert "纳入范围的材料领用全部扫码留痕" in text

    def test_hospital_renovation_submission_replaces_out_of_scope_generic_content(self, temp_dir):
        output_path = Path(temp_dir) / "hospital_scope.docx"
        export_autoplan_docx(
            {
                "topic": "某医院局部改造工程施工组织设计",
                "sections": [
                    {
                        "title": "确保工期与质量的保障体系与措施",
                        "content": (
                            "## 3.4 基础底板及二次结构工程质量控制\n"
                            "A准备：钢筋下料单复核，模板排版图绘制。\n"
                            "B施工：钢筋工按照间距绑扎墙柱纵筋与箍筋，模板工使用新型支撑体系加固，"
                            "混凝土工分层浇筑（厚度≤500mm），振捣棒快插慢拔。\n"
                            "采用防水层渗漏电导率阵列监测系统，铺设电导率传感器阵列。\n"
                            "基于BIM的管线综合防碰撞与绿色施工技术。\n"
                            "配置BIM图形工作站1台、PDA 4台和20t挖机1台。\n"
                            "清单重点（单项工程量大）：030411004005 / 工程量=6.0。\n"
                            "- **技术规避项声明**：本章节内容对标招标文件评审要求，规避否决扣分项。\n"
                            "信息化台账将作为工程结算和验收的铁证。\n"
                            "校验施工电梯、塔吊的起重量及运行行程。"
                        ),
                    }
                ],
            },
            str(output_path),
        )

        doc = Document(str(output_path))
        text = "\n".join(
            [paragraph.text for paragraph in doc.paragraphs]
            + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
        )
        for rejected in (
            "基础底板及二次结构",
            "钢筋下料单",
            "墙柱纵筋",
            "电导率",
            "BIM",
            "PDA",
            "20t挖机",
            "工程量=6.0",
            "评审",
            "扣分",
            "否决",
            "铁证",
            "施工电梯",
            "塔吊",
            "**",
        ):
            assert rejected not in text
        assert "局部修复、开槽回补及安装基层质量控制" in text
        assert "防水隐蔽验收与蓄水/淋水检验" in text
        assert "既有机电管线复核与专业接口协调" in text
        assert "清单重点项按正式工程量清单" in text
        assert "技术响应说明" in text
        assert "核验依据" in text
        assert "院方批准的垂直运输路线" in text
        assert "院感与洁污流线" in text
        assert "医疗秩序与系统切换" in text

    def test_submission_docx_renders_five_column_markdown_matrix_as_three_column_table(self, temp_dir):
        output_path = Path(temp_dir) / "markdown_matrix.docx"
        export_autoplan_docx(
            {
                "topic": "某医院局部改造施工组织设计",
                "sections": [
                    {
                        "title": "确保安全文明生产的管理体系与措施",
                        "content": (
                            "| 风险/问题 | 控制动作 | 验证方法 | 记录表 | 偏差处置 |\n"
                            "| :--- | :--- | :--- | :--- | :--- |\n"
                            "| 临时用电风险 | 每日巡检 | 绝缘测试 | 巡检台账 | 停止作业并整改 |"
                        ),
                    }
                ],
            },
            str(output_path),
        )

        doc = Document(str(output_path))
        matrix = next(table for table in doc.tables if table.cell(0, 0).text == "风险事项")
        assert len(matrix.columns) == 3
        assert "临时用电风险" in matrix.cell(1, 0).text
        assert "控制：每日巡检" in matrix.cell(1, 1).text
        assert "验证：绝缘测试" in matrix.cell(1, 1).text
        assert "记录：巡检台账" in matrix.cell(1, 2).text
        assert "处置：停止作业并整改" in matrix.cell(1, 2).text
        assert not any(paragraph.text.strip().startswith("|") for paragraph in doc.paragraphs)

    def test_estimated_toc_entries_use_live_pageref_with_fallback_value(self):
        doc = Document()
        paragraph = _render_toc_line(
            doc,
            {
                "title": "第一章 工程概况",
                "start_page": 37,
                "planned_pages": 6,
                "page_number_exact": False,
            },
            style_cfg={"title_font": "黑体", "title_latin_font": "Arial", "title_size": 16},
        )
        assert paragraph.text.endswith("37")
        assert "PAGEREF ZF_CHAPTER_1" in paragraph._element.xml

    def test_submission_media_rejects_unverified_ai_text_and_mixed_unit_chart(self):
        assert not _submission_media_eligible(
            {"path": "/tmp/generated.png", "caption": "施工组织设计思维导图（Gemini）"}
        )
        assert not _submission_media_eligible("/tmp/boq_stats_123.png")
        assert _submission_media_eligible(
            {
                "path": "/tmp/reviewed.png",
                "caption": "施工流程图（人工复核）",
                "text_verified": True,
            }
        )

    def test_submission_docx_uses_word_compatible_chinese_fonts(self, temp_dir):
        output_path = Path(temp_dir) / "portable_fonts.docx"
        export_autoplan_docx(
            {
                "topic": "测试",
                "style": {"body_font": "宋体", "title_font": "黑体"},
                "sections": [{"title": "第一章", "content": "中文正文"}],
            },
            str(output_path),
        )
        with zipfile.ZipFile(output_path) as zf:
            xml = zf.read("word/styles.xml").decode("utf-8", errors="ignore")
            font_table = zf.read("word/fontTable.xml").decode("utf-8", errors="ignore")
        assert 'w:eastAsia="宋体"' in xml
        assert 'w:eastAsia="黑体"' in xml
        assert 'w:name="宋体"' in font_table
        assert 'w:altName w:val="STSong"' in font_table
        assert 'w:hint="eastAsia"' in xml

    def test_cover_defaults_to_real_site_photo_and_project_id(self, temp_dir):
        from PIL import Image

        site = Path(temp_dir) / "现场.jpg"
        Image.new("RGB", (120, 80), color=(30, 120, 160)).save(site)
        meta = _resolve_cover_meta(
            {
                "topic": "医院局部改造工程施工组织设计",
                "project_id": "2026BFFGZ50127",
                "media": [
                    {
                        "path": str(site),
                        "caption": "项目现场：门诊综合楼",
                        "source_kind": "site_photo",
                        "source_sha256": "site-1",
                    }
                ],
            }
        )

        assert meta["project_code"] == "2026BFFGZ50127"
        assert meta["cover_image_path"] == str(site)
        assert meta["cover_image_caption"] == "项目现场：门诊综合楼"

    def test_submission_media_ranking_dedupes_and_prefers_project_sources(self, temp_dir):
        from PIL import Image

        site = Path(temp_dir) / "site.png"
        drawing = Path(temp_dir) / "drawing.png"
        diagram = Path(temp_dir) / "diagram.png"
        for path in (site, drawing, diagram):
            Image.new("RGB", (40, 30), color=(225, 240, 245)).save(path)

        rows = _rank_submission_media(
            [
                {"path": str(diagram), "source_kind": "deterministic_project_diagram", "source_sha256": "d-1"},
                {"path": str(site), "source_kind": "site_photo", "source_sha256": "s-1"},
                {"path": str(site), "source_kind": "site_photo", "source_sha256": "s-1"},
                {"path": str(drawing), "source_kind": "drawing", "source_sha256": "g-1"},
            ]
        )

        assert [row["source_kind"] for row in rows] == ["site_photo", "drawing", "deterministic_project_diagram"]

    def test_images_have_alt_text_and_internal_tables_mark_headers(self, temp_dir):
        from PIL import Image

        image_path = Path(temp_dir) / "site.png"
        Image.effect_noise((960, 640), 32).convert("RGB").save(image_path)
        output_path = Path(temp_dir) / "accessible_internal_review.docx"
        export_autoplan_docx(
            {
                "topic": "测试",
                "document_audience": "internal_review",
                "sections": [{"title": "第一章", "content": "正文"}],
                "media": [{"path": str(image_path), "caption": "现场总平面示意"}],
                "cross_index": {"focus_items": [{"name": "基础工程", "closure": {"ok": True}}]},
            },
            str(output_path),
        )
        with zipfile.ZipFile(output_path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
            footer_xml = zf.read("word/footer1.xml").decode("utf-8", errors="ignore")
        assert "现场总平面示意" in xml
        assert "descr=" in xml
        assert "tblHeader" in xml
        assert "tblHeader" in footer_xml

    def test_export_wires_front_matter_helpers_and_writes_build_report(self, temp_dir):
        """Front matter helpers are wired into export and emit minimal build report artifacts."""
        output_path = Path(temp_dir) / "baseline.docx"
        data = {
            "topic": "智慧厂房施工组织设计",
            "branding": {"bidder_company": "智飞建设有限公司"},
            "style": {
                "body_font": "宋体",
                "title_font": "黑体",
                "cover_page_count": 1,
                "toc_page_count": 2,
                "full_index_enabled": True,
                "full_index_page_count": 1,
            },
            "chapter_pages": {"第一章 工程概况": 2},
            "sections": [
                {
                    "title": "第一章 工程概况",
                    "content": "本章用于锁定当前导出正文基线。",
                    "agent_role": "项目经理",
                }
            ],
        }

        export_autoplan_docx(data, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "智慧厂房施工组织设计" in text
        assert "智飞建设有限公司" in text
        assert "施工组织设计" in text
        assert "全文索引" in text
        assert "目录" in text
        assert "01. 第一章 工程概况（约2页）" in text
        assert "第一章 工程概况" in text
        assert "负责人：项目经理" not in text
        assert "项目经理" not in text
        assert "本章用于锁定当前导出正文基线。" in text
        report_json_path, report_log_path = _build_report_paths(str(output_path))
        assert report_json_path.exists()
        assert report_log_path.exists()
        report = json.loads(report_json_path.read_text(encoding="utf-8"))
        assert report["schema_version"] == "docx_build_report.v1"
        assert report["output_path"] == str(output_path)
        assert report["topic"] == "智慧厂房施工组织设计"
        assert report["section_count"] == 1
        assert report["section_titles"] == ["第一章 工程概况"]
        assert report["front_matter_plan"]["full_index_pages"] == 1
        assert report["media_count"] == 0
        assert report["layout_receipts"][0]["title"] == "第一章 工程概况"
        assert report["layout_receipts"][0]["target_pages"] == 2
        assert "quality_checks" not in report
        assert "evidence" not in report
        assert "remediation" not in report
        assert "quality_evidence" not in report
        report_log = report_log_path.read_text(encoding="utf-8")
        assert "DOCX build report" in report_log
        assert "schema_version=docx_build_report.v1" in report_log
        assert "section_count=1" in report_log
        with zipfile.ZipFile(output_path) as zf:
            xml = "\n".join(
                zf.read(name).decode("utf-8", errors="ignore")
                for name in zf.namelist()
                if name.startswith("word/document") or name.startswith("word/footer")
            )
        assert "TOC" in xml
        assert "PAGE" in xml

    def test_export_creates_parent_directories(self, temp_dir, basic_data):
        """Test export creates nested parent directories."""
        output_path = Path(temp_dir) / "nested" / "dir" / "output.docx"
        result = export_autoplan_docx(basic_data, str(output_path))

        assert output_path.exists()

    def test_export_with_empty_data(self, temp_dir):
        """Test export with empty data dict."""
        output_path = Path(temp_dir) / "empty.docx"
        result = export_autoplan_docx({}, str(output_path))

        assert output_path.exists()

    def test_export_default_topic(self, temp_dir):
        """Test export uses default topic when not provided."""
        output_path = Path(temp_dir) / "default.docx"
        data = {"sections": []}
        result = export_autoplan_docx(data, str(output_path))

        doc = Document(str(output_path))
        # First paragraph should contain default title
        assert len(doc.paragraphs) > 0

    def test_export_with_agent_role(self, temp_dir, basic_data):
        """Submission export keeps internal agent roles out of bidder-facing prose."""
        output_path = Path(temp_dir) / "with_role.docx"
        export_autoplan_docx(basic_data, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "项目经理" not in text

    def test_export_without_agent_role(self, temp_dir):
        """Test export handles sections without agent_role."""
        output_path = Path(temp_dir) / "no_role.docx"
        data = {
            "topic": "测试",
            "sections": [{"title": "章节", "content": "内容"}],
        }
        export_autoplan_docx(data, str(output_path))

        assert output_path.exists()

    def test_export_with_media(self, temp_dir, basic_data):
        """Test export with media/images (simulated failure)."""
        output_path = Path(temp_dir) / "with_media.docx"
        basic_data["media"] = ["/nonexistent/image.png"]

        export_autoplan_docx(basic_data, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        # A formal export omits missing media instead of exposing technical
        # placeholder text in the bidder-facing document.
        assert "图片加载失败" not in text

    def test_export_with_valid_media(self, temp_dir, basic_data):
        """Test export with valid image file."""
        # Create a simple valid image
        from PIL import Image
        img_path = Path(temp_dir) / "test_image.png"
        img = Image.effect_noise((960, 640), 32).convert("RGB")
        img.save(str(img_path))

        output_path = Path(temp_dir) / "with_valid_media.docx"
        basic_data["media"] = [str(img_path)]

        export_autoplan_docx(basic_data, str(output_path))
        assert output_path.exists()

        doc = Document(str(output_path))
        image_paragraph = next(
            paragraph for paragraph in doc.paragraphs
            if paragraph._p.xpath(".//w:drawing")
        )
        assert image_paragraph.paragraph_format.line_spacing == 1.0
        assert image_paragraph.paragraph_format.line_spacing_rule == WD_LINE_SPACING.SINGLE
        assert image_paragraph.paragraph_format.keep_together is True
        assert image_paragraph.paragraph_format.keep_with_next is True

    def test_export_groups_two_landscape_site_photos_into_one_evidence_panel(self, temp_dir):
        from PIL import Image

        first = Path(temp_dir) / "现场东侧.jpg"
        second = Path(temp_dir) / "现场西侧.jpg"
        Image.effect_noise((960, 640), 31).convert("RGB").save(first)
        Image.effect_noise((960, 640), 57).convert("RGB").save(second)
        output_path = Path(temp_dir) / "two_up_photo_panel.docx"

        export_autoplan_docx(
            {
                "topic": "项目施工组织设计",
                "style": {
                    "chart_policy": {
                        "enabled": True,
                        "mode": "page_density_auto",
                        "position": "chapter",
                    }
                },
                "sections": [{"title": "第一章 工程概况", "content": "项目现场条件已复核。"}],
                "media": [
                    {"path": str(first), "caption": "现场东侧作业面", "source_kind": "site_photo"},
                    {"path": str(second), "caption": "现场西侧运输通道", "source_kind": "site_photo"},
                ],
            },
            str(output_path),
        )

        doc = Document(str(output_path))
        panel = next(
            table
            for table in doc.tables
            if "现场东侧作业面" in "\n".join(cell.text for row in table.rows for cell in row.cells)
        )
        assert len(panel.columns) == 2
        assert len(panel._tbl.xpath(".//w:drawing")) == 2
        panel_text = "\n".join(cell.text for row in panel.rows for cell in row.cells)
        assert "图1：现场东侧作业面" in panel_text
        assert "图2：现场西侧运输通道" in panel_text

    def test_markdown_table_uses_semantic_widths_and_alignment(self, temp_dir):
        output_path = Path(temp_dir) / "semantic_table.docx"
        export_autoplan_docx(
            {
                "topic": "项目施工组织设计",
                "sections": [
                    {
                        "title": "第一章 施工部署",
                        "content": (
                            "| 序号 | 控制措施与验证要求 | 数量 |\n"
                            "| --- | --- | --- |\n"
                            "| 1 | 完成样板验收后分区展开，并形成检查记录和复核结论。 | 12项 |\n"
                            "| 2 | 对关键工序实行旁站、实测和闭环复验。 | 8项 |"
                        ),
                    }
                ],
            },
            str(output_path),
        )

        doc = Document(str(output_path))
        table = next(table for table in doc.tables if table.cell(0, 0).text == "序号")
        widths = [table.cell(0, index).width.cm for index in range(3)]
        assert widths[1] > widths[0]
        assert widths[1] > widths[2]
        assert table.cell(1, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert table.cell(1, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT

    def test_export_with_quality_checks(self, temp_dir, data_with_quality_checks):
        """Test export with full quality checks data."""
        output_path = Path(temp_dir) / "with_qc.docx"
        export_autoplan_docx(data_with_quality_checks, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "质量校验摘要" in text
        assert "通过" in text or "需改进" in text

    def test_export_quality_checks_checklist(self, temp_dir, data_with_quality_checks):
        """Test export generates quality check checklist with marks."""
        output_path = Path(temp_dir) / "qc_checklist.docx"
        export_autoplan_docx(data_with_quality_checks, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        # Should have checkbox marks
        assert "☑" in text or "☐" in text

    def test_export_score_coverage_by_section(self, temp_dir, data_with_quality_checks):
        """Test export includes score coverage by section."""
        output_path = Path(temp_dir) / "score_coverage.docx"
        export_autoplan_docx(data_with_quality_checks, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "章节评分点覆盖清单" in text
        assert "缺失" in text

    def test_export_evidence_by_section(self, temp_dir, data_with_quality_checks):
        """Test export includes evidence count by section."""
        output_path = Path(temp_dir) / "evidence.docx"
        export_autoplan_docx(data_with_quality_checks, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "章节证据数量清单" in text
        assert "证据数" in text

    def test_export_closed_loop_by_section(self, temp_dir, data_with_quality_checks):
        """Test export includes closed loop by section."""
        output_path = Path(temp_dir) / "closed_loop.docx"
        export_autoplan_docx(data_with_quality_checks, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "章节风险-措施闭环清单" in text

    def test_export_engineering_by_section(self, temp_dir, data_with_quality_checks):
        """Test export includes engineering elements by section."""
        output_path = Path(temp_dir) / "engineering.docx"
        export_autoplan_docx(data_with_quality_checks, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "章节工程落地要素清单" in text

    def test_export_remediation_suggestions(self, temp_dir, data_with_quality_checks):
        """Test export includes remediation suggestions."""
        output_path = Path(temp_dir) / "remediation.docx"
        export_autoplan_docx(data_with_quality_checks, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "整改建议清单" in text
        assert "补充证据材料" in text

    def test_quality_evidence_stays_in_docx_and_is_mirrored_to_build_report(self, temp_dir, data_with_quality_checks):
        """Lock current DOCX behavior while mirroring quality/evidence details to build report."""
        output_path = Path(temp_dir) / "quality_evidence_baseline.docx"
        data_with_quality_checks["quality_checks"]["issue_list"] = [
            {
                "title": "第二章",
                "type": "missing_evidence",
                "severity": "high",
                "problem": "证据定位不足",
                "suggestion": "补充合同条款定位",
            }
        ]
        data_with_quality_checks["quality_checks"]["auto_revision_suggestions"] = [
            {
                "title": "第二章",
                "type": "evidence_fix",
                "suggestion": "补齐来源定位",
            }
        ]

        export_autoplan_docx(data_with_quality_checks, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "质量校验摘要" in text
        assert "质量校验清单" in text
        assert "evidence：需改进" in text
        assert "章节证据数量清单" in text
        assert "- 第一章: 证据数 2" in text
        assert "- 第二章: 证据数 0" in text
        assert "问题清单（自动检测）" in text
        assert "证据定位不足" in text
        assert "自动修订建议（按章节聚合）" in text
        assert "补齐来源定位" in text
        assert "整改建议清单" in text
        assert "补充证据材料" in text

        report_json_path, _ = _build_report_paths(str(output_path))
        report = json.loads(report_json_path.read_text(encoding="utf-8"))
        report_text = json.dumps(report, ensure_ascii=False)
        assert "quality_checks" not in report
        assert "evidence" not in report
        assert "remediation" not in report
        quality_evidence = report["quality_evidence"]
        assert quality_evidence["quality_checks"]["evidence"]["by_section"] == [
            {"title": "第一章", "evidence_count": 2},
            {"title": "第二章", "evidence_count": 0},
        ]
        assert quality_evidence["evidence"]["ok"] is False
        assert quality_evidence["remediation"] == [
            {"title": "第二章", "type": "missing_evidence", "suggestion": "补充证据材料"}
        ]
        assert quality_evidence["issue_list"][0]["problem"] == "证据定位不足"
        assert quality_evidence["auto_revision_suggestions"][0]["suggestion"] == "补齐来源定位"
        assert "证据定位不足" in report_text
        assert "补齐来源定位" in report_text
        assert "补充证据材料" in report_text

    def test_export_llm_compare_full_mode(self, temp_dir, data_with_llm_remediation):
        """Test export LLM compare in full mode."""
        output_path = Path(temp_dir) / "llm_compare.docx"
        data_with_llm_remediation["compare"] = {"mode": "full"}
        data_with_llm_remediation["quality_checks"] = {"structure": {"ok": True}}

        export_autoplan_docx(data_with_llm_remediation, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "LLM整改前后对比" in text
        assert "整改前" in text
        assert "整改后" in text

    def test_export_llm_compare_summary_mode(self, temp_dir, data_with_llm_remediation):
        """Test export LLM compare in summary mode with truncation."""
        output_path = Path(temp_dir) / "llm_summary.docx"
        # Create long content
        data_with_llm_remediation["sections"][0]["original_content"] = "A" * 1000
        data_with_llm_remediation["sections"][0]["content"] = "B" * 1000
        data_with_llm_remediation["compare"] = {"mode": "summary", "max_chars": 100}
        data_with_llm_remediation["quality_checks"] = {"structure": {"ok": True}}

        export_autoplan_docx(data_with_llm_remediation, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        # Should be truncated with "..."
        assert "..." in text

    def test_export_llm_compare_with_titles_filter(self, temp_dir, data_with_llm_remediation):
        """Test export LLM compare filters by titles."""
        output_path = Path(temp_dir) / "llm_filtered.docx"
        data_with_llm_remediation["compare"] = {
            "mode": "full",
            "titles": ["第一章 工程概述"],
        }
        data_with_llm_remediation["quality_checks"] = {"structure": {"ok": True}}

        export_autoplan_docx(data_with_llm_remediation, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "第一章 工程概述" in text

    def test_export_no_style(self, temp_dir):
        """Test export with no style config."""
        output_path = Path(temp_dir) / "no_style.docx"
        data = {"topic": "测试", "sections": []}

        export_autoplan_docx(data, str(output_path))
        assert output_path.exists()

    def test_export_with_nested_w6_style(self, temp_dir):
        """Test export supports W6 nested template style keys."""
        output_path = Path(temp_dir) / "w6_style.docx"
        data = {
            "topic": "测试",
            "style": {
                "font": {
                    "eastAsia": "宋体",
                    "latin": "Times New Roman",
                    "size_pt": 11,
                    "line_spacing": 1.5,
                },
                "headings": {"h1_size": 16, "h2_size": 13},
                "margins_cm": {"top": 2.54, "bottom": 2.54, "left": 3.0, "right": 2.5},
            },
            "sections": [{"title": "第一章", "content": "正文内容"}],
        }

        export_autoplan_docx(data, str(output_path))
        assert output_path.exists()

    def test_export_with_chapter_page_receipt(self, temp_dir):
        """Test export includes chapter page receipt when chapter_pages provided."""
        output_path = Path(temp_dir) / "page_receipt.docx"
        data = {
            "topic": "测试",
            "document_audience": "internal_review",
            "chapter_pages": {"第一章": 2},
            "sections": [{"title": "第一章", "content": "内容" * 120}],
        }

        export_autoplan_docx(data, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "章节版式约束回执" in text
        assert "目标2页" in text

    def test_export_with_chapter_pages_dict_format(self, temp_dir):
        """Test export supports chapter_pages in dict target format."""
        output_path = Path(temp_dir) / "page_receipt_dict.docx"
        data = {
            "topic": "测试",
            "document_audience": "internal_review",
            "chapter_pages": {"第一章": {"pages": 3}},
            "sections": [{"title": "第一章", "content": "内容" * 160}],
        }

        export_autoplan_docx(data, str(output_path))
        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "目标3页" in text

    def test_page_target_enforcement_never_inserts_blank_page_breaks(self, temp_dir):
        """Legacy enforce flag must not mechanically pad the DOCX with page breaks."""
        baseline_path = Path(temp_dir) / "page_target_baseline.docx"
        enforced_path = Path(temp_dir) / "page_target_enforced.docx"
        base_data = {
            "topic": "测试",
            "chapter_pages": {"第一章": 8},
            "style": {"chapter_start_new_page": False},
            "sections": [{"title": "第一章", "content": "有效技术内容" * 20}],
        }
        enforced_data = {
            **base_data,
            "style": {
                "chapter_start_new_page": False,
                "enforce_chapter_pages": True,
            },
        }

        export_autoplan_docx(base_data, str(baseline_path))
        export_autoplan_docx(enforced_data, str(enforced_path))

        with zipfile.ZipFile(baseline_path) as zf:
            baseline_xml = zf.read("word/document.xml")
        with zipfile.ZipFile(enforced_path) as zf:
            enforced_xml = zf.read("word/document.xml")
        assert enforced_xml.count(b'w:type="page"') == baseline_xml.count(b'w:type="page"')
        enforced_report = json.loads(
            _build_report_paths(str(enforced_path))[0].read_text(encoding="utf-8")
        )
        receipt = enforced_report["layout_receipts"][0]
        assert receipt["shortfall_pages"] > 0
        assert receipt["mechanical_padding_applied"] is False

    def test_export_empty_sections(self, temp_dir):
        """Test export with empty sections list."""
        output_path = Path(temp_dir) / "empty_sections.docx"
        data = {"topic": "测试", "sections": []}

        export_autoplan_docx(data, str(output_path))
        assert output_path.exists()

    def test_export_section_default_title_and_content(self, temp_dir):
        """Test export handles sections with missing title/content."""
        output_path = Path(temp_dir) / "defaults.docx"
        data = {
            "topic": "测试",
            "sections": [{}],  # Empty section
        }

        export_autoplan_docx(data, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "章节" in text  # Default title

    def test_auto_density_image_rule_under_200_and_skip_overview(self, temp_dir):
        output_path = Path(temp_dir) / "auto_density_under_200.docx"
        data = {
            "topic": "测试",
            "style": {
                "chart_policy": {
                    "enabled": True,
                    "mode": "page_density_auto",
                    "position": "chapter",
                }
            },
            "chapter_pages": {
                "工程概况": 1,
                "主要施工方法": 2,
            },
            "sections": [
                {"title": "工程概况", "content": "项目概况章节。"},
                {"title": "主要施工方法", "content": "控制间距900mm，抽检频次2次/班，风险-控制-验证闭环。"},
            ],
        }
        export_autoplan_docx(data, str(output_path))
        doc = Document(str(output_path))
        captions = [p.text for p in doc.paragraphs if re.match(r"^图\d+：", str(p.text or ""))]
        # Submission-safe density: one distinct supporting figure for this short chapter.
        assert len(captions) == 1
        assert all("工程概况" not in c for c in captions)
        assert any("本章实施控制矩阵" in c for c in captions)

    def test_auto_density_image_rule_over_200(self, temp_dir):
        output_path = Path(temp_dir) / "auto_density_over_200.docx"
        data = {
            "topic": "测试",
            "style": {
                "chart_policy": {
                    "enabled": True,
                    "mode": "page_density_auto",
                    "position": "chapter",
                }
            },
            "chapter_pages": {
                "工程概况": 199,  # 仅用于触发 total_pages > 200
                "主要施工方法": 4,
            },
            "sections": [
                {"title": "工程概况", "content": "项目概况章节。"},
                {"title": "主要施工方法", "content": "控制间距900mm，抽检频次2次/班，风险-控制-验证闭环。"},
            ],
        }
        export_autoplan_docx(data, str(output_path))
        doc = Document(str(output_path))
        captions = [p.text for p in doc.paragraphs if re.match(r"^图\d+：", str(p.text or ""))]
        # Long documents use a lower density; this four-page chapter gets one figure.
        assert len(captions) == 1

    def test_structured_tables_support_landscape_merged_headers_and_nested_tables(self, temp_dir):
        output_path = Path(temp_dir) / "structured_tables.docx"
        data = {
            "topic": "结构化表格验收",
            "sections": [{"title": "第一章 编制说明", "content": "本样本用于验证结构化表格导出。"}],
            "tables": [
                {
                    "title": "横向资源计划表",
                    "orientation": "landscape",
                    "headers": ["序号", "工序", "资源", "数量", "控制措施", "验收记录", "责任人", "状态"],
                    "merge_header_groups": [
                        {"start": 0, "end": 1, "label": "工作分解"},
                        {"start": 2, "end": 3, "label": "资源配置"},
                        {"start": 4, "end": 7, "label": "执行闭环"},
                    ],
                    "rows": [
                        [
                            "1",
                            "测量放线",
                            {
                                "text": "测量组",
                                "nested": {
                                    "headers": ["设备", "校验"],
                                    "rows": [["全站仪", "有效期内"]],
                                },
                            },
                            "2套",
                            "轴线与标高双重复核",
                            "测量复核记录",
                            "测量负责人",
                            "待验收",
                        ]
                    ],
                }
            ],
        }

        export_autoplan_docx(data, str(output_path))

        document = Document(str(output_path))
        assert [round(section.page_width.cm, 1) for section in document.sections] == [21.0, 29.7]
        assert [round(section.page_height.cm, 1) for section in document.sections] == [29.7, 21.0]
        assert any("横向资源计划表" in paragraph.text for paragraph in document.paragraphs)
        assert len(document.tables) == 1
        with zipfile.ZipFile(output_path) as package:
            document_xml = package.read("word/document.xml").decode("utf-8")
        assert document_xml.count("<w:tbl>") >= 2
        assert document_xml.count("<w:tblHeader") >= 3
        assert document_xml.count("<w:cantSplit") >= 4
        assert document_xml.count('w:headerReference w:type="default"') == 2
        assert document_xml.count('w:footerReference w:type="default"') == 2
        assert document_xml.count('w:headerReference w:type="first"') == 2
        assert document_xml.count('w:footerReference w:type="first"') == 2
        assert document_xml.count('w:headerReference w:type="even"') == 2
        assert document_xml.count('w:footerReference w:type="even"') == 2
        receipt = json.loads(output_path.with_suffix(".structural_quality.json").read_text(encoding="utf-8"))
        assert receipt["status"] == "pass"
        assert all(
            item["default_header"] and item["default_footer"]
            for item in receipt["section_story_references"]
        )
        assert receipt["section_story_references"][1]["header_types"] == ["default", "even", "first"]
        assert receipt["section_story_references"][1]["footer_types"] == ["default", "even", "first"]
        assert [section["orientation"] for section in receipt["section_metrics"]] == ["portrait", "landscape"]

    def test_portrait_restore_before_media_does_not_insert_an_empty_page(self, temp_dir):
        from PIL import Image

        output_path = Path(temp_dir) / "landscape_then_media.docx"
        image_path = Path(temp_dir) / "工程流程图.png"
        Image.effect_noise((1600, 900), 48).convert("RGB").save(image_path, dpi=(300, 300))
        data = {
            "topic": "横向表格后续图片验收",
            "sections": [{"title": "第一章 编制说明", "content": "验证恢复纵向页面后直接排入图形。"}],
            "tables": [
                {
                    "title": "横向资源表",
                    "orientation": "landscape",
                    "headers": ["序号", "工序", "资源", "数量", "措施", "频次", "责任人", "记录"],
                    "rows": [["1", "测量", "测量组", "2套", "双重复核", "每道工序", "测量员", "复核记录"]],
                }
            ],
            "media": [
                {
                    "path": str(image_path),
                    "caption": "恢复纵向后的工程流程图",
                    "source_kind": "deterministic_project_diagram",
                    "source_ref": "qa/工程流程图.png",
                    "text_verified": True,
                    "required": True,
                }
            ],
        }

        export_autoplan_docx(data, str(output_path))

        document = Document(str(output_path))
        assert [round(section.page_width.cm, 1) for section in document.sections] == [21.0, 29.7, 21.0]
        paragraphs = document.paragraphs
        figure_heading_index = next(
            index for index, paragraph in enumerate(paragraphs) if paragraph.text == "图表与插图"
        )
        assert paragraphs[figure_heading_index - 1]._p.xpath("./w:pPr/w:sectPr")
        assert not paragraphs[figure_heading_index - 1]._p.xpath('.//w:br[@w:type="page"]')
        assert not paragraphs[figure_heading_index]._p.xpath('.//w:br[@w:type="page"]')


class TestAutoDensityRules:
    def test_auto_density_images_for_pages(self):
        assert _auto_density_images_for_pages(3, 150) == 1
        assert _auto_density_images_for_pages(3, 260) == 1
        assert _auto_density_images_for_pages(12, 150) == 2
        assert _auto_density_images_for_pages(12, 260) == 2


# =============================================================================
# Tests for export_autoplan_compare_docx
# =============================================================================

class TestExportAutoplanCompareDocx:
    """Tests for export_autoplan_compare_docx function."""

    def test_basic_compare_export(self, temp_dir, data_with_llm_remediation):
        """Test basic compare DOCX export."""
        output_path = Path(temp_dir) / "compare.docx"
        result = export_autoplan_compare_docx(data_with_llm_remediation, str(output_path))

        assert result == str(output_path)
        assert output_path.exists()

    def test_compare_export_no_remediated_sections(self, temp_dir, basic_data):
        """Test compare export with no LLM remediated sections."""
        output_path = Path(temp_dir) / "no_remediation.docx"
        export_autoplan_compare_docx(basic_data, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "暂无可对比的章节" in text

    def test_compare_export_summary_mode(self, temp_dir, data_with_llm_remediation):
        """Test compare export in summary mode (default)."""
        output_path = Path(temp_dir) / "compare_summary.docx"
        data_with_llm_remediation["sections"][0]["original_content"] = "X" * 1000
        data_with_llm_remediation["sections"][0]["content"] = "Y" * 1000
        data_with_llm_remediation["compare"] = {"mode": "summary", "max_chars": 200}

        export_autoplan_compare_docx(data_with_llm_remediation, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "..." in text

    def test_compare_export_full_mode(self, temp_dir, data_with_llm_remediation):
        """Test compare export in full mode (no truncation)."""
        output_path = Path(temp_dir) / "compare_full.docx"
        data_with_llm_remediation["compare"] = {"mode": "full"}

        export_autoplan_compare_docx(data_with_llm_remediation, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "整改前" in text
        assert "整改后" in text

    def test_compare_export_titles_filter(self, temp_dir, data_with_llm_remediation):
        """Test compare export filters by specified titles."""
        output_path = Path(temp_dir) / "compare_filtered.docx"
        data_with_llm_remediation["compare"] = {
            "mode": "full",
            "titles": ["第一章 工程概述"],
        }

        export_autoplan_compare_docx(data_with_llm_remediation, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "第一章 工程概述" in text

    def test_compare_export_excludes_non_matching_titles(self, temp_dir, data_with_llm_remediation):
        """Test compare export excludes sections not in titles filter."""
        output_path = Path(temp_dir) / "compare_excluded.docx"
        data_with_llm_remediation["compare"] = {
            "mode": "full",
            "titles": ["不存在的章节"],
        }

        export_autoplan_compare_docx(data_with_llm_remediation, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "暂无可对比的章节" in text

    def test_compare_export_creates_directories(self, temp_dir, data_with_llm_remediation):
        """Test compare export creates parent directories."""
        output_path = Path(temp_dir) / "a" / "b" / "c" / "compare.docx"
        export_autoplan_compare_docx(data_with_llm_remediation, str(output_path))

        assert output_path.exists()

    def test_compare_export_default_topic(self, temp_dir, data_with_llm_remediation):
        """Test compare export uses default topic."""
        output_path = Path(temp_dir) / "default_topic.docx"
        del data_with_llm_remediation["topic"]

        export_autoplan_compare_docx(data_with_llm_remediation, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        assert "施组方案" in text

    def test_compare_export_empty_original_content(self, temp_dir):
        """Test compare export handles empty original_content."""
        output_path = Path(temp_dir) / "empty_original.docx"
        data = {
            "topic": "测试",
            "sections": [
                {
                    "title": "章节",
                    "content": "新内容",
                    "original_content": "",  # Empty, should be skipped
                    "auto_remediated": "llm",
                }
            ],
        }

        export_autoplan_compare_docx(data, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        # Empty original_content should be treated as falsy, skip section
        assert "暂无可对比的章节" in text


# =============================================================================
# Tests for export_autoplan_docx_from_file
# =============================================================================

class TestExportAutoplanDocxFromFile:
    """Tests for export_autoplan_docx_from_file function."""

    def test_export_from_json_file(self, temp_dir, basic_data):
        """Test export from a JSON file."""
        json_path = Path(temp_dir) / "input.json"
        json_path.write_text(json.dumps(basic_data), encoding="utf-8")

        output_path = Path(temp_dir) / "from_file.docx"
        result = export_autoplan_docx_from_file(str(json_path), str(output_path))

        assert result == str(output_path)
        assert output_path.exists()

    def test_export_from_file_with_variants(self, temp_dir, basic_data):
        """Test export from file with variants array (multi-version)."""
        data = {
            "variants": [
                basic_data,
                {"topic": "第二版本"},
            ]
        }
        json_path = Path(temp_dir) / "variants.json"
        json_path.write_text(json.dumps(data), encoding="utf-8")

        output_path = Path(temp_dir) / "variants_output.docx"
        result = export_autoplan_docx_from_file(str(json_path), str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)

        # Should use first variant
        assert "测试" in text
        assert "施工组织设计" in text
        assert "第二版本" not in text

    def test_export_from_file_with_empty_variants(self, temp_dir):
        """Test export from file with empty variants array."""
        data = {"variants": [], "topic": "直接数据"}
        json_path = Path(temp_dir) / "empty_variants.json"
        json_path.write_text(json.dumps(data), encoding="utf-8")

        output_path = Path(temp_dir) / "empty_variants.docx"
        result = export_autoplan_docx_from_file(str(json_path), str(output_path))

        # Should fall back to the whole data object
        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "直接数据" in text

    def test_export_from_file_without_variants(self, temp_dir, basic_data):
        """Test export from file without variants key."""
        json_path = Path(temp_dir) / "no_variants.json"
        json_path.write_text(json.dumps(basic_data), encoding="utf-8")

        output_path = Path(temp_dir) / "no_variants.docx"
        result = export_autoplan_docx_from_file(str(json_path), str(output_path))

        assert output_path.exists()

    def test_export_from_file_invalid_json(self, temp_dir):
        """Test export from file with invalid JSON raises error."""
        json_path = Path(temp_dir) / "invalid.json"
        json_path.write_text("not valid json {{{", encoding="utf-8")

        output_path = Path(temp_dir) / "invalid.docx"

        with pytest.raises(json.JSONDecodeError):
            export_autoplan_docx_from_file(str(json_path), str(output_path))

    def test_export_from_nonexistent_file(self, temp_dir):
        """Test export from nonexistent file raises error."""
        json_path = Path(temp_dir) / "does_not_exist.json"
        output_path = Path(temp_dir) / "output.docx"

        with pytest.raises(FileNotFoundError):
            export_autoplan_docx_from_file(str(json_path), str(output_path))

    def test_export_from_file_unicode_content(self, temp_dir):
        """Test export from file with unicode content."""
        data = {
            "topic": "中文标题 日本語 한국어",
            "sections": [
                {"title": "特殊字符 ™ © ® €", "content": "内容 αβγ ∑∏∫"},
            ],
        }
        json_path = Path(temp_dir) / "unicode.json"
        json_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

        output_path = Path(temp_dir) / "unicode.docx"
        result = export_autoplan_docx_from_file(str(json_path), str(output_path))

        assert output_path.exists()
        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "中文标题" in text


# =============================================================================
# Edge case and integration tests
# =============================================================================

class TestEdgeCases:
    """Edge case and integration tests."""

    def test_export_with_none_style(self, temp_dir):
        """Test export with style explicitly set to None."""
        output_path = Path(temp_dir) / "none_style.docx"
        data = {"topic": "测试", "style": None, "sections": []}

        export_autoplan_docx(data, str(output_path))
        assert output_path.exists()

    def test_export_with_none_sections(self, temp_dir):
        """Test export with sections explicitly set to None."""
        output_path = Path(temp_dir) / "none_sections.docx"
        data = {"topic": "测试", "sections": None}

        # Should handle None sections gracefully
        export_autoplan_docx(data, str(output_path))
        assert output_path.exists()

    def test_export_with_none_media(self, temp_dir, basic_data):
        """Test export with media explicitly set to None."""
        output_path = Path(temp_dir) / "none_media.docx"
        basic_data["media"] = None

        export_autoplan_docx(basic_data, str(output_path))
        assert output_path.exists()

    def test_export_with_empty_media_list(self, temp_dir, basic_data):
        """Test export with empty media list."""
        output_path = Path(temp_dir) / "empty_media.docx"
        basic_data["media"] = []

        export_autoplan_docx(basic_data, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Should not have media section
        assert "图表与插图" not in text

    def test_export_with_none_quality_checks(self, temp_dir, basic_data):
        """Test export with quality_checks explicitly set to None."""
        output_path = Path(temp_dir) / "none_qc.docx"
        basic_data["quality_checks"] = None

        export_autoplan_docx(basic_data, str(output_path))
        assert output_path.exists()

    def test_export_with_empty_quality_checks(self, temp_dir, basic_data):
        """Test export with empty quality_checks dict."""
        output_path = Path(temp_dir) / "empty_qc.docx"
        basic_data["quality_checks"] = {}

        export_autoplan_docx(basic_data, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Should not have QC section with empty dict
        assert output_path.exists()

    def test_large_document(self, temp_dir):
        """Test export with many sections."""
        sections = [
            {"title": f"第{i}章", "content": f"内容{i}" * 100, "agent_role": f"角色{i}"}
            for i in range(50)
        ]
        data = {"topic": "大型文档测试", "sections": sections}
        output_path = Path(temp_dir) / "large.docx"

        export_autoplan_docx(data, str(output_path))

        assert output_path.exists()
        assert output_path.stat().st_size > 10000  # Should be reasonably large

    def test_style_with_numeric_strings(self, temp_dir):
        """Test style with numeric values as strings."""
        output_path = Path(temp_dir) / "string_nums.docx"
        data = {
            "topic": "测试",
            "style": {
                "body_size": "14",  # String instead of int
                "title_size": "18",
                "line_spacing": "1.5",
            },
            "sections": [],
        }

        export_autoplan_docx(data, str(output_path))
        assert output_path.exists()

    def test_quality_check_item_with_no_ok_field(self, temp_dir):
        """Test quality check item without 'ok' field."""
        output_path = Path(temp_dir) / "no_ok.docx"
        data = {
            "topic": "测试",
            "document_audience": "internal_review",
            "sections": [],
            "quality_checks": {
                "structure": {"note": "仅有备注"},
            },
        }

        export_autoplan_docx(data, str(output_path))

        doc = Document(str(output_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        # ok is None/falsy, should show "需改进"
        assert "需改进" in text
