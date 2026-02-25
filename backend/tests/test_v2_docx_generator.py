from __future__ import annotations

from pathlib import Path

from docx import Document
from PIL import Image

from backend.zhifei_autoplan.v2.docx_generator import generate_v2_docx


def test_generate_v2_docx_with_visual_guardrails(tmp_path: Path) -> None:
    index_matrix = {
        "project_name": "测试项目",
        "index_matrix": [
            {"dimension": "质量", "keywords": ["质量", "验收"]},
            {"dimension": "安全", "keywords": ["安全", "防护"]},
        ],
    }
    sections = [
        {
            "title": "质量",
            "content": "执行质量控制，抽检频次2次/日，质量员复核。",
            "auto_generated_support": True,
            "graph_hit": {
                "title": "AUTO-QUALITY",
                "source_file": "self_healing_patch_nodes.json",
                "source_path": "/tmp/self_healing_patch_nodes.json",
                "resource_requirements": {"sampling_per_day": 2, "checker": "质量员"},
            },
        },
        {
            "title": "安全",
            "content": "执行安全控制，每班次检查2次，安全员复核。",
            "auto_generated_support": False,
            "graph_hit": {
                "title": "N-SAFE",
                "source_file": "kg.json",
                "source_path": "/tmp/kg.json",
                "resource_requirements": {"inspection_per_shift": 2, "checker": "安全员"},
            },
        },
    ]

    image_path = tmp_path / "visual.png"
    Image.new("RGB", (300, 180), color="#005BAC").save(image_path)
    visual_assets = [
        {
            "asset_id": "VIS-01",
            "title": "流程图",
            "caption": "流程图（GB/T 50104, CSCEC VI蓝绿灰, 仿宋）",
            "image_path": str(image_path),
        }
    ]

    out = tmp_path / "final.docx"
    result = generate_v2_docx(
        index_matrix=index_matrix,
        sections=sections,
        visual_assets=visual_assets,
        output_path=out,
    )

    assert result["ok"] is True
    assert out.exists()
    assert result["auto_generated_sections"] >= 1
    assert result["highlighted_paragraphs"] >= 1
    assert result["visual_assets_embedded"] == 1
    assert result["visual_assets_missing"] == 0

    doc = Document(str(out))
    all_text = "\n".join([p.text for p in doc.paragraphs])
    assert "1. 质量" in all_text
    assert "2. 安全" in all_text
    assert "附图与视觉生成" in all_text

    has_highlight = False
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color is not None:
                has_highlight = True
                break
        if has_highlight:
            break
    assert has_highlight is True
