from __future__ import annotations

import time
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _setup_doc_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _is_auto_generated_section(section: Dict[str, Any]) -> bool:
    if not isinstance(section, dict):
        return False
    if bool(section.get("auto_generated_support")):
        return True
    graph_hit = section.get("graph_hit")
    if not isinstance(graph_hit, dict):
        return False
    source_path = str(graph_hit.get("source_path") or "").lower()
    if "self_healing_patch_nodes" in source_path:
        return True
    snippet = str(graph_hit.get("snippet") or "").lower()
    return "is_auto_generated" in snippet


def _render_section_paragraph(
    doc: Document,
    *,
    text: str,
    highlight: bool,
) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(str(text or "").strip() or "（无内容）")
    if highlight:
        # Visual guardrail: mark sentences backed by self-healing auto-generated KG nodes.
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _resource_items(resource_requirements: Any) -> List[str]:
    if not isinstance(resource_requirements, dict):
        return []
    out: List[str] = []
    for key, value in resource_requirements.items():
        out.append(f"{key}: {value}")
    return out


def _evidence_grade_and_heat(section: Dict[str, Any]) -> Dict[str, str]:
    hit = section.get("graph_hit") if isinstance(section.get("graph_hit"), dict) else {}
    strength = hit.get("evidence_strength") if isinstance(hit.get("evidence_strength"), dict) else {}
    grade = str(strength.get("grade") or "").strip().upper()
    if not grade:
        completeness = hit.get("evidence_completeness") if isinstance(hit.get("evidence_completeness"), dict) else {}
        ratio = float(completeness.get("completeness_ratio") or 0.0)
        if ratio >= 0.85:
            grade = "A"
        elif ratio >= 0.70:
            grade = "B"
        elif ratio >= 0.55:
            grade = "C"
        else:
            grade = "D"
    heat = "high" if grade in {"A", "B"} else "medium" if grade == "C" else "low"
    return {"grade": grade, "heat": heat}


def _render_visual_assets(doc: Document, visual_assets: List[Dict[str, Any]]) -> Dict[str, int]:
    if not visual_assets:
        return {"embedded": 0, "missing": 0}
    embedded = 0
    missing = 0
    doc.add_page_break()
    doc.add_heading("附图与视觉生成", level=1)
    for asset in visual_assets:
        title = str(asset.get("title") or "自动生成图").strip()
        caption = str(asset.get("caption") or "").strip()
        image_path = str(asset.get("image_path") or "").strip()
        doc.add_paragraph(title)
        p = Path(image_path)
        if p.exists():
            doc.add_picture(str(p), width=Cm(15.5))
            embedded += 1
        else:
            doc.add_paragraph(f"图像缺失: {image_path or 'unknown'}")
            missing += 1
        if caption:
            doc.add_paragraph(caption)
    return {"embedded": embedded, "missing": missing}


def _render_evidence_heatmap(doc: Document, sections: List[Dict[str, Any]]) -> Dict[str, int]:
    if not sections:
        return {"rows": 0, "low_count": 0}
    doc.add_page_break()
    doc.add_heading("证据热力图", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "章节"
    hdr[1].text = "节点ID"
    hdr[2].text = "来源层级"
    hdr[3].text = "证据等级"
    hdr[4].text = "热度"
    low_count = 0
    for sec in sections:
        if not isinstance(sec, dict):
            continue
        hit = sec.get("graph_hit") if isinstance(sec.get("graph_hit"), dict) else {}
        heat = _evidence_grade_and_heat(sec)
        row = table.add_row().cells
        row[0].text = str(sec.get("title") or "")
        row[1].text = str(hit.get("node_id") or "")
        row[2].text = str(hit.get("source_hierarchy") or "")
        row[3].text = str(heat.get("grade") or "")
        row[4].text = str(heat.get("heat") or "")
        if heat.get("heat") == "low":
            low_count += 1
    return {"rows": len(sections), "low_count": low_count}


def generate_v2_docx(
    *,
    index_matrix: Dict[str, Any],
    sections: List[Dict[str, Any]],
    visual_assets: List[Dict[str, Any]] | None = None,
    output_path: Path | str,
    title_hint: str = "施工组织设计草案",
) -> Dict[str, Any]:
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _setup_doc_style(doc)

    title = str(index_matrix.get("project_name") or title_hint or "施工组织设计草案").strip()
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        f"生成时间: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime())} | "
        "标黄内容=AI自动补全参数(需人工复核)"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section_by_title: Dict[str, Dict[str, Any]] = {}
    for sec in sections:
        key = str(sec.get("title") or "").strip()
        if key and key not in section_by_title:
            section_by_title[key] = sec

    highlighted_paragraphs = 0
    auto_generated_sections = 0

    for idx, item in enumerate(index_matrix.get("index_matrix") or [], start=1):
        dimension = str(item.get("dimension") or f"章节{idx}").strip()
        keywords = "、".join([str(x) for x in (item.get("keywords") or [])[:10]])

        doc.add_heading(f"{idx}. {dimension}", level=1)
        doc.add_paragraph(f"响应关键词: {keywords}")

        sec = section_by_title.get(dimension) or {}
        highlight = _is_auto_generated_section(sec)
        evidence_heat = _evidence_grade_and_heat(sec)
        low_evidence = evidence_heat.get("heat") == "low"
        if highlight:
            auto_generated_sections += 1

        content = str(sec.get("content") or "").strip()
        _render_section_paragraph(doc, text=content or "该章节暂无可用内容。", highlight=highlight)
        if highlight:
            highlighted_paragraphs += 1

        graph_hit = sec.get("graph_hit") if isinstance(sec.get("graph_hit"), dict) else {}
        evidence_title = str(graph_hit.get("title") or "未命中")
        evidence_file = str(graph_hit.get("source_file") or "")
        evidence_para = doc.add_paragraph()
        evidence_run = evidence_para.add_run(f"证据节点: {evidence_title}  来源文件: {evidence_file}")
        if highlight:
            evidence_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            highlighted_paragraphs += 1
        elif low_evidence:
            evidence_run.font.highlight_color = WD_COLOR_INDEX.PINK
            highlighted_paragraphs += 1

        heat_para = doc.add_paragraph(
            f"证据强度: {evidence_heat.get('grade')} | 热度: {evidence_heat.get('heat')}"
        )
        if low_evidence:
            for run in heat_para.runs:
                run.font.highlight_color = WD_COLOR_INDEX.PINK
            highlighted_paragraphs += 1

        resources = _resource_items(graph_hit.get("resource_requirements"))
        if resources:
            doc.add_paragraph("参数清单:")
            for line in resources[:12]:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(line)
                if highlight:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    highlighted_paragraphs += 1

        if highlight:
            warning = doc.add_paragraph()
            warning_run = warning.add_run("AI审校信标: 本章节包含自愈引擎自动补全参数，请重点人工复核。")
            warning_run.bold = True
            warning_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            highlighted_paragraphs += 1

    heatmap_meta = _render_evidence_heatmap(doc, sections)
    visuals_meta = _render_visual_assets(doc, list(visual_assets or []))
    doc.save(str(out))
    return {
        "ok": True,
        "saved_at": str(out),
        "highlighted_paragraphs": highlighted_paragraphs,
        "auto_generated_sections": auto_generated_sections,
        "evidence_heatmap_rows": int(heatmap_meta["rows"]),
        "low_evidence_sections": int(heatmap_meta["low_count"]),
        "visual_assets_embedded": int(visuals_meta["embedded"]),
        "visual_assets_missing": int(visuals_meta["missing"]),
    }
