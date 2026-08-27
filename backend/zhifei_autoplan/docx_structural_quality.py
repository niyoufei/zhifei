from __future__ import annotations

import hashlib
import json
import os
import posixpath
import re
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


_A4_WIDTH_CM = 21.0
_A4_HEIGHT_CM = 29.7
_EMU_PER_CM = 360_000
_INTERNAL_LEAK_PATTERNS = (
    re.compile(r"\bmain:(?:anthropic|openai|google)\b", re.IGNORECASE),
    re.compile(r"\bfallback_[0-9]+\b", re.IGNORECASE),
    re.compile(r"\bjob_id\s*=", re.IGNORECASE),
    re.compile(r"\bquality_checks\b", re.IGNORECASE),
    re.compile(r"【(?:多Agent|系统全局指令|证据与追溯|章节结构蓝图|本地导出控制表)"),
)


class DocxStructuralQualityError(RuntimeError):
    """Raised when the final Word package violates the delivery contract."""

    def __init__(self, message: str, *, report: dict[str, Any] | None = None) -> None:
        super().__init__(message)
        self.report = report or {}


def _atomic_write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    handle = tempfile.NamedTemporaryFile(
        mode="w",
        encoding="utf-8",
        dir=str(path.parent),
        prefix=f".{path.name}.",
        suffix=".tmp",
        delete=False,
    )
    temp_path = Path(handle.name)
    try:
        with handle:
            json.dump(payload, handle, ensure_ascii=False, indent=2, sort_keys=True)
            handle.flush()
            os.fsync(handle.fileno())
        temp_path.chmod(0o600)
        os.replace(temp_path, path)
        path.chmod(0o600)
    finally:
        temp_path.unlink(missing_ok=True)


def _cm(value: Any) -> float:
    try:
        return round(float(value) / _EMU_PER_CM, 3)
    except (TypeError, ValueError):
        return 0.0


def _pt(value: Any) -> float | None:
    try:
        return round(float(value.pt), 2)
    except (AttributeError, TypeError, ValueError):
        return None


def _font_name(style: Any) -> str:
    try:
        rpr = style._element.rPr
        rfonts = rpr.rFonts if rpr is not None else None
        if rfonts is not None:
            for name in ("eastAsia", "ascii", "hAnsi"):
                value = rfonts.get(qn(f"w:{name}"))
                if value:
                    return str(value)
    except Exception:
        pass
    return str(getattr(getattr(style, "font", None), "name", "") or "")


def _normalise_font(value: Any) -> str:
    name = re.sub(r"\s+", "", str(value or "")).lower()
    aliases = {
        "宋体": "song",
        "simsun": "song",
        "stsong": "song",
        "songtisc": "song",
        "仿宋": "fangsong",
        "仿宋体": "fangsong",
        "stfangsong": "fangsong",
        "fangsong": "fangsong",
        "黑体": "heiti",
        "simhei": "heiti",
        "heitisc": "heiti",
        "stheiti": "heiti",
    }
    return aliases.get(name, name)


def _declared_font(value: Any) -> str:
    raw = str(value or "").strip()
    aliases = {
        "simsun": "宋体",
        "stsong": "宋体",
        "songti sc": "宋体",
        "simhei": "黑体",
        "heiti sc": "黑体",
        "stheiti": "黑体",
        "fangsong": "仿宋体",
        "stfangsong": "仿宋体",
        "仿宋": "仿宋体",
    }
    return aliases.get(raw.lower(), aliases.get(raw, raw))


def _stable_digest(payload: dict[str, Any]) -> str:
    material = dict(payload)
    for key in ("created_at", "receipt", "docx"):
        material.pop(key, None)
    encoded = json.dumps(material, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(encoded.encode("utf-8")).hexdigest()


def _read_package_xml(source: Path) -> dict[str, str]:
    with zipfile.ZipFile(source) as package:
        names = set(package.namelist())
        result: dict[str, str] = {}
        for name in (
            "word/document.xml",
            "word/settings.xml",
            "word/styles.xml",
            "word/fontTable.xml",
            "word/numbering.xml",
            "word/_rels/document.xml.rels",
            "docProps/core.xml",
            "docProps/app.xml",
            "[Content_Types].xml",
        ):
            if name in names:
                result[name] = package.read(name).decode("utf-8", errors="replace")
        footer_names = sorted(name for name in names if re.fullmatch(r"word/footer[0-9]+\.xml", name))
        result["word/footers.xml"] = "\n".join(
            package.read(name).decode("utf-8", errors="replace") for name in footer_names
        )
        result["package_names"] = "\n".join(sorted(names))
        return result


def _relationship_target(rels_name: str, target: str) -> str:
    if rels_name == "_rels/.rels":
        base = ""
    else:
        rels_dir = posixpath.dirname(rels_name)
        owner_dir = posixpath.dirname(rels_dir)
        base = owner_dir
    return posixpath.normpath(posixpath.join(base, str(target or ""))).lstrip("/")


def _audit_package_integrity(source: Path) -> dict[str, Any]:
    invalid_xml: list[str] = []
    duplicate_relationship_ids: list[dict[str, Any]] = []
    dangling_relationships: list[dict[str, Any]] = []
    duplicate_bookmark_ids: list[str] = []
    duplicate_bookmark_names: list[str] = []
    custom_parts: list[str] = []
    with zipfile.ZipFile(source) as package:
        names = set(package.namelist())
        custom_parts = sorted(
            name for name in names if name.startswith("customXml/") or name == "docProps/custom.xml"
        )
        parsed: dict[str, Any] = {}
        for name in sorted(names):
            if not (name.endswith(".xml") or name.endswith(".rels")):
                continue
            try:
                parsed[name] = etree.fromstring(package.read(name))
            except Exception:
                invalid_xml.append(name)
        for name, root in parsed.items():
            if not name.endswith(".rels"):
                continue
            seen: set[str] = set()
            duplicates: set[str] = set()
            for rel in root:
                rel_id = str(rel.get("Id") or "")
                if rel_id in seen:
                    duplicates.add(rel_id)
                seen.add(rel_id)
                if str(rel.get("TargetMode") or "").lower() == "external":
                    continue
                target = _relationship_target(name, str(rel.get("Target") or ""))
                if target and target not in names:
                    dangling_relationships.append({"part": name, "id": rel_id, "target": target})
            if duplicates:
                duplicate_relationship_ids.append({"part": name, "ids": sorted(duplicates)})
        document_root = parsed.get("word/document.xml")
        if document_root is not None:
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            ids: list[str] = []
            bookmark_names: list[str] = []
            for item in document_root.xpath("//w:bookmarkStart", namespaces=ns):
                ids.append(str(item.get(qn("w:id")) or ""))
                bookmark_names.append(str(item.get(qn("w:name")) or ""))
            duplicate_bookmark_ids = sorted({value for value in ids if value and ids.count(value) > 1})
            duplicate_bookmark_names = sorted(
                {value for value in bookmark_names if value and value != "_GoBack" and bookmark_names.count(value) > 1}
            )
    return {
        "invalid_xml": invalid_xml,
        "duplicate_relationship_ids": duplicate_relationship_ids,
        "dangling_relationships": dangling_relationships,
        "duplicate_bookmark_ids": duplicate_bookmark_ids,
        "duplicate_bookmark_names": duplicate_bookmark_names,
        "custom_parts": custom_parts,
    }


def _xml_text(root: Any, local_name: str) -> str:
    if root is None:
        return ""
    for element in root.iter():
        try:
            if etree.QName(element).localname == local_name:
                return str(element.text or "").strip()
        except Exception:
            continue
    return ""


def _field_present(xml: str, field_name: str) -> bool:
    name = re.escape(field_name)
    return bool(
        re.search(rf"w:instrText[^>]*>[^<]*\b{name}\b", xml, flags=re.IGNORECASE)
        or re.search(rf"w:fldSimple[^>]+w:instr=[\"'][^\"']*\b{name}\b", xml, flags=re.IGNORECASE)
    )


def _heading_level(style_name: str) -> int | None:
    match = re.search(r"(?:heading|标题)\s*([1-9])", str(style_name or ""), flags=re.IGNORECASE)
    return int(match.group(1)) if match else None


def audit_docx_structural_quality(
    docx_path: str | Path,
    *,
    expected_style: dict[str, Any] | None = None,
    figure_manifest: dict[str, Any] | None = None,
    require_heading_structure: bool = True,
    strict: bool = True,
    receipt_path: str | Path | None = None,
) -> dict[str, Any]:
    """Inspect the final DOCX package itself before it is exposed to a user.

    This gate is deliberately independent from the generator. It reads the
    saved package and verifies page geometry, named styles, Word fields,
    heading pagination, tables, figure delivery and bidder-facing cleanliness.
    """

    source = Path(docx_path)
    target_receipt = Path(receipt_path) if receipt_path else source.with_suffix(".structural_quality.json")
    hard_failures: list[dict[str, Any]] = []
    warnings: list[dict[str, Any]] = []

    if not source.is_file() or source.stat().st_size == 0:
        report = {
            "schema": "zhifei.docx_structural_quality.v1",
            "created_at": datetime.now(timezone.utc).isoformat(),
            "status": "blocked",
            "docx": str(source),
            "hard_failures": [{"code": "DOCX_MISSING_OR_EMPTY"}],
            "warnings": [],
        }
        report["decision_digest"] = _stable_digest(report)
        _atomic_write_json(target_receipt, report)
        report["receipt"] = str(target_receipt)
        if strict:
            raise DocxStructuralQualityError("最终 Word 不存在或为空", report=report)
        return report

    try:
        document = Document(str(source))
        package_xml = _read_package_xml(source)
        package_integrity = _audit_package_integrity(source)
    except Exception as exc:
        report = {
            "schema": "zhifei.docx_structural_quality.v1",
            "created_at": datetime.now(timezone.utc).isoformat(),
            "status": "blocked",
            "docx": str(source),
            "hard_failures": [{"code": "DOCX_UNREADABLE", "message": str(exc)[:300]}],
            "warnings": [],
        }
        report["decision_digest"] = _stable_digest(report)
        _atomic_write_json(target_receipt, report)
        report["receipt"] = str(target_receipt)
        if strict:
            raise DocxStructuralQualityError("最终 Word 包无法读取", report=report) from exc
        return report

    for code, key in (
        ("PACKAGE_XML_INVALID", "invalid_xml"),
        ("DUPLICATE_RELATIONSHIP_ID", "duplicate_relationship_ids"),
        ("DANGLING_RELATIONSHIP", "dangling_relationships"),
        ("DUPLICATE_BOOKMARK_ID", "duplicate_bookmark_ids"),
        ("DUPLICATE_BOOKMARK_NAME", "duplicate_bookmark_names"),
        ("SENSITIVE_CUSTOM_XML_REMAINS", "custom_parts"),
    ):
        items = package_integrity.get(key) or []
        if items:
            hard_failures.append({"code": code, "items": items[:40]})

    visible_paragraphs = [p for p in document.paragraphs if str(p.text or "").strip()]
    visible_text = "\n".join(str(p.text or "") for p in visible_paragraphs)
    visible_chars = len(re.sub(r"\s+", "", visible_text))
    # Structural QA also runs for small deterministic regression fixtures. The
    # content-quality gate enforces chapter substance separately; this layer
    # only rejects an effectively empty Word package.
    if visible_chars < 1:
        hard_failures.append({"code": "DOCUMENT_CONTENT_TOO_THIN", "visible_chars": visible_chars})

    section_metrics: list[dict[str, Any]] = []
    section_story_refs: list[dict[str, Any]] = []
    expected = dict(expected_style or {})
    margins = expected.get("margins_cm") if isinstance(expected.get("margins_cm"), dict) else {}
    margins = {
        "top": expected.get("margin_top_cm", margins.get("top", 2.5)),
        "right": expected.get("margin_right_cm", margins.get("right", 2.0)),
        "bottom": expected.get("margin_bottom_cm", margins.get("bottom", 2.0)),
        "left": expected.get("margin_left_cm", margins.get("left", 2.0)),
    }
    for index, section in enumerate(document.sections, start=1):
        section_properties = section._sectPr
        header_types = {
            str(node.get(qn("w:type")) or "default")
            for node in section_properties.findall(qn("w:headerReference"))
        }
        footer_types = {
            str(node.get(qn("w:type")) or "default")
            for node in section_properties.findall(qn("w:footerReference"))
        }
        story_ref = {
            "section": index,
            "header_types": sorted(header_types),
            "footer_types": sorted(footer_types),
            "default_header": "default" in header_types,
            "default_footer": "default" in footer_types,
        }
        section_story_refs.append(story_ref)
        if not story_ref["default_header"] or not story_ref["default_footer"]:
            hard_failures.append({"code": "SECTION_HEADER_FOOTER_REFERENCE_MISSING", **story_ref})
        if index > 1 and (header_types != {"default", "first", "even"} or footer_types != {"default", "first", "even"}):
            hard_failures.append({"code": "SECTION_HEADER_FOOTER_VARIANTS_INCOMPLETE", **story_ref})
        metric = {
            "section": index,
            "width_cm": _cm(section.page_width),
            "height_cm": _cm(section.page_height),
            "margins_cm": {
                "top": _cm(section.top_margin),
                "right": _cm(section.right_margin),
                "bottom": _cm(section.bottom_margin),
                "left": _cm(section.left_margin),
            },
        }
        section_metrics.append(metric)
        if str(expected.get("paper") or "A4").upper() == "A4":
            portrait_a4 = (
                abs(metric["width_cm"] - _A4_WIDTH_CM) <= 0.12
                and abs(metric["height_cm"] - _A4_HEIGHT_CM) <= 0.12
            )
            landscape_a4 = (
                abs(metric["width_cm"] - _A4_HEIGHT_CM) <= 0.12
                and abs(metric["height_cm"] - _A4_WIDTH_CM) <= 0.12
            )
            metric["orientation"] = "landscape" if landscape_a4 else "portrait"
            if not portrait_a4 and not landscape_a4:
                hard_failures.append({"code": "PAGE_SIZE_NOT_A4", **metric})
        for name, fallback in (("top", 2.5), ("right", 2.0), ("bottom", 2.0), ("left", 2.0)):
            expected_margin = float(margins.get(name, fallback))
            actual_margin = float(metric["margins_cm"][name])
            if abs(actual_margin - expected_margin) > 0.12:
                hard_failures.append(
                    {
                        "code": "MARGIN_MISMATCH",
                        "section": index,
                        "side": name,
                        "expected_cm": expected_margin,
                        "actual_cm": actual_margin,
                    }
                )

    try:
        normal = document.styles["Normal"]
        normal_font = _font_name(normal)
        normal_size = _pt(normal.font.size)
        normal_spacing = normal.paragraph_format.line_spacing
        normal_spacing_pt = _pt(normal_spacing)
    except Exception:
        normal_font = ""
        normal_size = None
        normal_spacing_pt = None
        hard_failures.append({"code": "NORMAL_STYLE_MISSING"})
    expected_font = _declared_font(expected.get("body_font") or "宋体")
    expected_size = float(expected.get("body_size_pt") or expected.get("body_size") or 14.0)
    expected_spacing_pt = expected.get("line_spacing_pt")
    if not normal_font:
        hard_failures.append({"code": "BODY_FONT_UNDEFINED"})
    elif str(normal_font).strip() != str(expected_font).strip():
        hard_failures.append(
            {"code": "BODY_EAST_ASIA_FONT_MISMATCH", "expected": expected_font, "actual": normal_font}
        )
    if normal_size is None or abs(normal_size - expected_size) > 0.2:
        hard_failures.append(
            {"code": "BODY_SIZE_MISMATCH", "expected_pt": expected_size, "actual_pt": normal_size}
        )
    if expected_spacing_pt is not None:
        expected_spacing = float(expected_spacing_pt)
        if normal_spacing_pt is None or abs(normal_spacing_pt - expected_spacing) > 0.3:
            hard_failures.append(
                {
                    "code": "LINE_SPACING_MISMATCH",
                    "expected_pt": expected_spacing,
                    "actual_pt": normal_spacing_pt,
                }
            )

    try:
        normal_p_pr = document.styles["Normal"]._element.pPr
        normal_ind = normal_p_pr.find(qn("w:ind")) if normal_p_pr is not None else None
        first_line_chars = normal_ind.get(qn("w:firstLineChars")) if normal_ind is not None else None
        normal_spacing_el = normal_p_pr.find(qn("w:spacing")) if normal_p_pr is not None else None
        before_twips = normal_spacing_el.get(qn("w:before")) if normal_spacing_el is not None else None
        after_twips = normal_spacing_el.get(qn("w:after")) if normal_spacing_el is not None else None
    except Exception:
        first_line_chars = None
        before_twips = None
        after_twips = None
    if str(first_line_chars or "") != "200":
        hard_failures.append({"code": "BODY_FIRST_LINE_CHARS_MISSING", "expected": "200", "actual": first_line_chars})
    if str(before_twips or "0") != "0" or str(after_twips or "0") != "0":
        hard_failures.append(
            {
                "code": "BODY_PARAGRAPH_SPACING_NONZERO",
                "before_twips": before_twips,
                "after_twips": after_twips,
            }
        )

    expected_title_font = _declared_font(expected.get("title_font") or "宋体")
    expected_title_size = float(expected.get("title_size_pt") or expected.get("title_size") or 16.0)
    expected_heading_sizes = {
        "Heading 1": float(expected.get("doc_title_size") or expected_title_size),
        "Heading 2": expected_title_size,
        "Heading 3": expected_title_size,
    }
    heading_style_issues: list[dict[str, Any]] = []
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        try:
            style = document.styles[style_name]
            font_name = _font_name(style)
            size = _pt(style.font.size)
            expected_heading_size = expected_heading_sizes[style_name]
            if font_name != expected_title_font or size is None or abs(size - expected_heading_size) > 0.2 or style.font.bold is not True:
                heading_style_issues.append(
                    {
                        "style": style_name,
                        "font": font_name,
                        "size_pt": size,
                        "expected_size_pt": expected_heading_size,
                        "bold": style.font.bold,
                    }
                )
        except Exception:
            heading_style_issues.append({"style": style_name, "missing": True})
    if heading_style_issues:
        hard_failures.append(
            {
                "code": "HEADING_STYLE_CONTRACT_MISMATCH",
                "expected_font": expected_title_font,
                "expected_sizes_pt": expected_heading_sizes,
                "items": heading_style_issues,
            }
        )
    try:
        caption = document.styles["Caption"]
        caption_font = _font_name(caption)
        caption_size = _pt(caption.font.size)
        if caption_font != expected_font or caption_size is None or abs(caption_size - expected_size) > 0.2:
            hard_failures.append(
                {
                    "code": "CAPTION_STYLE_CONTRACT_MISMATCH",
                    "expected_font": expected_font,
                    "expected_size_pt": expected_size,
                    "actual_font": caption_font,
                    "actual_size_pt": caption_size,
                }
            )
    except Exception:
        hard_failures.append({"code": "CAPTION_STYLE_MISSING"})

    document_xml = package_xml.get("word/document.xml", "")
    settings_xml = package_xml.get("word/settings.xml", "")
    styles_xml = package_xml.get("word/styles.xml", "")
    footers_xml = package_xml.get("word/footers.xml", "")
    if not _field_present(document_xml, "TOC"):
        hard_failures.append({"code": "TOC_FIELD_MISSING"})
    if not re.search(r"<w:updateFields\b[^>]*w:val=[\"'](?:true|1)[\"']", settings_xml):
        hard_failures.append({"code": "FIELD_UPDATE_DISABLED"})
    for field_name in ("PAGE", "NUMPAGES"):
        if not _field_present(footers_xml, field_name):
            hard_failures.append({"code": f"{field_name}_FIELD_MISSING"})
    if "<wp:anchor" in document_xml:
        hard_failures.append({"code": "FLOATING_IMAGE_ANCHOR_PRESENT"})
    if re.search(r"<w:(?:vanish|webHidden)\b", document_xml):
        hard_failures.append({"code": "HIDDEN_TEXT_REMAINS"})
    if re.search(r"<w:(?:ins|del|moveFrom|moveTo)\b", document_xml):
        hard_failures.append({"code": "TRACKED_CHANGES_REMAIN"})
    package_names = set(str(package_xml.get("package_names") or "").splitlines())
    if any(re.fullmatch(r"word/comments[0-9]*\.xml", name) for name in package_names):
        hard_failures.append({"code": "COMMENTS_REMAIN"})

    core_root = None
    app_root = None
    try:
        core_root = etree.fromstring(package_xml.get("docProps/core.xml", "").encode("utf-8"))
    except Exception:
        pass
    try:
        app_root = etree.fromstring(package_xml.get("docProps/app.xml", "").encode("utf-8"))
    except Exception:
        pass
    metadata_values = {
        "creator": _xml_text(core_root, "creator"),
        "lastModifiedBy": _xml_text(core_root, "lastModifiedBy"),
        "description": _xml_text(core_root, "description"),
        "company": _xml_text(app_root, "Company"),
        "manager": _xml_text(app_root, "Manager"),
    }
    metadata_values = {key: value for key, value in metadata_values.items() if value}
    if metadata_values:
        hard_failures.append({"code": "DOCUMENT_METADATA_REMAINS", "fields": metadata_values})

    try:
        styles_root = etree.fromstring(styles_xml.encode("utf-8"))
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        normal_nodes = styles_root.xpath("//w:style[@w:styleId='Normal']/w:rPr/w:rFonts", namespaces=ns)
        font_attrs = normal_nodes[0].attrib if normal_nodes else {}
        missing_font_attrs = [
            name for name in ("ascii", "hAnsi", "eastAsia", "cs") if not font_attrs.get(qn(f"w:{name}"))
        ]
        if missing_font_attrs:
            hard_failures.append({"code": "BODY_FONT_ATTRIBUTES_INCOMPLETE", "missing": missing_font_attrs})
    except Exception:
        hard_failures.append({"code": "STYLES_XML_UNREADABLE"})

    heading_count = 0
    hierarchy: list[int] = []
    unsafe_headings: list[dict[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        level = _heading_level(str(getattr(getattr(paragraph, "style", None), "name", "") or ""))
        if level is None:
            continue
        heading_count += 1
        hierarchy.append(level)
        fmt = paragraph.paragraph_format
        style_fmt = getattr(getattr(paragraph, "style", None), "paragraph_format", None)

        def _effective_bool(name: str) -> bool:
            direct = getattr(fmt, name, None)
            if direct is not None:
                return bool(direct)
            return bool(getattr(style_fmt, name, None)) if style_fmt is not None else False

        if not all(
            _effective_bool(name) for name in ("keep_with_next", "keep_together", "widow_control")
        ):
            unsafe_headings.append({"paragraph": index, "level": level, "text": paragraph.text[:80]})
    if require_heading_structure and heading_count == 0:
        hard_failures.append({"code": "HEADING_STRUCTURE_MISSING"})
    if unsafe_headings:
        hard_failures.append({"code": "HEADING_PAGINATION_UNSAFE", "items": unsafe_headings[:20]})
    jumps = [
        {"from": hierarchy[index - 1], "to": hierarchy[index], "position": index + 1}
        for index in range(1, len(hierarchy))
        if hierarchy[index] > hierarchy[index - 1] + 1
    ]
    if jumps:
        warnings.append({"code": "HEADING_HIERARCHY_JUMPS", "items": jumps[:20]})

    split_rows: list[dict[str, int]] = []
    headerless_tables: list[int] = []
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            tr_pr = row._tr.trPr
            if tr_pr is None or tr_pr.find(qn("w:cantSplit")) is None:
                split_rows.append({"table": table_index, "row": row_index})
        if len(table.rows) > 1:
            tr_pr = table.rows[0]._tr.trPr
            if tr_pr is None or tr_pr.find(qn("w:tblHeader")) is None:
                headerless_tables.append(table_index)
    if split_rows:
        hard_failures.append({"code": "TABLE_ROWS_CAN_SPLIT", "items": split_rows[:40]})
    if headerless_tables:
        warnings.append({"code": "TABLE_HEADER_NOT_REPEATED", "tables": headerless_tables})

    for pattern in _INTERNAL_LEAK_PATTERNS:
        match = pattern.search(visible_text)
        if match:
            hard_failures.append(
                {"code": "INTERNAL_IMPLEMENTATION_LEAK", "sample": match.group(0)[:100]}
            )
            break

    figure_summary: dict[str, Any] = {}
    if isinstance(figure_manifest, dict):
        figure_summary = {
            "delivery_allowed": bool(figure_manifest.get("delivery_allowed")),
            "figure_count": int(figure_manifest.get("figure_count") or 0),
            "decision_digest": str(figure_manifest.get("decision_digest") or ""),
        }
        if not figure_summary["delivery_allowed"]:
            hard_failures.append({"code": "FIGURE_DELIVERY_GATE_BLOCKED"})

    report = {
        "schema": "zhifei.docx_structural_quality.v1",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "status": "blocked" if hard_failures else "pass",
        "docx": str(source),
        "docx_sha256": hashlib.sha256(source.read_bytes()).hexdigest(),
        "visible_chars": visible_chars,
        "paragraph_count": len(document.paragraphs),
        "heading_count": heading_count,
        "table_count": len(document.tables),
        "section_metrics": section_metrics,
        "section_story_references": section_story_refs,
        "body_style": {
            "font": normal_font,
            "size_pt": normal_size,
            "line_spacing_pt": normal_spacing_pt,
            "first_line_chars": first_line_chars,
            "space_before_twips": before_twips,
            "space_after_twips": after_twips,
        },
        "word_fields": {
            "toc": _field_present(document_xml, "TOC"),
            "page": _field_present(footers_xml, "PAGE"),
            "numpages": _field_present(footers_xml, "NUMPAGES"),
            "update_on_open": bool(
                re.search(r"<w:updateFields\b[^>]*w:val=[\"'](?:true|1)[\"']", settings_xml)
            ),
        },
        "figure_delivery": figure_summary,
        "package_integrity": package_integrity,
        "hard_failures": hard_failures,
        "warnings": warnings,
    }
    report["decision_digest"] = _stable_digest(report)
    _atomic_write_json(target_receipt, report)
    report["receipt"] = str(target_receipt)
    if strict and report["status"] != "pass":
        codes = ", ".join(str(item.get("code")) for item in hard_failures)
        raise DocxStructuralQualityError(f"最终 Word 结构验收未通过：{codes}", report=report)
    return report
