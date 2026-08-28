from __future__ import annotations

import datetime as _dt
import copy
import json
import math
import os
import re
import tempfile
import zipfile
from pathlib import Path
from typing import Dict, Any, List

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from backend.zhifei_autoplan.media import generate_section_visuals
from backend.zhifei_autoplan.media_quality import (
    build_media_delivery_manifest,
    media_matches_chapter,
    validate_media_collection,
    validate_media_item,
    verify_docx_media_hashes,
)
from backend.zhifei_autoplan.docx_structural_quality import audit_docx_structural_quality
from backend.zhifei_autoplan.style_policy import resolve_line_spacing
from backend.zhifei_autoplan.terminology_guard import load_global_terminology, normalize_text_terminology

try:
    from backend.zhifei_autoplan.local_adapter_shim import validate_before_export as _local_adapter_validate_before_export
except Exception:
    _local_adapter_validate_before_export = None


_AUTOFIX_MARK_RE = re.compile(r"【自动补充】(?P<name>[^\n]{1,80}?)(?:：|:)")
_OVERVIEW_SECTION_RE = re.compile(r"(工程概况|项目概况)")
_TOC_CHAPTER_RE = re.compile(r"^第[一二三四五六七八九十百零0-9]+章")
_TOC_SECTION_RE = re.compile(r"^第[一二三四五六七八九十百零0-9]+节")
_TOC_CN_ITEM_RE = re.compile(r"^[一二三四五六七八九十百零]+[、.]")
_TOC_NUMERIC_ITEM_RE = re.compile(r"^\d+(?:\.\d+){1,3}")
_CN_DIGITS = {"0": "零", "1": "一", "2": "二", "3": "三", "4": "四", "5": "五", "6": "六", "7": "七", "8": "八", "9": "九"}
_GENERIC_COVER_IMAGE_STEMS = {
    "",
    "image",
    "img",
    "photo",
    "picture",
    "cover",
    "wechat image",
    "微信图片",
    "现场照片",
    "现场图",
    "现状",
}

_FONT_ALIASES = {
    # Store the portable Chinese family name in OOXML.  Platform-specific
    # alternatives are declared in fontTable.xml instead of silently replacing
    # w:eastAsia, so Windows/Word receives the required 宋体 declaration while
    # macOS/LibreOffice can still fall back to STSong for rendering.
    "宋体": "宋体",
    "simsun": "宋体",
    "stsong": "宋体",
    "songti sc": "宋体",
    "仿宋": "仿宋体",
    "仿宋体": "仿宋体",
    "fangsong": "仿宋体",
    "stfangsong": "仿宋体",
    "黑体": "黑体",
    "simhei": "黑体",
    "heiti sc": "黑体",
    "stheiti": "黑体",
}

_FONT_FALLBACKS = {
    "宋体": "STSong",
    "仿宋体": "STFangsong",
    "黑体": "Heiti SC",
}
_SUBMISSION_EVIDENCE_RE = re.compile(r"【证据\s*[:：][^】]{0,600}】")
_SUBMISSION_GRAPH_RE = re.compile(r"【(?:图谱节点|图谱经验值)\s*[:：][^】]{0,600}】")
_SUBMISSION_EXPERIENCE_RE = re.compile(
    r"[【（(]经验值(?:\s*[:：，,][^】）)]{0,200})?[】）)]"
)
_SUBMISSION_LOCATOR_RE = re.compile(
    r"[^\s，。；;：:【】]{1,180}#(?:p\d+_)?[0-9a-fA-F]{6,}@[0-9]+"
)
_SUBMISSION_INTERNAL_TOKENS = (
    "entity_master_key",
    "authority_rank",
    "formula_expression",
    "gemini_usefulness_score",
    "incremental_fingerprint",
    "is_auto_generated",
    "kg_dimension",
    "node_id",
    "professional_domain",
    "reference_standard_count",
    "reference_standard_primary",
    "safety_level",
    "source_hierarchy",
    "source_hierarchy_weight",
    "quality_checks",
    "auto_revision_suggestions",
    "span_start",
    "span_end",
)
_SUBMISSION_INTERNAL_PREFIXES = (
    "【多Agent",
    "【图谱节点绑定",
    "【证据与追溯",
    "【系统全局指令",
    "【篇幅约束",
    "【章节结构蓝图",
    "评分点命中关键词（用于本章覆盖校核）",
    "【消除空泛词】",
    "【核心结论证据补齐】",
    "【评分点覆盖建议】",
    "【证据可追溯定位】",
    "【数据一致性校核】",
    "【本地导出控制表】",
    "【排版及格式合规声明】",
    "【格式合规声明】",
    "【排版声明",
    "排版及格式合规声明",
    "格式合规声明",
    "评分点覆盖建议",
    "核心结论证据补齐",
    "证据可追溯定位",
    "证据标注",
    "评分维度回填",
    "消除空泛词",
    "本地导出控制表",
)
_SUBMISSION_INSTRUCTION_PREFIXES = (
    "将空泛词替换为",
    "对含“频次/阈值/时限/人数/型号/工期”等结论句逐条补",
    "示例：",
    "补充评分点覆盖：",
    "量化指标示例：",
    "本章至少保留 1 条带定位符的证据",
    "每条写清：动作+参数+频次+责任岗位+验收方法/阈值+记录表",
    "在每条关键结论句末追加",
    "至少 1 条/章",
    "至少1条/章",
)
_GENERIC_METRIC_TEMPLATE_PREFIX = (
    "量化指标：频次=2次/日（班前+收工）；阈值=偏差≤5mm；间距=1000mm；"
    "厚度=50mm；时长=4h/作业段；人数=8人/班；设备型号=20t挖机1台"
)
_SUBMISSION_IDENTIFIER_LABELS = {
    "inspection_batches": "已验收批次",
    "pass_rate_percent": "批次合格率",
    "total_batches": "总批次",
    "total_work_batches": "总施工批次",
    "emergency_response_minutes": "实际应急响应时长",
    "target_response_minutes": "目标应急响应时长",
    "defect_count": "缺陷数",
    "sample_count": "抽检样本数",
    "rework_volume": "返工工程量",
    "total_work_volume": "总工程量",
}
_RISK_TRIPLET_RE = re.compile(
    r"^(?:风险\s*[→-]\s*控制\s*[→-]\s*验证\s*[：:]\s*)?"
    r"(?:风险|风险点)\s*[：:]\s*(?P<risk>.*?)\s*[；;]\s*"
    r"(?:控制|控制措施)\s*[：:]\s*(?P<control>.*?)\s*[；;]\s*"
    r"(?:验证|验证方式)\s*[：:]\s*(?P<verification>.+)$"
)
_INLINE_MARKDOWN_RE = re.compile(r"(\*\*|__)(.+?)\1")
_MAX_SUBMISSION_IMAGES = 40
_INVALID_XML_10_RE = re.compile(
    "[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD\U00010000-\U0010FFFF]"
)

_SUBMISSION_HEADING_REPLACEMENTS = {
    "A. 组织与策划层 (Architecture & Planning)": "组织与策划",
    "B. 边界与防护层 (Boundary & Protection)": "作业边界与防护",
    "C. 核心执行层 (Core Implementation)": "施工执行与过程控制",
    "D. 监测与诊断层 (Detection & Diagnostics)": "监测、诊断与纠偏",
    "E. 应急与闭环层 (Emergency & Evaluation)": "应急处置与闭环",
    "A. 本章交付物与记录表": "交付成果与记录",
    "A. 交付物/记录表/验收点": "交付成果、记录与验收点",
    "B. 核心约束条件与重难点分析": "核心约束与重难点",
    "B. 约束条件": "施工约束与前置条件",
    "C. 执行步骤": "实施步骤",
    "D. 风险→控制→验证（闭环）": "风险控制与验证",
    "D. 风险→控制→验证（闭环管控）": "风险控制与验证",
    "E. 扣分项规避说明": "质量风险防控",
}


def _sanitize_xml_text(value: Any) -> str:
    """Remove characters that OOXML/XML 1.0 cannot represent."""
    return _INVALID_XML_10_RE.sub("", str(value or ""))


def _sanitize_docx_payload(value: Any) -> Any:
    """Return a structure-safe copy with every text value XML-compatible."""
    if isinstance(value, str):
        return _sanitize_xml_text(value)
    if isinstance(value, dict):
        return {key: _sanitize_docx_payload(item) for key, item in value.items()}
    if isinstance(value, list):
        return [_sanitize_docx_payload(item) for item in value]
    if isinstance(value, tuple):
        return tuple(_sanitize_docx_payload(item) for item in value)
    return value


def _local_adapter_export_error(export_kind: str, issues: Any) -> RuntimeError:
    payload = {
        "code": "LOCAL_ADAPTER_EXPORT_BLOCKED",
        "message": "输出未通过本地正式交付适配器门禁。",
        "status": "blocked",
        "export_allowed": False,
        "export_kind": export_kind,
        "issues": issues if isinstance(issues, list) else [{"code": "LOCAL_ADAPTER_EXPORT_BLOCKED", "message": str(issues)}],
    }
    return RuntimeError(json.dumps(payload, ensure_ascii=False))


def _require_local_adapter_export_allowed(data: Dict[str, Any] | None, export_kind: str) -> None:
    if _local_adapter_validate_before_export is None:
        raise _local_adapter_export_error(export_kind, [{"code": "ADAPTER_IMPORT_FAILURE", "message": "local adapter shim import failed"}])
    try:
        gate = _local_adapter_validate_before_export(data or {})
    except Exception as exc:
        raise _local_adapter_export_error(export_kind, [{"code": "ADAPTER_EXPORT_GATE_FAILURE", "message": repr(exc)}])
    if not gate.get("export_allowed"):
        raise _local_adapter_export_error(export_kind, gate.get("issues") or [{"code": "LOCAL_ADAPTER_EXPORT_BLOCKED", "message": "export blocked"}])


def _strip_internal_autofix_markers(text: str) -> str:
    """
    Internal remediation markers are useful for idempotency, but should not appear
    in the final DOCX deliverable.
    Keep the heading content, only remove the "自动补充" prefix.
    """
    s = str(text or "")
    s = _AUTOFIX_MARK_RE.sub(lambda m: f"【{m.group('name').strip()}】", s)
    s = s.replace("【自动补充】", "")
    return s


def _resolve_docx_font_name(value: Any) -> str:
    raw = str(value or "").strip()
    if not raw:
        return raw
    return _FONT_ALIASES.get(raw, _FONT_ALIASES.get(raw.lower(), raw))


def _set_first_line_chars(target: Any, chars: int) -> None:
    """Set Word's character-based first-line indent deterministically.

    ``python-docx`` exposes only distance-based indentation.  Chinese tender
    documents require ``w:firstLineChars=200`` (two ideographic characters),
    which remains stable when fonts are substituted across Word/WPS/macOS.
    """

    try:
        if hasattr(target, "_p"):
            p_pr = target._p.get_or_add_pPr()
        else:
            p_pr = target._element.get_or_add_pPr()
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            p_pr.append(ind)
        for name in ("firstLine", "hanging", "hangingChars"):
            ind.attrib.pop(qn(f"w:{name}"), None)
        ind.set(qn("w:firstLineChars"), str(max(0, int(chars))))
    except Exception:
        pass


def _iter_story_paragraphs(doc: Document):
    """Yield paragraphs from the body, tables, headers and footers once."""

    seen: set[int] = set()

    def _yield_container(container: Any):
        for paragraph in getattr(container, "paragraphs", ()):
            marker = id(paragraph._p)
            if marker not in seen:
                seen.add(marker)
                yield paragraph
        for table in getattr(container, "tables", ()):
            for row in table.rows:
                for cell in row.cells:
                    yield from _yield_container(cell)

    yield from _yield_container(doc)
    for section in doc.sections:
        for story in (section.header, section.first_page_header, section.even_page_header,
                      section.footer, section.first_page_footer, section.even_page_footer):
            yield from _yield_container(story)


def _enforce_chinese_paragraph_geometry(doc: Document) -> None:
    """Apply two-character indentation only to real body paragraphs.

    Headings, captions, lists, table cells and drawing host paragraphs must
    remain flush-left/centred.  Normal prose receives character-based
    indentation and exact zero before/after spacing.
    """

    def _inside_table(paragraph: Any) -> bool:
        node = paragraph._p.getparent()
        while node is not None:
            try:
                if node.tag == qn("w:tc"):
                    return True
            except Exception:
                pass
            node = node.getparent()
        return False

    for paragraph in _iter_story_paragraphs(doc):
        style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "")
        xml = paragraph._p.xml
        is_body = (
            style_name in {"Normal", "正文"}
            and not _inside_table(paragraph)
            and "<w:drawing" not in xml
            and "<w:pict" not in xml
            and "<w:fldChar" not in xml
            and str(paragraph.text or "").strip() != ""
        )
        _set_first_line_chars(paragraph, 200 if is_body else 0)


def _scrub_document_properties(doc: Document) -> None:
    """Clear authoring identity and machine metadata before package save."""

    try:
        props = doc.core_properties
        for name in (
            "author", "last_modified_by", "comments", "category", "content_status",
            "identifier", "keywords", "language", "subject", "title", "version",
        ):
            try:
                setattr(props, name, "")
            except Exception:
                pass
        for name in ("created", "modified", "last_printed"):
            try:
                setattr(props, name, None)
            except Exception:
                pass
        try:
            props.revision = 1
        except Exception:
            pass
    except Exception:
        pass


def _materialize_section_story_references(doc: Document) -> None:
    """Persist explicit default header/footer refs on every content section."""

    inherited_default: dict[str, Any] = {}
    for index, section in enumerate(doc.sections):
        if index > 0:
            section.different_first_page_header_footer = False
        for tag_name in ("headerReference", "footerReference"):
            default_nodes = [
                node
                for node in section._sectPr.findall(qn(f"w:{tag_name}"))
                if str(node.get(qn("w:type")) or "default") == "default"
            ]
            if index == 0 and default_nodes:
                inherited_default[tag_name] = default_nodes[0]
                continue
            prior = inherited_default.get(tag_name)
            if prior is None:
                continue
            for node in list(section._sectPr.findall(qn(f"w:{tag_name}"))):
                section._sectPr.remove(node)
            # Writer can select first/even page styles for continued landscape
            # tables even when Word would use the default story.  Point all
            # three variants at the verified non-empty default part.
            for story_type in ("even", "first", "default"):
                explicit = copy.deepcopy(prior)
                explicit.set(qn("w:type"), story_type)
                section._sectPr.insert(0, explicit)


def _secure_docx_package(path: str | Path) -> None:
    """Strip generated-package metadata/custom XML and add font fallbacks.

    The rewrite is atomic and preserves all unrelated ZIP members byte-for-byte.
    It also removes the default bibliography customXml part added by the
    python-docx template, preventing custom XML from becoming an unnoticed
    sensitive-data channel.
    """

    source = Path(path)
    handle = tempfile.NamedTemporaryFile(
        mode="wb", dir=str(source.parent), prefix=f".{source.name}.", suffix=".tmp", delete=False
    )
    temp_path = Path(handle.name)
    handle.close()
    try:
        with zipfile.ZipFile(source, "r") as reader, zipfile.ZipFile(temp_path, "w") as writer:
            for info in reader.infolist():
                name = info.filename
                if name.startswith("customXml/") or name == "docProps/custom.xml":
                    continue
                data = reader.read(name)
                if name.endswith(".rels") or name == "[Content_Types].xml":
                    try:
                        from lxml import etree

                        root = etree.fromstring(data)
                        if name.endswith(".rels"):
                            for rel in list(root):
                                target = str(rel.get("Target") or "")
                                rel_type = str(rel.get("Type") or "")
                                if "customXml" in rel_type or target.startswith("../customXml/") or target == "docProps/custom.xml":
                                    root.remove(rel)
                        else:
                            for child in list(root):
                                part_name = str(child.get("PartName") or "")
                                content_type = str(child.get("ContentType") or "")
                                if part_name.startswith("/customXml/") or part_name == "/docProps/custom.xml" or "custom-properties" in content_type:
                                    root.remove(child)
                        data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                    except Exception:
                        pass
                elif name == "docProps/core.xml":
                    try:
                        from lxml import etree

                        root = etree.fromstring(data)
                        for child in root:
                            local = etree.QName(child).localname
                            if local in {
                                "creator", "lastModifiedBy", "description", "keywords", "subject",
                                "title", "category", "contentStatus", "identifier", "language", "version",
                                "created", "modified", "lastPrinted",
                            }:
                                child.text = ""
                        data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                    except Exception:
                        pass
                elif name == "docProps/app.xml":
                    try:
                        from lxml import etree

                        root = etree.fromstring(data)
                        for child in root:
                            if etree.QName(child).localname in {"Company", "Manager"}:
                                child.text = ""
                        data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                    except Exception:
                        pass
                elif name == "word/fontTable.xml":
                    try:
                        from lxml import etree

                        root = etree.fromstring(data)
                        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                        for declared, fallback in _FONT_FALLBACKS.items():
                            font = next(
                                (item for item in root.findall(f"{{{ns}}}font") if item.get(f"{{{ns}}}name") == declared),
                                None,
                            )
                            if font is None:
                                font = etree.SubElement(root, f"{{{ns}}}font")
                                font.set(f"{{{ns}}}name", declared)
                            alt = font.find(f"{{{ns}}}altName")
                            if alt is None:
                                alt = etree.SubElement(font, f"{{{ns}}}altName")
                            alt.set(f"{{{ns}}}val", fallback)
                        data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                    except Exception:
                        pass
                writer.writestr(info, data)
        os.replace(temp_path, source)
    finally:
        temp_path.unlink(missing_ok=True)


def _save_docx_secure(doc: Document, output_path: str | Path) -> None:
    _enforce_chinese_paragraph_geometry(doc)
    _scrub_document_properties(doc)
    _materialize_section_story_references(doc)
    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target))
    _secure_docx_package(target)


def _rewrite_submission_notation(value: str) -> str:
    """Translate machine-oriented formulas into bidder-facing engineering prose."""
    text = str(value or "")
    formula_replacements = (
        (
            r"inspection_batches\s*\*\s*pass_rate_percent\s*/\s*max\s*\(\s*(?:total_batches|total_work_batches)\s*,\s*1\s*\)",
            "已验收批次乘以批次合格率，再除以总批次",
        ),
        (
            r"emergency_response_minutes\s*/\s*max\s*\(\s*target_response_minutes\s*,\s*1\s*\)",
            "实际应急响应时长与目标应急响应时长的比值",
        ),
        (
            r"100\s*-\s*defect_count\s*\*\s*100\s*/\s*max\s*\(\s*sample_count\s*,\s*1\s*\)",
            "抽检合格率〔（抽检样本数－缺陷数）÷抽检样本数×100%〕",
        ),
        (
            r"rework_volume\s*\*\s*100\s*/\s*max\s*\(\s*total_work_volume\s*,\s*1\s*\)",
            "返工率〔返工工程量÷总工程量×100%〕",
        ),
        (
            r"high_risk_tasks\s*\*\s*100\s*/\s*(?:max\s*\(\s*)?total_tasks(?:\s*,\s*1\s*\))?",
            "高风险任务数÷任务总数×100%",
        ),
        (
            r"delay_hours\s*\*\s*100\s*/\s*(?:max\s*\(\s*)?planned_hours(?:\s*,\s*1\s*\))?",
            "累计延误时长÷计划时长×100%",
        ),
    )
    for pattern, replacement in formula_replacements:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    for identifier, label in _SUBMISSION_IDENTIFIER_LABELS.items():
        text = re.sub(rf"\b{re.escape(identifier)}\b", label, text, flags=re.IGNORECASE)
    text = re.sub(r"\bmax\s*\(([^,()]+),\s*1\s*\)", r"\1（最小按1计）", text, flags=re.IGNORECASE)
    text = text.replace("`", "").replace("->", "→")
    return text


def _missing_parameter_keys(value: Any) -> set[str]:
    if not isinstance(value, dict):
        return set()
    missing = value.get("missing")
    if not isinstance(missing, list):
        return set()
    return {
        str(item.get("key") or "").strip()
        for item in missing
        if isinstance(item, dict) and str(item.get("key") or "").strip()
    }


def _rewrite_unverified_submission_parameters(text: str, missing_parameters: Any) -> str:
    """Replace guessed planning values when the source packet says they are missing.

    The prose remains usable, but the submission no longer states model/default
    values as if they had been confirmed by the tender documents.
    """
    keys = _missing_parameter_keys(missing_parameters)
    value = str(text or "")
    value = value.replace(_GENERIC_METRIC_TEMPLATE_PREFIX + "。", "")
    value = value.replace(_GENERIC_METRIC_TEMPLATE_PREFIX, "")
    if "总工期" in keys:
        unverified_durations: set[str] = set()
        duration_patterns = (
            r"总工期[^\n，。；;]{0,24}?(\d+(?:\.\d+)?)\s*天",
            r"(\d+(?:\.\d+)?)\s*天[^\n，。；;]{0,24}?总工期",
            r"总进度计划网络图（(\d+(?:\.\d+)?)\s*天控制口径）",
        )
        for pattern in duration_patterns:
            unverified_durations.update(re.findall(pattern, value))
        value = re.sub(
            r"(?:本工程)?(?:计划)?总工期\s*(?:[=：:]|为)?\s*\d+(?:\.\d+)?\s*天",
            "本工程总工期以招标文件及经批准的总进度计划为准",
            value,
        )
        for duration in sorted(unverified_durations, key=len, reverse=True):
            value = re.sub(
                rf"(?<![0-9.]){re.escape(duration)}\s*天(?![0-9])",
                "经批准计划确定的工期",
                value,
            )
    if "资源峰值" in keys:
        value = re.sub(
            r"(?:锁定)?资源峰值\s*(?:[=：:]|为)?\s*\d+(?:\.\d+)?\s*人(?:（当量）)?",
            "各阶段资源按经批准的进度计划和实际工作面动态配置",
            value,
        )
    if "关键线路间隔" in keys:
        value = re.sub(
            r"关键线路间隔\s*(?:[=：:]|为)?\s*\d+(?:\.\d+)?\s*天",
            "关键线路节拍以经批准的进度计划为准",
            value,
        )
    if "风险检查频次" in keys:
        value = value.replace("频次=2次/日（班前+收工）", "按班前检查、过程巡检和收工复核执行")
        value = value.replace("2次/日（班前+收工）", "班前检查、过程巡检和收工复核")
        value = re.sub(
            r"(?:风险检查)?频次\s*[=：:]\s*2\s*次\s*/\s*日",
            "检查频次按班前、过程巡检和收工复核执行",
            value,
        )
        value = re.sub(
            r"(?<![0-9])2\s*次\s*/\s*日(?![0-9])",
            "按班前、过程巡检和收工复核",
            value,
        )
    if "质量阈值" in keys:
        value = value.replace("偏差≤5mm", "允许偏差按设计文件和适用验收标准执行")
        value = re.sub(
            r"(?:一次验收)?(?:质量)?合格率(?:阈值)?(?:要求)?\s*(?:设定为|控制为)?\s*[=：:]?\s*≥?\s*98%",
            "验收合格标准按招标文件、设计文件及适用规范执行",
            value,
        )
        value = re.sub(
            r"一次验收通过率(?:阈值)?(?:设置|要求|控制在)?\s*[=：:]?\s*≥?\s*95%",
            "一次验收通过率按经批准的项目质量目标考核",
            value,
        )
        value = re.sub(
            r"扫码准确率\s*≥?\s*98%",
            "扫码信息与实物及台账逐项核对",
            value,
        )
    if "偏差处置时限" in keys:
        value = value.replace("偏差处置时限≤4h", "偏差应立即隔离并在复验合格后销项")
        value = re.sub(
            r"(?:处置)?时限\s*[=：:]?\s*≤?\s*\d+(?:\.\d+)?\s*h\s*内?",
            "发现偏差后立即隔离整改，复验合格后销项",
            value,
            flags=re.IGNORECASE,
        )
        value = re.sub(
            r"偏差应立即隔离并在复验合格后销项内到场",
            "立即启动备用人员调配，并依据经批准的进度计划调整作业安排",
            value,
        )
    if "资源峰值" in keys:
        value = re.sub(
            r"作业人数\s*[（(]\s*\d+\s*人\s*/\s*班\s*[）)]",
            "作业人数按工作面和经批准的劳动力计划配置",
            value,
        )
        value = re.sub(
            r"人数\s*[=：:]\s*\d+\s*人\s*/\s*班",
            "人数按工作面和经批准的劳动力计划配置",
            value,
        )
    value = re.sub(
        r"防锈漆厚度\s*[=：:]\s*50\s*mm(?:（[^）]*）)?",
        "防锈漆干膜厚度按设计文件、产品体系和适用验收标准检查",
        value,
    )
    # Remove generation-template language and generic equipment defaults that
    # are not suitable for a hospital partial-renovation submission.
    value = value.replace("A/B/C/D/E标准控制工序", "准备、复核、材料、施工、验收五步控制工序")
    value = value.replace("A/B/C/D/E结构", "五步闭环")
    value = value.replace("A/B/C/D/E", "五步闭环")
    value = value.replace(
        "A（准备）/B（施工）/C（检查）/D（验收）/E（成品保护/后期处理）五步工序法",
        "准备、施工、检查、验收、成品保护五步闭环工序法",
    )
    value = re.sub(r"[（(](?:模版|模板)A[）)]", "", value)
    value = value.replace("五层级（五步闭环）", "准备、复核、材料、施工、验收五步闭环")
    value = re.sub(
        r"防尘网覆盖厚度控制在50mm",
        "采用阻燃密目网连续覆盖并压边固定",
        value,
    )
    value = re.sub(
        r"裸土及渣土使用密目网覆盖，厚度控制50mm以保持水土",
        "裸土及渣土采用阻燃密目网连续覆盖并压边固定",
        value,
    )
    value = re.sub(
        r"(?:铺设砂浆|干硬性水泥砂浆打底|底层铺设)厚度[=：:]?50mm(?:（含找平层）)?",
        "找平层及结合层厚度按设计构造、样板确认和标高复核结果控制",
        value,
    )
    value = value.replace(
        "干硬性水泥砂浆打底（厚度=50mm）",
        "干硬性水泥砂浆找平及结合层厚度按设计构造、样板确认和标高复核结果控制",
    )
    value = re.sub(
        r"细石混凝土保护层厚度[=：:]?50mm",
        "细石混凝土保护层厚度按设计构造和适用验收标准控制",
        value,
    )
    value = re.sub(
        r"蓄水深度不低于厚度[=：:]?50mm",
        "蓄水深度和持续时间按设计文件及适用验收标准执行",
        value,
    )
    value = re.sub(
        r"设备型号[=：:]?20t挖机1台(?:（量化默认值）|（按工序替换）)?",
        "与作业面及院区交通条件相匹配的小型开挖设备",
        value,
    )
    value = value.replace("20t挖机1台", "与作业面及院区交通条件相匹配的小型开挖设备")
    value = value.replace("20t挖机", "与作业面及院区交通条件相匹配的小型开挖设备")
    value = value.replace("50m车载泵", "按专项方案选配的混凝土输送设备")
    value = re.sub(
        r"(?:连续作业)?时长[=：:]?4h/作业段",
        "作业时长按班次计划与职业健康要求执行",
        value,
    )
    value = re.sub(
        r"时长[（(]?4h/作业段[）)]?",
        "作业时长按班次计划与职业健康要求执行",
        value,
    )
    value = value.replace("单次搅拌量满足4h/作业段施工", "单次搅拌量与当班可使用量匹配")
    value = value.replace("照片上传系统延迟：时长≤4h/作业段", "照片须随检验批及时上传并完成关联")
    value = re.sub(
        r"总人数配置满足8人/班的作业分组",
        "班组人数按工作面和经批准的劳动力计划配置",
        value,
    )
    value = re.sub(
        r"扫码领料覆盖率≥95%",
        "纳入范围的材料领用全部扫码留痕",
        value,
    )
    value = value.replace(
        "整改偏差应立即隔离并在复验合格后销项/作业段",
        "发现不合格时立即隔离整改，复验合格后销项",
    )
    value = re.sub(
        r"进度纠偏行动需在偏差应立即隔离并在复验合格后销项内下达指令",
        "发现进度偏差后立即下达纠偏指令",
        value,
    )
    value = value.replace(
        "时限：偏差应立即隔离并在复验合格后销项",
        "处置：立即隔离整改，复验合格后销项",
    )
    value = value.replace(
        "限定偏差应立即隔离并在复验合格后销项",
        "执行立即隔离整改、复验合格后销项的闭环要求",
    )
    value = value.replace(
        "执行偏差应立即隔离并在复验合格后销项的刚性闭环",
        "立即隔离整改并在复验合格后销项",
    )
    value = value.replace(
        "要求备用人员在立即启动备用人员调配，并依据经批准的进度计划调整作业安排",
        "立即启动备用人员调配，并依据经批准的进度计划调整作业安排",
    )
    value = value.replace(
        "在规定偏差应立即隔离并在复验合格后销项内进行局部复涂",
        "立即进行局部复涂",
    )
    value = value.replace(
        "系统于偏差应立即隔离并在复验合格后销项内自动报警",
        "系统在出现异常时自动报警",
    )
    value = value.replace(
        "间距=1000mm；厚度=50mm",
        "间距和厚度按设计图纸、材料体系及适用规范确定",
    )
    value = value.replace(
        "在找平层上铺设电导率传感网格（间距=1000mm布置）",
        "在找平层上按经批准的专项方案铺设电导率传感网格",
    )
    value = value.replace(
        "电工施放线缆，间距=1000mm绑扎固定",
        "电工施放线缆，并按设计及适用规范要求分段绑扎固定",
    )
    value = value.replace(
        "预设间距=1000mm、吊杆直径8mm等核心检验参数",
        "预设吊杆间距不大于1000mm、吊杆直径8mm等核心检验参数",
    )
    value = value.replace(
        "综合布线强弱电缆保持间距=1000mm以上",
        "综合布线强弱电缆保持设计及适用规范要求的隔离距离",
    )
    value = value.replace(
        "以间距=1000mm为节点进行布线",
        "按经批准专项方案确定的网格间距布线",
    )
    value = value.replace(
        "；与作业面及院区交通条件相匹配的小型开挖设备。",
        "；设备按具体工序和现场条件配置。",
    )
    value = value.replace(
        "设备/工具=与作业面及院区交通条件相匹配的小型开挖设备（按工序替换）",
        "设备/工具=按具体工序和现场条件配置",
    )
    value = value.replace(
        "数据超标即时停止土方作业，加派与作业面及院区交通条件相匹配的小型开挖设备协同覆盖防尘网，增设雾炮车（时限：30min内）",
        "数据超标时立即停止土方作业，组织人员覆盖裸土并启动雾炮喷淋，复核达标后恢复作业",
    )
    value = value.replace(
        "针对高风险任务触发安全响应控制补强节点",
        "针对高风险任务增配安全监督人员并提高过程巡检频次",
    )
    value = value.replace(
        "安全响应控制补强节点",
        "安全风险预警与处置机制",
    )
    value = value.replace(
        "进度响应控制补强节点",
        "进度偏差预警与纠偏机制",
    )
    value = value.replace(
        "质量响应控制补强节点",
        "质量偏差复核与闭环机制",
    )
    value = value.replace(
        "环保响应控制补强节点",
        "环境指标预警与处置机制",
    )
    value = value.replace(
        "清单-工序-资源映射补强节点",
        "清单、工序与资源匹配复核机制",
    )
    value = value.replace(
        "清单-工序-资源映射",
        "清单与工序资源匹配",
    )
    value = value.replace(
        "触发“清单、工序与资源匹配复核机制”评估实体返工率。",
        "施工过程中按检验批记录实体质量，并对不合格项复验销项。",
    )
    value = value.replace(
        "证据未给出红线边界及夜间禁行具体要求 → 编制口径：按合肥市中心城区管控标准，"
        "暂定晚22:00-早6:00禁止重型渣土车及材料运输车通行，场区内部划定4m宽单向循环通道1处 "
        "→ 需澄清项：具体交通开口坐标及夜间施工许可办理权限归属。",
        "施工前联合建设单位复核改造区域红线、材料运输时段、院区通行路线和交通开口；"
        "未取得书面确认前，不采用临时假定作为施工依据。",
    )
    return value


def _rewrite_hospital_renovation_scope(text: str, project_topic: Any) -> str:
    """Keep a hospital partial-renovation submission inside its evidenced scope.

    Model drafts sometimes import generic new-build crews, heavy equipment,
    experimental sensors, or bare BoQ numbers into a renovation narrative.  For
    this project type those claims are either unsupported or actively distract
    from the operational hospital constraints.  Replace them with conventional,
    verifiable renovation controls instead of silently presenting them as facts.
    """
    topic = str(project_topic or "")
    if "医院" not in topic or "改造" not in topic:
        return str(text or "")

    value = str(text or "")
    replacements = (
        (
            "防水层渗漏电导率阵列监测（新工艺）",
            "防水隐蔽验收与蓄水/淋水检验",
        ),
        (
            "防水层渗漏电导率阵列监测系统",
            "防水隐蔽验收与蓄水/淋水检验制度",
        ),
        (
            "防水层渗漏电导率阵列监测",
            "防水隐蔽验收与蓄水/淋水检验",
        ),
        (
            "基于BIM的管线综合防碰撞与绿色施工技术",
            "既有机电管线复核与专业接口协调",
        ),
        (
            "基础底板及二次结构工程质量控制",
            "局部修复、开槽回补及安装基层质量控制",
        ),
        (
            "防爆工业级手持扫码终端机（PDA）4台，热转印工业条码打印机2台，二维码耗材50卷",
            "经项目批准的移动扫码终端和标签打印设备，数量按材料收发工作量配置",
        ),
        (
            "防爆防摔三防平板电脑6台（配置微距防抖摄像头），携带RTK精准定位模块的劳保安全帽8顶，红外测距仪6台",
            "项目受控移动终端、激光测量工具和影像采集设备，数量按作业面配置",
        ),
        (
            "测量工1人/班，模板工（吊顶安装）4人/班，防腐工2人/班，质检员1人",
            "测量放线、吊顶安装、防腐和质量检查人员按作业面及经批准的劳动力计划配置",
        ),
        (
            "测量工、钢筋工、模板工、混凝土工、架子工、防水工、电工、焊工、管道工、起重信号司索工、机械设备操作工等11类技术工种",
            "装饰装修工、防水工、电工、弱电安装工、管道工、通风空调工、焊工及测量放线人员等与本项目范围相匹配的专业工种",
        ),
        (
            "测量工、钢筋工、模板工、混凝土工、架子工、防水工、电工、焊工、管道工、起重信号司索工、机械设备操作工",
            "装饰装修工、防水工、电工、弱电安装工、管道工、通风空调工、焊工及测量放线人员",
        ),
        (
            "地下结构、主体框架、二次结构、机电智能化安装、装饰装修等五个控制节点",
            "区域移交、拆改与基层处理、机电管线安装、隐蔽验收、装饰收口和系统联调等控制节点",
        ),
        (
            "配置人数按工作面和经批准的劳动力计划配置的专业攻坚组（含模板工、钢筋工、混凝土工等），实行双班轮替作业（16小时工作制），保证工序连续",
            "按区域开放条件和经批准的劳动力计划配置专业班组，并依据院方确认的时段组织错峰施工",
        ),
        (
            "钢筋下料单复核，模板排版图绘制",
            "复核拆改边界、基层条件、既有管线位置和修复做法，完成样板确认",
        ),
        (
            "钢筋工按照间距绑扎墙柱纵筋与箍筋，模板工使用新型支撑体系加固，混凝土工分层浇筑（厚度≤500mm），振捣棒快插慢拔",
            "清理松散基层，按设计材料分层修补找平；机电开槽封堵和安装基层按隐蔽验收结果施工",
        ),
        (
            "混凝土初凝后覆盖薄膜保湿养护，时长不少于14天",
            "修复区域按材料体系要求养护并设置成品保护，达到后续工序条件后办理移交",
        ),
        (
            "配置起重信号司索工及特种设备操作工，持证上岗率100%",
            "涉及登高、动火和临时用电的人员按作业类别持有效证件上岗",
        ),
        (
            "针对基础、主体、二次结构、屋面与防渗、装饰及智能化等核心工序",
            "针对拆改、基层修复、防水、装饰装修、机电安装及智能化等核心工序",
        ),
        (
            "完全覆盖基础、主体、二次结构、屋面与防渗及装饰、智能化",
            "覆盖拆改与基层修复、防水、装饰装修、机电安装及智能化",
        ),
        (
            "土方、吊装、机电安装阶段",
            "室外局部开挖、拆改、材料转运和机电安装阶段",
        ),
        (
            "如提升机、与作业面及院区交通条件相匹配的小型开挖设备等",
            "如小型运输、切割、钻孔、登高及测试设备等",
        ),
        (
            "带GPS/BIM坐标的三维影像资料",
            "带施工部位、房间/轴线、标高和时间信息的影像资料",
        ),
        (
            "配置BIM坐标插件的移动终端",
            "配置部位标识和时间水印的项目受控移动终端",
        ),
        (
            "扫码出库指令与BIM模型用量强制比对",
            "扫码出库指令与经批准的材料计划用量逐项核对",
        ),
        (
            "BIM导出的无碰撞排布图及精准下料清单",
            "经会审确认的综合管线排布图及材料计划",
        ),
        (
            "利用BIM平台建立地下一层及地上建筑高精度三维模型",
            "汇总各专业图纸和现场复测成果，形成综合管线叠图及接口问题清单",
        ),
        (
            "BIM图形工作站1台",
            "综合管线复核所需的受控办公终端和制图软件",
        ),
        (
            "BIM工程师1人",
            "机电专业工程师按接口复核工作量配置",
        ),
        (
            "现场实测标高输入系统，系统自动计算并标注",
            "现场实测标高与经会审确认的综合管线排布图逐项核对，标注",
        ),
        (
            "指导模板工调整主龙骨吊放位置",
            "指导吊顶安装人员调整主龙骨和吊点位置",
        ),
        (
            "模板工安装拉爆螺丝",
            "吊顶安装人员按批准做法安装后置锚固件",
        ),
        (
            "模板工（吊顶安装）",
            "吊顶安装工",
        ),
        (
            "测量工佩戴RTK定位安全帽，手持平板",
            "施工员使用项目受控移动终端",
        ),
        (
            "质量员手持防爆平板电脑至点位",
            "质量员使用项目受控移动终端到点检查",
        ),
        (
            "一键上传云端服务器",
            "上传至项目受控资料库",
        ),
        (
            "同步至云端归档",
            "同步至项目受控资料库归档",
        ),
        (
            "20t挖机1台",
            "与作业面、地下管线和院区通行条件相匹配的小型开挖设备",
        ),
        (
            "50m车载泵",
            "按具体修复工序和现场条件选配的施工设备",
        ),
        (
            "雾炮机",
            "局部吸尘与移动雾化降尘设备",
        ),
        (
            "雾炮车",
            "移动雾化降尘设备",
        ),
        (
            "雾炮设备",
            "移动雾化降尘设备",
        ),
        (
            "BIM模型",
            "经会审确认的综合管线排布图",
        ),
        (
            "BIM碰撞检测",
            "专业接口复核",
        ),
        (
            "BIM出图",
            "经会审确认的综合管线排布图",
        ),
        (
            "BIM图纸",
            "综合管线排布图",
        ),
        (
            "BIM坐标",
            "施工部位标识",
        ),
        (
            "BIM",
            "综合管线复核",
        ),
        (
            "PDA",
            "移动扫码终端",
        ),
        (
            "GPS定位与标高时间戳",
            "施工部位、标高和时间标识",
        ),
        (
            "GPS/综合管线复核",
            "施工部位",
        ),
    )
    for source, target in replacements:
        value = value.replace(source, target)

    # Replace unsupported sensor claims with conventional, auditable waterproof
    # acceptance steps.  These sentences are intentionally conservative.
    value = re.sub(
        r"在找平层上(?:按经批准的专项方案)?铺设电导率传感网格[^。]*。?",
        "防水施工前复核基层、附加层、搭接和收头条件，形成隐蔽验收记录。",
        value,
    )
    value = re.sub(
        r"(?:敷设|铺设|网格化铺设)[^。；\n]*电导率(?:传感器)?阵列[^。；\n]*[。；]?",
        "对基层、附加层、搭接、收头及穿墙管根逐项旁站检查并留痕。",
        value,
    )
    value = re.sub(
        r"防水隐蔽验收与蓄水/淋水检验(?:数据反馈单|数据报告|日志)",
        "防水隐蔽验收记录和蓄水/淋水检验记录",
        value,
    )
    value = value.replace("电导率异常报警复核 (渗漏电流≤30mA，实时监测)", "渗漏点复核（检查节点、搭接、收头和穿墙管根）")
    value = value.replace("智慧监测日志 / 渗漏坐标、电流值、处理状态", "防水检查台账 / 检查部位、渗漏点、处理状态")
    value = value.replace("启动局部揭开修补，复测电导率", "局部揭开修补，并重新执行蓄水/淋水检验")
    value = re.sub(
        r"(?:系统|后台)[^。；\n]*(?:电导率|介电常数)[^。；\n]*[。；]?",
        "蓄水/淋水期间分区巡查，发现渗漏立即标记部位并启动修补复验。",
        value,
    )
    value = re.sub(
        r"(?:电导率|渗漏电流)[^。；\n]*(?:阈值|报警|基线)[^。；\n]*[。；]?",
        "防水检验按设计文件、适用验收标准和批准方案判定，渗漏点修补后重新检验。",
        value,
    )
    value = re.sub(
        r"(?:执行)?(?:24|48)小时蓄水/淋水试验",
        "执行蓄水/淋水试验，持续时间按适用验收标准和批准方案确定",
        value,
    )
    value = re.sub(
        r"(?:注入|蓄水深度为?)\s*20mm[^。；\n]*",
        "蓄水深度和持续时间按适用验收标准及批准方案执行",
        value,
    )
    value = re.sub(
        r"PM10(?:浓度)?\s*(?:≤|≥|=)?\s*\d+(?:\.\d+)?\s*(?:ug/m3|μg/m³)",
        "现场扬尘指标",
        value,
        flags=re.IGNORECASE,
    )
    value = value.replace("PM2.5及PM10报警联动控制", "粉尘监测、局部吸尘和雾化降尘联动控制")
    value = value.replace("PM10数值", "粉尘监测数据")
    value = value.replace("PM10监测仪读数", "粉尘监测数据")
    value = value.replace("PM10浓度", "粉尘监测数据")
    value = value.replace("PM10", "粉尘")
    value = value.replace("雾炮", "移动雾化降尘设备")
    value = value.replace("GPS定位", "施工部位标识")
    value = value.replace("工程云端管理平台", "项目受控资料库")
    value = value.replace("项目云盘", "项目受控资料库")
    value = value.replace("云端", "项目受控资料库")
    value = value.replace(
        "交付清单驱动物：防水基层隐蔽验收单、执行蓄水/淋水试验，持续时间按适用验收标准和批准方案确定记录、电导率阵列监测数据反馈单",
        "交付清单：防水基层隐蔽验收单、蓄水/淋水检验记录、渗漏修补复验记录",
    )
    value = value.replace(
        "作业：分层涂刷防水层，敷设渗漏电导率传感阵列，施工细石混凝土保护层厚度按设计构造和适用验收标准控制",
        "作业：按批准的材料体系分层施工防水层，管根、阴阳角、收头和搭接部位逐项旁站；保护层按设计构造和适用验收标准施工",
    )
    value = value.replace(
        "检查验收：封堵下水口，蓄水深度和持续时间按适用验收标准及批准方案执行；作业时长按班次计划与职业健康要求执行电导率监测",
        "检查验收：封堵排水口，蓄水深度和持续时间按适用验收标准及批准方案执行；发现渗漏后标记部位、修补并重新检验",
    )
    value = value.replace(
        "《防水施工与电导率监测日志》（记录厚度、电阻率变化峰值）",
        "《防水施工与蓄水/淋水检验记录》（记录施工部位、节点做法、检验时段和处理结果）",
    )
    value = value.replace(
        "执行48小时闭水试验，结合电导率数据双重判定",
        "执行蓄水/淋水试验，持续时间和判定标准按适用验收标准及批准方案确定",
    )
    value = value.replace(
        "依托电导率阵列监测技术，实时监控水分子穿透防水层后的介电常数变化",
        "依托隐蔽验收、蓄水/淋水检验和分区巡查，对管根、收头、搭接和节点部位重点复核",
    )
    value = value.replace(
        "在防水层下方网格化铺设电导率传感阵列，按经批准专项方案确定的网格间距布线，接入智慧工地监控后台",
        "防水施工前复核基层、附加层、搭接、收头及穿墙管根，过程影像和验收结果录入项目受控资料库",
    )
    value = value.replace(
        "防水层下方网格化铺设渗漏电导率传感器阵列",
        "基层、附加层、搭接、收头及穿墙管根逐项检查并留痕",
    )
    value = value.replace("电导率阵列监测数据报告", "防水隐蔽验收和蓄水/淋水检验记录")
    value = value.replace("电导率监测数据反馈单", "蓄水/淋水检验记录")
    value = value.replace("电导率监测", "蓄水/淋水检验")
    value = value.replace("渗漏电导率传感器阵列", "防水节点检查清单")
    value = value.replace("渗漏电导率传感阵列", "防水节点检查清单")
    value = value.replace("电导率阵列监测技术", "防水隐蔽验收与蓄水/淋水检验制度")
    value = value.replace("电导率数据", "蓄水/淋水检验记录")
    value = value.replace(
        "起重信号司索工及各专业工种按资源映射节点流水施工",
        "各专业班组按经批准的穿插计划和区域移交条件流水施工",
    )
    value = value.replace(
        "起重信号司索工及机械设备操作工凭证上岗，执行“定人、定机、定岗”制度",
        "涉及设备操作的人员按设备类别持有效证件上岗，执行定人、定机、定岗管理",
    )
    value = re.sub(
        r"技术工种配置：测量工1人/班；钢筋工2人/班；模板工2人/班；混凝土工2人/班；"
        r"架子工2人/班；电工1人/班；焊工1人/班；起重信号司索工1人/班（起重作业时）[。.]?",
        "技术工种配置：装饰装修、防水、电气、弱电、管道、通风空调、焊接及测量放线人员按作业面和经批准的劳动力计划配置。",
        value,
    )
    value = value.replace(
        "防水工、钢筋工、电工作业过程中的隐蔽工程节点",
        "防水、电气、管道和装饰安装作业过程中的隐蔽工程节点",
    )
    value = value.replace(
        "数据超标时立即停止土方作业，组织人员覆盖裸土并启动移动雾化降尘设备喷淋，复核达标后恢复作业",
        "粉尘指标异常时立即停止室外开挖或拆除作业，组织局部吸尘、覆盖和雾化降尘，复核达标后恢复作业",
    )
    value = value.replace(
        "结合工程量及造价清单重点（如：9.1工程量=6.0、030411004005合价=17197.0等）",
        "结合正式工程量清单和图纸确定的重点材料、设备与施工部位",
    )
    value = re.sub(
        r"本项目涉及门诊综合楼[^。\n]*?重点管控造价高、单价高的设备与材料，包含清单项：[^。\n]*。",
        "本项目涉及门诊综合楼和运动训练楼的弱电智能化系统，重点管控综合布线、网络、机房、广播、呼叫等系统的设备材料、接口条件和联调记录。",
        value,
    )
    value = re.sub(
        r"检查验收与资料归档：技术人员登录系统后台[^。\n]*图纸工程量[^。\n]*。[^。\n]*。",
        "检查验收与资料归档：材料员、施工员和质量员核对领料数量、安装部位、剩余退库和检验批记录，形成可追溯的材料流向台账。",
        value,
    )
    value = re.sub(
        r"涉及高单价材料清单：[^。\n]*。",
        "需重点管控的材料和设备按正式工程量清单、设计文件及采购计划确定。",
        value,
    )
    value = re.sub(
        r"高价材料管控项：[^。\n]*。",
        "需重点管控的材料和设备建立批次、进场报验、领用部位和退库记录。",
        value,
    )
    value = re.sub(
        r"依据图纸清单及高价材料项（[^）]*），编制物资采购计划。采购比价=≥3家/批次",
        "依据图纸、正式工程量清单和经批准的材料计划编制采购计划，询比价和审批按招标要求及公司制度执行",
        value,
    )
    value = re.sub(
        r"高价值智能化线缆（如清单[^）]*）",
        "需重点管控的智能化设备和线缆",
        value,
    )
    value = re.sub(
        r"由技术负责人组织测量工、机电工程师核对智能化16个系统与原建筑结构图纸，标定全部碰撞点；"
        r"建立高价值清单项（[^）]*）采购台账",
        "由技术负责人组织测量放线和机电专业人员核对智能化系统与既有建筑图纸、现场条件及装饰标高，形成接口问题清单；建立重点材料和设备采购台账",
        value,
    )
    value = re.sub(
        r"按照领料单提取[^。；\n]*，利用施工电梯运至各楼层指定作业区，堆载高度受控",
        "按照领料单配发经报验合格的管材、线缆、装饰材料和设备，沿院方确认的运输路线运至指定作业区，并控制临时堆载",
        value,
    )
    value = re.sub(
        r"适用条件：针对清单中合价超1万的[^。\n]*。",
        "适用条件：需重点管控的智能化设备、线缆、管材和装饰材料。",
        value,
    )
    # The final submission must read as a construction plan, not as an AI
    # scoring memo.  Replace review-language, invented formulas and generic
    # new-build logistics with hospital-renovation controls that can actually
    # be implemented and accepted on site.
    polished_replacements = (
        (
            "《项目特征分析及清单量价评估报告》",
            "《项目特征与施工界面复核记录》《场地与运输路线查勘记录》《重点材料设备采购与报验台账》",
        ),
        ("引发重大偏差扣分", "形成重大质量安全风险"),
        ("规避否决扣分项", "控制关键质量、安全和进度风险"),
        (
            "依托工程项目响应控制补强机制，对质量、安全、进度目标设定数字化红线管控，控制关键质量、安全和进度风险：",
            "建立质量、安全和进度预警与纠偏机制，对关键工序、接口和院区运行风险实施闭环控制：",
        ),
        (
            "控制：计算高风险任务占比指标（公式：高风险任务数÷任务总数×100%），当指标过高时增加安全监督岗。执行应急响应时间考核（公式：实际应急响应时长与目标应急响应时长的比值）。现场布置防尘喷雾，裸土及渣土采用阻燃密目网连续覆盖并压边固定。",
            "控制：高处、动火、临时用电、拆除和吊装等高风险作业严格执行作业许可、班前交底、过程监护和完工销项；出现异常立即停工、隔离并按批准的应急预案处置。室外局部开挖和拆除作业采用局部吸尘、覆盖与雾化降尘。",
        ),
        (
            "验证：专职安全员每班开展隐患排查（检查频次按班前、过程巡检和收工复核执行），应急演练频次=1次/季度，维持零安全生产事故记录。",
            "验证：专职安全员开展班前检查、过程巡检和收工复核；应急演练按施工阶段和批准的应急预案组织，问题整改后复验销项。",
        ),
        (
            "控制：应用进度延误率模型（公式：累计延误时长÷计划时长×100%）逐日核算偏差。当延误超前置预警线，立即触发资源动态调配，增加夜班作业时段或投入备用设备。",
            "控制：计划人员每日比对实际完成量与经批准的计划，发现偏差立即分析接口、材料、人员和工作面原因；仅在院方批准的作业时段内调整穿插顺序和资源配置。",
        ),
        (
            "控制：严密监测返工体量比率（公式：返工率〔返工工程量÷总工程量×100%〕）。石材缝隙、天棚标高控制及系统电缆绝缘测试由专人盯控，工序资源按清单映射匹配（例如 重点清单项 对应专项熟练技术工种）。",
            "控制：按检验批统计返工原因；石材铺贴、吊顶标高和系统电缆绝缘测试由专人复核，正式清单项、施工部位、工序、资源和验收资料逐项对应。",
        ),
        (
            "若系统功能测试不达标，构成质量验收重大偏差。执行设备材料进场100%报验，不合格品即刻退场，杜绝使用劣质材料导致否决项触发。",
            "系统功能测试或材料报验未通过时不得进入下一工序；不合格材料隔离标识，并按审批流程退换处理。",
        ),
        (
            "吊顶拉爆螺丝松动及石材空鼓视为安全与质量双重否决项。执行100%全覆盖敲击排查，消除维修改造工程常见的坠落与返工风险。",
            "吊顶后置锚固件松动和石材空鼓属于重点质量安全风险；按检验批和施工区域逐项检查，问题整改复验后方可移交。",
        ),
        (
            "防止材料溯源不清引发的进度与质量纠纷，信息化台账将作为工程结算和验收的铁证，杜绝审计核减风险。",
            "材料批次、报验、领用部位和退库记录形成可追溯台账，作为工程验收和结算的核验依据。",
        ),
        (
            "控制：启用进度响应控制补强算法（累计延误时长÷计划时长×100%），动态调整关键线路缓冲期。",
            "控制：每日核对关键线路和区域移交条件，按批准的预警标准分析偏差并调整穿插顺序、材料到货和人员配置。",
        ),
        (
            "超期交付为合同重大违约扣分项。将各阶段资源按经批准的进度计划和实际工作面动态配置按网格化分解至各作业段，保障日均产值底线，规避延期罚款。",
            "将各阶段资源按经批准的进度计划、院区开放条件和实际工作面配置到各作业段，持续控制合同履约风险。",
        ),
        (
            "安全事故是项目管理的绝对红线与否决项。本措施彻底消除重大安全与环保违规隐患，实现零事故目标，打造绿色文明标准工地。",
            "安全、院感、消防和环保要求是项目管理底线；各类风险按作业许可、过程检查、整改复验和资料归档闭环控制。",
        ),
        (
            "本章详细阐述本项目（含门诊综合楼改造、运动训练楼智能化及装饰改造）中拟采用的四新技术与信息化管理技术。各项工艺和技术均围绕项目清单重点项与图纸规范编制，所有参数设定与操作规程均经过校验，无重大偏差，规避废标与扣分项。施工本工程总工期以招标文件及经批准的总进度计划为准，各阶段资源按经批准的进度计划和实际工作面动态配置，关键线路节拍以经批准的进度计划为准。",
            "本章围绕门诊综合楼和运动训练楼的智能化、装饰及相关改造工作，选择可实施、可验证、可移交的新技术。工期、资源和关键线路均以招标文件及经批准的计划为准。",
        ),
        (
            "控制：在验收APP中配置“强制全景+微距必拍点”算法。对于门诊大厅、发热门诊等核心改造区域，必须完整上传“吊点焊接局部图”“防锈漆涂装细节图”“防火涂料三遍对比图”，集齐后方可激活下方“提交隐蔽验收”按钮，执行防呆锁死控制。",
            "控制：验收程序将全景、局部和尺量影像设置为隐蔽验收必备资料；资料齐全并经复核后方可提交验收。",
        ),
        (
            "测量复核与碰撞检测：将墙上1m高水平控制线作为标高基准，载入模型进行软硬碰撞检测，系统自动生成干涉点清单。综合管线复核工程师通过调整管线路由，消除所有碰撞点。",
            "测量与接口复核：以现场水平控制线和复测标高为基准，将各专业图纸叠合并与现场条件核对，形成干涉点清单；经专业会审确认后调整管线路由和末端位置。",
        ),
        ("视为安全与质量双重否决项", "属于重点质量安全风险"),
        ("补强机制", "预警与纠偏机制"),
        ("补强算法", "预警与纠偏机制"),
        ("清单映射", "清单对应关系"),
        ("三维映射数据库", "材料安装部位台账"),
        ("进度模型", "进度计划"),
        ("扣分项", "合同与技术风险"),
        ("否决项", "关键风险"),
        ("铁证", "核验依据"),
        ("评审要求", "合同与技术要求"),
        (
            "本章节内容对标招标文件合同与技术要求，控制关键质量、安全和进度风险。",
            "本节明确合同、图纸、规范和审批流程的执行要求。",
        ),
        ("智能电导率阵列监测预警", "防水隐蔽验收、蓄水/淋水检验和渗漏巡查"),
        ("发热门诊或运动训练楼", "门诊综合楼或运动训练楼"),
        ("门诊大厅、发热门诊等核心改造区域", "门诊综合楼和运动训练楼的重点改造区域"),
        ("施工电梯、塔吊的起重量及运行行程", "院方批准的垂直运输路线及升降平台等设备的额定载荷和作业范围"),
        ("超过一定规模的吊顶标高定位、外脚手架", "高处、动火、临时用电、拆除和吊装（如涉及）"),
        ("防水工+瓦工", "石材铺装工+质量员"),
        ("系统台账数据与实物相符率100%，台账抽查频次=1次/周。防范高价材料流失或用错批次。", "系统台账与实物逐项核对，按周抽查并留痕，防范重点材料错领、错用或去向不清。"),
        ("所有隐蔽工程数字影像留存率100%，支撑项目终验追溯。", "隐蔽工程按检验批留存部位清晰、时间可追溯的影像资料，支撑过程验收和竣工移交。"),
        ("超期交付为合同重大违约合同与技术风险。", "进度偏差可能形成合同履约风险。"),
    )
    for source, target in polished_replacements:
        value = value.replace(source, target)

    value = re.sub(
        r"机械设备配置：[^。\n]*。",
        "机械设备按具体工序和作业面配置，重点采用带集尘切割或钻孔设备、移动吸尘器、液压升降平台、激光测量工具、光纤熔接与测试设备及石材切割设备；室外局部开挖确需机械时，选用与地下管线和院区通行条件相匹配的小型设备。",
        value,
    )
    value = re.sub(
        r"本章节内容对标招标文件[^。\n]*：",
        "本节明确合同、图纸、规范和审批流程的执行要求：",
        value,
    )
    value = re.sub(
        r"全面规避废标与重大偏差：[^\n]*",
        "开工前核对招标目录、技术条款和施工范围；涉及设计或范围变化时先履行书面确认，不降低质量与安全标准。",
        value,
    )
    value = re.sub(
        r"技术规避项声明：[^\n]*",
        "技术响应说明：施工前核对图纸、现场和材料设备接口，按批准流程解决差异；质量、安全、进度和环保要求均纳入检查与整改闭环。",
        value,
    )
    value = re.sub(
        r"基于本项目本工程总工期[^\n]*",
        "本项目在运营医院环境内实施，安全、院感、医疗秩序、进度和环保要求相互制约；项目部按准备、复核、材料、施工、验收五步闭环组织各作业区。",
        value,
    )
    value = re.sub(
        r"应用进度延误率模型[^。\n]*逐日核算偏差。[^。\n]*。",
        "计划人员每日比对实际完成量与经批准的计划，发现偏差立即分析接口、材料、人员和工作面原因；仅在院方批准的作业时段内调整穿插顺序和资源配置。",
        value,
    )
    value = re.sub(
        r"计算高风险任务占比指标[^\n]*",
        "高处、动火、临时用电、拆除和吊装等高风险作业严格执行作业许可、班前交底、过程监护和完工销项；出现异常立即停工、隔离并按批准的应急预案处置。",
        value,
    )
    value = re.sub(
        r"严密监测返工体量比率[^\n]*",
        "按检验批统计返工原因；石材铺贴、吊顶标高和系统电缆绝缘测试由专人复核，正式清单项、施工部位、工序、资源和验收资料逐项对应。",
        value,
    )
    value = re.sub(
        r"启用进度响应控制[^\n]*",
        "每日核对关键线路和区域移交条件，按批准的预警标准分析偏差并调整穿插顺序、材料到货和人员配置。",
        value,
    )
    value = re.sub(
        r"在验收APP中配置[^。\n]*算法。[^\n]*防呆锁死控制。",
        "验收程序将全景、局部和尺量影像设置为隐蔽验收必备资料；资料齐全并经复核后方可提交验收。",
        value,
    )
    value = re.sub(
        r"测量复核与碰撞检测[^：:\n]*[：:][^\n]*",
        "测量与接口复核：以现场水平控制线和复测标高为基准，将各专业图纸叠合并与现场条件核对，形成干涉点清单；经专业会审确认后调整管线路由和末端位置。",
        value,
    )
    value = re.sub(
        r"技术规避项声明[^：:\n]*[：:][^\n]*",
        "技术响应说明：施工前核对图纸、现场和材料设备接口，按批准流程解决差异；质量、安全、进度和环保要求均纳入检查与整改闭环。",
        value,
    )
    value = re.sub(
        r"基于本项目[^\n]*规模特点[^\n]*",
        "本项目在运营医院环境内实施，安全、院感、医疗秩序、进度和环保要求相互制约；项目部按准备、复核、材料、施工、验收五步闭环组织各作业区。",
        value,
    )
    value = re.sub(
        r"粉尘阈值\s*(?:[=：]\s*)?(?:≤\s*)?\d+(?:\.\d+)?\s*(?:ug/m3|μg/m³)?",
        "粉尘控制指标执行属地环保要求和院方管理规定",
        value,
        flags=re.IGNORECASE,
    )
    value = re.sub(
        r"昼间噪声阈值\s*(?:[=：]\s*)?≤?\s*\d+\s*dB[^。；\n]*夜间(?:噪声阈值)?\s*(?:[=：]\s*)?≤?\s*\d+\s*dB",
        "施工噪声执行属地标准和院方批准的作业时段",
        value,
        flags=re.IGNORECASE,
    )
    value = re.sub(r"(?:检查|抽检|台账抽查|应急演练|上传)频次\s*=\s*1次/季度", "实施安排按施工阶段和批准的专项计划", value)
    value = re.sub(r"(?:检查|抽检|台账抽查|上传)频次\s*=\s*1次/周", "按周抽查并留痕", value)
    value = re.sub(r"洒水车每日按频次\s*=\s*4次/日", "道路降尘频次根据作业强度和粉尘监测结果调整", value)
    value = re.sub(r"(?:质量)?阈值\s*=\s*", "", value)
    value = re.sub(r"抽检频次\s*=\s*", "抽检频次按", value)
    value = re.sub(r"频次\s*=\s*", "频次按", value)
    value = re.sub(r"(?:材料)?采购比价\s*=\s*≥?3家/批次", "材料询比价和审批按招标要求及公司制度执行", value)
    value = value.replace("材料采购比价≥3家/批次", "材料询比价和审批按招标要求及公司制度执行")
    value = value.replace("采购比价≥3家/批次", "材料询比价和审批按招标要求及公司制度执行")
    value = value.replace("效率映射关联", "人员投入与完成量对应关系")
    value = value.replace("施工班组通过移动端“扫码领料”，数据自动关联WBS进度计划，扣减工程量清单库存额度。", "施工班组通过移动端扫码领料，领料数量、材料批次、安装部位和退库记录与材料台账关联。")
    value = value.replace("合规管理与合同与技术风险显式规避说明", "合同与技术标准响应管理")
    value = value.replace("合规管理与扣分项显式规避说明", "合同与技术标准响应管理")
    value = value.replace("材料进场抽检频次按每100m2 1次", "材料进场按批次报验和抽检")
    value = value.replace("系统联调时测量信噪比及衰减度，抽检频次按每100m2 1次", "系统联调按子系统、回路和点位测量信噪比及衰减度")
    value = value.replace("当周计划滞后超过20%时，触发进度响应控制预警与纠偏机制", "达到经批准的进度预警条件时，立即启动偏差分析和纠偏")
    value = value.replace("水土保持沉淀池清理记录", "建筑垃圾分类清运和施工废水处置记录")
    value = value.replace("本章证据及验收要求", "本章控制要求和适用验收标准")
    value = value.replace("节点映射监测100%覆盖", "节点检查记录按检验批完整留存")
    value = value.replace("保证最终项目整体质量等级达到合格标准要求，杜绝重大偏差出现", "确保实体质量和资料满足合同、设计及适用验收标准，问题整改复验后闭环")
    value = value.replace("各类特种作业（架子工、焊工等）必须100%持证", "高处、焊接、临时用电等特种作业人员必须持有效证件")
    value = value.replace("土方开挖与裸露区域100%覆盖防尘网", "室外局部开挖和裸露区域连续覆盖防尘网并压边固定")
    value = value.replace("应急响应时间（实际应急响应时长）达标率100%", "异常情况按批准的应急预案及时响应并形成处置记录")
    value = value.replace("管线空间坐标标高与经会审确认的综合管线排布图一致率：100%", "管线标高、路由和末端位置经复核后应与会审确认的综合管线排布图一致")
    value = value.replace("按到货批次100%实施", "按到货批次和施工区域实施")
    value = value.replace("涂布率达100%", "石材六面防护液涂布完整")
    value = value.replace("每作业面隐蔽前100%全数检查", "每个作业面隐蔽前逐段检查")
    value = value.replace("防水隐蔽验收与蓄水/淋水检验制度实时在线", "防水隐蔽验收及蓄水/淋水检验记录完整")
    value = value.replace("杜绝逾期风险", "控制工期履约风险")
    value = value.replace("杜绝水泥砂浆泛碱及变色", "降低水泥砂浆泛碱和石材变色风险")
    value = value.replace("照片上传率100%", "每个检验批的规定影像资料完整")
    value = value.replace("物流数据与领料数据相符率100%", "物流、领料和安装部位记录一致")
    value = value.replace("扫码识别率100%，实物库存与系统台账误差率0", "扫码信息可识别，实物库存、领用和退库记录与系统台账一致")
    value = value.replace("扫码出库准确率100%", "扫码出库记录与领料单、实物和安装部位一致")
    value = value.replace("隐蔽部位影像覆盖率100%", "隐蔽部位按检验批留存规定影像")
    value = value.replace("批次字段齐全率=100%", "规定批次字段齐全")
    value = value.replace("影像覆盖率=100%", "规定影像资料齐全")
    value = value.replace("执行项目安全等级要求，杜绝重大安全事故，实现文明施工", "执行项目安全管理要求，控制重大安全风险并落实文明施工")
    value = value.replace("关键标准定位：杜绝水泥砂浆泛碱", "关键标准：降低水泥砂浆泛碱风险")
    value = value.replace("杜绝水泥砂浆泛碱", "降低水泥砂浆泛碱风险")
    value = value.replace("执行“六个百分百”要求，裸土及易扬尘物料100%覆盖", "涉及室外土方或易扬尘物料时，按属地扬尘治理要求连续覆盖并压边固定")
    value = value.replace("不合格批次=100%隔离并在24h内退换", "不合格批次立即隔离，并按审批流程退换")
    value = value.replace("台账字段齐全率=100%+上传频次按1次/日", "规定台账字段齐全并及时上传")
    value = value.replace("抽查覆盖率=100%，记录=《资料台账》；偏差处置：缺项≤24h补齐并复核关闭", "按计划抽查并记录于《资料台账》；缺项及时补齐，复核后关闭")
    value = value.replace("巡检记录齐全率=100%，违章=0次/日", "巡检记录齐全，违章行为立即纠正")
    value = value.replace("试跳记录齐全率=100%，带病运行=0次", "漏电保护试跳记录齐全，不合格设备不得运行")
    value = value.replace("围挡闭合率=100%+道路硬化=100%+喷淋=按班前、过程巡检和收工复核+车辆冲洗=1次/车+覆盖=100%", "围挡或隔离完整，产尘作业同步吸尘或喷雾，车辆按规定冲洗，裸露物料连续覆盖")
    value = value.replace("现场扬尘指标（监测=1次/日），投诉=0次/周", "产尘作业期间监测粉尘并记录投诉处置")
    value = value.replace("粉尘超限≤15min启动加密喷淋，2h内复测达标", "粉尘指标异常时立即停止产尘作业、加强吸尘或喷雾，复测合格后恢复")
    value = re.sub(
        r"材料进场时启动采购比价机制[^。\n]*。质检员按每100m²\s*1次的抽检频次对外护套、线径及防火涂料理化指标进行检验，[^。\n]*。",
        "材料进场按批次报验，重点核对外护套、线径、防火涂料合格证和复验资料；抽检方法与频次按设计、适用标准和经批准的检验计划执行。",
        value,
    )
    value = re.sub(
        r"技术负责人每工作段验证综合管线排布图与实体的一致性[^\n]*",
        "技术负责人按作业区复核综合管线排布图与实体的一致性；粉尘监测、吸尘和雾化降尘措施在产尘作业期间同步检查，应急演练按施工阶段和批准的专项计划组织。",
        value,
    )
    value = value.replace("环境监测系统24小时实时传输", "产尘作业期间监测并记录")
    # Last-resort vocabulary guard for bidder-facing hospital submissions.
    value = value.replace("否决", "关键风险")
    value = value.replace("扣分", "履约风险")
    # Normalize split or unsplit bare BoQ identifiers only after the surrounding
    # prose has been rewritten. A code without an item name is not useful to a
    # reviewer and must not masquerade as project-specific detail.
    value = re.sub(r"\b(0\d{8})\s*\n\s*(\d{3,5})\b", r"\1\2", value)
    value = re.sub(r"\b(WB\d{7})\s*\n\s*(\d{3,6})\b", r"\1\2", value, flags=re.IGNORECASE)
    value = re.sub(r"\b(?:0\d{8,}|WB\d{8,})\b", "重点清单项", value, flags=re.IGNORECASE)

    # Bare BoQ codes and price/quantity fragments without item names are not
    # intelligible tender content. Replace the first occurrence in each chapter
    # with a usable control statement and remove duplicates.
    cleaned: List[str] = []
    boq_notice_added = False
    for raw_line in value.splitlines():
        line = raw_line
        if re.match(r"^\s*(?:[-*+]\s*)?清单项：\s*(?:重点清单项|[A-Z0-9]+)\s*$", line):
            continue
        if re.search(r"清单重点（|(?:工程量|合价|单价)\s*[=：]", line):
            if not boq_notice_added:
                cleaned.append(
                    "清单重点项按正式工程量清单建立项目名称、施工部位、工序、资源和验收资料的对应关系；"
                    "清单与图纸不一致时先履行书面核验程序。"
                )
                boq_notice_added = True
            continue
        cleaned.append(line)
    return "\n".join(cleaned)


def _normalize_submission_heading(value: str) -> str:
    heading = re.sub(r"\s+", " ", str(value or "").strip())
    heading = re.sub(r"\s*[\(（]方案版本\s*[:：]\s*v?\d+[\)）]\s*$", "", heading, flags=re.IGNORECASE)
    if heading in _SUBMISSION_HEADING_REPLACEMENTS:
        return _SUBMISSION_HEADING_REPLACEMENTS[heading]
    heading = re.sub(r"^([A-E])\.\s*", "", heading)
    heading = re.sub(r"\s*\([A-Za-z][A-Za-z &/\-]+\)\s*$", "", heading)
    return heading.strip()


def _submission_heading_key(value: str) -> str:
    normalized = _normalize_submission_heading(value)
    normalized = re.sub(r"（如有）$", "", normalized)
    normalized = re.sub(r"^\s*第[\u4e00-\u9fa50-9]+章\s*", "", normalized)
    normalized = re.sub(r"[\s（）()\-_:：、。]", "", normalized)
    return normalized.lower()


def _sanitize_submission_text(
    value: Any,
    *,
    field_name: str,
    missing_parameters: Any = None,
    project_topic: Any = None,
) -> str:
    """Return bidder-facing prose and fail closed on non-text section payloads.

    The generation pipeline intentionally keeps machine locators and graph/agent
    metadata for quality analysis.  Those values belong in the JSON/XLSX review
    artifacts, never in the submission DOCX.
    """
    if value is None:
        return ""
    if not isinstance(value, str):
        raise _local_adapter_export_error(
            "docx",
            [
                {
                    "code": "SUBMISSION_TEXT_TYPE_INVALID",
                    "field": field_name,
                    "message": "submission DOCX accepts text section content only",
                }
            ],
        )

    text = _rewrite_submission_notation(_strip_internal_autofix_markers(value).replace("\x00", ""))
    text = _rewrite_unverified_submission_parameters(text, missing_parameters)
    text = _rewrite_hospital_renovation_scope(text, project_topic)
    text = _SUBMISSION_EVIDENCE_RE.sub("", text)
    text = _SUBMISSION_GRAPH_RE.sub("", text)
    text = _SUBMISSION_EXPERIENCE_RE.sub("", text)
    text = _SUBMISSION_LOCATOR_RE.sub("资料依据", text)

    clean_lines: List[str] = []
    seen_lines: set[str] = set()
    for raw_line in text.splitlines():
        line = raw_line.strip()
        # A scoped phrase replacement may consume the closing Markdown marker
        # of a bold lead while leaving its opening marker. Do not leak that
        # authoring syntax into the submitted Word document.
        if line.count("**") == 1:
            line = line.replace("**", "", 1).lstrip()
        if not line:
            clean_lines.append("")
            continue
        if line.startswith("【证据摘要】"):
            # Legacy provider-failure payloads appended the raw evidence dump at
            # the end of a chapter. The remainder is review-only material.
            break
        lowered = line.lower()
        semantic_line = re.sub(
            r"^(?:#{1,6}\s+|[-*+]\s+|\d+[.)、]\s*)",
            "",
            line,
        ).strip()
        semantic_line = re.sub(r"^\*\*(.*?)\*\*$", r"\1", semantic_line).strip()
        semantic_key = semantic_line.strip("【】").strip()
        if any(semantic_line.startswith(prefix) or semantic_key.startswith(prefix) for prefix in _SUBMISSION_INTERNAL_PREFIXES):
            continue
        if any(semantic_line.startswith(prefix) for prefix in _SUBMISSION_INSTRUCTION_PREFIXES):
            continue
        if semantic_line.startswith(_GENERIC_METRIC_TEMPLATE_PREFIX):
            continue
        if semantic_line.startswith("量化指标：") and sum(
            marker in semantic_line
            for marker in ("频次=", "阈值=", "间距=", "厚度=", "时长=", "人数=", "设备型号=")
        ) >= 4:
            continue
        if semantic_line in {"---", "***", "___", "风险→控制→验证"}:
            continue
        if re.search(r"纸张\s*A4|正文字体|页边距|排版导致的", semantic_line) and (
            "排版" in semantic_line
            or "字号" in semantic_line
            or "行距" in semantic_line
            or "正文字体" in semantic_line
            or "页边距" in semantic_line
        ):
            continue
        if any(token.lower() in lowered for token in _SUBMISSION_INTERNAL_TOKENS):
            continue
        if "zf-kg-" in lowered or "knowledge_graph" in lowered or "graph_nodes" in lowered:
            continue
        if ("负责人：" in line or "负责人=" in line) and ("逻辑模" in line or line.startswith("【范围】")):
            continue
        if "冲突值" in line and "{" in line and "}" in line:
            line = "资料之间如有差异，以招标文件、答疑文件及经批准的设计文件为准，并履行书面核验程序。"
        # Raw mapping dumps are never bidder-facing prose.
        if line.startswith(("{", "[")) and line.endswith(("}", "]")) and ":" in line:
            continue
        # Retain the relevant construction step while removing model/template
        # outline letters that have no place in the submitted document.
        # The outline letter may sit inside Markdown emphasis, for example
        # ``- **A. 施工准备**``. Remove only that internal letter marker while
        # preserving the bidder-facing title and its emphasis.
        if semantic_line not in _SUBMISSION_HEADING_REPLACEMENTS:
            line = re.sub(
                r"^((?:(?:#{1,6}|[-*+])\s+)?(?:\*\*|__)?)"
                r"[A-E](?:[.．、]\s*|\s*)(?=[\u4e00-\u9fff])",
                r"\1",
                line,
            )
        line = re.sub(r"^((?:[-*+]\s+)?)[A-E](?=[\u4e00-\u9fff]{2,12}[：:])", r"\1", line)
        line = re.sub(r"\s+([，。；：])", r"\1", line).strip()
        if line and line not in {"-", "•"}:
            duplicate_key = re.sub(r"\s+", "", line)
            if len(duplicate_key) >= 12 and duplicate_key in seen_lines:
                continue
            if len(duplicate_key) >= 12:
                seen_lines.add(duplicate_key)
            clean_lines.append(line)

    # Collapse excessive blank lines while preserving intentional paragraph breaks.
    collapsed: List[str] = []
    for line in clean_lines:
        if not line and (not collapsed or not collapsed[-1]):
            continue
        collapsed.append(line)
    result = "\n".join(collapsed).strip()
    remaining = [
        token
        for token in ("entity_master_key", "incremental_fingerprint", "【多Agent", "【图谱节点", "$.章节")
        if token.lower() in result.lower()
    ]
    if remaining:
        raise _local_adapter_export_error(
            "docx",
            [
                {
                    "code": "SUBMISSION_INTERNAL_METADATA_REMAINS",
                    "field": field_name,
                    "tokens": remaining,
                    "message": "internal metadata remains after submission sanitization",
                }
            ],
        )
    return result


def _prepare_submission_sections(
    raw_sections: Any,
    *,
    internal_review: bool,
    missing_parameters: Any = None,
    project_topic: Any = None,
) -> List[Dict[str, Any]]:
    if raw_sections is None:
        return []
    if not isinstance(raw_sections, list):
        raise _local_adapter_export_error(
            "docx",
            [{"code": "SUBMISSION_SECTIONS_TYPE_INVALID", "message": "sections must be a list"}],
        )
    prepared: List[Dict[str, Any]] = []
    for idx, raw in enumerate(raw_sections):
        if not isinstance(raw, dict):
            raise _local_adapter_export_error(
                "docx",
                [{"code": "SUBMISSION_SECTION_TYPE_INVALID", "index": idx, "message": "section must be an object"}],
            )
        section = dict(raw)
        title = section.get("title") or f"章节{idx + 1}"
        if not isinstance(title, str):
            raise _local_adapter_export_error(
                "docx",
                [{"code": "SUBMISSION_TITLE_TYPE_INVALID", "index": idx, "message": "section title must be text"}],
            )
        section["title"] = title.strip() or f"章节{idx + 1}"
        if not internal_review:
            section["content"] = _sanitize_submission_text(
                section.get("content"),
                field_name=f"sections[{idx}].content",
                missing_parameters=missing_parameters,
                project_topic=project_topic,
            )
        prepared.append(section)
    return prepared


def _to_float(v: Any, default: float) -> float:
    try:
        return float(v)
    except Exception:
        return default


def _to_int(v: Any, default: int) -> int:
    try:
        return int(float(v))
    except Exception:
        return default


def _to_bool(v: Any, default: bool = False) -> bool:
    if isinstance(v, bool):
        return v
    if isinstance(v, (int, float)):
        return bool(v)
    if isinstance(v, str):
        s = v.strip().lower()
        if s in {"1", "true", "yes", "y", "on"}:
            return True
        if s in {"0", "false", "no", "n", "off"}:
            return False
    return default


def _merge_style(base: Dict[str, Any], override: Dict[str, Any]) -> Dict[str, Any]:
    merged = dict(base or {})
    for k, v in (override or {}).items():
        if isinstance(v, dict) and isinstance(merged.get(k), dict):
            item = dict(merged.get(k) or {})
            item.update(v)
            merged[k] = item
        else:
            merged[k] = v
    return merged


def _resolve_alignment(v: Any):
    if not isinstance(v, str):
        return None
    key = v.strip().lower()
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(key)


def _normalize_style(style: Dict[str, Any]) -> Dict[str, Any]:
    style = style if isinstance(style, dict) else {}
    font_cfg = style.get("font") if isinstance(style.get("font"), dict) else {}
    headings_cfg = style.get("headings") if isinstance(style.get("headings"), dict) else {}
    margins_cfg = style.get("margins_cm") if isinstance(style.get("margins_cm"), dict) else {}

    body_font = (
        style.get("body_font")
        or (style.get("font") if isinstance(style.get("font"), str) else None)
        or font_cfg.get("eastAsia")
        or "宋体"
    )
    body_latin_font = style.get("body_latin_font") or font_cfg.get("latin") or body_font
    body_size = _to_float(
        style.get("body_size") or style.get("font_size") or font_cfg.get("size_pt") or 14,
        14.0,
    )
    # A tender-specified multiple spacing must not inherit the 22 pt fallback.
    # Use 22 pt only when neither spacing mode is explicitly configured.
    line_spacing, line_spacing_pt = resolve_line_spacing(style, default_pt=22.0)

    title_font = style.get("title_font") or headings_cfg.get("eastAsia") or body_font
    title_latin_font = style.get("title_latin_font") or headings_cfg.get("latin") or body_latin_font
    title_size = _to_float(style.get("title_size") or headings_cfg.get("h2_size") or max(body_size + 2, 14), 14.0)
    doc_title_size = _to_float(
        style.get("doc_title_size") or headings_cfg.get("h1_size") or title_size or 16,
        16.0,
    )

    margins_list = style.get("margins") if isinstance(style.get("margins"), list) else []
    top = _to_float(margins_cfg.get("top"), _to_float(margins_list[0] if len(margins_list) > 0 else 2.5, 2.5))
    right = _to_float(margins_cfg.get("right"), _to_float(margins_list[1] if len(margins_list) > 1 else 2.0, 2.0))
    bottom = _to_float(margins_cfg.get("bottom"), _to_float(margins_list[2] if len(margins_list) > 2 else 2.0, 2.0))
    left = _to_float(margins_cfg.get("left"), _to_float(margins_list[3] if len(margins_list) > 3 else 2.0, 2.0))

    def _hex_color(value: Any, default: str) -> str:
        cleaned = re.sub(r"[^0-9A-Fa-f]", "", str(value or ""))
        return cleaned.upper() if len(cleaned) == 6 else default

    palette_raw = style.get("palette") if isinstance(style.get("palette"), dict) else {}
    palette = {
        # Deep blue/teal and cool gray are deliberately restrained: they match
        # the visual language shared by the four reviewed benchmark tenders
        # without copying any one bidder's brand identity.
        "accent": _hex_color(palette_raw.get("accent") or style.get("accent_color"), "0F5966"),
        "accent_dark": _hex_color(palette_raw.get("accent_dark"), "103B52"),
        "accent_light": _hex_color(palette_raw.get("accent_light"), "EAF2F5"),
        "signal": _hex_color(palette_raw.get("signal"), "D9792B"),
        "signal_light": _hex_color(palette_raw.get("signal_light"), "FFF2E8"),
        "border": _hex_color(palette_raw.get("border"), "AFC4CE"),
        "table_header": _hex_color(palette_raw.get("table_header"), "103B52"),
        "table_band": _hex_color(palette_raw.get("table_band"), "F5F8FA"),
        "muted": _hex_color(palette_raw.get("muted"), "53656E"),
    }

    return {
        "paper": str(style.get("paper") or style.get("paper_size") or "A4"),
        "body_font": body_font,
        "body_latin_font": body_latin_font,
        "body_size": max(8.0, min(22.0, body_size)),
        "title_font": title_font,
        "title_latin_font": title_latin_font,
        "title_size": max(10.0, min(28.0, title_size)),
        "doc_title_size": max(12.0, min(36.0, doc_title_size)),
        "line_spacing": max(1.0, min(2.5, line_spacing)),
        "line_spacing_pt": line_spacing_pt,
        "first_line_indent_cm": max(0.0, _to_float(style.get("first_line_indent_cm"), 0.74)),
        "body_align": _resolve_alignment(style.get("body_align")) or WD_ALIGN_PARAGRAPH.JUSTIFY,
        "title_align": _resolve_alignment(style.get("title_align")),
        "margins_cm": {
            "top": max(0.5, top),
            "right": max(0.5, right),
            "bottom": max(0.5, bottom),
            "left": max(0.5, left),
        },
        "header_distance_cm": max(0.3, _to_float(style.get("header_distance_cm"), 0.65)),
        "footer_distance_cm": max(0.3, _to_float(style.get("footer_distance_cm"), 0.65)),
        "palette": palette,
        "chapter_start_new_page": _to_bool(style.get("chapter_start_new_page"), True),
        "enforce_chapter_pages": _to_bool(style.get("enforce_chapter_pages"), False),
    }


def _apply_page_setup(doc: Document, style_cfg: Dict[str, Any]):
    paper = str(style_cfg.get("paper") or "A4").upper()
    margins = style_cfg.get("margins_cm") or {}
    for section in doc.sections:
        if paper == "A4":
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        section.top_margin = Cm(_to_float(margins.get("top"), 2.5))
        section.right_margin = Cm(_to_float(margins.get("right"), 2.0))
        section.bottom_margin = Cm(_to_float(margins.get("bottom"), 2.0))
        section.left_margin = Cm(_to_float(margins.get("left"), 2.0))
        section.header_distance = Cm(_to_float(style_cfg.get("header_distance_cm"), 0.65))
        section.footer_distance = Cm(_to_float(style_cfg.get("footer_distance_cm"), 0.65))


def _set_run_font(run, east_font: str, latin_font: str, size_pt: float):
    east_font = _resolve_docx_font_name(east_font)
    latin_font = _resolve_docx_font_name(latin_font)
    run.font.name = latin_font or east_font
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    if east_font:
        rfonts.set(qn("w:eastAsia"), east_font)
        rfonts.set(qn("w:hint"), "eastAsia")
    if latin_font:
        rfonts.set(qn("w:ascii"), latin_font)
        rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:cs"), east_font or latin_font)


def _set_style_font(style, *, east_font: str, latin_font: str, size_pt: float | None = None) -> None:
    east_font = _resolve_docx_font_name(east_font)
    latin_font = _resolve_docx_font_name(latin_font)
    style.font.name = latin_font or east_font
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    try:
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        if east_font:
            rfonts.set(qn("w:eastAsia"), east_font)
            rfonts.set(qn("w:hint"), "eastAsia")
        if latin_font:
            rfonts.set(qn("w:ascii"), latin_font)
            rfonts.set(qn("w:hAnsi"), latin_font)
        rfonts.set(qn("w:cs"), east_font or latin_font)
    except Exception:
        pass


def _configure_professional_named_styles(doc: Document, style_cfg: Dict[str, Any]) -> None:
    """Normalize built-in Word styles used by captions, lists and the TOC.

    Keeping these as named styles (rather than one-off paragraph formatting)
    makes the export easier to revise in Word while preserving a consistent
    technical-bid visual system.
    """
    cfg = _normalize_style(style_cfg or {})
    palette = cfg["palette"]
    try:
        caption = doc.styles["Caption"]
        _set_style_font(
            caption,
            east_font=cfg["body_font"],
            latin_font=cfg["body_latin_font"],
            size_pt=cfg["body_size"],
        )
        caption.font.italic = False
        caption.font.color.rgb = RGBColor.from_string(palette["muted"])
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.first_line_indent = Cm(0)
        caption.paragraph_format.line_spacing = 1.0
        caption.paragraph_format.space_before = Pt(2)
        caption.paragraph_format.space_after = Pt(8)
        caption.paragraph_format.keep_together = True
        caption.paragraph_format.keep_with_next = False
        caption.paragraph_format.widow_control = True
    except Exception:
        pass
    try:
        bullet = doc.styles["List Bullet"]
        _set_style_font(
            bullet,
            east_font=cfg["body_font"],
            latin_font=cfg["body_latin_font"],
            size_pt=cfg["body_size"],
        )
        bullet.paragraph_format.left_indent = Cm(0.74)
        bullet.paragraph_format.first_line_indent = Cm(-0.37)
        bullet.paragraph_format.space_after = Pt(3)
        bullet.paragraph_format.widow_control = True
    except Exception:
        pass
    for style_name, left_cm, size_pt in (
        ("TOC 1", 0.0, cfg["body_size"]),
        ("TOC 2", 0.55, cfg["body_size"]),
        ("TOC 3", 1.1, cfg["body_size"]),
    ):
        try:
            toc_style = doc.styles[style_name]
            _set_style_font(
                toc_style,
                east_font=cfg["body_font"],
                latin_font=cfg["body_latin_font"],
                size_pt=size_pt,
            )
            toc_style.paragraph_format.left_indent = Cm(left_cm)
            toc_style.paragraph_format.first_line_indent = Cm(0)
            toc_style.paragraph_format.space_after = Pt(2)
            toc_style.paragraph_format.widow_control = True
        except Exception:
            pass


def _set_inline_shape_alt_text(shape: Any, *, title: str, description: str) -> None:
    try:
        doc_pr = shape._inline.docPr
        doc_pr.set("title", str(title or "插图"))
        doc_pr.set("descr", str(description or title or "插图"))
    except Exception:
        pass


def _configure_inline_image_paragraph(
    paragraph: Any,
    *,
    space_before_pt: float = 6.0,
    space_after_pt: float = 4.0,
    keep_with_next: bool = True,
) -> None:
    """Keep inline images independent from the tender body's fixed line grid.

    Word and LibreOffice can clip a drawing to a thin horizontal strip when its
    host paragraph inherits an exact 22 pt line height.  Every formal image,
    including cover/header/footer assets, therefore gets a single-line host
    paragraph with explicit spacing and zero indentation.
    """
    try:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.right_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(float(space_before_pt or 0))
        paragraph.paragraph_format.space_after = Pt(float(space_after_pt or 0))
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = bool(keep_with_next)
    except Exception:
        pass


def _mark_table_header_row(table: Any) -> None:
    try:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)
    except Exception:
        pass


def _topic_to_cover_project_name(topic: Any) -> str:
    raw = str(topic or "").strip()
    if not raw:
        return ""
    for suffix in ("施工组织设计方案", "施工组织设计", "施组方案"):
        if raw.endswith(suffix):
            raw = raw[: -len(suffix)].strip()
            break
    # A parser can occasionally retain only the tail of a long tender title
    # (for example, a closing parenthesis with no opening parenthesis).  A
    # professional cover must never print that fragment as if it were the full
    # project name.  Explicit project_name metadata still takes precedence in
    # _resolve_cover_meta.
    if raw.count("）") > raw.count("（") or raw.count(")") > raw.count("("):
        return ""
    return raw


def _to_cn_month(month: int) -> str:
    month = max(1, min(12, int(month or 1)))
    if month < 10:
        return _CN_DIGITS[str(month)]
    if month == 10:
        return "十"
    return "十" + _CN_DIGITS[str(month % 10)]


def _format_cover_year_month(dt: _dt.datetime | None = None) -> str:
    current = dt or _dt.datetime.now()
    year_cn = "".join(_CN_DIGITS.get(ch, ch) for ch in f"{int(current.year):04d}")
    return f"{year_cn}年{_to_cn_month(int(current.month))}月"


def _normalize_front_matter_page_mode(value: Any) -> str:
    return "exclude" if str(value or "").strip().lower() == "exclude" else "include"


def _normalize_full_index_enabled(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return bool(value)
    return str(value or "").strip().lower() in {"1", "true", "yes", "y", "on", "enable", "enabled"}


def _resolve_front_matter_plan(
    *,
    style_raw: Dict[str, Any],
    data: Dict[str, Any],
    body_pages_estimate: int,
) -> Dict[str, Any]:
    style = style_raw if isinstance(style_raw, dict) else {}
    payload = data if isinstance(data, dict) else {}
    cover_pages = max(1, _to_int(style.get("cover_page_count"), 1))
    toc_pages = max(1, _to_int(style.get("toc_page_count"), 1))
    configured_index_pages = max(1, _to_int(style.get("full_index_page_count"), 1))
    full_index_enabled = _normalize_full_index_enabled(style.get("full_index_enabled"))
    count_mode = _normalize_front_matter_page_mode(style.get("front_matter_page_mode"))
    document_total_target = max(
        1,
        _to_int(style.get("document_total_pages_target"), 0)
        or _to_int(payload.get("total_pages_target"), 0)
        or int(body_pages_estimate or 1),
    )
    full_index_pages = configured_index_pages if full_index_enabled else 0
    actual_front_matter_pages = cover_pages + toc_pages + full_index_pages
    effective_document_pages = (
        document_total_target
        if count_mode == "include"
        else document_total_target + actual_front_matter_pages
    )
    return {
        "cover_pages": cover_pages,
        "toc_pages": toc_pages,
        "configured_index_pages": configured_index_pages,
        "full_index_enabled": full_index_enabled,
        "count_mode": count_mode,
        "full_index_pages": full_index_pages,
        "actual_front_matter_pages": actual_front_matter_pages,
        "document_total_target": document_total_target,
        "effective_document_pages": effective_document_pages,
    }


def _insert_full_index_page(
    doc: Document,
    apply_paragraph,
    *,
    topic: str,
    sections: List[Dict[str, Any]],
    chapter_pages: Dict[str, Any] | None,
    effective_document_pages: int,
    index_entries: List[Dict[str, Any]] | None = None,
) -> None:
    heading = doc.add_heading("全文索引", level=1)
    apply_paragraph(heading, is_title=True)
    intro = doc.add_paragraph(
        f"{str(topic or '施工组织设计').strip()}；"
        f"章节数={len(sections or [])}；"
        f"成品预计总页数={max(1, int(effective_document_pages or 1))}页。"
    )
    apply_paragraph(intro)

    normalized_entries = [item for item in (index_entries or []) if isinstance(item, dict)]
    if normalized_entries:
        for idx, item in enumerate(normalized_entries, start=1):
            title = str(item.get("title") or f"章节{idx}").strip() or f"章节{idx}"
            summary = str(item.get("summary") or "").strip()
            planned_pages = _to_int(item.get("planned_pages"), 0)
            line = summary or f"{idx:02d}. {title}"
            if planned_pages and "约" not in line:
                line += f"（约{int(planned_pages)}页）"
            paragraph = doc.add_paragraph(line)
            apply_paragraph(paragraph)
        doc.add_page_break()
        return

    if not sections:
        note = doc.add_paragraph("当前无可索引章节。")
        apply_paragraph(note)
        doc.add_page_break()
        return

    for idx, section in enumerate(sections or [], start=1):
        title = str((section or {}).get("title") or f"章节{idx}").strip() or f"章节{idx}"
        planned_pages = _extract_chapter_page_target(chapter_pages or {}, title)
        line = f"{idx:02d}. {title}"
        if planned_pages:
            line += f"（约{int(planned_pages)}页）"
        paragraph = doc.add_paragraph(line)
        apply_paragraph(paragraph)
    doc.add_page_break()


def _build_static_toc_entries(
    *,
    sections: List[Dict[str, Any]],
    section_pages: List[int],
    front_matter_plan: Dict[str, Any],
) -> List[Dict[str, Any]]:
    plan = front_matter_plan if isinstance(front_matter_plan, dict) else {}
    cover_pages = max(1, _to_int(plan.get("cover_pages"), 1))
    toc_pages = max(1, _to_int(plan.get("toc_pages"), 1))
    full_index_pages = max(0, _to_int(plan.get("full_index_pages"), 0))
    current_page = cover_pages + full_index_pages + toc_pages + 1
    entries: List[Dict[str, Any]] = []
    for idx, sec in enumerate(sections or []):
        title = str((sec or {}).get("title") or f"章节{idx + 1}").strip() or f"章节{idx + 1}"
        planned_pages = max(1, _to_int(section_pages[idx] if idx < len(section_pages) else 1, 1))
        entries.append(
            {
                "order": idx + 1,
                "title": title,
                "start_page": current_page,
                "planned_pages": planned_pages,
                "page_number_exact": False,
            }
        )
        current_page += planned_pages
    return entries


def _paginate_toc_entries(entries: List[Dict[str, Any]], toc_pages: int) -> List[List[Dict[str, Any]]]:
    page_count = max(1, int(toc_pages or 1))
    chunks: List[List[Dict[str, Any]]] = []
    cursor = 0
    total = len(entries or [])
    for page_idx in range(page_count):
        remaining_pages = page_count - page_idx
        remaining_entries = total - cursor
        if remaining_entries <= 0:
            chunks.append([])
            continue
        take = max(1, math.ceil(remaining_entries / max(1, remaining_pages)))
        chunks.append((entries or [])[cursor: cursor + take])
        cursor += take
    return chunks


def _infer_toc_level(entry: Dict[str, Any]) -> int:
    raw = _to_int(entry.get("level"), 0) if isinstance(entry, dict) else 0
    if raw > 0:
        return min(3, raw)
    title = re.sub(r"\s+", " ", str((entry or {}).get("title") or "").strip())
    if _TOC_CHAPTER_RE.match(title):
        return 1
    if _TOC_SECTION_RE.match(title):
        return 2
    if _TOC_CN_ITEM_RE.match(title) or _TOC_NUMERIC_ITEM_RE.match(title):
        return 3
    return 1


def _format_toc_display_title(title: str) -> str:
    cleaned = re.sub(r"\s+", " ", str(title or "").strip())
    for pattern in (_TOC_CHAPTER_RE, _TOC_SECTION_RE):
        match = pattern.match(cleaned)
        if match:
            prefix = match.group(0)
            suffix = cleaned[len(prefix):].strip(" 、.")
            return f"{prefix}、{suffix}" if suffix else prefix
    return cleaned


def _append_field_run(paragraph, instruction: str, result_text: str | None = None):
    run = paragraph.add_run()
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = str(instruction or "")
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_separate)
    if result_text is not None:
        result = OxmlElement("w:t")
        result.text = str(result_text)
        r.append(result)
    r.append(fld_end)
    return run


def _add_paragraph_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    """Bind a stable Word bookmark to a chapter heading for live TOC page refs."""
    safe_name = re.sub(r"[^A-Za-z0-9_]", "_", str(name or "")) or "ZF_CHAPTER"
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(int(bookmark_id)))
    start.set(qn("w:name"), safe_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(int(bookmark_id)))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _hide_paragraph(paragraph) -> None:
    if not getattr(paragraph, "runs", None):
        paragraph.add_run()
    for run in paragraph.runs:
        try:
            rpr = run._element.get_or_add_rPr()
            if rpr.find(qn("w:vanish")) is None:
                rpr.append(OxmlElement("w:vanish"))
        except Exception:
            continue


def _toc_entry_style(style_cfg: Dict[str, Any], level: int) -> Dict[str, Any]:
    style = style_cfg if isinstance(style_cfg, dict) else {}
    body_font = str(style.get("body_font") or "宋体")
    body_latin = str(style.get("body_latin_font") or body_font)
    title_font = str(style.get("title_font") or body_font)
    title_latin = str(style.get("title_latin_font") or body_latin)
    body_size = _to_float(style.get("body_size"), 14.0)
    title_size = _to_float(style.get("title_size"), body_size + 2.0)
    palette = _normalize_style(style).get("palette") or {}
    accent_rgb = tuple(int(str(palette.get("accent") or "0F5966")[i : i + 2], 16) for i in (0, 2, 4))
    if level <= 1:
        return {
            "font_east": title_font,
            "font_latin": title_latin,
            "size_pt": title_size,
            "bold": True,
            "left_indent_cm": 0.0,
            "color_rgb": tuple(int(str(palette.get("accent_dark") or "103B52")[i : i + 2], 16) for i in (0, 2, 4)),
        }
    if level == 2:
        return {
            "font_east": body_font,
            "font_latin": body_latin,
            "size_pt": body_size,
            "bold": False,
            "left_indent_cm": 1.0,
            "color_rgb": accent_rgb,
        }
    return {
        "font_east": body_font,
        "font_latin": body_latin,
        "size_pt": body_size,
        "bold": False,
        "left_indent_cm": 2.0,
        "color_rgb": (0, 0, 0),
    }


def _render_toc_line(
    doc: Document,
    entry: Dict[str, Any],
    *,
    style_cfg: Dict[str, Any],
):
    level = _infer_toc_level(entry)
    line_cfg = _toc_entry_style(style_cfg, level)
    title = _format_toc_display_title(str((entry or {}).get("title") or "章节"))
    page_number = int(_to_int((entry or {}).get("start_page"), 1) or 1)
    paragraph = doc.add_paragraph()
    try:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(float(line_cfg["left_indent_cm"]))
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(22)
        p_pr = paragraph._p.get_or_add_pPr()
        tabs = p_pr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            p_pr.append(tabs)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "8850")
        tabs.append(tab)
    except Exception:
        pass
    suffix = str(page_number)
    title_run = paragraph.add_run(title)
    tab_run = paragraph.add_run("\t")
    order = max(1, _to_int((entry or {}).get("order"), 1))
    bookmark_name = f"ZF_CHAPTER_{order}"
    suffix_run = _append_field_run(
        paragraph,
        f"PAGEREF {bookmark_name} \\h",
        result_text=suffix,
    )
    for run in (title_run, tab_run, suffix_run):
        _set_run_font(run, str(line_cfg["font_east"]), str(line_cfg["font_latin"]), float(line_cfg["size_pt"]))
        try:
            run.bold = bool(line_cfg["bold"])
            color = tuple(line_cfg.get("color_rgb") or (0, 0, 0))
            run.font.color.rgb = RGBColor(int(color[0]), int(color[1]), int(color[2]))
        except Exception:
            pass
    return paragraph


def _insert_auto_toc(
    doc: Document,
    apply_paragraph,
    *,
    style_cfg: Dict[str, Any],
    toc_pages: int = 1,
    toc_entries: List[Dict[str, Any]] | None = None,
) -> None:
    style = style_cfg if isinstance(style_cfg, dict) else {}
    title_font = str(style.get("title_font") or "宋体")
    title_latin = str(style.get("title_latin_font") or style.get("body_latin_font") or title_font)
    title_size = _to_float(style.get("doc_title_size"), 16.0)
    heading = doc.add_paragraph()
    try:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.first_line_indent = Cm(0)
        heading.paragraph_format.space_before = Pt(6)
        heading.paragraph_format.space_after = Pt(10)
        heading.paragraph_format.line_spacing = Pt(22)
    except Exception:
        pass
    title_run = heading.add_run("目录")
    _set_run_font(title_run, title_font, title_latin, title_size)
    try:
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(15, 89, 102)
    except Exception:
        pass

    # A visible live TOC is the sole directory representation.  The former
    # hidden TOC plus hand-built duplicate directory left sensitive hidden text
    # and produced two directories after Word refreshed fields.
    field_paragraph = doc.add_paragraph()
    _append_field_run(field_paragraph, 'TOC \\o "1-3" \\h \\z \\u')
    apply_paragraph(field_paragraph)
    try:
        field_paragraph.paragraph_format.first_line_indent = Cm(0)
        field_paragraph.paragraph_format.space_before = Pt(0)
        field_paragraph.paragraph_format.space_after = Pt(0)
    except Exception:
        pass
    doc.add_page_break()


def _usable_page_width_cm(doc: Document) -> float:
    try:
        section = doc.sections[-1]
        width_emu = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
        return max(8.0, float(width_emu) / 360000.0)
    except Exception:
        return 17.0


def _clear_block_container(container) -> None:
    try:
        for paragraph in list(container.paragraphs):
            element = paragraph._element
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
    except Exception:
        pass
    try:
        for table in list(container.tables):
            element = table._element
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
    except Exception:
        pass


def _set_cell_width(cell, width_cm: float) -> None:
    try:
        cell.width = Cm(float(width_cm))
    except Exception:
        pass
    try:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.first_child_found_in("w:tcW")
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(int(float(width_cm) * 567)))
        tc_w.set(qn("w:type"), "dxa")
    except Exception:
        pass


def _set_cell_shading(cell, fill: str) -> None:
    try:
        tc_pr = cell._tc.get_or_add_tcPr()
        for child in list(tc_pr):
            if child.tag == qn("w:shd"):
                tc_pr.remove(child)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), str(fill or "").strip() or "FFFFFF")
        tc_pr.append(shd)
    except Exception:
        pass


def _set_paragraph_shading(paragraph: Any, fill: str) -> None:
    try:
        p_pr = paragraph._p.get_or_add_pPr()
        for child in list(p_pr):
            if child.tag == qn("w:shd"):
                p_pr.remove(child)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), str(fill or "").strip() or "FFFFFF")
        p_pr.append(shd)
    except Exception:
        pass


def _set_paragraph_border(paragraph: Any, *, edge: str, color: str, sz: int = 8, space: int = 3) -> None:
    """Add one deterministic paragraph rule without relying on a Word theme."""
    try:
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        current = p_bdr.find(qn(f"w:{edge}"))
        if current is None:
            current = OxmlElement(f"w:{edge}")
            p_bdr.append(current)
        current.set(qn("w:val"), "single")
        current.set(qn("w:sz"), str(int(sz)))
        current.set(qn("w:space"), str(int(space)))
        current.set(qn("w:color"), str(color or "AFC4CE"))
    except Exception:
        pass


def _set_cell_margins(cell: Any, *, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    try:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
            node = tc_mar.find(qn(f"w:{name}"))
            if node is None:
                node = OxmlElement(f"w:{name}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(int(value)))
            node.set(qn("w:type"), "dxa")
    except Exception:
        pass


def _prevent_table_row_split(row: Any) -> None:
    try:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
    except Exception:
        pass


def _style_professional_table(table: Any, style_cfg: Dict[str, Any] | None = None) -> None:
    """Apply the restrained blue-gray tender table treatment used by exports."""
    cfg = _normalize_style(style_cfg or {})
    palette = cfg["palette"]
    try:
        table.style = "Table Grid"
        table.autofit = False
    except Exception:
        pass
    _mark_table_header_row(table)
    for row_index, row in enumerate(table.rows):
        _prevent_table_row_split(row)
        for cell in row.cells:
            _set_cell_margins(cell)
            if row_index == 0:
                _set_cell_shading(cell, palette["table_header"])
            elif row_index % 2 == 0:
                _set_cell_shading(cell, palette["table_band"])
            _set_cell_border(
                cell,
                top={"color": palette["border"], "sz": 5},
                bottom={"color": palette["border"], "sz": 5},
                start={"color": palette["border"], "sz": 5},
                end={"color": palette["border"], "sz": 5},
                insideH={"color": palette["border"], "sz": 4},
                insideV={"color": palette["border"], "sz": 4},
            )
            for paragraph in cell.paragraphs:
                try:
                    paragraph.paragraph_format.keep_together = True
                    paragraph.paragraph_format.widow_control = True
                except Exception:
                    pass


def _table_text_weight(value: Any) -> float:
    """Approximate printed width, counting CJK glyphs as full-width."""
    text = re.sub(r"\s+", " ", str(value or "").strip())
    if not text:
        return 0.0
    weight = 0.0
    for char in text:
        weight += 1.0 if "\u2e80" <= char <= "\uffff" else 0.56
    return weight


def _suggest_table_column_widths(
    headers: List[str],
    rows: List[List[str]],
    *,
    total_width_cm: float = 16.2,
) -> tuple[float, ...]:
    """Allocate A4 table width by semantic content instead of equal columns."""
    if not headers:
        return (float(total_width_cm),)
    narrative_tokens = ("措施", "内容", "说明", "分析", "要求", "控制", "验证", "处置", "记录", "描述", "建议", "章节")
    compact_tokens = ("序号", "编号", "单位", "数量", "页码", "状态", "等级", "日期", "时间", "单价", "合价")
    raw_weights: List[float] = []
    for col, header in enumerate(headers):
        samples = [str(header or "")]
        for row in rows[:24]:
            if col < len(row):
                samples.append(str(row[col] or ""))
        measured = sorted(min(54.0, _table_text_weight(value)) for value in samples if str(value).strip())
        representative = measured[min(len(measured) - 1, max(0, int(len(measured) * 0.78)))] if measured else 2.0
        weight = max(2.0, _table_text_weight(header) * 1.15, representative)
        header_text = str(header or "")
        if any(token in header_text for token in narrative_tokens):
            weight *= 1.35
        if any(token in header_text for token in compact_tokens):
            weight = min(weight, 5.2)
        raw_weights.append(max(2.0, min(weight, 34.0)))

    min_width = 1.55 if len(headers) >= 4 else 2.0
    widths = [max(min_width, total_width_cm * weight / max(1.0, sum(raw_weights))) for weight in raw_weights]
    # Re-normalize after minimum-width clamping.
    scale = float(total_width_cm) / max(0.1, sum(widths))
    widths = [round(width * scale, 2) for width in widths]
    widths[-1] = round(widths[-1] + (float(total_width_cm) - sum(widths)), 2)
    return tuple(widths)


def _table_column_prefers_center(header: str, values: List[str]) -> bool:
    header_text = str(header or "").strip()
    if any(token in header_text for token in ("序号", "编号", "单位", "数量", "页码", "状态", "等级", "日期", "时间", "单价", "合价", "比例")):
        return True
    nonempty = [str(value or "").strip() for value in values if str(value or "").strip()]
    if not nonempty:
        return False
    numeric_like = re.compile(r"^[\d\s.,%‰+\-—/年月日时分秒万元m²㎡m³³]+$", re.IGNORECASE)
    return all(_table_text_weight(value) <= 12 and numeric_like.fullmatch(value) for value in nonempty)


def _apply_semantic_table_layout(
    table: Any,
    headers: List[str],
    rows: List[List[str]],
    style_cfg: Dict[str, Any] | None = None,
) -> None:
    """Apply print-oriented alignment, type size and first-column emphasis."""
    cfg = _normalize_style(style_cfg or {})
    palette = cfg["palette"]
    if not headers:
        return
    columns = len(headers)
    centered = [
        _table_column_prefers_center(
            headers[col],
            [row[col] if col < len(row) else "" for row in rows],
        )
        for col in range(columns)
    ]
    first_values = [row[0] for row in rows if row]
    emphasize_first = bool(first_values) and max((_table_text_weight(value) for value in first_values), default=0) <= 18
    for row_index, row in enumerate(table.rows):
        for col, cell in enumerate(row.cells):
            try:
                cell.vertical_alignment = (
                    WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    if row_index == 0 or (col < len(centered) and centered[col])
                    else WD_CELL_VERTICAL_ALIGNMENT.TOP
                )
            except Exception:
                pass
            if row_index > 0 and col == 0 and emphasize_first:
                _set_cell_shading(cell, palette["accent_light"])
            for paragraph in cell.paragraphs:
                try:
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER
                        if row_index == 0 or (col < len(centered) and centered[col]) or (col == 0 and emphasize_first)
                        else WD_ALIGN_PARAGRAPH.LEFT
                    )
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.keep_together = True
                    paragraph.paragraph_format.widow_control = True
                except Exception:
                    pass
                for run in paragraph.runs:
                    _set_run_font(run, cfg["body_font"], cfg["body_latin_font"], cfg["body_size"])
                    if row_index == 0 or (row_index > 0 and col == 0 and emphasize_first):
                        run.bold = True
                    if row_index == 0:
                        try:
                            run.font.color.rgb = RGBColor.from_string("FFFFFF")
                        except Exception:
                            pass


def _set_inline_shape_border(shape: Any, *, color: str = "AFC4CE", width: int = 12700) -> None:
    """Give an inserted picture a thin neutral keyline for clean print output."""
    try:
        graphic_data = shape._inline.graphic.graphicData
        pic = next((child for child in graphic_data.iter() if child.tag == qn("pic:pic")), None)
        if pic is None:
            return
        sp_pr = pic.find(qn("pic:spPr"))
        if sp_pr is None:
            return
        old = sp_pr.find(qn("a:ln"))
        if old is not None:
            sp_pr.remove(old)
        line = OxmlElement("a:ln")
        line.set("w", str(int(width)))
        solid = OxmlElement("a:solidFill")
        rgb = OxmlElement("a:srgbClr")
        rgb.set("val", str(color or "AFC4CE"))
        solid.append(rgb)
        line.append(solid)
        sp_pr.append(line)
    except Exception:
        pass


def _image_width_for_docx(path: Any, *, source_kind: str = "") -> float:
    """Choose a print-safe width from the actual aspect ratio."""
    try:
        from PIL import Image

        with Image.open(str(path)) as image:
            width, height = image.size
        ratio = float(width) / max(1.0, float(height))
        if ratio < 0.82:
            return 9.6
        if ratio < 1.15:
            return 11.2
        if ratio > 2.15:
            return 15.6
        if ratio > 1.55:
            return 15.0
        return 13.8
    except Exception:
        return 14.5 if str(source_kind or "").lower() in {"site_photo", "drawing", "deterministic_project_diagram"} else 13.8


def _image_aspect_ratio(path: Any) -> float | None:
    try:
        from PIL import Image

        with Image.open(str(path)) as image:
            width, height = image.size
        return float(width) / max(1.0, float(height))
    except Exception:
        return None


def _media_pair_eligible(items: List[Any]) -> bool:
    """Return true for two compatible landscape project photos/drawings."""
    if len(items) != 2:
        return False
    kinds: List[str] = []
    for item in items:
        if not isinstance(item, dict):
            return False
        path = str(item.get("path") or "").strip()
        kind = str(item.get("source_kind") or "").strip().lower()
        if kind not in {"site_photo", "drawing"} or not path or not Path(path).is_file():
            return False
        ratio = _image_aspect_ratio(path)
        if ratio is None or ratio < 1.12 or ratio > 2.25:
            return False
        kinds.append(kind)
    # A photo pair or drawing pair reads as one coherent evidence panel.
    return kinds[0] == kinds[1]


def _enable_field_updates(doc: Document) -> None:
    """Ask Word-compatible editors to refresh TOC/PAGE fields on open."""
    try:
        settings = doc.settings._element
        current = settings.find(qn("w:updateFields"))
        if current is None:
            current = OxmlElement("w:updateFields")
            settings.append(current)
        current.set(qn("w:val"), "true")
    except Exception:
        pass


def _set_cell_border(cell, **kwargs) -> None:
    try:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        for edge, cfg in kwargs.items():
            edge_el = tc_borders.find(qn(f"w:{edge}"))
            if edge_el is None:
                edge_el = OxmlElement(f"w:{edge}")
                tc_borders.append(edge_el)
            edge_el.set(qn("w:val"), str(cfg.get("val", "single")))
            edge_el.set(qn("w:sz"), str(cfg.get("sz", 6)))
            edge_el.set(qn("w:space"), str(cfg.get("space", 0)))
            edge_el.set(qn("w:color"), str(cfg.get("color", "D9EAF0")))
    except Exception:
        pass


def _set_table_all_borders(table, *, color: str, sz: int = 6, top: bool = False, bottom: bool = False) -> None:
    for row in table.rows:
        for cell in row.cells:
            edges: Dict[str, Dict[str, Any]] = {}
            if top:
                edges["top"] = {"color": color, "sz": sz}
            if bottom:
                edges["bottom"] = {"color": color, "sz": sz}
            if edges:
                _set_cell_border(cell, **edges)


def _apply_footer_page_numbers(
    doc: Document,
    style_cfg: Dict[str, Any],
    *,
    bidder_company: str,
    logo_path: str | None,
    document_label: str = "",
) -> None:
    font_east = str((style_cfg or {}).get("body_font") or "宋体")
    font_latin = str((style_cfg or {}).get("body_latin_font") or font_east)
    palette = (_normalize_style(style_cfg or {}).get("palette") or {})
    company = str(bidder_company or "").strip()
    logo = str(logo_path or "").strip()
    if logo and not Path(logo).exists():
        logo = ""
    usable_width = _usable_page_width_cm(doc)
    for section in doc.sections:
        try:
            section.different_first_page_header_footer = True
        except Exception:
            pass
        try:
            footer = section.footer
            footer.is_linked_to_previous = False
        except Exception:
            continue
        _clear_block_container(footer)
        table = footer.add_table(rows=1, cols=2, width=Cm(usable_width))
        _mark_table_header_row(table)
        try:
            table.autofit = False
        except Exception:
            pass
        _set_table_all_borders(table, color=str(palette.get("border") or "AFC4CE"), sz=6, top=True)
        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        _set_cell_width(left_cell, usable_width * 0.72)
        _set_cell_width(right_cell, usable_width * 0.28)
        try:
            left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        except Exception:
            pass
        p_left = left_cell.paragraphs[0] if left_cell.paragraphs else left_cell.add_paragraph()
        p_right = right_cell.paragraphs[0] if right_cell.paragraphs else right_cell.add_paragraph()
        try:
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_left.paragraph_format.space_before = Pt(4)
            p_left.paragraph_format.space_after = Pt(0)
            p_right.paragraph_format.space_before = Pt(4)
            p_right.paragraph_format.space_after = Pt(0)
        except Exception:
            pass
        if logo:
            try:
                shape = p_left.add_run().add_picture(logo, width=Cm(1.0))
                _set_inline_shape_alt_text(shape, title="投标人标识", description="投标人企业标识")
                _configure_inline_image_paragraph(
                    p_left,
                    space_before_pt=4,
                    space_after_pt=0,
                    keep_with_next=False,
                )
                p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            except Exception:
                pass
        if company:
            run_company = p_left.add_run(f" {company}" if logo else company)
            _set_run_font(run_company, font_east, font_latin, 12.0)
            try:
                run_company.bold = True
            except Exception:
                pass
        elif str(document_label or "").strip():
            run_label = p_left.add_run(str(document_label).strip())
            _set_run_font(run_label, font_east, font_latin, 14.0)
            try:
                run_label.font.color.rgb = RGBColor(15, 89, 102)
            except Exception:
                pass
        prefix = p_right.add_run("第 ")
        _set_run_font(prefix, font_east, font_latin, 14.0)
        _append_field_run(p_right, "PAGE")
        middle = p_right.add_run(" 页 / 共 ")
        _set_run_font(middle, font_east, font_latin, 14.0)
        _append_field_run(p_right, "NUMPAGES")
        suffix = p_right.add_run(" 页")
        _set_run_font(suffix, font_east, font_latin, 14.0)
        for run in p_right.runs:
            _set_run_font(run, font_east, font_latin, 14.0)
            try:
                run.font.color.rgb = RGBColor(83, 101, 110)
            except Exception:
                pass


def _style_cover_paragraph(
    paragraph,
    *,
    east_font: str,
    latin_font: str,
    size_pt: float,
    text: Any = "",
    bold: bool = False,
    color_rgb: tuple[int, int, int] | None = None,
    space_before_pt: float = 0.0,
    space_after_pt: float = 0.0,
    line_spacing_pt: float | None = None,
):
    try:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_before = Pt(float(space_before_pt or 0))
        paragraph.paragraph_format.space_after = Pt(float(space_after_pt or 0))
        if line_spacing_pt is not None:
            paragraph.paragraph_format.line_spacing = Pt(float(line_spacing_pt))
    except Exception:
        pass
    run = paragraph.add_run(str(text or ""))
    _set_run_font(run, str(east_font or "宋体"), str(latin_font or east_font or "宋体"), float(size_pt or 12))
    try:
        run.bold = bool(bold)
        if color_rgb is not None:
            run.font.color.rgb = RGBColor(int(color_rgb[0]), int(color_rgb[1]), int(color_rgb[2]))
    except Exception:
        pass
    return run


def _cover_image_caption(project_name: Any, filename: Any, source_hint: Any = "") -> str:
    name = str(project_name or "").strip()
    source = str(source_hint or "").strip().lower()
    stem = Path(str(filename or "")).stem.strip()
    stem = re.sub(r"[_\-]+", " ", stem).strip()
    stem = re.sub(r"(?:\(|（)\d+(?:\)|）)$", "", stem).strip()
    stem = re.sub(r"^微信图片\s*\d{6,}$", "微信图片", stem).strip()
    stem_key = stem.lower()

    if stem and stem_key not in _GENERIC_COVER_IMAGE_STEMS:
        label = stem
    elif source in {"site_photo", "site", "scene", "现场", "现场照片"}:
        label = "现场实景图"
    else:
        label = "项目效果图"
    return f"{name} · {label}" if name else label


def _submission_media_eligible(item: Any) -> bool:
    """Reject unverified text-bearing figures from bidder-facing exports.

    AI image generators are not reliable at rendering Chinese labels, and the
    legacy BoQ overview mixed incomparable units on one axis.  Both are useful
    as internal previews but must not be inserted into a formal submission
    unless a caller explicitly marks the figure as text-verified.
    """
    if isinstance(item, dict):
        caption = str(item.get("caption") or "").strip().lower()
        path = str(item.get("path") or "").strip().lower()
        source_kind = str(item.get("source_kind") or "").strip().lower()
        text_verified = item.get("text_verified") is True
    else:
        caption = ""
        path = str(item or "").strip().lower()
        source_kind = ""
        text_verified = False
    if source_kind == "logo" or "投标单位logo" in caption.replace(" ", ""):
        return False
    if "gemini" in caption or "ai生成" in caption or "ai 生成" in caption:
        return text_verified
    if source_kind in {"external_ai", "ai", "generated_ai", "outline_mindmap"}:
        return text_verified
    if "boq统计概览" in caption.replace(" ", "") or "boq_stats_" in path:
        return False
    return True


def _rank_submission_media(items: Any) -> List[Any]:
    """Deduplicate and rank real project evidence ahead of generated figures."""
    ranked: List[tuple[int, int, Any]] = []
    seen: set[str] = set()
    priorities = {
        "site_photo": 0,
        "drawing": 1,
        "deterministic_project_diagram": 2,
    }
    for index, item in enumerate(items or []):
        if not _submission_media_eligible(item):
            continue
        if isinstance(item, dict):
            path = str(item.get("path") or "").strip()
            source_kind = str(item.get("source_kind") or "").strip().lower()
            identity = str(item.get("source_sha256") or "").strip() or path
        else:
            path = str(item or "").strip()
            source_kind = ""
            identity = path
        if not path or not Path(path).exists() or not identity or identity in seen:
            continue
        seen.add(identity)
        ranked.append((priorities.get(source_kind, 3), index, item))
    ranked.sort(key=lambda entry: (entry[0], entry[1]))
    return [entry[2] for entry in ranked]


def _resolve_cover_meta(data: Dict[str, Any] | None) -> Dict[str, Any]:
    raw = data if isinstance(data, dict) else {}
    branding = raw.get("branding") if isinstance(raw.get("branding"), dict) else {}
    topic = str(raw.get("topic") or "施工组织设计").strip() or "施工组织设计"
    project_name = str(raw.get("project_name") or "").strip() or _topic_to_cover_project_name(topic)
    project_code = str(raw.get("project_code") or raw.get("project_id") or "").strip()
    bidder_company = str(branding.get("bidder_company") or raw.get("bidder_company") or "").strip()

    logo_path = str(branding.get("logo_path") or raw.get("logo_path") or "").strip()
    if logo_path and not Path(logo_path).exists():
        logo_path = ""
    cover_image_path = str(raw.get("cover_image_path") or branding.get("cover_image_path") or "").strip()
    if cover_image_path and not Path(cover_image_path).exists():
        cover_image_path = ""

    cover_source_hint = "site_photo"
    cover_source_caption = ""
    if not cover_image_path:
        for media_item in _rank_submission_media(raw.get("media") or []):
            if not isinstance(media_item, dict):
                continue
            if str(media_item.get("source_kind") or "").strip().lower() != "site_photo":
                continue
            candidate = str(media_item.get("path") or "").strip()
            if candidate and Path(candidate).exists():
                cover_image_path = candidate
                cover_source_hint = "site_photo"
                cover_source_caption = str(media_item.get("caption") or "").strip()
                break

    cover_image_caption = str(raw.get("cover_image_caption") or branding.get("cover_image_caption") or "").strip()
    if cover_image_path and not cover_image_caption:
        cover_image_caption = cover_source_caption or _cover_image_caption(
            project_name,
            Path(cover_image_path).name,
            cover_source_hint,
        )

    return {
        "project_id": str(raw.get("project_id") or branding.get("project_id") or "").strip(),
        "project_name": project_name,
        "project_code": project_code,
        "topic": topic,
        "cover_title": str(raw.get("cover_title") or "施工组织设计").strip() or "施工组织设计",
        "cover_image_path": cover_image_path,
        "cover_image_caption": cover_image_caption,
        "bidder_company": bidder_company,
        "logo_path": logo_path,
        "issue_year_month": str(raw.get("issue_year_month") or branding.get("issue_year_month") or "").strip()
        or _format_cover_year_month(),
    }


def _insert_cover_page(doc: Document, style_cfg: Dict[str, Any], cover_meta: Dict[str, Any] | None) -> None:
    cfg = _normalize_style(style_cfg or {})
    meta = cover_meta if isinstance(cover_meta, dict) else {}
    title_font = str(cfg.get("title_font") or "黑体")
    title_latin = str(cfg.get("title_latin_font") or cfg.get("body_latin_font") or title_font)
    body_font = str(cfg.get("body_font") or "宋体")
    body_latin = str(cfg.get("body_latin_font") or body_font)
    accent = (15, 89, 102)

    project_name = str(meta.get("project_name") or "").strip()
    project_code = str(meta.get("project_code") or "").strip()
    cover_title = str(meta.get("cover_title") or "施工组织设计").strip() or "施工组织设计"
    bidder_company = str(meta.get("bidder_company") or "").strip()
    issue_year_month = str(meta.get("issue_year_month") or "").strip()

    kicker = doc.add_paragraph()
    _style_cover_paragraph(
        kicker,
        east_font=title_font,
        latin_font=title_latin,
        size_pt=12,
        text="技术文件",
        bold=True,
        color_rgb=accent,
        space_before_pt=18,
        space_after_pt=12,
        line_spacing_pt=18,
    )
    _set_paragraph_border(
        kicker,
        edge="bottom",
        color=str(cfg["palette"]["border"]),
        sz=8,
        space=6,
    )

    if project_name:
        _style_cover_paragraph(
            doc.add_paragraph(),
            east_font=title_font,
            latin_font=title_latin,
            size_pt=max(float(cfg.get("doc_title_size") or 20), 20.0),
            text=project_name,
            bold=True,
            color_rgb=accent,
            space_before_pt=18,
            space_after_pt=10,
            line_spacing_pt=30,
        )
    if project_code:
        _style_cover_paragraph(
            doc.add_paragraph(),
            east_font=body_font,
            latin_font=body_latin,
            size_pt=12,
            text=f"项目编号：{project_code}",
            space_after_pt=18,
        )
    _style_cover_paragraph(
        doc.add_paragraph(),
        east_font=title_font,
        latin_font=title_latin,
        size_pt=28,
        text=cover_title,
        bold=True,
        space_before_pt=18,
        space_after_pt=24,
        line_spacing_pt=36,
    )

    cover_image_path = str(meta.get("cover_image_path") or "").strip()
    cover_image_caption = str(meta.get("cover_image_caption") or "").strip()
    if cover_image_path and Path(cover_image_path).exists():
        try:
            shape = doc.add_picture(cover_image_path, width=Cm(14.2))
            _set_inline_shape_alt_text(
                shape,
                title="封面项目图片",
                description=cover_image_caption or "项目现场或效果图",
            )
            _set_inline_shape_border(shape, color=str(cfg["palette"]["border"]), width=12700)
            _configure_inline_image_paragraph(
                doc.paragraphs[-1],
                space_before_pt=8,
                space_after_pt=4,
                keep_with_next=bool(cover_image_caption),
            )
        except Exception:
            pass
    if cover_image_caption:
        _style_cover_paragraph(
            doc.add_paragraph(),
            east_font=body_font,
            latin_font=body_latin,
            size_pt=cfg["body_size"],
            text=cover_image_caption,
            color_rgb=(90, 98, 102),
            space_before_pt=4,
            space_after_pt=20,
        )

    logo_path = str(meta.get("logo_path") or "").strip()
    if logo_path and Path(logo_path).exists():
        try:
            shape = doc.add_picture(logo_path, width=Cm(2.4))
            _set_inline_shape_alt_text(shape, title="投标人标识", description="投标人企业标识")
            _configure_inline_image_paragraph(
                doc.paragraphs[-1],
                space_before_pt=12,
                space_after_pt=4,
                keep_with_next=True,
            )
        except Exception:
            pass
    if bidder_company:
        _style_cover_paragraph(
            doc.add_paragraph(),
            east_font=body_font,
            latin_font=body_latin,
            size_pt=14,
            text=bidder_company,
            bold=True,
            space_before_pt=18,
            space_after_pt=8,
        )
    if issue_year_month:
        _style_cover_paragraph(
            doc.add_paragraph(),
            east_font=body_font,
            latin_font=body_latin,
            size_pt=12,
            text=issue_year_month,
            space_after_pt=12,
        )
    doc.add_page_break()


def _apply_style(doc: Document, style: Dict[str, Any]):
    cfg = _normalize_style(style)
    _configure_professional_named_styles(doc, cfg)
    try:
        st = doc.styles["Normal"]
        _set_style_font(
            st,
            east_font=cfg["body_font"],
            latin_font=cfg["body_latin_font"],
            size_pt=cfg["body_size"],
        )
        if cfg.get("line_spacing_pt") is not None:
            st.paragraph_format.line_spacing = Pt(cfg["line_spacing_pt"])
        else:
            st.paragraph_format.line_spacing = cfg["line_spacing"]
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)
        st.paragraph_format.alignment = cfg["body_align"]
        st.paragraph_format.first_line_indent = Cm(cfg["first_line_indent_cm"])
        _set_first_line_chars(st, 200)
        st.paragraph_format.widow_control = True
    except Exception:
        pass
    heading_sizes = {
        "Heading 1": cfg["doc_title_size"],
        "Heading 2": cfg["title_size"],
        "Heading 3": cfg["title_size"],
    }
    palette = cfg["palette"]
    heading_colors = {
        "Heading 1": RGBColor.from_string(palette["accent_dark"]),
        "Heading 2": RGBColor.from_string(palette["accent"]),
        "Heading 3": RGBColor.from_string(palette["accent_dark"]),
    }
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        try:
            heading_style = doc.styles[style_name]
            _set_style_font(
                heading_style,
                east_font=cfg["title_font"],
                latin_font=cfg["title_latin_font"],
                size_pt=heading_sizes[style_name],
            )
            heading_style.font.bold = True
            heading_style.font.color.rgb = heading_colors[style_name]
            heading_style.paragraph_format.keep_with_next = True
            heading_style.paragraph_format.keep_together = True
            heading_style.paragraph_format.page_break_before = False
            heading_style.paragraph_format.widow_control = True
            heading_style.paragraph_format.first_line_indent = Cm(0)
            heading_style.paragraph_format.space_before = Pt(12 if style_name == "Heading 1" else 8)
            heading_style.paragraph_format.space_after = Pt(6)
        except Exception:
            pass

    def apply_paragraph(p, is_title: bool = False):
        font_east = cfg["title_font"] if is_title else cfg["body_font"]
        font_latin = cfg["title_latin_font"] if is_title else cfg["body_latin_font"]
        size = cfg["title_size"] if is_title else cfg["body_size"]
        try:
            if cfg.get("line_spacing_pt") is not None:
                p.paragraph_format.line_spacing = Pt(cfg["line_spacing_pt"])
            else:
                p.paragraph_format.line_spacing = cfg["line_spacing"]
            p.paragraph_format.widow_control = True
            p.paragraph_format.space_before = Pt(0 if not is_title else 8)
            p.paragraph_format.space_after = Pt(0 if not is_title else 6)
            if not is_title and cfg["first_line_indent_cm"] > 0:
                _set_first_line_chars(p, 200)
            align = cfg["title_align"] if is_title else cfg["body_align"]
            if align is not None:
                p.paragraph_format.alignment = align
        except Exception:
            pass
        for r in p.runs:
            _set_run_font(r, font_east, font_latin, size)
        if is_title:
            try:
                style_name = str(getattr(getattr(p, "style", None), "name", "") or "")
                if style_name == "Heading 1":
                    # Benchmark technical bids consistently use a centered
                    # chapter opener with a restrained rule, while lower-level
                    # headings remain left aligned for fast scanning.
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _set_paragraph_border(p, edge="bottom", color=palette["accent"], sz=10, space=6)
                    p.paragraph_format.left_indent = Cm(0)
                    p.paragraph_format.right_indent = Cm(0)
                    p.paragraph_format.space_before = Pt(14)
                    p.paragraph_format.space_after = Pt(10)
                elif style_name == "Heading 2":
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    _set_paragraph_border(p, edge="bottom", color=palette["border"], sz=6, space=4)
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after = Pt(6)
                elif style_name == "Heading 3":
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(4)
            except Exception:
                pass

    return apply_paragraph


def _apply_branding_header(doc: Document, style_cfg: Dict[str, Any], *, topic: str, bidder_company: str, logo_path: str | None):
    """
    Best-effort brand output:
    - Put company name (and logo if available) into page header.
    This is intentionally lightweight: it should not affect tender-driven chapter structure.
    """
    company = str(bidder_company or "").strip()
    logo = str(logo_path or "").strip()
    if logo and not Path(logo).exists():
        logo = ""
    font_east = str(style_cfg.get("body_font") or "宋体")
    font_latin = str(style_cfg.get("body_latin_font") or font_east)
    size_pt = 9.0
    palette = (_normalize_style(style_cfg or {}).get("palette") or {})

    for sec in doc.sections:
        try:
            header = sec.header
        except Exception:
            continue
        try:
            # Avoid duplicate branding if caller exports multiple times into same doc instance.
            if getattr(header, "_zf_branding", False):
                continue
            setattr(header, "_zf_branding", True)
        except Exception:
            pass

        try:
            _clear_block_container(header)
            usable_width = _usable_page_width_cm(doc)
            table = header.add_table(rows=1, cols=2, width=Cm(usable_width))
            _mark_table_header_row(table)
            try:
                table.autofit = True
            except Exception:
                pass
            cell_logo = table.cell(0, 0)
            cell_text = table.cell(0, 1)
            _set_cell_width(cell_logo, usable_width * 0.62)
            _set_cell_width(cell_text, usable_width * 0.38)
            _set_table_all_borders(table, color=str(palette.get("border") or "AFC4CE"), sz=6, bottom=True)

            # Logo (left)
            if logo:
                p0 = cell_logo.paragraphs[0] if cell_logo.paragraphs else cell_logo.add_paragraph()
                try:
                    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
                except Exception:
                    pass
                try:
                    r0 = p0.add_run()
                    shape = r0.add_picture(logo, width=Cm(2.0))
                    _set_inline_shape_alt_text(shape, title="投标人标识", description="投标人企业标识")
                    _configure_inline_image_paragraph(
                        p0,
                        space_before_pt=0,
                        space_after_pt=3,
                        keep_with_next=False,
                    )
                    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
                except Exception:
                    pass

            p0 = cell_logo.paragraphs[0] if cell_logo.paragraphs else cell_logo.add_paragraph()
            try:
                p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p0.paragraph_format.space_before = Pt(0)
                p0.paragraph_format.space_after = Pt(3)
            except Exception:
                pass
            left_text = company or (_topic_to_cover_project_name(topic) or str(topic or "施工组织设计"))
            if left_text:
                r0_text = p0.add_run(f" {left_text}" if logo else left_text)
                _set_run_font(r0_text, font_east, font_latin, size_pt)
                try:
                    r0_text.font.color.rgb = RGBColor(15, 89, 102)
                    r0_text.bold = True
                except Exception:
                    pass

            # Document label (right)
            p1 = cell_text.paragraphs[0] if cell_text.paragraphs else cell_text.add_paragraph()
            try:
                p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            except Exception:
                pass
            header_text = "施工组织设计"
            if header_text:
                r1 = p1.add_run(header_text)
                try:
                    _set_run_font(r1, font_east, font_latin, size_pt)
                    r1.font.color.rgb = RGBColor(83, 101, 110)
                except Exception:
                    pass
        except Exception:
            continue


def _extract_chapter_page_target(chapter_pages: Dict[str, Any], title: str) -> int | None:
    if not isinstance(chapter_pages, dict):
        return None
    raw = chapter_pages.get(title)
    if raw is None:
        return None
    if isinstance(raw, dict):
        raw = (
            raw.get("target")
            or raw.get("pages")
            or raw.get("page_target")
            or raw.get("count")
        )
    target = _to_int(raw, 0)
    return target if target > 0 else None


def _estimate_chars_per_page(style_cfg: Dict[str, Any]) -> int:
    body_size = _to_float(style_cfg.get("body_size"), 14.0)
    line_spacing = _to_float(style_cfg.get("line_spacing"), 1.5)
    line_spacing_pt = style_cfg.get("line_spacing_pt")
    if line_spacing_pt is not None:
        try:
            line_spacing = max(1.0, min(3.0, float(line_spacing_pt) / max(8.0, body_size)))
        except Exception:
            pass
    margins = style_cfg.get("margins_cm") or {}
    left = _to_float(margins.get("left"), 2.0)
    right = _to_float(margins.get("right"), 2.0)
    top = _to_float(margins.get("top"), 2.5)
    bottom = _to_float(margins.get("bottom"), 2.0)

    width_factor = 5.5 / max(2.0, left + right)
    height_factor = 5.0 / max(2.0, top + bottom)
    margin_factor = max(0.75, min(1.25, (width_factor + height_factor) / 2.0))
    est = int(900 * (12.0 / max(8.0, body_size)) * (1.5 / max(1.0, line_spacing)) * margin_factor)
    return max(350, min(1800, est))


def _estimate_content_pages(content: str, chars_per_page: int) -> int:
    text = (content or "").replace("\n", "")
    if not text:
        return 1
    return max(1, math.ceil(len(text) / max(1, chars_per_page)))


def _is_overview_section(title: str) -> bool:
    return bool(_OVERVIEW_SECTION_RE.search(str(title or "").strip()))


def _auto_density_images_for_pages(chapter_pages: int, total_pages: int) -> int:
    """
    Submission-safe image density policy.

    Figures are supporting evidence, not page filler.  Keep at most two distinct
    figures per chapter and reduce density further for very long documents.
    """
    cp = max(0, int(chapter_pages or 0))
    tp = max(0, int(total_pages or 0))
    if cp <= 0:
        return 0
    divisor = 4 if tp <= 200 else 6
    return min(2, max(1, math.ceil(cp / divisor)))


def _append_inline_markdown_runs(paragraph: Any, text: str) -> None:
    """Render the small Markdown subset produced by generation without leaking markers."""
    value = str(text or "")
    cursor = 0
    for match in _INLINE_MARKDOWN_RE.finditer(value):
        if match.start() > cursor:
            paragraph.add_run(value[cursor: match.start()])
        run = paragraph.add_run(match.group(2))
        run.bold = True
        cursor = match.end()
    if cursor < len(value):
        paragraph.add_run(value[cursor:])


def _normalize_risk_triplet_line(value: str) -> Dict[str, str] | None:
    line = re.sub(r"^(?:[-*•·]|\d+[.)、])\s*", "", str(value or "").strip())
    line = re.sub(r"\*\*|__", "", line).strip()
    match = _RISK_TRIPLET_RE.match(line)
    if not match:
        return None
    result = {key: str(val or "").strip() for key, val in match.groupdict().items()}
    return result if all(result.values()) else None


def _risk_triplet_signature(triplet: Dict[str, str]) -> tuple[str, str, str]:
    """Build a stable signature for suppressing repeated generic control cards."""
    values = []
    for key in ("risk", "control", "verification"):
        value = str(triplet.get(key) or "").strip()
        value = re.sub(r"(?:资料依据】)+$", "", value).strip()
        value = re.sub(r"\s+", "", value)
        values.append(value)
    return tuple(values)


def _append_risk_triplet_table(
    doc: Document,
    apply_paragraph: Any,
    triplet: Dict[str, str],
    style_cfg: Dict[str, Any] | None = None,
) -> None:
    table = doc.add_table(rows=2, cols=3)
    try:
        table.style = "Table Grid"
        table.autofit = False
    except Exception:
        pass
    _mark_table_header_row(table)
    headers = ("风险", "控制措施", "验证与留痕")
    values = (triplet["risk"], triplet["control"], triplet["verification"])
    widths = (4.7, 6.2, 5.1)
    for index, header in enumerate(headers):
        header_cell = table.rows[0].cells[index]
        value_cell = table.rows[1].cells[index]
        _set_cell_width(header_cell, widths[index])
        _set_cell_width(value_cell, widths[index])
        _set_cell_shading(header_cell, "D9F2FA")
        header_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        hp = header_cell.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(header)
        hr.bold = True
        apply_paragraph(hp)
        hp.paragraph_format.first_line_indent = Cm(0)
        vp = value_cell.paragraphs[0]
        _append_inline_markdown_runs(vp, values[index])
        apply_paragraph(vp)
        vp.paragraph_format.first_line_indent = Cm(0)
        vp.paragraph_format.space_after = Pt(0)
    _style_professional_table(table, style_cfg)
    _apply_semantic_table_layout(table, list(headers), [list(values)], style_cfg)


def _project_chapter_lead(topic: str, chapter_title: str) -> str:
    """Return a concise bidder-facing lead for the known project context."""
    if "医院" not in str(topic or "") or "改造" not in str(topic or ""):
        return ""
    title = str(chapter_title or "")
    project_name = _topic_to_cover_project_name(topic) or "本工程"
    if "整体理解" in title or "工程概况" in title:
        return (
            f"{project_name}按医院局部改造场景组织实施。施工部署以改造区域边界、专业接口协同、"
            "既有设施保护、院感边界、医疗秩序和交付验证为主线；具体作业时段、通行边界及临时切换方案在开工前由建设单位确认。"
        )
    if "重点难点" in title:
        return (
            "本章将装饰装修、弱电智能化、防水防渗及既有设施保护纳入同一个接口台账，"
            "按“前置复核—样板确认—过程旁站—隐蔽验收—成品保护”形成闭环，减少拆改返工和医疗环境扰动。"
        )
    if "新技术" in title or "新工艺" in title:
        return (
            "拟采用的新技术以可实施、可验证和可移交为原则。每项技术均须经专项交底、小范围试用和建设单位确认，"
            "不改变设计功能，不以技术展示代替工序质量验收。"
        )
    if "工期与质量" in title:
        return (
            "进度管理以招标文件和经批准的总进度计划为基准，将区域移交、材料到场、专业穿插、隐蔽验收与系统联调作为关键控制点。"
            "质量管理坚持先样板、后展开，上道工序未验收不转入下道工序。"
        )
    if "人、材、机" in title:
        return (
            "人、材、机配置随区域开放条件和工序节拍动态调整，重点保障装饰、防水、电气与智能化专业的接口配合。"
            "所有材料和设备实行进场报验、批次追溯和领用去向记录，未验收或不合格品不得投入使用。"
        )
    if "安全文明" in title:
        return (
            "安全文明施工除完成常规高处、临电、动火和机械作业控制外，重点管控医院改造期间的院感边界、洁污分流、粉尘噪声、人员通行和成品保护。"
            "异常情况先停止作业、隔离现场，经复核确认后再恢复施工。"
        )
    return ""


def _append_hospital_renovation_response_table(
    doc: Document,
    apply_paragraph: Any,
    style_cfg: Dict[str, Any] | None = None,
) -> None:
    rows = (
        (
            "改造区域与医疗环境边界",
            "开工前确认封闭范围、人员及材料通行路径，对粉尘、噪声和成品保护实施分区控制。",
            "区域移交单、边界巡检记录、恢复确认记录",
        ),
        (
            "装饰、防水与智能化接口",
            "通过图纸会审、样板先行和隐蔽验收统一标高、点位、收口及系统切换条件。",
            "图纸会审记录、样板验收记录、隐蔽工程验收记录",
        ),
        (
            "既有设施保护与临时切换",
            "实行先调查、后作业；涉及停用或切换时编制专项操作顺序和回退措施，未经确认不擅自操作。",
            "切换审批记录、功能测试记录、恢复确认单",
        ),
        (
            "清单、图纸与方案一致性",
            "建立清单项、设计位置、施工工序和验收资料的对应关系，发现差异先履行书面核验程序。",
            "清单对照表、设计变更/洽商记录、检验批资料",
        ),
        (
            "院感与洁污流线",
            "如邻近在用诊疗区域，采用连续封闭隔离、局部吸尘、门口粘尘和密闭转运；保洁消杀及垃圾转运路线由院方确认。",
            "隔离验收记录、保洁消杀记录、建筑垃圾转运记录",
        ),
        (
            "医疗秩序与系统切换",
            "噪声、停水停电、消防及弱电系统切换实行预约审批和书面告知；先验证回退条件，再按批准窗口实施。",
            "作业窗口确认单、切换审批单、功能测试及恢复记录",
        ),
    )
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    table.autofit = False
    _mark_table_header_row(table)
    headers = ("项目约束", "施工组织响应", "验证与留痕")
    widths = (4.0, 7.6, 4.6)
    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        _set_cell_width(cell, widths[col])
        _set_cell_shading(cell, "D9F2FA")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        apply_paragraph(paragraph)
        paragraph.paragraph_format.first_line_indent = Cm(0)
    for row_index, row in enumerate(rows, 1):
        for col, value in enumerate(row):
            cell = table.rows[row_index].cells[col]
            _set_cell_width(cell, widths[col])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_index % 2 == 0:
                _set_cell_shading(cell, "F4FAFC")
            paragraph = cell.paragraphs[0]
            _append_inline_markdown_runs(paragraph, value)
            apply_paragraph(paragraph)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_after = Pt(0)
    _style_professional_table(table, style_cfg)
    _apply_semantic_table_layout(table, list(headers), [list(row) for row in rows], style_cfg)


def _split_markdown_table_row(value: str) -> List[str]:
    line = str(value or "").strip().strip("|")
    return [re.sub(r"\\\|", "|", cell).strip() for cell in line.split("|")]


def _is_markdown_table_separator(value: str) -> bool:
    cells = _split_markdown_table_row(value)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def _append_submission_markdown_table(
    doc: Document,
    apply_paragraph: Any,
    headers: List[str],
    rows: List[List[str]],
    style_cfg: Dict[str, Any] | None = None,
) -> None:
    """Render generated Markdown matrices as readable Word tables.

    Five-column safety matrices are condensed into three wider columns so the
    result remains legible on A4 portrait pages.
    """
    normalized_headers = [re.sub(r"\*\*", "", str(item or "")).strip() for item in headers]
    normalized_rows = [
        [re.sub(r"\*\*", "", str(item or "")).strip() for item in row]
        for row in rows
    ]
    if len(normalized_headers) == 5:
        table_headers = ["风险事项", "控制与验证", "记录与处置"]
        table_rows: List[List[str]] = []
        for row in normalized_rows:
            padded = (row + [""] * 5)[:5]
            table_rows.append(
                [
                    padded[0],
                    f"控制：{padded[1]}\n验证：{padded[2]}",
                    f"记录：{padded[3]}\n处置：{padded[4]}",
                ]
            )
        widths = (3.8, 7.0, 5.4)
    else:
        column_count = max(1, min(6, len(normalized_headers)))
        table_headers = (normalized_headers + [f"字段{idx + 1}" for idx in range(column_count)])[:column_count]
        table_rows = [(row + [""] * column_count)[:column_count] for row in normalized_rows]
        widths = _suggest_table_column_widths(table_headers, table_rows, total_width_cm=16.2)

    table = doc.add_table(rows=1 + len(table_rows), cols=len(table_headers))
    table.style = "Table Grid"
    table.autofit = False
    _mark_table_header_row(table)
    for col, header in enumerate(table_headers):
        cell = table.rows[0].cells[col]
        _set_cell_width(cell, widths[col])
        _set_cell_shading(cell, "D9F2FA")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        apply_paragraph(paragraph)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        for run in paragraph.runs:
            run.font.size = Pt(style_cfg.get("body_size") or 14.0)
    for row_index, row in enumerate(table_rows, 1):
        for col, value in enumerate(row):
            cell = table.rows[row_index].cells[col]
            _set_cell_width(cell, widths[col])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_index % 2 == 0:
                _set_cell_shading(cell, "F4FAFC")
            paragraph = cell.paragraphs[0]
            for part_index, part in enumerate(str(value or "").splitlines()):
                if part_index:
                    paragraph.add_run().add_break()
                _append_inline_markdown_runs(paragraph, part)
            apply_paragraph(paragraph)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(style_cfg.get("body_size") or 14.0)
    _style_professional_table(table, style_cfg)
    _apply_semantic_table_layout(table, table_headers, table_rows, style_cfg)


def _set_a4_section_geometry(
    section: Any,
    *,
    landscape: bool,
    style_cfg: Dict[str, Any],
    previous_section: Any | None = None,
) -> None:
    # LibreOffice does not reliably resolve Word's implicit "linked to
    # previous" header/footer when a new section only contains page geometry.
    # Reuse the explicit default story relationships so every orientation
    # change retains the branding header and PAGE/NUMPAGES footer.
    if previous_section is not None:
        for tag_name in ("headerReference", "footerReference"):
            for node in list(section._sectPr.findall(qn(f"w:{tag_name}"))):
                section._sectPr.remove(node)
            for node in previous_section._sectPr.findall(qn(f"w:{tag_name}")):
                if str(node.get(qn("w:type")) or "default") == "default":
                    section._sectPr.insert(0, copy.deepcopy(node))
        section.different_first_page_header_footer = False
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    section.page_width = Cm(29.7 if landscape else 21.0)
    section.page_height = Cm(21.0 if landscape else 29.7)
    section.top_margin = Cm(float(style_cfg.get("margin_top_cm") or 2.5))
    section.right_margin = Cm(float(style_cfg.get("margin_right_cm") or 2.0))
    section.bottom_margin = Cm(float(style_cfg.get("margin_bottom_cm") or 2.0))
    section.left_margin = Cm(float(style_cfg.get("margin_left_cm") or 2.0))


def _append_nested_table(
    cell: Any,
    nested: Dict[str, Any],
    *,
    style_cfg: Dict[str, Any],
) -> None:
    headers = [str(value or "").strip() for value in (nested.get("headers") or [])]
    raw_rows = nested.get("rows") or []
    if not headers or not isinstance(raw_rows, list):
        return
    rows = [
        [str(value or "").strip() for value in row]
        for row in raw_rows
        if isinstance(row, (list, tuple))
    ]
    rows = [(row + [""] * len(headers))[: len(headers)] for row in rows]
    table = cell.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col, header in enumerate(headers):
        table.rows[0].cells[col].text = header
    for row_index, row in enumerate(rows, start=1):
        for col, value in enumerate(row):
            table.rows[row_index].cells[col].text = value
    _style_professional_table(table, style_cfg)
    _apply_semantic_table_layout(table, headers, rows, style_cfg)


def _append_structured_tables(
    doc: Document,
    apply_paragraph: Any,
    definitions: Any,
    style_cfg: Dict[str, Any],
    *,
    restore_portrait: bool = True,
) -> bool:
    """Render explicit tender tables, including wide/merged/nested variants."""

    if definitions in (None, []):
        return False
    if not isinstance(definitions, list):
        raise ValueError("tables must be a list")
    current_landscape = False
    for table_index, definition in enumerate(definitions, start=1):
        if not isinstance(definition, dict):
            raise ValueError(f"tables[{table_index - 1}] must be an object")
        headers = [str(value or "").strip() for value in (definition.get("headers") or [])]
        raw_rows = definition.get("rows") or []
        if not headers or not isinstance(raw_rows, list):
            continue
        column_count = min(12, len(headers))
        headers = headers[:column_count]
        rows: List[List[Any]] = []
        for raw_row in raw_rows:
            if not isinstance(raw_row, (list, tuple)):
                continue
            rows.append((list(raw_row) + [""] * column_count)[:column_count])

        landscape = str(definition.get("orientation") or definition.get("layout") or "").strip().lower() in {
            "landscape",
            "horizontal",
            "横向",
        }
        if landscape != current_landscape:
            previous_section = doc.sections[-1]
            new_section = doc.add_section(WD_SECTION.NEW_PAGE)
            _set_a4_section_geometry(
                new_section,
                landscape=landscape,
                style_cfg=style_cfg,
                previous_section=previous_section,
            )
            current_landscape = landscape

        title = str(definition.get("title") or f"附表{table_index}").strip()
        if title:
            heading = doc.add_heading(title, level=2)
            apply_paragraph(heading, is_title=True)

        merge_groups = definition.get("merge_header_groups") or []
        has_group_header = isinstance(merge_groups, list) and any(isinstance(item, dict) for item in merge_groups)
        header_rows = 2 if has_group_header else 1
        table = doc.add_table(rows=header_rows + len(rows), cols=column_count)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        total_width_cm = 25.7 if landscape else 16.2
        text_rows = [
            [str(value.get("text") or "") if isinstance(value, dict) else str(value or "") for value in row]
            for row in rows
        ]
        widths = _suggest_table_column_widths(headers, text_rows, total_width_cm=total_width_cm)

        if has_group_header:
            covered: set[int] = set()
            for group in merge_groups:
                if not isinstance(group, dict):
                    continue
                start = max(0, min(column_count - 1, _to_int(group.get("start"), 0)))
                end = max(start, min(column_count - 1, _to_int(group.get("end"), start)))
                merged = table.rows[0].cells[start]
                if end > start:
                    merged = merged.merge(table.rows[0].cells[end])
                merged.text = str(group.get("label") or "").strip()
                covered.update(range(start, end + 1))
            for col in range(column_count):
                if col not in covered:
                    table.rows[0].cells[col].text = headers[col]
            _mark_table_header_row(table)
            second_tr_pr = table.rows[1]._tr.get_or_add_trPr()
            if second_tr_pr.find(qn("w:tblHeader")) is None:
                second_tr_pr.append(OxmlElement("w:tblHeader"))

        header_row = table.rows[header_rows - 1]
        for col, header in enumerate(headers):
            header_row.cells[col].text = header
            _set_cell_width(header_row.cells[col], widths[col])

        for row_index, row in enumerate(rows, start=header_rows):
            for col, raw_value in enumerate(row):
                cell = table.rows[row_index].cells[col]
                _set_cell_width(cell, widths[col])
                if isinstance(raw_value, dict):
                    cell.text = str(raw_value.get("text") or "").strip()
                    nested = raw_value.get("nested")
                    if isinstance(nested, dict):
                        _append_nested_table(cell, nested, style_cfg=style_cfg)
                else:
                    cell.text = str(raw_value or "")

        _style_professional_table(table, style_cfg)
        _apply_semantic_table_layout(table, headers, text_rows, style_cfg)
        if has_group_header:
            palette = _normalize_style(style_cfg)["palette"]
            for group_row_index in (0, 1):
                for cell in table.rows[group_row_index].cells:
                    _set_cell_shading(cell, palette["table_header"])
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.paragraph_format.first_line_indent = Cm(0)
                        for run in paragraph.runs:
                            _set_run_font(run, style_cfg["body_font"], style_cfg["body_latin_font"], style_cfg["body_size"])
                            run.bold = True
                            run.font.color.rgb = RGBColor.from_string("FFFFFF")

    if current_landscape and restore_portrait:
        previous_section = doc.sections[-1]
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        _set_a4_section_geometry(
            new_section,
            landscape=False,
            style_cfg=style_cfg,
            previous_section=previous_section,
        )
        return True
    return False


def _append_submission_content(
    doc: Document,
    apply_paragraph: Any,
    content: str,
    *,
    seen_risk_triplets: set[tuple[str, str, str]] | None = None,
    chapter_title: str = "",
    style_cfg: Dict[str, Any] | None = None,
) -> None:
    """Render plain generated text as semantic DOCX paragraphs.

    The old exporter put an entire chapter into one paragraph with soft line
    breaks, which produced unstable pagination and inaccessible structure.
    """
    lines = str(content or "").splitlines() or [""]
    line_index = 0
    while line_index < len(lines):
        raw = lines[line_index]
        line = raw.strip()
        if (
            line.startswith("|")
            and line_index + 1 < len(lines)
            and _is_markdown_table_separator(lines[line_index + 1])
        ):
            headers = _split_markdown_table_row(line)
            table_rows: List[List[str]] = []
            line_index += 2
            while line_index < len(lines) and lines[line_index].strip().startswith("|"):
                table_rows.append(_split_markdown_table_row(lines[line_index]))
                line_index += 1
            if headers and table_rows:
                _append_submission_markdown_table(doc, apply_paragraph, headers, table_rows, style_cfg)
            continue
        line_index += 1
        if not line:
            continue
        triplet = _normalize_risk_triplet_line(line)
        if triplet:
            signature = _risk_triplet_signature(triplet)
            if seen_risk_triplets is not None:
                if signature in seen_risk_triplets:
                    continue
                seen_risk_triplets.add(signature)
            _append_risk_triplet_table(doc, apply_paragraph, triplet, style_cfg)
            continue
        markdown_heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        bracket_heading = re.match(r"^【([^】]{2,48})】\s*$", line)
        if markdown_heading or bracket_heading:
            if markdown_heading:
                heading_text = _normalize_submission_heading(markdown_heading.group(2))
                level = 2 if len(markdown_heading.group(1)) <= 2 else 3
            else:
                heading_text = _normalize_submission_heading(bracket_heading.group(1))
                level = 2
            if not heading_text or (
                chapter_title and _submission_heading_key(heading_text) == _submission_heading_key(chapter_title)
            ):
                continue
            p = doc.add_heading(heading_text, level=level)
            apply_paragraph(p, is_title=True)
            continue
        bold_heading = re.match(r"^\*\*(.{2,100})\*\*$", line)
        if bold_heading:
            heading_text = _normalize_submission_heading(bold_heading.group(1))
            if chapter_title and _submission_heading_key(heading_text) == _submission_heading_key(chapter_title):
                continue
            level = 3 if re.match(r"^(?:\d+\.\d+|[A-E]\.)", bold_heading.group(1).strip()) else 2
            p = doc.add_heading(heading_text, level=level)
            apply_paragraph(p, is_title=True)
            continue
        bracket_lead = re.match(r"^【([^】]{2,48})】\s*(.+)$", line)
        if bracket_lead:
            p = doc.add_paragraph()
            p.add_run(f"{bracket_lead.group(1).strip()}：").bold = True
            _append_inline_markdown_runs(p, bracket_lead.group(2).strip())
            apply_paragraph(p)
            continue
        if re.match(r"^(?:[-•·]|\*(?!\*)|\d+[.)、])\s*", line):
            cleaned = re.sub(r"^(?:[-•·]|\*(?!\*)|\d+[.)、])\s*", "", line).strip()
            try:
                p = doc.add_paragraph(style="List Bullet")
                _append_inline_markdown_runs(p, cleaned)
            except Exception:
                p = doc.add_paragraph()
                _append_inline_markdown_runs(p, f"• {cleaned}")
            apply_paragraph(p)
            try:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.left_indent = Cm(0.74)
                p.paragraph_format.space_after = Pt(2)
            except Exception:
                pass
            continue
        p = doc.add_paragraph()
        _append_inline_markdown_runs(p, line)
        apply_paragraph(p)


def _build_report_paths(output_path: str) -> tuple[Path, Path]:
    out = Path(output_path)
    return out.with_suffix(".build_report.json"), out.with_suffix(".build_report.log")


def _json_safe_report_value(value: Any) -> Any:
    try:
        return json.loads(json.dumps(value, ensure_ascii=False, default=str))
    except Exception:
        return str(value)


def _build_quality_evidence_report(quality_checks: Any) -> Dict[str, Any]:
    if not isinstance(quality_checks, dict) or not quality_checks:
        return {}
    return {
        "quality_checks": _json_safe_report_value(quality_checks),
        "evidence": _json_safe_report_value(quality_checks.get("evidence") or {}),
        "evidence_quality": _json_safe_report_value(quality_checks.get("evidence_quality") or {}),
        "evidence_traceability": _json_safe_report_value(quality_checks.get("evidence_traceability") or {}),
        "drawing_evidence": _json_safe_report_value(quality_checks.get("drawing_evidence") or {}),
        "standard_evidence": _json_safe_report_value(quality_checks.get("standard_evidence") or {}),
        "remediation": _json_safe_report_value(quality_checks.get("remediation") or []),
        "issue_list": _json_safe_report_value(quality_checks.get("issue_list") or []),
        "auto_revision_suggestions": _json_safe_report_value(quality_checks.get("auto_revision_suggestions") or []),
    }


def _write_docx_build_report(
    output_path: str,
    *,
    topic: str,
    sections: List[Dict[str, Any]],
    front_matter_plan: Dict[str, Any],
    layout_receipts: List[Dict[str, Any]],
    media_count: int,
    quality_checks: Any = None,
    figure_quality: Any = None,
    structural_quality: Any = None,
) -> None:
    _require_local_adapter_export_allowed({"sections": sections, "quality_checks": quality_checks}, "docx_build_report")
    report_json_path, report_log_path = _build_report_paths(output_path)
    report_json_path.parent.mkdir(parents=True, exist_ok=True)
    section_titles = [
        str((sec or {}).get("title") or f"章节{idx + 1}").strip() or f"章节{idx + 1}"
        for idx, sec in enumerate(sections or [])
    ]
    report = {
        "schema_version": "docx_build_report.v1",
        "output_path": str(output_path),
        "topic": str(topic or ""),
        "section_count": len(section_titles),
        "section_titles": section_titles,
        "front_matter_plan": dict(front_matter_plan or {}),
        "layout_receipts": list(layout_receipts or []),
        "media_count": int(media_count or 0),
    }
    quality_evidence = _build_quality_evidence_report(quality_checks)
    if quality_evidence:
        report["quality_evidence"] = quality_evidence
    if figure_quality:
        report["figure_quality"] = _json_safe_report_value(figure_quality)
    if structural_quality:
        report["structural_quality"] = _json_safe_report_value(structural_quality)
    report_json_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    report_log_path.write_text(
        "\n".join(
            [
                "DOCX build report",
                f"schema_version={report['schema_version']}",
                f"output_path={report['output_path']}",
                f"topic={report['topic']}",
                f"section_count={report['section_count']}",
                f"media_count={report['media_count']}",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def _append_quality_evidence_appendix(
    doc: Document,
    apply_paragraph: Any,
    quality_checks: Any,
    *,
    sections: List[Dict[str, Any]],
    compare_cfg: Dict[str, Any],
) -> None:
    qc = quality_checks or {}
    if not qc:
        return

    doc.add_page_break()
    hq = doc.add_heading("质量校验摘要", level=1)
    apply_paragraph(hq, is_title=True)
    quality_keys = (
        "structure",
        "score_coverage",
        "closed_loop",
        "engineering",
        "risk_triplet",
        "qse_closed_loop",
        "logic_template_adherence",
        "quantitative",
        "vague_terms",
        "officialese",
        "repetition_control",
        "content_specificity",
        "content_density",
        "consistency",
        "boq_focus_coverage",
        "boq_focus_item_closure",
        "boq_focus_item_typed_evidence",
        "required_topics",
        "required_topics_detail",
        "trade_names",
        "evidence",
        "evidence_quality",
        "evidence_traceability",
        "drawing_evidence",
        "standard_evidence",
        "template_style",
    )
    for key in quality_keys:
        item = qc.get(key) or {}
        ok = item.get("ok")
        p = doc.add_paragraph(f"{key}：{'通过' if ok else '需改进'}")
        apply_paragraph(p)
        for k, v in item.items():
            if k == "ok":
                continue
            p2 = doc.add_paragraph(f"- {k}: {v}")
            apply_paragraph(p2)

    hqc = doc.add_heading("质量校验清单", level=2)
    apply_paragraph(hqc, is_title=True)
    for key in quality_keys:
        item = qc.get(key) or {}
        ok = item.get("ok")
        mark = "☑" if ok else "☐"
        p = doc.add_paragraph(f"{mark} {key}")
        apply_paragraph(p)
        details = {k: v for k, v in item.items() if k != "ok"}
        if details and not ok:
            for k, v in details.items():
                p2 = doc.add_paragraph(f"  - {k}: {v}")
                apply_paragraph(p2)

    by_section = qc.get("score_coverage_by_section") or []
    if by_section:
        hsc = doc.add_heading("章节评分点覆盖清单", level=2)
        apply_paragraph(hsc, is_title=True)
        for sec in by_section:
            title = sec.get("title") or "章节"
            ok = sec.get("ok")
            mark = "☑" if ok else "☐"
            p = doc.add_paragraph(f"{mark} {title}")
            apply_paragraph(p)
            if not ok:
                for miss in sec.get("missing", []):
                    p2 = doc.add_paragraph(f"  - 缺失：{miss.get('dimension')} / {miss.get('keywords')}")
                    apply_paragraph(p2)

    by_evidence = (qc.get("evidence") or {}).get("by_section") or []
    if by_evidence:
        hce = doc.add_heading("章节证据数量清单", level=2)
        apply_paragraph(hce, is_title=True)
        for sec in by_evidence:
            title = sec.get("title") or "章节"
            cnt = sec.get("evidence_count")
            p = doc.add_paragraph(f"- {title}: 证据数 {cnt}")
            apply_paragraph(p)

    issues = qc.get("issue_list") or []
    hi = doc.add_heading("问题清单（自动检测）", level=2)
    apply_paragraph(hi, is_title=True)
    if issues:
        for it in issues[:200]:
            sev = it.get("severity") or "medium"
            title = it.get("title") or "章节"
            typ = it.get("type") or "issue"
            prob = it.get("problem") or ""
            sugg = it.get("suggestion") or ""
            p = doc.add_paragraph(f"- [{sev}] {title} / {typ}: {prob}")
            apply_paragraph(p)
            if sugg:
                p2 = doc.add_paragraph(f"  建议：{sugg}")
                apply_paragraph(p2)
    else:
        p = doc.add_paragraph("- 无")
        apply_paragraph(p)

    recs = qc.get("auto_revision_suggestions") or []
    hr = doc.add_heading("自动修订建议（按章节聚合）", level=2)
    apply_paragraph(hr, is_title=True)
    if recs:
        seen = set()
        for r in recs[:300]:
            title = r.get("title") or "章节"
            typ = r.get("type") or "issue"
            sugg = r.get("suggestion") or ""
            key = (title, typ, sugg)
            if key in seen:
                continue
            seen.add(key)
            p = doc.add_paragraph(f"- {title} / {typ}: {sugg}")
            apply_paragraph(p)
    else:
        p = doc.add_paragraph("- 无")
        apply_paragraph(p)

    by_closed = qc.get("closed_loop_by_section") or []
    if by_closed:
        hcl = doc.add_heading("章节风险-措施闭环清单", level=2)
        apply_paragraph(hcl, is_title=True)
        for sec in by_closed:
            title = sec.get("title") or "章节"
            ok = sec.get("ok")
            mark = "☑" if ok else "☐"
            p = doc.add_paragraph(f"{mark} {title}（风险: {sec.get('has_risk')} / 措施: {sec.get('has_measure')}）")
            apply_paragraph(p)

    by_eng = qc.get("engineering_by_section") or []
    if by_eng:
        heg = doc.add_heading("章节工程落地要素清单", level=2)
        apply_paragraph(heg, is_title=True)
        for sec in by_eng:
            title = sec.get("title") or "章节"
            ok = sec.get("ok")
            mark = "☑" if ok else "☐"
            missing = sec.get("missing") or []
            p = doc.add_paragraph(f"{mark} {title}（缺失: {missing}）")
            apply_paragraph(p)

    remediation = qc.get("remediation") or []
    if remediation:
        hrs = doc.add_heading("整改建议清单", level=2)
        apply_paragraph(hrs, is_title=True)
        for rec in remediation:
            title = rec.get("title") or "章节"
            rtype = rec.get("type") or "issue"
            suggestion = rec.get("suggestion") or ""
            p = doc.add_paragraph(f"- {title} / {rtype}: {suggestion}")
            apply_paragraph(p)

    issue_list = qc.get("issue_list") or []
    if issue_list:
        hi = doc.add_heading("问题清单", level=2)
        apply_paragraph(hi, is_title=True)
        for it in issue_list:
            title = it.get("title") or "章节"
            itype = it.get("type") or "issue"
            sev = it.get("severity") or "medium"
            prob = it.get("problem") or ""
            p = doc.add_paragraph(f"- [{sev}] {title} / {itype}: {prob}")
            apply_paragraph(p)

    auto_revision = qc.get("auto_revision_suggestions") or []
    if auto_revision:
        ha = doc.add_heading("自动修订建议", level=2)
        apply_paragraph(ha, is_title=True)
        for rec in auto_revision:
            title = rec.get("title") or "章节"
            rtype = rec.get("type") or "issue"
            suggestion = rec.get("suggestion") or ""
            p = doc.add_paragraph(f"- {title} / {rtype}: {suggestion}")
            apply_paragraph(p)

    has_compare = False
    for sec in sections:
        if sec.get("auto_remediated") == "llm" and sec.get("original_content"):
            has_compare = True
            break
    if has_compare:
        hcp = doc.add_heading("LLM整改前后对比", level=2)
        apply_paragraph(hcp, is_title=True)
        mode = compare_cfg.get("mode", "full")
        max_chars = _to_int(compare_cfg.get("max_chars"), 800)
        titles_filter = compare_cfg.get("titles")
        for sec in sections:
            if sec.get("auto_remediated") != "llm" or not sec.get("original_content"):
                continue
            if isinstance(titles_filter, list) and sec.get("title") not in titles_filter:
                continue
            title = sec.get("title") or "章节"
            h3 = doc.add_heading(title, level=3)
            apply_paragraph(h3, is_title=True)
            p1 = doc.add_paragraph("整改前：")
            apply_paragraph(p1)
            before = sec.get("original_content") or ""
            after = sec.get("content") or ""
            if mode == "summary":
                before = before[:max_chars] + ("..." if len(before) > max_chars else "")
                after = after[:max_chars] + ("..." if len(after) > max_chars else "")
            p2 = doc.add_paragraph(before)
            apply_paragraph(p2)
            p3 = doc.add_paragraph("整改后：")
            apply_paragraph(p3)
            p4 = doc.add_paragraph(after)
            apply_paragraph(p4)


def export_autoplan_docx(data: Dict[str, Any], output_path: str) -> str:
    data = _sanitize_docx_payload(data)
    _require_local_adapter_export_allowed(data, "docx")
    style_raw = data.get("style") or {}
    style_cfg = _normalize_style(style_raw)
    chapter_pages = data.get("chapter_pages") or {}
    chapter_styles = style_raw.get("chapter_styles") if isinstance(style_raw, dict) else {}

    doc = Document()
    _apply_page_setup(doc, style_cfg)
    apply_paragraph = _apply_style(doc, style_raw)

    topic = data.get("topic") or "施组方案"
    cover_meta = _resolve_cover_meta(data)
    bidder_company = str(cover_meta.get("bidder_company") or "").strip()
    logo_path = cover_meta.get("logo_path")

    def _brand_image_with_logo(src_path: str) -> str:
        """
        Add a small logo corner mark to an image (best-effort) to unify brand output.
        Returns the branded image path or the original path when branding is not possible.
        """
        try:
            if not isinstance(logo_path, str) or not logo_path.strip():
                return src_path
            if not src_path or str(src_path) == str(logo_path):
                return src_path
            sp = Path(str(src_path))
            lp = Path(str(logo_path))
            if not sp.exists() or not sp.is_file() or not lp.exists() or not lp.is_file():
                return src_path

            # Create a stable output name next to the source image (avoid touching originals).
            out = sp.with_name(f"{sp.stem}_brand.png")
            if out.exists() and out.is_file():
                return str(out)

            from PIL import Image

            with Image.open(sp) as base_im:
                base = base_im.convert("RGBA")
                with Image.open(lp) as logo_im:
                    logo = logo_im.convert("RGBA")
                    # Scale logo: up to 12% width, capped.
                    max_w = min(220, max(80, int(base.width * 0.12)))
                    if logo.width > max_w:
                        hh = int(logo.height * (max_w / max(1, logo.width)))
                        logo = logo.resize((max_w, max(1, hh)))
                    margin = max(12, int(base.width * 0.015))
                    x0 = max(0, base.width - logo.width - margin)
                    y0 = max(0, margin)
                    base.alpha_composite(logo, (x0, y0))
                # Save as PNG for docx compatibility.
                base.convert("RGB").save(out, format="PNG")
            return str(out)
        except Exception:
            return src_path

    layout_receipts = []
    internal_review = str(data.get("document_audience") or "").strip().lower() == "internal_review"
    sections = _prepare_submission_sections(
        data.get("sections"),
        internal_review=internal_review,
        missing_parameters=data.get("missing_parameters"),
        project_topic=data.get("topic"),
    )
    # The submission sanitizer can remove internal-only evidence tails that made
    # the raw payload appear complete.  Re-run the build-report gate against the
    # exact bidder-facing sections before any DOCX bytes are saved.  This keeps a
    # blocked export from leaving behind a plausible-looking partial document.
    _require_local_adapter_export_allowed(
        {"sections": sections, "quality_checks": data.get("quality_checks")},
        "docx_build_report",
    )
    terminology_entries = load_global_terminology()
    raw_media = list(data.get("media") or [])
    required_prefilter_failures = []
    for raw_item in raw_media:
        if not isinstance(raw_item, dict) or not raw_item.get("required"):
            continue
        if not _submission_media_eligible(raw_item):
            required_prefilter_failures.append(
                {"item": dict(raw_item), "reason": ["required_image_not_eligible_for_submission"]}
            )
            continue
        required_path = Path(str(raw_item.get("path") or raw_item.get("source_path") or "")).expanduser()
        if not required_path.exists() or not required_path.is_file():
            required_prefilter_failures.append(
                {"item": dict(raw_item), "reason": ["image_file_missing"]}
            )
    ranked_media = _rank_submission_media(raw_media)
    media_quality = validate_media_collection(ranked_media)
    if required_prefilter_failures:
        media_quality["required_failures"] = [
            *(media_quality.get("required_failures") or []),
            *required_prefilter_failures,
        ]
        media_quality["rejected"] = [
            *(media_quality.get("rejected") or []),
            *required_prefilter_failures,
        ]
        media_quality["rejected_count"] = len(media_quality["rejected"])
        media_quality["status"] = "blocked"
    if media_quality.get("status") == "blocked":
        raise RuntimeError(
            json.dumps(
                {
                    "status": "blocked",
                    "export_allowed": False,
                    "export_kind": "docx_figure_quality",
                    "issues": [
                        {
                            "code": "REQUIRED_FIGURE_REJECTED",
                            "reason": list(entry.get("reason") or []),
                            "caption": str((entry.get("item") or {}).get("caption") or ""),
                        }
                        for entry in media_quality.get("required_failures") or []
                    ],
                },
                ensure_ascii=False,
            )
        )
    media_all = list(media_quality.get("accepted") or [])
    chart_policy = style_raw.get("chart_policy") if isinstance(style_raw, dict) and isinstance(style_raw.get("chart_policy"), dict) else {}
    chart_enabled = _to_bool(chart_policy.get("enabled"), True)
    chart_mode = str(chart_policy.get("mode") or "").strip().lower()
    chart_position = str(chart_policy.get("position") or "end").strip().lower()  # end|chapter
    chart_every_n = max(1, _to_int(chart_policy.get("every_n_chapters"), 2))
    chart_mode_auto_density = chart_mode in {"page_density_auto", "page_density", "auto_density"}
    if not chart_enabled:
        chart_mode_auto_density = False

    # Planned total pages are used for automatic image density policy.
    def _effective_pages_for_section(title: str, content_doc: str, section_style_cfg: Dict[str, Any]) -> int:
        t = _extract_chapter_page_target(chapter_pages, title)
        if t:
            return int(t)
        chars_per_page = _estimate_chars_per_page(section_style_cfg)
        return _estimate_content_pages(content_doc, chars_per_page)

    total_planned_pages = 0
    section_pages: List[int] = []
    for sec in sections:
        title = str(sec.get("title") or "章节")
        content = _strip_internal_autofix_markers(sec.get("content") or "")
        pages = _effective_pages_for_section(title, content, style_cfg)
        section_pages.append(pages)
        total_planned_pages += pages
    total_planned_pages = max(1, int(total_planned_pages))

    front_matter_plan = _resolve_front_matter_plan(
        style_raw=style_raw,
        data=data,
        body_pages_estimate=total_planned_pages,
    )
    _apply_branding_header(doc, style_cfg, topic=str(topic), bidder_company=bidder_company, logo_path=logo_path)
    _apply_footer_page_numbers(
        doc,
        style_cfg,
        bidder_company=bidder_company,
        logo_path=str(logo_path or ""),
        document_label=str(cover_meta.get("project_name") or topic or "施工组织设计"),
    )
    _insert_cover_page(doc, style_cfg, cover_meta)
    if int(front_matter_plan.get("full_index_pages") or 0) > 0:
        _insert_full_index_page(
            doc,
            apply_paragraph,
            topic=str(topic),
            sections=sections,
            chapter_pages=chapter_pages,
            effective_document_pages=int(front_matter_plan.get("effective_document_pages") or total_planned_pages),
        )
    _insert_auto_toc(
        doc,
        apply_paragraph,
        style_cfg=style_cfg,
        toc_pages=int(front_matter_plan.get("toc_pages") or 1),
        toc_entries=_build_static_toc_entries(
            sections=sections,
            section_pages=section_pages,
            front_matter_plan=front_matter_plan,
        ),
    )

    media_index = 0
    chapter_media_started = False
    seen_risk_triplets: set[tuple[str, str, str]] = set()
    consumed_media: set[str] = set()
    inserted_media_hashes: list[str] = []
    figure_insertions: list[Dict[str, Any]] = []
    figure_failures: list[Dict[str, Any]] = []

    def _media_identity(item: Any) -> str:
        if isinstance(item, dict):
            return str(
                item.get("asset_sha256")
                or item.get("source_sha256")
                or item.get("path")
                or ""
            ).strip()
        return str(item or "").strip()

    def _available_media_for_chapter(chapter_title: str, *, is_context: bool) -> List[Any]:
        return [
            item
            for item in media_all
            if _media_identity(item) not in consumed_media
            and media_matches_chapter(
                item,
                chapter_title,
                allow_unbound_project_source=bool(is_context),
            )
        ]

    def _media_parts(item: Any) -> tuple[str, str, str]:
        path = item.get("path") if isinstance(item, dict) else item
        caption = item.get("caption") if isinstance(item, dict) else None
        source_kind = (
            str(item.get("source_kind") or "local_asset").strip().lower()
            if isinstance(item, dict)
            else "local_asset"
        )
        if not caption:
            try:
                name = Path(str(path)).name
            except Exception:
                name = str(path)
            caption = "BoQ 统计概览" if "boq_stats_" in str(name) else name
        return str(path or ""), str(caption or "施工图表"), source_kind

    def _media_source_ref(item: Any, path: str) -> str:
        if isinstance(item, dict):
            explicit = str(
                item.get("source_ref")
                or item.get("source_filename")
                or item.get("original_filename")
                or ""
            ).strip()
            if explicit:
                return explicit
        return Path(str(path or "")).name

    def _append_media_item(item: Any, *, chapter_title: str = "") -> bool:
        nonlocal media_index
        if media_index >= _MAX_SUBMISSION_IMAGES:
            failure = {
                "caption": str((item or {}).get("caption") or "") if isinstance(item, dict) else "",
                "chapter_title": chapter_title,
                "reason": ["formal_figure_limit_reached"],
                "required": bool(isinstance(item, dict) and item.get("required")),
            }
            figure_failures.append(failure)
            if failure["required"]:
                raise RuntimeError(json.dumps({"status": "blocked", "issue": failure}, ensure_ascii=False))
            return False
        path, caption, source_kind = _media_parts(item)
        source_ref = _media_source_ref(item, path)
        source_receipt = validate_media_item(item, chapter_title=chapter_title)
        if not source_receipt.get("ok"):
            failure = {
                "caption": caption,
                "chapter_title": chapter_title,
                "reason": list(source_receipt.get("errors") or []),
                "required": bool(isinstance(item, dict) and item.get("required")),
            }
            figure_failures.append(failure)
            if failure["required"]:
                raise RuntimeError(json.dumps({"status": "blocked", "issue": failure}, ensure_ascii=False))
            return False
        try:
            branded_path = None if source_kind in {"site_photo", "drawing"} else _brand_image_with_logo(str(path))
            path_to_add = branded_path or str(path)
            insert_receipt = validate_media_item(
                {
                    "path": path_to_add,
                    "caption": caption,
                    "required": bool(isinstance(item, dict) and item.get("required")),
                },
                chapter_title=chapter_title,
            )
            if not insert_receipt.get("ok"):
                failure = {
                    "caption": caption,
                    "chapter_title": chapter_title,
                    "reason": list(insert_receipt.get("errors") or []),
                    "required": bool(isinstance(item, dict) and item.get("required")),
                }
                figure_failures.append(failure)
                if failure["required"]:
                    raise RuntimeError(json.dumps({"status": "blocked", "issue": failure}, ensure_ascii=False))
                return False
            image_width = _image_width_for_docx(path_to_add, source_kind=source_kind)
            shape = doc.add_picture(str(path_to_add), width=Cm(image_width))
            _set_inline_shape_alt_text(
                shape,
                title=str(caption or "施工图表"),
                description=f"与施工组织设计正文配套的图表：{str(caption or '施工图表')}",
            )
            _set_inline_shape_border(shape, color=str(style_cfg["palette"]["border"]), width=12700)
            try:
                image_paragraph = doc.paragraphs[-1]
                _configure_inline_image_paragraph(image_paragraph)
            except Exception:
                pass
            media_index += 1
            pc = doc.add_paragraph(f"图{media_index}：{caption}", style="Caption")
            apply_paragraph(pc)
            try:
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pc.paragraph_format.first_line_indent = Cm(0)
                pc.paragraph_format.line_spacing = 1.0
                pc.paragraph_format.space_before = Pt(0)
                pc.paragraph_format.space_after = Pt(8)
                pc.paragraph_format.keep_together = True
                for run in pc.runs:
                    _set_run_font(run, style_cfg["body_font"], style_cfg["body_latin_font"], style_cfg["body_size"])
                    run.font.color.rgb = RGBColor(24, 82, 112)
            except Exception:
                pass
            asset_hash = str(insert_receipt.get("sha256") or "")
            if asset_hash:
                inserted_media_hashes.append(asset_hash)
            figure_insertions.append(
                {
                    "figure_number": media_index,
                    "caption": caption,
                    "chapter_title": chapter_title,
                    "source_kind": source_kind,
                    "source_ref": source_ref,
                    "required": bool(isinstance(item, dict) and item.get("required")),
                    "asset_sha256": asset_hash,
                    "width_px": int(insert_receipt.get("width_px") or 0),
                    "height_px": int(insert_receipt.get("height_px") or 0),
                    "effective_dpi": insert_receipt.get("effective_dpi"),
                }
            )
            return True
        except Exception as exc:
            failure = {
                "caption": caption,
                "chapter_title": chapter_title,
                "reason": ["docx_image_insert_failed"],
                "error_type": type(exc).__name__,
                "required": bool(isinstance(item, dict) and item.get("required")),
            }
            figure_failures.append(failure)
            if failure["required"]:
                raise
            return False

    def _append_media_pair(items: List[Any], *, chapter_title: str = "") -> bool:
        """Render two landscape project photos as one professional evidence panel."""
        nonlocal media_index
        if media_index + 2 > _MAX_SUBMISSION_IMAGES:
            required_items = [item for item in items if isinstance(item, dict) and item.get("required")]
            if required_items:
                failure = {
                    "caption": "、".join(str(item.get("caption") or "") for item in required_items),
                    "chapter_title": chapter_title,
                    "reason": ["formal_figure_limit_reached"],
                    "required": True,
                }
                figure_failures.append(failure)
                raise RuntimeError(json.dumps({"status": "blocked", "issue": failure}, ensure_ascii=False))
            return False
        if not _media_pair_eligible(items):
            return False
        table = None
        pair_receipts: list[tuple[str, str, str, str, bool, Dict[str, Any]]] = []
        start_media_index = media_index
        staged_hashes: list[str] = []
        staged_insertions: list[Dict[str, Any]] = []
        try:
            for item in items:
                path, caption, source_kind = _media_parts(item)
                source_ref = _media_source_ref(item, path)
                required = bool(isinstance(item, dict) and item.get("required"))
                receipt = validate_media_item(item, chapter_title=chapter_title, insert_width_cm=7.45)
                if not receipt.get("ok"):
                    failure = {
                        "caption": caption,
                        "chapter_title": chapter_title,
                        "reason": list(receipt.get("errors") or []),
                        "required": bool(isinstance(item, dict) and item.get("required")),
                    }
                    figure_failures.append(failure)
                    if failure["required"]:
                        raise RuntimeError(json.dumps({"status": "blocked", "issue": failure}, ensure_ascii=False))
                    return False
                pair_receipts.append((path, caption, source_kind, source_ref, required, receipt))
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _prevent_table_row_split(table.rows[0])
            cell_width = 8.0
            first_figure_number = media_index + 1
            for col, (path, caption, source_kind, source_ref, required, receipt) in enumerate(pair_receipts):
                cell = table.rows[0].cells[col]
                _set_cell_width(cell, cell_width)
                _set_cell_margins(cell, top=80, start=85, bottom=75, end=85)
                _set_cell_border(
                    cell,
                    top={"color": style_cfg["palette"]["border"], "sz": 5},
                    bottom={"color": style_cfg["palette"]["border"], "sz": 5},
                    start={"color": style_cfg["palette"]["border"], "sz": 5},
                    end={"color": style_cfg["palette"]["border"], "sz": 5},
                )
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                image_paragraph = cell.paragraphs[0]
                _configure_inline_image_paragraph(
                    image_paragraph,
                    space_before_pt=0,
                    space_after_pt=3,
                    keep_with_next=True,
                )
                shape = image_paragraph.add_run().add_picture(path, width=Cm(7.45))
                _set_inline_shape_alt_text(
                    shape,
                    title=caption,
                    description=f"与施工组织设计正文配套的项目证据图片：{caption}",
                )
                _set_inline_shape_border(shape, color=str(style_cfg["palette"]["border"]), width=12700)
                media_index += 1
                caption_paragraph = cell.add_paragraph(f"图{media_index}：{caption}", style="Caption")
                caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_paragraph.paragraph_format.first_line_indent = Cm(0)
                caption_paragraph.paragraph_format.line_spacing = 1.0
                caption_paragraph.paragraph_format.space_before = Pt(0)
                caption_paragraph.paragraph_format.space_after = Pt(2)
                caption_paragraph.paragraph_format.keep_together = True
                for run in caption_paragraph.runs:
                    _set_run_font(run, style_cfg["body_font"], style_cfg["body_latin_font"], style_cfg["body_size"])
                    run.font.color.rgb = RGBColor.from_string(style_cfg["palette"]["muted"])
                asset_hash = str(receipt.get("sha256") or "")
                if asset_hash:
                    staged_hashes.append(asset_hash)
                staged_insertions.append(
                    {
                        "figure_number": first_figure_number + col,
                        "caption": caption,
                        "chapter_title": chapter_title,
                        "source_kind": source_kind,
                        "source_ref": source_ref,
                        "required": required,
                        "asset_sha256": asset_hash,
                        "width_px": int(receipt.get("width_px") or 0),
                        "height_px": int(receipt.get("height_px") or 0),
                        "effective_dpi": receipt.get("effective_dpi"),
                    }
                )
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(2)
            spacer.paragraph_format.line_spacing = 1.0
            inserted_media_hashes.extend(staged_hashes)
            figure_insertions.extend(staged_insertions)
            return True
        except Exception:
            media_index = start_media_index
            if table is not None:
                try:
                    parent = table._element.getparent()
                    if parent is not None:
                        parent.remove(table._element)
                except Exception:
                    pass
            if any(bool(isinstance(item, dict) and item.get("required")) for item in items):
                raise
            return False

    # 目录/章节内容
    for idx, sec in enumerate(sections):
        title = sec.get("title") or "章节"
        content = sec.get("content") or ""
        content_doc = _strip_internal_autofix_markers(content)
        try:
            content_doc, _ = normalize_text_terminology(content_doc, terminology_entries)
        except Exception:
            pass
        apply_this = apply_paragraph
        section_style_cfg = style_cfg
        if isinstance(chapter_styles, dict) and isinstance(chapter_styles.get(title), dict):
            merged_style = _merge_style(style_raw, chapter_styles.get(title) or {})
            apply_this = _apply_style(doc, merged_style)
            section_style_cfg = _normalize_style(merged_style)

        chapter_heading = doc.add_heading(title, level=1)
        apply_this(chapter_heading, is_title=True)
        _add_paragraph_bookmark(chapter_heading, f"ZF_CHAPTER_{idx + 1}", 1000 + idx)
        if idx > 0 and style_cfg.get("chapter_start_new_page"):
            # Put the page-start requirement on the heading itself.  An empty
            # page-break paragraph can combine with a template's heading rule
            # and yield a nearly blank sheet in Word/LibreOffice.
            chapter_heading.paragraph_format.page_break_before = True
        if internal_review and sec.get("agent_role"):
            p = doc.add_paragraph(f"内部责任角色：{sec.get('agent_role')}")
            apply_this(p)
        if not internal_review:
            chapter_lead = _project_chapter_lead(str(topic), str(title))
            if chapter_lead:
                lead = doc.add_paragraph()
                lead.add_run("本章实施主线：").bold = True
                lead.add_run(chapter_lead)
                apply_this(lead)
                try:
                    _set_paragraph_shading(lead, "EEF3F8")
                    lead.paragraph_format.left_indent = Cm(0.3)
                    lead.paragraph_format.right_indent = Cm(0.3)
                    lead.paragraph_format.first_line_indent = Cm(0)
                    lead.paragraph_format.space_before = Pt(4)
                    lead.paragraph_format.space_after = Pt(8)
                except Exception:
                    pass
            if idx == 0 and "医院" in str(topic) and "改造" in str(topic):
                response_heading = doc.add_heading("医院局部改造施工组织响应", level=2)
                apply_this(response_heading, is_title=True)
                _append_hospital_renovation_response_table(doc, apply_this, section_style_cfg)
        _append_submission_content(
            doc,
            apply_this,
            content_doc,
            seen_risk_triplets=seen_risk_triplets,
            chapter_title=str(title),
            style_cfg=section_style_cfg,
        )

        target_pages = _extract_chapter_page_target(chapter_pages, title)
        if target_pages:
            chars_per_page = _estimate_chars_per_page(section_style_cfg)
            estimated_pages = _estimate_content_pages(content_doc, chars_per_page)
            delta = estimated_pages - target_pages
            layout_receipts.append(
                {
                    "title": title,
                    "target_pages": target_pages,
                    "estimated_pages": estimated_pages,
                    "delta": delta,
                    "shortfall_pages": max(0, target_pages - estimated_pages),
                    "mechanical_padding_applied": False,
                }
            )
        else:
            chars_per_page = _estimate_chars_per_page(section_style_cfg)
            estimated_pages = _estimate_content_pages(content_doc, chars_per_page)

        # Auto density policy: the first chapter receives real project context;
        # later chapters receive at most one project-specific generated figure.
        if chart_mode_auto_density and chart_position in {"chapter", "per_chapter", "by_chapter"}:
            is_context_chapter = idx == 0 or _is_overview_section(title)
            chapter_candidates = _available_media_for_chapter(str(title), is_context=is_context_chapter)
            if is_context_chapter:
                remaining_sources = sum(
                    1
                    for pending in chapter_candidates
                    if isinstance(pending, dict)
                    and str(pending.get("source_kind") or "").strip().lower() in {"site_photo", "drawing"}
                )
                need_images = min(2, remaining_sources, max(0, _MAX_SUBMISSION_IMAGES - media_index))
            else:
                effective_pages = target_pages if target_pages else estimated_pages
                need_images = min(
                    _auto_density_images_for_pages(effective_pages, total_planned_pages),
                    max(0, _MAX_SUBMISSION_IMAGES - media_index),
                )
            if need_images > 0:
                if not chapter_media_started:
                    heading_text = "项目现场与图纸依据" if is_context_chapter else "施工组织图示"
                    hmc = doc.add_heading(heading_text, level=2)
                    apply_paragraph(hmc, is_title=True)
                    chapter_media_started = True
                chapter_items: List[Any] = []
                chapter_items.extend(chapter_candidates[:need_images])
                shortfall = need_images - len(chapter_items)
                if shortfall > 0 and not is_context_chapter:
                    generated = generate_section_visuals(
                        title=title,
                        content=content_doc,
                        image_count=min(shortfall, 1),
                        include_mindmap=False,
                    )
                    for generated_item in (generated or [])[:1]:
                        item = dict(generated_item) if isinstance(generated_item, dict) else {"path": str(generated_item or "")}
                        item["chapter_scope"] = [str(title)]
                        item.setdefault("semantic_terms", [str(title)])
                        item.setdefault("source_kind", "deterministic_project_diagram")
                        chapter_items.append(item)
                selected_items = chapter_items[:need_images]
                if _append_media_pair(selected_items, chapter_title=str(title)):
                    consumed_media.update(_media_identity(item) for item in selected_items)
                else:
                    for item in selected_items:
                        if _append_media_item(item, chapter_title=str(title)):
                            consumed_media.add(_media_identity(item))
        # Legacy chapter frequency policy (backward compatibility).
        elif media_all and chart_enabled and chart_position in {"chapter", "per_chapter", "by_chapter"}:
            chapter_candidates = _available_media_for_chapter(
                str(title),
                is_context=bool(idx == 0 or _is_overview_section(title)),
            )
            if ((idx + 1) % chart_every_n == 0) and chapter_candidates:
                if not chapter_media_started:
                    hmc = doc.add_heading("图表与插图（按章节分布）", level=2)
                    apply_paragraph(hmc, is_title=True)
                    chapter_media_started = True
                item = chapter_candidates[0]
                if _append_media_item(item, chapter_title=str(title)):
                    consumed_media.add(_media_identity(item))

    tables_restored_portrait = _append_structured_tables(
        doc,
        apply_paragraph,
        data.get("tables"),
        style_cfg,
        restore_portrait=bool(
            (media_all and not chart_mode_auto_density)
            or (
                internal_review
                and any(
                    data.get(name)
                    for name in ("drawing_index", "standard_index", "cross_index", "param_trace", "quality_checks")
                )
            )
        ),
    )

    # 图纸证据索引（可追溯）
    drawing_index = data.get("drawing_index") or {}
    if internal_review and isinstance(drawing_index, dict) and (drawing_index.get("drawings") or drawing_index.get("chapter_bindings")):
        doc.add_page_break()
        hd = doc.add_heading("图纸证据索引（自动生成）", level=1)
        apply_paragraph(hd, is_title=True)

        drawings = drawing_index.get("drawings") or []
        if drawings:
            h1 = doc.add_heading("图纸清单", level=2)
            apply_paragraph(h1, is_title=True)
            for d in drawings[:30]:
                fn = d.get("filename") or ""
                pages = d.get("pages")
                kws = d.get("keywords") or []
                line = f"- {fn}"
                if pages:
                    line += f"（页数={pages}）"
                if kws:
                    line += f"；关键词={'、'.join([str(x) for x in kws[:8] if str(x).strip()])}"
                p = doc.add_paragraph(line)
                apply_paragraph(p)

        binds = drawing_index.get("chapter_bindings") or []
        if binds:
            h2 = doc.add_heading("章节-图纸绑定", level=2)
            apply_paragraph(h2, is_title=True)
            for b in binds[:40]:
                ch = b.get("chapter") or ""
                loc = b.get("locator") or ""
                p = doc.add_paragraph(f"- {ch} -> {loc}")
                apply_paragraph(p)

    # 企业标准证据索引（可追溯）
    standard_index = data.get("standard_index") or {}
    if internal_review and isinstance(standard_index, dict) and (standard_index.get("standards") or standard_index.get("chapter_bindings")):
        doc.add_page_break()
        hd = doc.add_heading("企业标准证据索引（自动生成）", level=1)
        apply_paragraph(hd, is_title=True)

        standards = standard_index.get("standards") or []
        if standards:
            h1 = doc.add_heading("标准文件清单", level=2)
            apply_paragraph(h1, is_title=True)
            for d in standards[:40]:
                fn = d.get("filename") or ""
                pages = d.get("pages")
                kws = d.get("keywords") or []
                line = f"- {fn}"
                if pages:
                    line += f"（页数={pages}）"
                if kws:
                    line += f"；关键词={'、'.join([str(x) for x in kws[:8] if str(x).strip()])}"
                p = doc.add_paragraph(line)
                apply_paragraph(p)

        binds = standard_index.get("chapter_bindings") or []
        if binds:
            h2 = doc.add_heading("章节-标准绑定", level=2)
            apply_paragraph(h2, is_title=True)
            for b in binds[:60]:
                ch = b.get("chapter") or ""
                loc = b.get("locator") or ""
                p = doc.add_paragraph(f"- {ch} -> {loc}")
                apply_paragraph(p)

    # 重点项证据闭环索引（BoQ focus cross-index）
    cross_index = data.get("cross_index") or {}
    if internal_review and isinstance(cross_index, dict) and isinstance(cross_index.get("focus_items"), list) and cross_index.get("focus_items"):
        doc.add_page_break()
        hd = doc.add_heading("重点项证据闭环索引（自动生成）", level=1)
        apply_paragraph(hd, is_title=True)

        fc = int(cross_index.get("focus_count") or 0)
        mc = int(cross_index.get("mentioned_count") or 0)
        okc = int(cross_index.get("closed_ok_count") or 0)
        md = int(cross_index.get("missing_drawing_locator_count") or 0)
        ms = int(cross_index.get("missing_standard_locator_count") or 0)
        p0 = doc.add_paragraph(f"重点项={fc}；出现={mc}；闭环OK={okc}；缺图纸定位={md}；缺标准定位={ms}。")
        apply_paragraph(p0)
        p1 = doc.add_paragraph("闭环OK判定：同一章节内同时满足=量化（含单位数值）+风险→控制→验证+证据定位。")
        apply_paragraph(p1)

        items = cross_index.get("focus_items") or []
        # Keep table compact to avoid layout explosion on A4.
        table = doc.add_table(rows=1, cols=9)
        _mark_table_header_row(table)
        hdr = table.rows[0].cells
        headers = ["清单项", "类别/工序", "工程量", "单价", "合价", "落位章节", "图纸定位", "标准定位", "闭环"]
        for i, h in enumerate(headers):
            try:
                hdr[i].text = h
            except Exception:
                pass

        def _fmt_num(v: Any) -> str:
            try:
                if v is None:
                    return ""
                f = float(v)
                if abs(f - int(f)) < 1e-6:
                    return str(int(f))
                return f"{f:.4g}"
            except Exception:
                return str(v)

        for it in items[:24]:
            if not isinstance(it, dict):
                continue
            name = str(it.get("name") or "").strip()
            cats = it.get("categories") or []
            if isinstance(cats, list):
                cats = "、".join([str(x) for x in cats if str(x).strip()][:4])
            else:
                cats = str(cats) if cats else ""
            proc = str(it.get("process_name") or "").strip()
            cat_proc = cats
            if proc:
                cat_proc = (cat_proc + ("；" if cat_proc else "") + f"工序={proc}").strip()

            qty = _fmt_num(it.get("quantity"))
            unit = str(it.get("unit") or "").strip()
            qty_disp = (qty + unit).strip()
            up = _fmt_num(it.get("unit_price"))
            tp = _fmt_num(it.get("total_price"))
            ch = str(it.get("chapter") or "").strip()
            dwg = str(it.get("drawing_locator") or "").strip()
            std = str(it.get("standard_locator") or "").strip()
            clo = it.get("closure") if isinstance(it.get("closure"), dict) else {}
            clo_ok = bool(clo.get("ok"))
            missing = clo.get("missing_parts") or []
            if isinstance(missing, list):
                missing = [str(x) for x in missing if str(x).strip()]
            else:
                missing = []
            clo_disp = "OK" if clo_ok else ("缺:" + ",".join(missing) if missing else "缺")

            row = table.add_row().cells
            vals = [name, cat_proc, qty_disp, up, tp, ch, dwg, std, clo_disp]
            for i, v in enumerate(vals):
                try:
                    row[i].text = str(v or "")
                except Exception:
                    pass
        _style_professional_table(table, style_cfg)

    # 可编辑参数影响回执（参数键 -> 出现位置/影响章节）
    param_trace = data.get("param_trace") or {}
    receipt = param_trace.get("receipt") if isinstance(param_trace, dict) else None
    if internal_review and isinstance(receipt, dict) and isinstance(receipt.get("keys"), dict) and receipt.get("keys"):
        doc.add_page_break()
        hd = doc.add_heading("可编辑参数影响回执（自动生成）", level=1)
        apply_paragraph(hd, is_title=True)
        ver = str(receipt.get("version") or "").strip()
        keys = receipt.get("keys") or {}
        try:
            impacted = set()
            for _, item in keys.items():
                for t in (item or {}).get("impacted_chapters") or []:
                    if str(t).strip():
                        impacted.add(str(t).strip())
            impacted_count = len(impacted)
        except Exception:
            impacted_count = 0
        p0 = doc.add_paragraph(
            f"参数版本={ver or '-'}；参数键={len(keys)}；影响章节数={impacted_count}。"
        )
        apply_paragraph(p0)

        table = doc.add_table(rows=1, cols=3)
        _mark_table_header_row(table)
        hdr = table.rows[0].cells
        hdr[0].text = "参数键"
        hdr[1].text = "当前值"
        hdr[2].text = "影响章节"

        # Keep compact to reduce DOCX size. Keys are already limited (quant_defaults + boq_focus_card).
        for k in sorted(keys.keys())[:60]:
            item = keys.get(k) or {}
            val = str(item.get("value") or "").strip()
            chs = item.get("impacted_chapters") or []
            if isinstance(chs, list):
                chs = "；".join([str(x).strip() for x in chs if str(x).strip()][:10])
            else:
                chs = str(chs) if chs else ""
            row = table.add_row().cells
            row[0].text = str(k)
            row[1].text = val
            row[2].text = chs
        _style_professional_table(table, style_cfg)

    # Remaining chart/images (default: append at end, or chapter mode leftover).
    remaining_media = (
        []
        if chart_mode_auto_density
        else [item for item in media_all if _media_identity(item) not in consumed_media]
    )
    if remaining_media:
        # add_section(NEW_PAGE) already moved the cursor to a fresh portrait
        # page after a trailing landscape table.  Adding another page break at
        # that exact point creates a branded but otherwise empty page.  Public
        # exports have no intervening annexes, so the figure heading can safely
        # occupy the restored page.  Internal-review annexes do intervene and
        # retain the explicit page break below.
        if not (tables_restored_portrait and not internal_review):
            doc.add_page_break()
        hm = doc.add_heading("图表与插图", level=1)
        apply_paragraph(hm, is_title=True)
        for item in remaining_media:
            if _append_media_item(item, chapter_title="图表与插图"):
                consumed_media.add(_media_identity(item))

    # 章节版式回执
    if internal_review and layout_receipts:
        doc.add_page_break()
        hl = doc.add_heading("章节版式约束回执", level=1)
        apply_paragraph(hl, is_title=True)
        for item in layout_receipts:
            title = item["title"]
            target = item["target_pages"]
            estimated = item["estimated_pages"]
            delta = item["delta"]
            if delta == 0:
                status = "达成"
            elif delta > 0:
                status = f"超出{abs(delta)}页"
            else:
                status = f"不足{abs(delta)}页"
            p = doc.add_paragraph(f"- {title}: 目标{target}页，估算{estimated}页（{status}）")
            apply_paragraph(p)

    if internal_review:
        _append_quality_evidence_appendix(
            doc,
            apply_paragraph,
            data.get("quality_checks"),
            sections=sections,
            compare_cfg=data.get("compare") or {},
        )

    _enable_field_updates(doc)
    _save_docx_secure(doc, output_path)
    embedded_media_verification = verify_docx_media_hashes(output_path, inserted_media_hashes)
    source_media_summary = {
            "accepted_count": int(media_quality.get("accepted_count") or 0),
            "rejected_count": int(media_quality.get("rejected_count") or 0),
            "rejected": [
                {
                    "caption": str((entry.get("item") or {}).get("caption") or ""),
                    "reason": list(entry.get("reason") or []),
                    "required": bool((entry.get("item") or {}).get("required")),
                }
                for entry in media_quality.get("rejected") or []
            ],
    }
    figure_quality_report = build_media_delivery_manifest(
        source_media=source_media_summary,
        insertions=figure_insertions,
        insertion_failures=figure_failures,
        embedded_media_verification=embedded_media_verification,
    )
    figure_manifest_path = Path(output_path).with_suffix(".figure_manifest.json")
    figure_manifest_path.write_text(
        json.dumps(figure_quality_report, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    if not figure_quality_report.get("delivery_allowed"):
        raise RuntimeError(
            json.dumps(
                {
                    "status": "blocked",
                    "export_allowed": False,
                    "export_kind": "docx_figure_delivery_manifest",
                    "decision_digest": figure_quality_report.get("decision_digest"),
                    "issues": figure_quality_report.get("issues") or [],
                },
                ensure_ascii=False,
            )
        )
    structural_quality_report = audit_docx_structural_quality(
        output_path,
        expected_style=style_cfg,
        figure_manifest=figure_quality_report,
        require_heading_structure=bool(sections),
        strict=True,
    )
    _write_docx_build_report(
        output_path,
        topic=str(topic),
        sections=sections,
        front_matter_plan=front_matter_plan,
        layout_receipts=layout_receipts,
        media_count=media_index,
        quality_checks=data.get("quality_checks"),
        figure_quality=figure_quality_report,
        structural_quality=structural_quality_report,
    )
    return output_path


def export_autoplan_compare_docx(data: Dict[str, Any], output_path: str) -> str:
    data = _sanitize_docx_payload(data)
    _require_local_adapter_export_allowed(data, "compare_docx")
    style_raw = data.get("style") or {}
    style_cfg = _normalize_style(style_raw)
    doc = Document()
    _apply_page_setup(doc, style_cfg)
    apply_paragraph = _apply_style(doc, style_raw)

    h = doc.add_heading((data.get("topic") or "施组方案") + " - 整改对比", level=1)
    apply_paragraph(h, is_title=True)

    compare_cfg = data.get("compare") or {}
    mode = compare_cfg.get("mode", "summary")
    max_chars = _to_int(compare_cfg.get("max_chars"), 800)
    titles_filter = compare_cfg.get("titles")

    has_any = False
    for sec in data.get("sections", []):
        if sec.get("auto_remediated") != "llm" or not sec.get("original_content"):
            continue
        if isinstance(titles_filter, list) and sec.get("title") not in titles_filter:
            continue
        has_any = True
        title = sec.get("title") or "章节"
        h2 = doc.add_heading(title, level=2)
        apply_paragraph(h2, is_title=True)
        before = sec.get("original_content") or ""
        after = sec.get("content") or ""
        if mode == "summary":
            before = before[:max_chars] + ("..." if len(before) > max_chars else "")
            after = after[:max_chars] + ("..." if len(after) > max_chars else "")
        p1 = doc.add_paragraph("整改前：")
        apply_paragraph(p1)
        p2 = doc.add_paragraph(before)
        apply_paragraph(p2)
        p3 = doc.add_paragraph("整改后：")
        apply_paragraph(p3)
        p4 = doc.add_paragraph(after)
        apply_paragraph(p4)

    if not has_any:
        p = doc.add_paragraph("暂无可对比的章节（需使用 LLM 整改且保留 original_content）。")
        apply_paragraph(p)

    _enable_field_updates(doc)
    _save_docx_secure(doc, output_path)
    return output_path


def export_autoplan_focus_xlsx(data: Dict[str, Any], output_path: str) -> str:
    _require_local_adapter_export_allowed(data, "focus_xlsx")
    """
    Export a reviewer-friendly XLSX for BoQ focus closure:
    - focus items cross-index (chapter + drawing/std locator + closure gaps)
    - issue list + auto revision suggestions (for quick triage)
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except Exception:
        return ""

    cross_index = data.get("cross_index") if isinstance(data.get("cross_index"), dict) else {}
    focus_items = cross_index.get("focus_items") if isinstance(cross_index.get("focus_items"), list) else []
    qc = data.get("quality_checks") if isinstance(data.get("quality_checks"), dict) else {}
    issues = qc.get("issue_list") if isinstance(qc.get("issue_list"), list) else []
    recs = qc.get("auto_revision_suggestions") if isinstance(qc.get("auto_revision_suggestions"), list) else []
    param_trace = data.get("param_trace") if isinstance(data.get("param_trace"), dict) else {}
    plan_consistency = data.get("plan_consistency") if isinstance(data.get("plan_consistency"), dict) else {}
    cpm_receipt = plan_consistency.get("cpm") if isinstance(plan_consistency.get("cpm"), dict) else {}
    boq_focus = data.get("boq_focus") if isinstance(data.get("boq_focus"), dict) else {}
    four_new_recs = boq_focus.get("four_new_recommendations") if isinstance(boq_focus.get("four_new_recommendations"), list) else []
    variant_similarity = data.get("variant_similarity") if isinstance(data.get("variant_similarity"), dict) else {}
    drawing_index = data.get("drawing_index") if isinstance(data.get("drawing_index"), dict) else {}
    standard_index = data.get("standard_index") if isinstance(data.get("standard_index"), dict) else {}
    branding = data.get("branding") if isinstance(data.get("branding"), dict) else {}

    drawings = drawing_index.get("drawings") if isinstance(drawing_index.get("drawings"), list) else []
    drawing_binds = drawing_index.get("chapter_bindings") if isinstance(drawing_index.get("chapter_bindings"), list) else []
    standards = standard_index.get("standards") if isinstance(standard_index.get("standards"), list) else []
    standard_binds = standard_index.get("chapter_bindings") if isinstance(standard_index.get("chapter_bindings"), list) else []
    has_drawing_index = bool(drawings or drawing_binds)
    has_standard_index = bool(standards or standard_binds)

    # Parse structured focus cards from section text so reviewers can verify quant + triplet + typed evidence.
    focus_cards = []
    try:
        from backend.zhifei_autoplan.focus_card_parser import extract_focus_cards

        focus_cards = extract_focus_cards(data.get("sections") or [])
    except Exception:
        focus_cards = []

    # Param trace may exist even when focus items are empty.
    receipt = param_trace.get("receipt") if isinstance(param_trace.get("receipt"), dict) else {}
    receipt_keys = receipt.get("keys") if isinstance(receipt.get("keys"), dict) else {}

    has_plan_consistency = bool(plan_consistency and (plan_consistency.get("canonical") or plan_consistency.get("changed")))
    has_cpm_receipt = bool(cpm_receipt and (cpm_receipt.get("computed") or cpm_receipt.get("activities")))
    has_variant_similarity = bool(variant_similarity and int(variant_similarity.get("variant_count") or 0) >= 2)

    # Sync variant-similarity findings into issues/revision sheets so reviewers don't need JSON.
    issues = list(issues) if isinstance(issues, list) else []
    recs = list(recs) if isinstance(recs, list) else []
    if has_variant_similarity:
        try:
            chapter_thr = float(variant_similarity.get("chapter_threshold") or 0.90)
        except Exception:
            chapter_thr = 0.90
        strict_flagged = variant_similarity.get("flagged") if isinstance(variant_similarity.get("flagged"), list) else []
        if strict_flagged and (variant_similarity.get("ok") is False):
            for f in strict_flagged[:24]:
                if not isinstance(f, dict):
                    continue
                title = str(f.get("title") or "").strip() or "章节"
                pair = str(f.get("pair") or "").strip() or "pair"
                sim = f.get("similarity")
                thr = f.get("threshold") if f.get("threshold") is not None else chapter_thr
                problem = f"多方案相似度过高：{pair}={sim}（阈值={thr}）。"
                suggestion = (
                    "不改招标目录，重写本章章内逻辑：A=交付物/约束/步骤/闭环；"
                    "B=工序流程/控制点表/资源节拍；C=指标矩阵/人机料法环/闭环分组；"
                    "用短句卡片表达，避免长段复述。"
                )
                issues.append(
                    {
                        "severity": "high",
                        "title": title,
                        "type": "variant_diversity_gap",
                        "problem": problem,
                        "suggestion": suggestion,
                    }
                )
                recs.append({"title": title, "type": "variant_diversity_gap", "suggestion": suggestion})

        relaxed_flagged = (
            variant_similarity.get("relaxed_flagged")
            if isinstance(variant_similarity.get("relaxed_flagged"), list)
            else []
        )
        # Informational: relaxed chapters are excluded from gate, but still shown for reviewers.
        for f in relaxed_flagged[:12]:
            if not isinstance(f, dict):
                continue
            title = str(f.get("title") or "").strip() or "章节"
            pair = str(f.get("pair") or "").strip() or "pair"
            sim = f.get("similarity")
            thr = f.get("threshold")
            problem = f"多方案相似度偏高（白名单/放宽章节）：{pair}={sim}（放宽阈值={thr}）。"
            suggestion = "可选优化：用“表/矩阵/闭环卡片”替代叙述段，保留事实数据与证据定位。"
            issues.append(
                {
                    "severity": "low",
                    "title": title,
                    "type": "variant_diversity_relaxed_note",
                    "problem": problem,
                    "suggestion": suggestion,
                }
            )

    # If chapters were auto-fixed for diversity, record as low-severity reviewer notes.
    try:
        secs = data.get("sections") if isinstance(data.get("sections"), list) else []
        notes = []
        for s in secs:
            if not isinstance(s, dict):
                continue
            if str(s.get("auto_remediated") or "") != "diversity_autofix":
                continue
            title = str(s.get("title") or "").strip()
            if not title:
                continue
            tid = str(s.get("logic_template_id") or "").strip().upper()
            dom = str(s.get("chapter_domain") or "").strip().lower()
            notes.append((title, tid, dom))
        seen = set()
        for title, tid, dom in notes[:24]:
            key = (title, tid, dom)
            if key in seen:
                continue
            seen.add(key)
            issues.append(
                {
                    "severity": "low",
                    "title": title,
                    "type": "variant_diversity_autofixed",
                    "problem": f"本章已执行多方案差异化结构重排（模版={tid or '-'}；domain={dom or '-'}）。",
                    "suggestion": "复核：清单重点项/量化指标/风险→控制→验证闭环/证据定位是否仍匹配本章内容。",
                }
            )
    except Exception:
        pass

    # Deduplicate (keeps sheets readable).
    try:
        dedup = {}
        for it in issues:
            if not isinstance(it, dict):
                continue
            key = (it.get("title"), it.get("type"), it.get("problem"))
            dedup[key] = it
        issues = list(dedup.values())
    except Exception:
        pass
    try:
        dedup2 = {}
        for it in recs:
            if not isinstance(it, dict):
                continue
            key = (it.get("title"), it.get("type"), it.get("suggestion"))
            dedup2[key] = it
        recs = list(dedup2.values())
    except Exception:
        pass

    # If nothing to export, do not create a file.
    if not (
        focus_items
        or issues
        or recs
        or focus_cards
        or four_new_recs
        or receipt_keys
        or has_plan_consistency
        or has_cpm_receipt
        or has_variant_similarity
        or has_drawing_index
        or has_standard_index
    ):
        return ""

    def _fmt_num(v: Any) -> Any:
        try:
            if v is None:
                return ""
            f = float(v)
            if abs(f - int(f)) < 1e-9:
                return int(f)
            return f
        except Exception:
            return str(v) if v is not None else ""

    def _join_list(v: Any, sep: str = "；", limit: int = 10) -> str:
        if isinstance(v, list):
            parts = [str(x).strip() for x in v if str(x).strip()]
            return sep.join(parts[: max(0, int(limit or 0))])
        return str(v).strip() if isinstance(v, str) else ""

    def _clean_val(v: Any) -> str:
        s = str(v or "").strip()
        # Focus-card values often end with '。' from sentence punctuation.
        return s.rstrip("。.;；")

    cards_by_key = {}
    cards_by_name = {}
    try:
        import re as _re

        def _norm_name(s: Any) -> str:
            t = str(s or "").strip()
            t = _re.sub(r"[\\s\\u3000]+", "", t)
            t = t.strip("，,。．.；;：:、")
            return t

        def _pick_better(a: dict | None, b: dict) -> dict:
            if not a:
                return b
            ae = len(a.get("evidence_sources") or [])
            be = len(b.get("evidence_sources") or [])
            if be != ae:
                return b if be > ae else a
            ar = len(str(a.get("raw") or ""))
            br = len(str(b.get("raw") or ""))
            return b if br > ar else a

        for c in focus_cards or []:
            if not isinstance(c, dict):
                continue
            ch = str(c.get("chapter") or "").strip()
            name = _norm_name(c.get("name"))
            if not name:
                continue
            key = (ch, name)
            cards_by_key[key] = _pick_better(cards_by_key.get(key), c)
            cards_by_name[name] = _pick_better(cards_by_name.get(name), c)
    except Exception:
        cards_by_key = {}
        cards_by_name = {}

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()

    hdr_font = Font(bold=True)
    hdr_fill = PatternFill("solid", fgColor="F2F2F2")
    warn_fill = PatternFill("solid", fgColor="FFF2CC")
    wrap = Alignment(wrap_text=True, vertical="top")

    def _write(ws, r: int, c: int, v: Any, header: bool = False):
        cell = ws.cell(row=r, column=c, value=v)
        cell.alignment = wrap
        if header:
            cell.font = hdr_font
            cell.fill = hdr_fill
        return cell

    def _set_width(ws, col: int, width: float):
        try:
            ws.column_dimensions[get_column_letter(int(col))].width = float(width)
        except Exception:
            pass

    # Summary sheet
    ws = wb.active
    ws.title = "summary"
    row = 1
    logic = data.get("logic_template") if isinstance(data.get("logic_template"), dict) else {}
    logic_id = str(logic.get("id") or "").strip()
    logic_name = str(logic.get("name") or "").strip()
    summary_pairs = [
        ("topic", data.get("topic") or ""),
        ("project_id", (cross_index.get("project_id") or data.get("project_id") or "")),
        ("logic_template_id", logic_id),
        ("logic_template_name", logic_name),
        ("bidder_company", str(branding.get("bidder_company") or "")),
        ("bidder_domain", str(branding.get("bidder_domain") or "")),
        ("logo_path", str(branding.get("logo_path") or "")),
        ("focus_count", cross_index.get("focus_count") or 0),
        ("mentioned_count", cross_index.get("mentioned_count") or 0),
        ("closed_ok_count", cross_index.get("closed_ok_count") or 0),
        ("missing_drawing_locator_count", cross_index.get("missing_drawing_locator_count") or 0),
        ("missing_standard_locator_count", cross_index.get("missing_standard_locator_count") or 0),
        ("focus_cards_count", len(focus_cards or [])),
        ("four_new_recommendations_count", len(four_new_recs or [])),
        ("drawing_files_count", len(drawings or [])),
        ("drawing_bindings_count", len(drawing_binds or [])),
        ("standard_files_count", len(standards or [])),
        ("standard_bindings_count", len(standard_binds or [])),
    ]
    if has_variant_similarity:
        summary_pairs.extend(
            [
                ("variant_similarity_ok", bool(variant_similarity.get("ok"))),
                ("variant_similarity_avg_max_strict", variant_similarity.get("avg_max_similarity")),
                ("variant_similarity_avg_max_all", variant_similarity.get("avg_max_similarity_all")),
                ("variant_similarity_flagged_count", variant_similarity.get("flagged_count")),
                ("variant_similarity_relaxed_flagged_count", variant_similarity.get("relaxed_flagged_count")),
            ]
        )
    if has_cpm_receipt:
        computed = cpm_receipt.get("computed") if isinstance(cpm_receipt.get("computed"), dict) else {}
        cpm_conflicts = cpm_receipt.get("conflicts") if isinstance(cpm_receipt.get("conflicts"), list) else []
        summary_pairs.extend(
            [
                ("cpm_ok", bool(cpm_receipt.get("ok"))),
                ("cpm_algorithm", str(cpm_receipt.get("algorithm") or "")),
                ("cpm_duration_days", computed.get("project_duration_days")),
                ("cpm_resource_peak", computed.get("resource_peak")),
                ("cpm_critical_interval_days", computed.get("critical_interval_days")),
                ("cpm_conflict_count", len(cpm_conflicts)),
            ]
        )

    for k, v in summary_pairs:
        _write(ws, row, 1, str(k), header=True)
        _write(ws, row, 2, v if not isinstance(v, (dict, list)) else _join_list(v))
        row += 1
    _set_width(ws, 1, 28)
    _set_width(ws, 2, 80)

    # Variant similarity (cross-variant anti-paraphrase report)
    if has_variant_similarity:
        ws = wb.create_sheet("variant_similarity")
        row = 1
        for k, v in [
            ("ok", bool(variant_similarity.get("ok"))),
            ("variant_count", int(variant_similarity.get("variant_count") or 0)),
            ("avg_max_similarity_strict", variant_similarity.get("avg_max_similarity")),
            ("avg_max_similarity_all", variant_similarity.get("avg_max_similarity_all")),
            ("strict_chapter_count", variant_similarity.get("strict_chapter_count")),
            ("relaxed_chapter_count", variant_similarity.get("relaxed_chapter_count")),
            ("chapter_threshold", variant_similarity.get("chapter_threshold")),
            ("relaxed_chapter_threshold", variant_similarity.get("relaxed_chapter_threshold")),
            ("overall_threshold", variant_similarity.get("overall_threshold")),
            ("min_chars", variant_similarity.get("min_chars")),
            ("flagged_count", variant_similarity.get("flagged_count")),
            ("relaxed_flagged_count", variant_similarity.get("relaxed_flagged_count")),
        ]:
            _write(ws, row, 1, str(k), header=True)
            _write(ws, row, 2, v)
            row += 1
        row += 1

        vcount = int(variant_similarity.get("variant_count") or 0)
        pairs = []
        for i in range(1, vcount + 1):
            for j in range(i + 1, vcount + 1):
                pairs.append((i, j))

        headers = ["title", "relaxed", "threshold", "max_pair", "max_combined"]
        for i in range(1, vcount + 1):
            headers.append(f"len_v{i}")
        for i, j in pairs:
            headers.extend([f"v{i}_v{j}_combined", f"v{i}_v{j}_jaccard3", f"v{i}_v{j}_cosine2"])
        for c, h in enumerate(headers, start=1):
            _write(ws, row, c, h, header=True)
        row += 1

        chapter_thr = float(variant_similarity.get("chapter_threshold") or 0.90)
        for rec in (variant_similarity.get("by_chapter") or [])[:300]:
            if not isinstance(rec, dict):
                continue
            title = str(rec.get("title") or "")
            is_relaxed = bool(rec.get("relaxed"))
            thr_used = rec.get("threshold")
            max_pair = str(rec.get("max_pair") or "")
            max_combined = rec.get("max_combined")
            lens = rec.get("lens") if isinstance(rec.get("lens"), list) else []
            values = [title, "Y" if is_relaxed else "N", _fmt_num(thr_used), max_pair, _fmt_num(max_combined)]
            for i in range(1, vcount + 1):
                values.append(_fmt_num(lens[i - 1] if i - 1 < len(lens) else ""))
            for i, j in pairs:
                key = f"v{i}_v{j}"
                ps = rec.get(key) if isinstance(rec.get(key), dict) else {}
                values.extend([_fmt_num(ps.get("combined")), _fmt_num(ps.get("jaccard3")), _fmt_num(ps.get("cosine2"))])
            for c, v in enumerate(values, start=1):
                cell = _write(ws, row, c, v)
                try:
                    # Highlight only strict chapters; relaxed chapters are informational.
                    thr = float(thr_used or chapter_thr)
                    if (not is_relaxed) and float(max_combined or 0.0) >= thr:
                        cell.fill = warn_fill
                except Exception:
                    pass
            row += 1

        _set_width(ws, 1, 22)
        _set_width(ws, 2, 10)
        _set_width(ws, 3, 12)
        _set_width(ws, 4, 10)
        _set_width(ws, 5, 12)
        base = 6
        for i in range(1, vcount + 1):
            _set_width(ws, base + i - 1, 10)
        # Pair columns
        for c in range(base + vcount, base + vcount + (len(pairs) * 3)):
            _set_width(ws, c, 14)

    # Drawings index (files + chapter bindings)
    if has_drawing_index:
        if drawings:
            ws = wb.create_sheet("drawings")
            headers = [
                "filename",
                "sha256",
                "pages",
                "keywords",
                "topology_nodes",
                "topology_edges",
                "topology_components",
                "topology_endpoints",
                "topology_trunk_length",
                "topology_suggested_flow_segments",
                "topology_confidence",
                "preview",
            ]
            for c, h in enumerate(headers, start=1):
                _write(ws, 1, c, h, header=True)
            for r, d in enumerate(drawings[:300], start=2):
                if not isinstance(d, dict):
                    continue
                kws = _join_list(d.get("keywords") or [], sep="、", limit=12)
                topo = d.get("topology") if isinstance(d.get("topology"), dict) else {}
                vals = [
                    str(d.get("filename") or ""),
                    str(d.get("sha256") or "")[:12],
                    _fmt_num(d.get("pages")),
                    kws,
                    _fmt_num(topo.get("nodes_count")),
                    _fmt_num(topo.get("edges_count")),
                    _fmt_num(topo.get("components_count")),
                    _fmt_num(topo.get("endpoint_count")),
                    _fmt_num(topo.get("trunk_length")),
                    _fmt_num(topo.get("suggested_flow_segments")),
                    str(topo.get("topology_confidence") or ""),
                    str(d.get("preview") or ""),
                ]
                for c, v in enumerate(vals, start=1):
                    _write(ws, r, c, v)
            _set_width(ws, 1, 54)
            _set_width(ws, 2, 18)
            _set_width(ws, 3, 10)
            _set_width(ws, 4, 48)
            _set_width(ws, 5, 12)
            _set_width(ws, 6, 12)
            _set_width(ws, 7, 12)
            _set_width(ws, 8, 12)
            _set_width(ws, 9, 14)
            _set_width(ws, 10, 16)
            _set_width(ws, 11, 14)
            _set_width(ws, 12, 78)

        if drawing_binds:
            ws = wb.create_sheet("drawing_bindings")
            headers = ["chapter", "locator", "filename", "page", "offset", "snippet"]
            for c, h in enumerate(headers, start=1):
                _write(ws, 1, c, h, header=True)
            for r, b in enumerate(drawing_binds[:600], start=2):
                if not isinstance(b, dict):
                    continue
                vals = [
                    str(b.get("chapter") or ""),
                    str(b.get("locator") or ""),
                    str(b.get("filename") or ""),
                    _fmt_num(b.get("page")),
                    _fmt_num(b.get("offset")),
                    str(b.get("snippet") or ""),
                ]
                for c, v in enumerate(vals, start=1):
                    _write(ws, r, c, v)
            _set_width(ws, 1, 28)
            _set_width(ws, 2, 46)
            _set_width(ws, 3, 54)
            _set_width(ws, 4, 10)
            _set_width(ws, 5, 12)
            _set_width(ws, 6, 90)

    # Enterprise standards index (files + chapter bindings)
    if has_standard_index:
        if standards:
            ws = wb.create_sheet("standards")
            headers = ["filename", "sha256", "pages", "keywords"]
            for c, h in enumerate(headers, start=1):
                _write(ws, 1, c, h, header=True)
            for r, d in enumerate(standards[:400], start=2):
                if not isinstance(d, dict):
                    continue
                kws = _join_list(d.get("keywords") or [], sep="、", limit=12)
                vals = [
                    str(d.get("filename") or ""),
                    str(d.get("sha256") or "")[:12],
                    _fmt_num(d.get("pages")),
                    kws,
                ]
                for c, v in enumerate(vals, start=1):
                    _write(ws, r, c, v)
            _set_width(ws, 1, 58)
            _set_width(ws, 2, 18)
            _set_width(ws, 3, 10)
            _set_width(ws, 4, 80)

        if standard_binds:
            ws = wb.create_sheet("standard_bindings")
            headers = ["chapter", "locator", "filename", "page", "offset", "snippet"]
            for c, h in enumerate(headers, start=1):
                _write(ws, 1, c, h, header=True)
            for r, b in enumerate(standard_binds[:800], start=2):
                if not isinstance(b, dict):
                    continue
                vals = [
                    str(b.get("chapter") or ""),
                    str(b.get("locator") or ""),
                    str(b.get("filename") or ""),
                    _fmt_num(b.get("page")),
                    _fmt_num(b.get("offset")),
                    str(b.get("snippet") or ""),
                ]
                for c, v in enumerate(vals, start=1):
                    _write(ws, r, c, v)
            _set_width(ws, 1, 28)
            _set_width(ws, 2, 46)
            _set_width(ws, 3, 58)
            _set_width(ws, 4, 10)
            _set_width(ws, 5, 12)
            _set_width(ws, 6, 90)

    # Param trace (editable parameter -> impacted chapters)
    if receipt_keys:
        ws = wb.create_sheet("param_trace")
        headers = [
            "param_key",
            "value",
            "impacted_chapters",
            "placeholder_hits",
            "value_hits",
            "placeholder_occurrences",
            "value_occurrences",
        ]
        for c, h in enumerate(headers, start=1):
            _write(ws, 1, c, h, header=True)

        def _occ_list(arr: Any, limit: int = 24) -> str:
            if not isinstance(arr, list):
                return ""
            out = []
            for it in arr[: max(1, int(limit or 0))]:
                if not isinstance(it, dict):
                    continue
                t = str(it.get("title") or "").strip()
                off = it.get("offset")
                if t and off is not None:
                    out.append(f"{t}@{off}")
            return "；".join(out)

        for r, (k, item) in enumerate(sorted(receipt_keys.items(), key=lambda x: x[0]), start=2):
            if not isinstance(item, dict):
                continue
            impacted = _join_list(item.get("impacted_chapters") or [], sep="；", limit=40)
            ph = item.get("placeholder_occurrences") or []
            vh = item.get("value_occurrences") or []
            values = [
                str(k),
                str(item.get("value") or ""),
                impacted,
                len(ph) if isinstance(ph, list) else 0,
                len(vh) if isinstance(vh, list) else 0,
                _occ_list(ph),
                _occ_list(vh),
            ]
            for c, v in enumerate(values, start=1):
                _write(ws, r, c, v)
        _set_width(ws, 1, 32)
        _set_width(ws, 2, 28)
        _set_width(ws, 3, 60)
        _set_width(ws, 4, 14)
        _set_width(ws, 5, 14)
        _set_width(ws, 6, 80)
        _set_width(ws, 7, 80)

    # Plan consistency receipt (工期/资源峰值/关键线路间隔)
    if plan_consistency and (plan_consistency.get("canonical") or plan_consistency.get("changed")):
        ws = wb.create_sheet("plan_consistency")
        _write(ws, 1, 1, "metric", header=True)
        _write(ws, 1, 2, "canonical_value", header=True)
        row = 2
        can = plan_consistency.get("canonical") if isinstance(plan_consistency.get("canonical"), dict) else {}
        for k, v in (can or {}).items():
            _write(ws, row, 1, str(k))
            _write(ws, row, 2, str(v))
            row += 1
        row += 1
        _write(ws, row, 1, "changed_chapters", header=True)
        changed = plan_consistency.get("changed") if isinstance(plan_consistency.get("changed"), list) else []
        ch_titles = [str(it.get("title") or "").strip() for it in changed if isinstance(it, dict) and str(it.get("title") or "").strip()]
        _write(ws, row, 2, "；".join(ch_titles[:60]))
        _set_width(ws, 1, 24)
        _set_width(ws, 2, 90)

    # CPM deterministic schedule receipt (NetworkX DAG + critical path)
    if has_cpm_receipt:
        ws = wb.create_sheet("cpm")
        row = 1
        computed = cpm_receipt.get("computed") if isinstance(cpm_receipt.get("computed"), dict) else {}
        mentioned = cpm_receipt.get("mentioned") if isinstance(cpm_receipt.get("mentioned"), dict) else {}
        graph = cpm_receipt.get("graph") if isinstance(cpm_receipt.get("graph"), dict) else {}
        critical_path = cpm_receipt.get("critical_path") if isinstance(cpm_receipt.get("critical_path"), list) else []
        conflict_rows = cpm_receipt.get("conflicts") if isinstance(cpm_receipt.get("conflicts"), list) else []
        for k, v in [
            ("ok", bool(cpm_receipt.get("ok"))),
            ("algorithm", str(cpm_receipt.get("algorithm") or "")),
            ("mentioned_工期", mentioned.get("工期")),
            ("mentioned_资源峰值", mentioned.get("资源峰值")),
            ("mentioned_关键线路间隔", mentioned.get("关键线路间隔")),
            ("computed_project_duration_days", computed.get("project_duration_days")),
            ("computed_resource_peak", computed.get("resource_peak")),
            ("computed_critical_interval_days", computed.get("critical_interval_days")),
            ("critical_path", " -> ".join([str(x) for x in critical_path[:40]])),
            ("graph_node_count", graph.get("node_count")),
            ("graph_edge_count", graph.get("edge_count")),
            ("cycle_edges_removed", _join_list(graph.get("cycle_edges_removed") or [], sep="；", limit=20)),
            ("conflict_count", len(conflict_rows)),
        ]:
            _write(ws, row, 1, str(k), header=True)
            _write(ws, row, 2, v if not isinstance(v, (dict, list)) else _join_list(v))
            row += 1
        row += 1
        _write(ws, row, 1, "metric", header=True)
        _write(ws, row, 2, "mentioned", header=True)
        _write(ws, row, 3, "computed", header=True)
        _write(ws, row, 4, "tolerance", header=True)
        _write(ws, row, 5, "delta", header=True)
        row += 1
        for c in conflict_rows[:60]:
            if not isinstance(c, dict):
                continue
            _write(ws, row, 1, str(c.get("metric") or ""))
            _write(ws, row, 2, _fmt_num(c.get("mentioned")))
            _write(ws, row, 3, _fmt_num(c.get("computed")))
            _write(ws, row, 4, _fmt_num(c.get("tolerance")))
            _write(ws, row, 5, _fmt_num(c.get("delta")))
            row += 1
        row += 1
        _write(ws, row, 1, "activity_id", header=True)
        _write(ws, row, 2, "name", header=True)
        _write(ws, row, 3, "deps", header=True)
        _write(ws, row, 4, "duration_days", header=True)
        _write(ws, row, 5, "resource_units", header=True)
        _write(ws, row, 6, "ES", header=True)
        _write(ws, row, 7, "EF", header=True)
        _write(ws, row, 8, "LS", header=True)
        _write(ws, row, 9, "LF", header=True)
        _write(ws, row, 10, "total_float", header=True)
        _write(ws, row, 11, "critical", header=True)
        row += 1
        for act in (cpm_receipt.get("activities") or [])[:600]:
            if not isinstance(act, dict):
                continue
            vals = [
                str(act.get("id") or ""),
                str(act.get("name") or ""),
                _join_list(act.get("deps") or [], sep="、", limit=20),
                _fmt_num(act.get("duration_days")),
                _fmt_num(act.get("resource_units")),
                _fmt_num(act.get("es")),
                _fmt_num(act.get("ef")),
                _fmt_num(act.get("ls")),
                _fmt_num(act.get("lf")),
                _fmt_num(act.get("total_float")),
                "Y" if act.get("critical") else "N",
            ]
            for c, v in enumerate(vals, start=1):
                _write(ws, row, c, v)
            row += 1
        _set_width(ws, 1, 18)
        _set_width(ws, 2, 28)
        _set_width(ws, 3, 22)
        _set_width(ws, 4, 12)
        _set_width(ws, 5, 12)
        _set_width(ws, 6, 10)
        _set_width(ws, 7, 10)
        _set_width(ws, 8, 10)
        _set_width(ws, 9, 10)
        _set_width(ws, 10, 12)
        _set_width(ws, 11, 10)

    # Four-new recommendations (editable library + BoQ/process matching)
    if isinstance(four_new_recs, list) and four_new_recs:
        ws = wb.create_sheet("four_new")
        headers = ["category", "name", "project_type", "project_types", "score", "matched_keywords", "trades", "roles", "acceptance_preview"]
        for c, h in enumerate(headers, start=1):
            _write(ws, 1, c, h, header=True)
        for r, it in enumerate(four_new_recs[:24], start=2):
            if not isinstance(it, dict):
                continue
            cat = str(it.get("category") or "").strip()
            name = str(it.get("name") or "").strip()
            pt = str(it.get("project_type") or "").strip()
            pts = _join_list(it.get("project_types") or [], sep="、", limit=8)
            score = it.get("score")
            matched = _join_list(it.get("matched") or [], sep="、", limit=12)
            trades = _join_list(it.get("trades") or [], sep="、", limit=12)
            roles = _join_list(it.get("roles") or [], sep="、", limit=12)
            acc = _join_list(it.get("acceptance") or [], sep="；", limit=2)
            vals = [cat, name, pt, pts, _fmt_num(score), matched, trades, roles, acc]
            for c, v in enumerate(vals, start=1):
                _write(ws, r, c, v)
        _set_width(ws, 1, 16)
        _set_width(ws, 2, 58)
        _set_width(ws, 3, 12)
        _set_width(ws, 4, 24)
        _set_width(ws, 5, 10)
        _set_width(ws, 6, 40)
        _set_width(ws, 7, 28)
        _set_width(ws, 8, 28)
        _set_width(ws, 9, 68)

    # Focus index
    if focus_items:
        ws = wb.create_sheet("focus_index")
        headers = [
            "清单项",
            "类别",
            "工序",
            "清单编码",
            "工程量",
            "单位",
            "单价",
            "合价",
            "落位章节",
            "图纸定位",
            "标准定位",
            "闭环OK",
            "缺口",
            "flags",
            "近邻证据",
            "卡-工程量",
            "卡-单价",
            "卡-合价",
            "卡-频次",
            "卡-阈值",
            "卡-间距",
            "卡-厚度",
            "卡-时长",
            "卡-人数",
            "卡-设备型号",
            "卡-图纸定位",
            "卡-标准引用",
            "卡-风险",
            "卡-控制",
            "卡-验证",
            "卡-证据来源",
        ]
        for c, h in enumerate(headers, start=1):
            _write(ws, 1, c, h, header=True)

        for r, it in enumerate(focus_items[:300], start=2):
            if not isinstance(it, dict):
                continue
            clo = it.get("closure") if isinstance(it.get("closure"), dict) else {}
            # Find the best-matching control card by (chapter, name), fallback to name-only.
            card = None
            try:
                name_norm = _norm_name(it.get("name"))
                ch_norm = str(it.get("chapter") or "").strip()
                card = cards_by_key.get((ch_norm, name_norm)) or cards_by_name.get(name_norm)
            except Exception:
                card = None
            head = card.get("head") if isinstance(card, dict) and isinstance(card.get("head"), dict) else {}
            quant = card.get("quant") if isinstance(card, dict) and isinstance(card.get("quant"), dict) else {}

            values = [
                str(it.get("name") or ""),
                _join_list(it.get("categories") or [], sep="、", limit=8),
                str(it.get("process_name") or ""),
                str(it.get("boq_code") or ""),
                _fmt_num(it.get("quantity")),
                str(it.get("unit") or ""),
                _fmt_num(it.get("unit_price")),
                _fmt_num(it.get("total_price")),
                str(it.get("chapter") or ""),
                str(it.get("drawing_locator") or ""),
                str(it.get("standard_locator") or ""),
                "OK" if clo.get("ok") else "NO",
                _join_list(clo.get("missing_parts") or [], sep=",", limit=8),
                _join_list(it.get("flags") or [], sep="；", limit=8),
                _join_list(it.get("evidence_locators_near") or [], sep="；", limit=8),
                _clean_val(head.get("工程量") or ""),
                _clean_val(head.get("单价") or ""),
                _clean_val(head.get("合价") or ""),
                _clean_val(quant.get("频次") or ""),
                _clean_val(quant.get("阈值") or ""),
                _clean_val(quant.get("间距") or ""),
                _clean_val(quant.get("厚度") or ""),
                _clean_val(quant.get("时长") or ""),
                _clean_val(quant.get("人数") or ""),
                _clean_val(quant.get("设备型号") or ""),
                _clean_val((card or {}).get("drawing_locator") or ""),
                _clean_val((card or {}).get("standard_locator") or ""),
                _clean_val((card or {}).get("risk") or ""),
                _clean_val((card or {}).get("control") or ""),
                _clean_val((card or {}).get("verify") or ""),
                _join_list((card or {}).get("evidence_sources") or [], sep="；", limit=10),
            ]
            for c, v in enumerate(values, start=1):
                _write(ws, r, c, v)

        # Column widths (best-effort)
        _set_width(ws, 1, 26)
        _set_width(ws, 2, 18)
        _set_width(ws, 3, 18)
        for c in range(4, 9):
            _set_width(ws, c, 12)
        _set_width(ws, 9, 28)
        _set_width(ws, 10, 44)
        _set_width(ws, 11, 44)
        _set_width(ws, 12, 12)
        _set_width(ws, 13, 12)
        _set_width(ws, 14, 34)
        _set_width(ws, 15, 34)
        for c in range(16, 19):
            _set_width(ws, c, 14)
        for c in range(19, 26):
            _set_width(ws, c, 18)
        _set_width(ws, 26, 44)
        _set_width(ws, 27, 44)
        for c in range(28, 32):
            _set_width(ws, c, 44)

    # Focus cards (raw structured export)
    if focus_cards:
        ws = wb.create_sheet("focus_cards")
        headers = [
            "chapter",
            "name",
            "工程量",
            "单价",
            "合价",
            "图纸定位",
            "标准引用",
            "频次",
            "阈值",
            "间距",
            "厚度",
            "时长",
            "人数",
            "设备型号",
            "风险",
            "控制",
            "验证",
            "evidence_sources",
            "raw",
        ]
        for c, h in enumerate(headers, start=1):
            _write(ws, 1, c, h, header=True)
        for r, card in enumerate((focus_cards or [])[:600], start=2):
            if not isinstance(card, dict):
                continue
            head = card.get("head") if isinstance(card.get("head"), dict) else {}
            quant = card.get("quant") if isinstance(card.get("quant"), dict) else {}
            values = [
                str(card.get("chapter") or ""),
                str(card.get("name") or ""),
                _clean_val(head.get("工程量") or ""),
                _clean_val(head.get("单价") or ""),
                _clean_val(head.get("合价") or ""),
                _clean_val(card.get("drawing_locator") or ""),
                _clean_val(card.get("standard_locator") or ""),
                _clean_val(quant.get("频次") or ""),
                _clean_val(quant.get("阈值") or ""),
                _clean_val(quant.get("间距") or ""),
                _clean_val(quant.get("厚度") or ""),
                _clean_val(quant.get("时长") or ""),
                _clean_val(quant.get("人数") or ""),
                _clean_val(quant.get("设备型号") or ""),
                _clean_val(card.get("risk") or ""),
                _clean_val(card.get("control") or ""),
                _clean_val(card.get("verify") or ""),
                _join_list(card.get("evidence_sources") or [], sep="；", limit=12),
                str(card.get("raw") or ""),
            ]
            for c, v in enumerate(values, start=1):
                _write(ws, r, c, v)
        _set_width(ws, 1, 28)
        _set_width(ws, 2, 28)
        for c in range(3, 8):
            _set_width(ws, c, 22)
        for c in range(8, 15):
            _set_width(ws, c, 18)
        for c in range(15, 20):
            _set_width(ws, c, 70)

    # Issues
    if issues:
        ws = wb.create_sheet("issues")
        headers = ["severity", "title", "type", "problem", "suggestion"]
        for c, h in enumerate(headers, start=1):
            _write(ws, 1, c, h, header=True)
        for r, it in enumerate(issues[:800], start=2):
            if not isinstance(it, dict):
                continue
            values = [
                str(it.get("severity") or ""),
                str(it.get("title") or ""),
                str(it.get("type") or ""),
                str(it.get("problem") or ""),
                str(it.get("suggestion") or ""),
            ]
            for c, v in enumerate(values, start=1):
                _write(ws, r, c, v)
        _set_width(ws, 1, 18)
        _set_width(ws, 2, 22)
        _set_width(ws, 3, 18)
        _set_width(ws, 4, 90)
        _set_width(ws, 5, 90)

    # Auto revision suggestions (chapter-aggregated)
    if recs:
        ws = wb.create_sheet("auto_revision")
        headers = ["title", "type", "suggestion"]
        for c, h in enumerate(headers, start=1):
            _write(ws, 1, c, h, header=True)
        for r, it in enumerate(recs[:1000], start=2):
            if not isinstance(it, dict):
                continue
            values = [str(it.get("title") or ""), str(it.get("type") or ""), str(it.get("suggestion") or "")]
            for c, v in enumerate(values, start=1):
                _write(ws, r, c, v)
        _set_width(ws, 1, 22)
        _set_width(ws, 2, 22)
        _set_width(ws, 3, 110)

    wb.save(str(output_path))
    return str(output_path)


def export_autoplan_docx_from_file(json_path: str, output_path: str) -> str:
    data = json.loads(Path(json_path).read_text(encoding="utf-8"))
    # 若为多版本，默认导出第一个版本
    if isinstance(data, dict) and isinstance(data.get("variants"), list) and data["variants"]:
        return export_autoplan_docx(data["variants"][0], output_path)
    return export_autoplan_docx(data, output_path)


def export_scoring_evidence_overview_xlsx(data: Dict[str, Any], output_path: str) -> str:
    _require_local_adapter_export_allowed(data, "scoring_evidence_overview_xlsx")
    """
    评分点覆盖与证据引用总览
    - 评分点覆盖: score_mapping item_cards
    - 段落证据链: evidence_tracking rows
    - 评分点×证据矩阵: 便于评审定位“第几页用哪条证据响应了哪个评分点”
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except Exception:
        return ""

    evidence_tracking = data.get("evidence_tracking") if isinstance(data.get("evidence_tracking"), dict) else {}
    rows = evidence_tracking.get("rows") if isinstance(evidence_tracking.get("rows"), list) else []
    summary = evidence_tracking.get("summary") if isinstance(evidence_tracking.get("summary"), dict) else {}
    score_mapping = data.get("score_mapping") if isinstance(data.get("score_mapping"), dict) else {}
    item_cards = score_mapping.get("item_cards") if isinstance(score_mapping.get("item_cards"), list) else []
    high_risk_items = score_mapping.get("high_risk_items") if isinstance(score_mapping.get("high_risk_items"), list) else []

    def _s(v: Any) -> str:
        return str(v or "").strip()

    def _join_list(v: Any, sep: str = "；", limit: int = 20) -> str:
        if isinstance(v, list):
            out = [_s(x) for x in v if _s(x)]
            return sep.join(out[: max(1, int(limit or 1))])
        return _s(v)

    def _sheet_set_width(ws, col: int, width: float) -> None:
        try:
            ws.column_dimensions[get_column_letter(int(col))].width = float(width)
        except Exception:
            pass

    def _infer_src_type(src: str, typed: Dict[str, Any]) -> str:
        s = _s(src).lower()
        if not s:
            return ""
        graph_nodes = [_s(x) for x in (typed.get("graph_nodes") or []) if _s(x)]
        drawing_refs = [_s(x) for x in (typed.get("drawing_refs") or []) if _s(x)]
        standard_refs = [_s(x) for x in (typed.get("standard_refs") or []) if _s(x)]
        if src in drawing_refs:
            return "drawing_dxf"
        if src in standard_refs:
            return "standard"
        if src in graph_nodes or s.startswith("图谱节点:"):
            return "graph"
        return "other"

    wb = Workbook()
    hdr_font = Font(bold=True)
    hdr_fill = PatternFill("solid", fgColor="F2F2F2")
    wrap = Alignment(wrap_text=True, vertical="top")

    def _write(ws, r: int, c: int, v: Any, *, header: bool = False):
        cell = ws.cell(row=r, column=c, value=v)
        cell.alignment = wrap
        if header:
            cell.font = hdr_font
            cell.fill = hdr_fill
        return cell

    # 1) summary
    ws = wb.active
    ws.title = "summary"
    meta_rows = [
        ("topic", _s(data.get("topic"))),
        ("project_id", _s(data.get("project_id"))),
        ("paragraph_count", int(summary.get("paragraph_count") or 0)),
        ("score_point_bound_rows", int(summary.get("score_point_bound_rows") or 0)),
        ("evidence_bound_rows", int(summary.get("evidence_bound_rows") or 0)),
        ("traceable_locator_rows", int(summary.get("traceable_locator_rows") or 0)),
        ("score_item_count", len(item_cards)),
        ("high_risk_item_count", len(high_risk_items)),
    ]
    for i, (k, v) in enumerate(meta_rows, start=1):
        _write(ws, i, 1, k, header=True)
        _write(ws, i, 2, v)
    _sheet_set_width(ws, 1, 32)
    _sheet_set_width(ws, 2, 80)

    # 2) 评分点覆盖
    ws = wb.create_sheet("score_items")
    headers = [
        "item_id",
        "dimension",
        "keywords",
        "matched_keywords",
        "missing_keywords",
        "coverage_ratio",
        "estimated_score",
        "deduction_risk",
        "matched_sections",
    ]
    for c, h in enumerate(headers, start=1):
        _write(ws, 1, c, h, header=True)
    for r, it in enumerate(item_cards[:800], start=2):
        if not isinstance(it, dict):
            continue
        values = [
            _s(it.get("item_id")),
            _s(it.get("dimension")),
            _join_list(it.get("keywords"), sep="、", limit=50),
            _join_list(it.get("matched_keywords"), sep="、", limit=50),
            _join_list(it.get("missing_keywords"), sep="、", limit=50),
            it.get("coverage_ratio"),
            it.get("estimated_score"),
            it.get("deduction_risk"),
            _join_list(it.get("matched_sections"), sep="；", limit=50),
        ]
        for c, v in enumerate(values, start=1):
            _write(ws, r, c, v)
    _sheet_set_width(ws, 1, 14)
    _sheet_set_width(ws, 2, 18)
    _sheet_set_width(ws, 3, 42)
    _sheet_set_width(ws, 4, 42)
    _sheet_set_width(ws, 5, 42)
    _sheet_set_width(ws, 6, 14)
    _sheet_set_width(ws, 7, 14)
    _sheet_set_width(ws, 8, 14)
    _sheet_set_width(ws, 9, 40)

    # 3) 评分点 × 证据矩阵
    ws = wb.create_sheet("score_evidence_matrix")
    headers = [
        "score_rule_id",
        "score_dimension",
        "matched_keywords",
        "page_estimate",
        "section_title",
        "paragraph_id",
        "evidence_source",
        "evidence_type",
        "response_excerpt",
    ]
    for c, h in enumerate(headers, start=1):
        _write(ws, 1, c, h, header=True)

    row_no = 2
    for rec in rows[:8000]:
        if not isinstance(rec, dict):
            continue
        score_hits = rec.get("tender_score_points") if isinstance(rec.get("tender_score_points"), list) else []
        score_hits = score_hits or [{"rule_id": "", "dimension": "", "matched_keywords": []}]
        sources = rec.get("evidence_sources") if isinstance(rec.get("evidence_sources"), list) else []
        sources = sources or [""]
        typed = rec.get("evidence_typed") if isinstance(rec.get("evidence_typed"), dict) else {}
        excerpt = _s(rec.get("system_response"))[:400]
        for sp in score_hits:
            if not isinstance(sp, dict):
                continue
            for src in sources:
                values = [
                    _s(sp.get("rule_id")),
                    _s(sp.get("dimension")),
                    _join_list(sp.get("matched_keywords"), sep="、", limit=20),
                    rec.get("page_estimate"),
                    _s(rec.get("section_title")),
                    _s(rec.get("paragraph_id")),
                    _s(src),
                    _infer_src_type(_s(src), typed),
                    excerpt,
                ]
                for c, v in enumerate(values, start=1):
                    _write(ws, row_no, c, v)
                row_no += 1

    _sheet_set_width(ws, 1, 16)
    _sheet_set_width(ws, 2, 20)
    _sheet_set_width(ws, 3, 32)
    _sheet_set_width(ws, 4, 12)
    _sheet_set_width(ws, 5, 28)
    _sheet_set_width(ws, 6, 16)
    _sheet_set_width(ws, 7, 68)
    _sheet_set_width(ws, 8, 16)
    _sheet_set_width(ws, 9, 90)

    # 4) 段落证据链明细
    ws = wb.create_sheet("paragraph_evidence")
    headers = [
        "paragraph_id",
        "section_title",
        "paragraph_index",
        "page_estimate",
        "tender_score_points",
        "evidence_sources",
        "graph_nodes",
        "drawing_refs",
        "standard_refs",
        "other_refs",
        "system_response",
    ]
    for c, h in enumerate(headers, start=1):
        _write(ws, 1, c, h, header=True)
    for r, rec in enumerate(rows[:4000], start=2):
        if not isinstance(rec, dict):
            continue
        typed = rec.get("evidence_typed") if isinstance(rec.get("evidence_typed"), dict) else {}
        score_points = rec.get("tender_score_points") if isinstance(rec.get("tender_score_points"), list) else []
        score_text = []
        for sp in score_points:
            if not isinstance(sp, dict):
                continue
            score_text.append(f"{_s(sp.get('rule_id'))}|{_s(sp.get('dimension'))}|{_join_list(sp.get('matched_keywords'), sep='、', limit=10)}")
        values = [
            _s(rec.get("paragraph_id")),
            _s(rec.get("section_title")),
            rec.get("paragraph_index"),
            rec.get("page_estimate"),
            _join_list(score_text, sep="\n", limit=100),
            _join_list(rec.get("evidence_sources"), sep="\n", limit=100),
            _join_list(typed.get("graph_nodes"), sep="\n", limit=100),
            _join_list(typed.get("drawing_refs"), sep="\n", limit=100),
            _join_list(typed.get("standard_refs"), sep="\n", limit=100),
            _join_list(typed.get("other_refs"), sep="\n", limit=100),
            _s(rec.get("system_response"))[:800],
        ]
        for c, v in enumerate(values, start=1):
            _write(ws, r, c, v)
    _sheet_set_width(ws, 1, 16)
    _sheet_set_width(ws, 2, 24)
    _sheet_set_width(ws, 3, 10)
    _sheet_set_width(ws, 4, 10)
    _sheet_set_width(ws, 5, 48)
    _sheet_set_width(ws, 6, 48)
    _sheet_set_width(ws, 7, 36)
    _sheet_set_width(ws, 8, 36)
    _sheet_set_width(ws, 9, 36)
    _sheet_set_width(ws, 10, 36)
    _sheet_set_width(ws, 11, 90)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(output_path))
    return str(output_path)


def export_expert_review_brief_docx(data: Dict[str, Any], output_path: str) -> str:
    data = _sanitize_docx_payload(data)
    _require_local_adapter_export_allowed(data, "expert_review_brief_docx")
    """
    10%专家复核提要版
    仅保留: 工期关键节点、资源峰值、重大安质风控、加分策略触发摘要。
    """
    style_raw = data.get("style") if isinstance(data.get("style"), dict) else {}
    style_cfg = _normalize_style(style_raw)
    doc = Document()
    _apply_page_setup(doc, style_cfg)
    apply_paragraph = _apply_style(doc, style_raw)

    topic = str(data.get("topic") or "施工组织设计").strip()
    h1 = doc.add_heading(f"{topic} - 专家复核提要版", level=1)
    apply_paragraph(h1, is_title=True)
    p_intro = doc.add_paragraph("本稿为高浓度复核摘要，仅保留关键指标、闭环措施与加分策略触发信息。")
    apply_paragraph(p_intro)

    sections = data.get("sections") if isinstance(data.get("sections"), list) else []
    quality_checks = data.get("quality_checks") if isinstance(data.get("quality_checks"), dict) else {}
    plan_consistency = data.get("plan_consistency") if isinstance(data.get("plan_consistency"), dict) else {}
    boq_wbs_cpm = data.get("boq_wbs_cpm") if isinstance(data.get("boq_wbs_cpm"), dict) else {}
    cpm = plan_consistency.get("cpm") if isinstance(plan_consistency.get("cpm"), dict) else {}
    cpm_summary = boq_wbs_cpm.get("summary") if isinstance(boq_wbs_cpm.get("summary"), dict) else {}
    cpm_critical_names = [str(x).strip() for x in (cpm_summary.get("critical_path_names") or []) if str(x).strip()]
    if not cpm_critical_names:
        cpm_critical_names = [str(x).strip() for x in (cpm.get("critical_path") or []) if str(x).strip()]

    duration_days = cpm_summary.get("project_duration_days")
    if duration_days in (None, ""):
        duration_days = ((cpm.get("computed") or {}).get("project_duration_days") if isinstance(cpm.get("computed"), dict) else None)
    resource_peak = cpm_summary.get("resource_peak")
    if resource_peak in (None, ""):
        resource_peak = ((cpm.get("computed") or {}).get("resource_peak") if isinstance(cpm.get("computed"), dict) else None)
    critical_gap = cpm_summary.get("critical_interval_days")
    if critical_gap in (None, ""):
        critical_gap = ((cpm.get("computed") or {}).get("critical_interval_days") if isinstance(cpm.get("computed"), dict) else None)

    h = doc.add_heading("1. 核心工期网络节点", level=2)
    apply_paragraph(h, is_title=True)
    for line in [
        f"- 项目总工期(天): {duration_days if duration_days not in (None, '') else '未提取'}",
        f"- 关键线路最小间隔(天): {critical_gap if critical_gap not in (None, '') else '未提取'}",
        f"- 关键线路节点: {' -> '.join(cpm_critical_names[:20]) if cpm_critical_names else '未提取'}",
    ]:
        p = doc.add_paragraph(line)
        apply_paragraph(p)

    h = doc.add_heading("2. 资源投入峰值", level=2)
    apply_paragraph(h, is_title=True)
    for line in [
        f"- 资源峰值: {resource_peak if resource_peak not in (None, '') else '未提取'}",
        f"- WBS工序数量: {len(boq_wbs_cpm.get('wbs') or []) if isinstance(boq_wbs_cpm, dict) else 0}",
    ]:
        p = doc.add_paragraph(line)
        apply_paragraph(p)

    h = doc.add_heading("3. 重大安全/质量风控闭环", level=2)
    apply_paragraph(h, is_title=True)
    issue_list = quality_checks.get("issue_list") if isinstance(quality_checks.get("issue_list"), list) else []
    picked = []
    for it in issue_list:
        if not isinstance(it, dict):
            continue
        sev = str(it.get("severity") or "").strip().lower()
        txt = f"{it.get('title') or '章节'}|{it.get('type') or ''}|{it.get('problem') or ''}|{it.get('suggestion') or ''}"
        if sev in {"high", "critical"} or ("安全" in txt) or ("质量" in txt) or ("风险" in txt):
            picked.append(it)
    if not picked:
        p = doc.add_paragraph("- 当前质量审计未检出高风险问题。")
        apply_paragraph(p)
    else:
        for it in picked[:12]:
            line = (
                f"- [{it.get('severity') or 'medium'}] {it.get('title') or '章节'}: "
                f"{it.get('problem') or ''}；建议: {it.get('suggestion') or ''}"
            )
            p = doc.add_paragraph(line)
            apply_paragraph(p)

    h = doc.add_heading("4. qt_score_booster 加分策略触发", level=2)
    apply_paragraph(h, is_title=True)
    booster_hits = []
    # 1) 从章节内容中捕捉显式触发
    for sec in sections:
        if not isinstance(sec, dict):
            continue
        title = str(sec.get("title") or "").strip() or "章节"
        txt = str(sec.get("content") or "")
        if "qt_score_booster" in txt or "加分项" in txt or "加分策略" in txt:
            booster_hits.append(f"{title}: 命中章节文本策略标记")
        m = re.findall(r"(?:\+|加分)\s*(\d+(?:\.\d+)?)\s*分", txt)
        if m:
            booster_hits.append(f"{title}: 检测到分值表达 {','.join(m[:3])} 分")
    # 2) 从图谱调度摘要补充
    multi_agent = data.get("multi_agent") if isinstance(data.get("multi_agent"), dict) else {}
    selected_graphs = multi_agent.get("selected_graphs") if isinstance(multi_agent.get("selected_graphs"), list) else []
    for g in selected_graphs:
        if not isinstance(g, dict):
            continue
        booster = g.get("qt_score_booster") if isinstance(g.get("qt_score_booster"), dict) else {}
        if booster:
            gname = str(g.get("graph_name") or g.get("filename") or "图谱").strip()
            booster_hits.append(
                f"{gname}: {str(booster.get('strategy') or booster.get('score_weight') or booster.get('weight') or '已配置加分策略')}"
            )
    # 去重
    dedup = []
    for x in booster_hits:
        if x not in dedup:
            dedup.append(x)
    if dedup:
        for ln in dedup[:16]:
            p = doc.add_paragraph(f"- {ln}")
            apply_paragraph(p)
    else:
        p = doc.add_paragraph("- 未检测到显式 qt_score_booster 触发项（建议在图谱节点或章节中补充加分策略表达）。")
        apply_paragraph(p)

    # 10%章节摘录（供总工快速定位）
    h = doc.add_heading("5. 10%章节快速摘录", level=2)
    apply_paragraph(h, is_title=True)
    if not sections:
        p = doc.add_paragraph("- 无章节数据。")
        apply_paragraph(p)
    else:
        ranked = []
        for sec in sections:
            if not isinstance(sec, dict):
                continue
            title = str(sec.get("title") or "").strip()
            content = str(sec.get("content") or "")
            score = 0
            for kw, w in [
                ("工期", 3),
                ("进度", 3),
                ("关键线路", 4),
                ("资源", 3),
                ("质量", 3),
                ("安全", 3),
                ("风险", 3),
                ("控制", 2),
                ("验证", 2),
            ]:
                if kw in content or kw in title:
                    score += w
            score += min(6, int(len(content) / 800))
            ranked.append((score, title, content))
        ranked.sort(key=lambda x: x[0], reverse=True)
        keep = max(1, min(12, math.ceil(len(ranked) * 0.10)))
        for _, title, content in ranked[:keep]:
            p1 = doc.add_paragraph(f"- {title}")
            apply_paragraph(p1)
            p2 = doc.add_paragraph((content[:360] + "...") if len(content) > 360 else content)
            apply_paragraph(p2)

    _enable_field_updates(doc)
    _save_docx_secure(doc, output_path)
    return str(output_path)
