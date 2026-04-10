from __future__ import annotations

import json
import re
from typing import Any, Dict, List

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from backend.zhifei_autoplan.style_policy import normalize_cn_font_name


DEFAULT_FORMAT_CONFIG: Dict[str, Any] = {
    "body_font": "宋体",
    "title_font": "黑体",
    "body_size_pt": 14.0,
    "title_size_pt": 16.0,
    "line_spacing_pt": 22.0,
    "margins_cm": {
        "top": 2.5,
        "right": 2.0,
        "bottom": 2.0,
        "left": 2.0,
    },
    "cover_page_count": 1,
    "toc_page_count": 2,
    "front_matter_page_mode": "include",
}

_EVIDENCE_TAG_RE = re.compile(r"【证据:(.+?)】")
_INLINE_INTERNAL_TAG_RE = re.compile(r"【(?:图谱节点|经验值|图谱经验值):[^】]+】")
_WRAPPED_HEADING_RE = re.compile(r"^【([^】]{1,60})】[:：]?\s*")
_KV_TOKEN_RE = re.compile(r"([^=：:；;，,\s]{1,20})\s*(?:=|：|:)\s*([^；;，,\n]+)")
_RISK_ROW_RE = re.compile(
    r"(?:^[-*•]?\s*)?风险[:：]\s*(?P<risk>.+?)\s*[；;]\s*控制[:：]\s*(?P<control>.+?)\s*[；;]\s*验证[:：]\s*(?P<verify>.+)$"
)

_SYSTEM_ONLY_LABELS = {
    "范围",
    "系统全局指令",
    "图谱节点绑定",
    "图谱逻辑节点（必须绑定）",
    "多agent",
    "多agent协作",
    "章节结构蓝图",
    "证据摘要",
    "证据与追溯",
    "可编辑参数（优先采用；若招标/图纸/清单有明确要求，则以证据为准）",
    "可编辑参数",
    "知识图谱证据",
    "招标/清单/图纸证据",
    "合规检查要点",
    "文风硬约束（必须）",
    "编制要求",
    "权重与扣分项",
    "系统全局指令（必须无条件执行）",
    "青天适配硬约束（本章必须执行）",
}
_TABLE_ONLY_LABELS = {
    "风险→控制→验证",
    "风险-控制-验证",
    "风险控制验证表",
    "资源-工序耦合表",
    "资源-工序耦合",
}
_VISIBLE_LABEL_ALIASES = {
    "适用范围与关键参数": "适用范围与关键参数",
    "重点难点/风险点及措施": "重点难点及控制措施",
    "重点难点/风险措施": "重点难点及控制措施",
    "工序流程": "施工工序流程",
    "步骤控制点（量化）": "关键控制指标",
    "控制指标矩阵": "关键控制指标",
    "人机料法环落地": "人机料法环控制要点",
    "监管红线清单": "监管红线清单",
    "岗位联签链": "岗位联签链",
    "闭环时限表": "闭环时限要求",
    "接口冲突清单": "接口冲突清单",
    "关键路径纠偏卡": "关键路径纠偏要求",
    "实施场景卡片": "实施场景",
    "参数对照表": "参数控制要点",
    "验收样表": "验收记录要点",
    "区域网格": "区域网格划分",
    "班组行为清单": "班组行为要求",
    "红黄牌处置": "红黄牌处置",
    "复核与销项": "复核与销项",
    "清单重点项": "清单重点项",
}
_RESOURCE_KEYS = {"班组人数", "人数", "设备", "设备型号", "班组", "机械", "资源", "工器具"}
_CONTROL_KEYS = {
    "频次",
    "阈值",
    "间距",
    "厚度",
    "时长",
    "节拍",
    "抽检",
    "抽检频次",
    "采购比价",
    "合格率阈值",
    "一次验收通过率",
    "台账抽查频次",
    "应急演练频次",
}
_SCaffold_LINE_PATTERNS = [
    re.compile(r"^\s*角色定位[:：]"),
    re.compile(r"^\s*章节标题[:：]"),
    re.compile(r"^\s*方案版本[:：]"),
    re.compile(r"^\s*输出要求[:：]"),
    re.compile(r"^\s*样本[:：]"),
    re.compile(r"^\s*\d+\)\s"),
    re.compile(r"^\s*provider\s*[:=]"),
    re.compile(r"^\s*model\s*[:=]"),
    re.compile(r"^\s*constraint_log\s*[:=]"),
]


def _is_number(value: Any) -> bool:
    try:
        float(value)
        return True
    except Exception:
        return False


def _to_float(value: Any) -> float | None:
    try:
        return float(value)
    except Exception:
        return None


def _clean_scalar(value: Any) -> Any:
    if value is None:
        return None
    if isinstance(value, (str, int, float, bool)):
        return value
    return json.loads(json.dumps(value, ensure_ascii=False, default=str))


def _style_scalar(style: Dict[str, Any], key: str, nested_key: str | None = None) -> Any:
    if style.get(key) is not None:
        return style.get(key)
    if nested_key and isinstance(style.get("font"), dict) and style["font"].get(nested_key) is not None:
        return style["font"].get(nested_key)
    return None


def build_bidding_format_config_from_style(style: Dict[str, Any] | None) -> Dict[str, Any]:
    raw = style if isinstance(style, dict) else {}
    font_cfg = raw.get("font") if isinstance(raw.get("font"), dict) else {}
    headings_cfg = raw.get("headings") if isinstance(raw.get("headings"), dict) else {}
    margins_cfg = raw.get("margins_cm") if isinstance(raw.get("margins_cm"), dict) else {}
    margins_list = raw.get("margins") if isinstance(raw.get("margins"), list) else []

    def _margin_value(name: str, index: int) -> Any:
        if margins_cfg.get(name) is not None:
            return margins_cfg.get(name)
        if index < len(margins_list):
            return margins_list[index]
        return None

    return {
        "body_font": (
            normalize_cn_font_name(
            raw.get("body_font")
            or (raw.get("font") if isinstance(raw.get("font"), str) else None)
            or font_cfg.get("eastAsia")
            )
            or None
        ),
        "title_font": normalize_cn_font_name(raw.get("title_font") or headings_cfg.get("eastAsia")) or None,
        "body_size_pt": raw.get("body_size") or raw.get("font_size") or font_cfg.get("size_pt"),
        "title_size_pt": raw.get("title_size") or headings_cfg.get("h2_size") or headings_cfg.get("h1_size"),
        "line_spacing_pt": raw.get("line_spacing_pt") or font_cfg.get("line_spacing_pt"),
        "margins_cm": {
            "top": _margin_value("top", 0),
            "right": _margin_value("right", 1),
            "bottom": _margin_value("bottom", 2),
            "left": _margin_value("left", 3),
        },
        "cover_page_count": raw.get("cover_page_count"),
        "toc_page_count": raw.get("toc_page_count"),
        "front_matter_page_mode": raw.get("front_matter_page_mode"),
        "document_total_pages_target": raw.get("document_total_pages_target"),
    }


def merge_style_with_bidding_fallback(
    *,
    user_style: Dict[str, Any] | None,
    bidding_format_config: Dict[str, Any] | None,
) -> Dict[str, Any]:
    user = dict(user_style or {})
    bidding = bidding_format_config if isinstance(bidding_format_config, dict) else {}

    merged = dict(user)
    bidding_body_font = normalize_cn_font_name(bidding.get("body_font"))
    bidding_title_font = normalize_cn_font_name(bidding.get("title_font"))
    user_body_font = normalize_cn_font_name(user.get("body_font"))
    user_title_font = normalize_cn_font_name(user.get("title_font"))
    merged["body_font"] = bidding_body_font or user_body_font or DEFAULT_FORMAT_CONFIG["body_font"]
    merged["title_font"] = bidding_title_font or user_title_font or DEFAULT_FORMAT_CONFIG["title_font"]
    merged["body_size"] = (
        bidding.get("body_size_pt")
        or user.get("body_size")
        or user.get("font_size")
        or DEFAULT_FORMAT_CONFIG["body_size_pt"]
    )
    merged["title_size"] = bidding.get("title_size_pt") or user.get("title_size") or DEFAULT_FORMAT_CONFIG["title_size_pt"]
    merged["line_spacing_pt"] = (
        bidding.get("line_spacing_pt") or user.get("line_spacing_pt") or DEFAULT_FORMAT_CONFIG["line_spacing_pt"]
    )

    merged_margins = dict(DEFAULT_FORMAT_CONFIG["margins_cm"])
    if isinstance(user.get("margins_cm"), dict):
        for key, value in user["margins_cm"].items():
            if value is not None:
                merged_margins[key] = value
    if isinstance(bidding.get("margins_cm"), dict):
        for key, value in bidding["margins_cm"].items():
            if value is not None:
                merged_margins[key] = value
    merged["margins_cm"] = merged_margins

    for key in ("cover_page_count", "toc_page_count", "front_matter_page_mode", "document_total_pages_target"):
        if bidding.get(key) is not None:
            merged[key] = bidding.get(key)
        elif user.get(key) is not None:
            merged[key] = user.get(key)
        elif key in DEFAULT_FORMAT_CONFIG:
            merged[key] = DEFAULT_FORMAT_CONFIG[key]
    return merged


def _strip_line_prefix(line: str) -> tuple[str, str]:
    match = _WRAPPED_HEADING_RE.match(line.strip())
    if not match:
        return "", line.strip()
    return match.group(1).strip(), line.strip()[match.end():].strip()


def _normalize_label(label: str) -> str:
    return str(label or "").strip().lower()


def is_system_scaffold_line(line: str) -> bool:
    raw = str(line or "").strip()
    if not raw:
        return False
    label, rest = _strip_line_prefix(raw)
    if _normalize_label(label) in {_normalize_label(x) for x in _SYSTEM_ONLY_LABELS}:
        return True
    if raw.startswith("【系统") or raw.startswith("【图谱") or raw.startswith("【证据摘要】"):
        return True
    if raw.startswith("{") and raw.endswith("}"):
        return True
    if raw.startswith("[") and raw.endswith("]") and any(token in raw for token in ('"title"', '"problem"', '"suggestion"')):
        return True
    return any(p.search(raw) for p in _SCaffold_LINE_PATTERNS)


def _strip_inline_internal_tags(text: str) -> str:
    return _INLINE_INTERNAL_TAG_RE.sub("", str(text or ""))


def _strip_evidence_tags(text: str) -> tuple[str, List[str]]:
    anchors = [m.group(1).strip() for m in _EVIDENCE_TAG_RE.finditer(str(text or "")) if m.group(1).strip()]
    cleaned = _EVIDENCE_TAG_RE.sub("", str(text or ""))
    return cleaned, anchors


def _cleanup_spacing(text: str) -> str:
    s = str(text or "")
    s = re.sub(r"[ \t]{2,}", " ", s)
    s = re.sub(r"[；;，,。]{2,}", lambda m: m.group(0)[0], s)
    s = re.sub(r"\s+([，。；：:])", r"\1", s)
    s = re.sub(r"([（(])\s+", r"\1", s)
    s = re.sub(r"\s+([）)])", r"\1", s)
    return s.strip(" \t\r\n-•*")


def _parse_kv_pairs(text: str) -> List[tuple[str, str]]:
    pairs = []
    for key, value in _KV_TOKEN_RE.findall(str(text or "")):
        k = str(key or "").strip()
        v = str(value or "").strip().strip("。")
        if not k or not v:
            continue
        pairs.append((k, v))
    return pairs


def _naturalize_pair(key: str, value: str) -> str:
    v = str(value or "").strip().rstrip("。")
    normalized = re.sub(r"^偏差[≤<=]*", "", v)
    mapping = {
        "频次": f"巡检频次控制为{v}",
        "抽检频次": f"抽检频次控制为{v}",
        "台账抽查频次": f"台账抽查频次控制为{v}",
        "应急演练频次": f"应急演练按{v}组织",
        "阈值": f"关键偏差控制在{normalized}以内" if normalized else f"关键偏差控制按{v}执行",
        "合格率阈值": f"合格率控制在{v}以上",
        "一次验收通过率": f"一次验收通过率控制在{v}以上",
        "间距": f"关键间距按{v}控制",
        "厚度": f"关键厚度按{v}控制",
        "时长": f"单个作业段控制在{v}内完成",
        "节拍": f"单个工序节拍控制为{v}",
        "人数": f"现场安排{v}组织施工",
        "班组人数": f"单班配置{v}",
        "设备型号": f"主要设备采用{v}",
        "设备": f"主要设备采用{v}",
        "采购比价": f"材料采购执行{v}比价机制",
    }
    if key in mapping:
        return mapping[key]
    return f"{key}按{v}执行"


def naturalize_machine_text(text: str) -> str:
    raw = _cleanup_spacing(_strip_inline_internal_tags(str(text or "")))
    if not raw:
        return ""
    label, rest = _strip_line_prefix(raw)
    pairs = _parse_kv_pairs(rest or raw)
    if len(pairs) < 2:
        return raw
    pair_chars = sum(len(k) + len(v) for k, v in pairs)
    if pair_chars < max(8, int(len(rest or raw) * 0.45)):
        return raw
    sentences = [_naturalize_pair(key, value) for key, value in pairs]
    if label == "量化指标":
        return "本章量化控制要求如下：" + "，".join(sentences) + "。"
    if label in {"参数对照表", "控制指标矩阵"}:
        return "本章控制参数如下：" + "，".join(sentences) + "。"
    if label:
        return f"{label}：" + "，".join(sentences) + "。"
    return "本章控制要求如下：" + "，".join(sentences) + "。"


def sanitize_delivery_line(line: str) -> Dict[str, Any]:
    raw = str(line or "").strip()
    if not raw:
        return {"visible": "", "anchors": [], "stripped": False, "reason": "empty"}
    if is_system_scaffold_line(raw):
        return {"visible": "", "anchors": [], "stripped": True, "reason": "system_scaffold", "raw": raw}

    label, rest = _strip_line_prefix(raw)
    if _normalize_label(label) in {_normalize_label(x) for x in _TABLE_ONLY_LABELS} and not rest:
        return {"visible": "", "anchors": [], "stripped": True, "reason": "table_marker", "raw": raw}

    visible, anchors = _strip_evidence_tags(raw)
    visible = _strip_inline_internal_tags(visible)
    label, rest = _strip_line_prefix(visible)
    if label in {"量化指标", "控制指标矩阵", "参数对照表", "步骤控制点（量化）"}:
        visible = naturalize_machine_text(raw)
        visible, extra_anchors = _strip_evidence_tags(visible)
        anchors.extend(extra_anchors)
        visible = _cleanup_spacing(visible)
        if not visible:
            return {"visible": "", "anchors": anchors, "stripped": True, "reason": "empty_after_cleanup", "raw": raw}
        return {"visible": visible, "anchors": anchors, "stripped": False}
    if label and _normalize_label(label) in {_normalize_label(x) for x in _SYSTEM_ONLY_LABELS}:
        return {"visible": "", "anchors": anchors, "stripped": True, "reason": "system_heading", "raw": raw}
    if label and label in _VISIBLE_LABEL_ALIASES and not rest:
        visible = _VISIBLE_LABEL_ALIASES[label]
    elif label and label not in _VISIBLE_LABEL_ALIASES and rest:
        visible = f"{label}：{rest}"
    elif label in _VISIBLE_LABEL_ALIASES and rest:
        visible = f"{_VISIBLE_LABEL_ALIASES[label]}：{rest}"

    visible = naturalize_machine_text(visible)
    visible = _cleanup_spacing(visible)
    if not visible:
        return {"visible": "", "anchors": anchors, "stripped": True, "reason": "empty_after_cleanup", "raw": raw}
    return {"visible": visible, "anchors": anchors, "stripped": False}


def _risk_table_from_lines(lines: List[str]) -> tuple[List[Dict[str, Any]], int]:
    rows: List[Dict[str, Any]] = []
    consumed = 0
    for raw in lines:
        line = str(raw or "").strip()
        if not line:
            break
        visible, anchors = _strip_evidence_tags(_strip_inline_internal_tags(line))
        match = _RISK_ROW_RE.match(visible.strip())
        if not match:
            if rows:
                break
            return [], 0
        rows.append(
            {
                "cells": [
                    _cleanup_spacing(match.group("risk")),
                    _cleanup_spacing(match.group("control")),
                    _cleanup_spacing(match.group("verify")),
                ],
                "anchors": anchors,
            }
        )
        consumed += 1
    return rows, consumed


def _resource_table_from_lines(lines: List[str]) -> tuple[List[Dict[str, Any]], int]:
    rows: List[Dict[str, Any]] = []
    consumed = 0
    for raw in lines:
        line = str(raw or "").strip()
        if not line:
            break
        visible, anchors = _strip_evidence_tags(_strip_inline_internal_tags(line))
        pairs = _parse_kv_pairs(visible)
        if not pairs:
            if rows:
                break
            return [], 0
        process = ""
        resource_parts: List[str] = []
        control_parts: List[str] = []
        for key, value in pairs:
            if key == "工序" and not process:
                process = value
            elif key in _RESOURCE_KEYS:
                resource_parts.append(_naturalize_pair(key, value))
            else:
                control_parts.append(_naturalize_pair(key, value))
        if not process:
            if rows:
                break
            return [], 0
        rows.append(
            {
                "cells": [
                    _cleanup_spacing(process),
                    "；".join(resource_parts) if resource_parts else "按章节资源计划配置。",
                    "；".join(control_parts) if control_parts else "按工序控制要求执行。",
                ],
                "anchors": anchors,
            }
        )
        consumed += 1
    return rows, consumed


def prepare_delivery_render(text: str) -> Dict[str, Any]:
    source_lines = str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: List[Dict[str, Any]] = []
    stripped_items: List[Dict[str, Any]] = []
    all_anchors: List[str] = []
    idx = 0
    while idx < len(source_lines):
        raw = source_lines[idx]
        line = str(raw or "").strip()
        if not line:
            idx += 1
            continue

        label, rest = _strip_line_prefix(line)
        normalized_label = _normalize_label(label)
        if normalized_label in {_normalize_label("风险→控制→验证"), _normalize_label("风险-控制-验证"), _normalize_label("风险控制验证表")}:
            rows, consumed = _risk_table_from_lines(source_lines[idx + 1 :])
            stripped_items.append({"type": "marker", "raw": line, "reason": "risk_table_marker"})
            if rows:
                all_anchors.extend(anchor for row in rows for anchor in row.get("anchors") or [])
                blocks.append(
                    {
                        "type": "table",
                        "title": "风险控制验证表",
                        "headers": ["风险源", "控制措施", "验证方式"],
                        "rows": rows,
                    }
                )
            idx += max(1, consumed + 1)
            continue
        if normalized_label in {_normalize_label("资源-工序耦合表"), _normalize_label("资源-工序耦合")}:
            rows, consumed = _resource_table_from_lines(source_lines[idx + 1 :])
            stripped_items.append({"type": "marker", "raw": line, "reason": "resource_table_marker"})
            if rows:
                all_anchors.extend(anchor for row in rows for anchor in row.get("anchors") or [])
                blocks.append(
                    {
                        "type": "table",
                        "title": "资源-工序耦合表",
                        "headers": ["工序", "资源配置", "控制要求"],
                        "rows": rows,
                    }
                )
            idx += max(1, consumed + 1)
            continue

        cleaned = sanitize_delivery_line(line)
        if cleaned.get("stripped"):
            stripped_items.append({"type": "line", "raw": line, "reason": cleaned.get("reason")})
            idx += 1
            continue
        anchors = cleaned.get("anchors") or []
        all_anchors.extend(anchors)
        blocks.append({"type": "paragraph", "text": cleaned.get("visible") or "", "anchors": anchors})
        idx += 1

    visible_lines = []
    for block in blocks:
        if block["type"] == "paragraph":
            visible_lines.append(block["text"])
        elif block["type"] == "table":
            visible_lines.append(block["title"])
            for row in block.get("rows") or []:
                visible_lines.append(" ".join([str(cell or "").strip() for cell in row.get("cells") or [] if str(cell or "").strip()]))
    return {
        "blocks": blocks,
        "visible_text": "\n".join([line for line in visible_lines if line]).strip(),
        "stripped_items": stripped_items,
        "hidden_evidence_anchors": all_anchors,
    }


def _append_hidden_run(paragraph, text: str) -> None:
    if not str(text or "").strip():
        return
    run = paragraph.add_run(" " + str(text).strip())
    try:
        run.font.hidden = True
    except Exception:
        pass
    rpr = run._element.get_or_add_rPr()
    vanish = OxmlElement("w:vanish")
    rpr.append(vanish)


def render_delivery_blocks(doc: Document, apply_paragraph, prepared: Dict[str, Any]) -> Dict[str, Any]:
    rendered_tables = 0
    hidden_anchor_count = 0
    for block in prepared.get("blocks") or []:
        if block.get("type") == "paragraph":
            p = doc.add_paragraph(block.get("text") or "")
            apply_paragraph(p)
            anchors = [str(x).strip() for x in (block.get("anchors") or []) if str(x).strip()]
            if anchors:
                _append_hidden_run(p, "；".join([f"证据:{x}" for x in anchors]))
                hidden_anchor_count += len(anchors)
            continue
        if block.get("type") == "table":
            title = str(block.get("title") or "").strip()
            if title:
                hp = doc.add_paragraph(title)
                apply_paragraph(hp, is_title=True)
            headers = [str(x or "").strip() for x in (block.get("headers") or [])]
            table = doc.add_table(rows=1, cols=max(1, len(headers)))
            for idx, value in enumerate(headers):
                table.rows[0].cells[idx].text = value
                for p in table.rows[0].cells[idx].paragraphs:
                    apply_paragraph(p, is_title=True)
            for row in block.get("rows") or []:
                cells = table.add_row().cells
                for idx, value in enumerate(row.get("cells") or []):
                    if idx < len(cells):
                        cells[idx].text = str(value or "")
                        for p in cells[idx].paragraphs:
                            apply_paragraph(p)
                anchors = [str(x).strip() for x in (row.get("anchors") or []) if str(x).strip()]
                if anchors and len(cells) > 0:
                    _append_hidden_run(cells[-1].paragraphs[0], "；".join([f"证据:{x}" for x in anchors]))
                    hidden_anchor_count += len(anchors)
            rendered_tables += 1
    return {
        "table_count": rendered_tables,
        "hidden_anchor_count": hidden_anchor_count,
        "visible_text": prepared.get("visible_text") or "",
        "stripped_items": prepared.get("stripped_items") or [],
    }


def collect_delivery_report_payload(*, section_reports: List[Dict[str, Any]], internal_payload: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "section_render_reports": [
            {
                "title": str(item.get("title") or ""),
                "table_count": int(item.get("table_count") or 0),
                "hidden_anchor_count": int(item.get("hidden_anchor_count") or 0),
                "visible_length": len(str(item.get("visible_text") or "")),
                "stripped_items": [_clean_scalar(x) for x in (item.get("stripped_items") or [])],
            }
            for item in (section_reports or [])
        ],
        "internal_payload": _clean_scalar(internal_payload) or {},
    }
