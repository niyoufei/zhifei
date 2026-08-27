from __future__ import annotations

import asyncio
import concurrent.futures
import hashlib
import hmac
import json
import logging
import multiprocessing
import os
import re
import shutil
import subprocess
import tempfile
import threading
import time
from collections.abc import Callable
from concurrent.futures.process import BrokenProcessPool
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Annotated, Any

from fastapi import APIRouter, File, HTTPException, UploadFile
from pypdf import PdfReader

from backend.zhifei_autoplan.case_library_service import CASE_LIBRARY_SCOPE
from backend.zhifei_autoplan.image_library import (
    IMAGE_LIBRARY_SCOPE,
    normalize_text_list,
)
from backend.zhifei_autoplan.ingest_tags import (
    classify_document_tags,
    normalize_source_hint,
)
from backend.zhifei_autoplan.job_store import (
    JobLeaseLostError,
    acquire_job_lease,
    create_job,
    get_job,
    heartbeat_job,
    job_lease_active,
    merge_job,
    run_with_job_lease,
    transition_job,
)
from backend.zhifei_autoplan.local_job_queue import submit_local_job
from backend.zhifei_autoplan.project_types import normalize_project_type
from backend.zhifei_autoplan.runtime_events import append_runtime_event

logger = logging.getLogger(__name__)

try:
    from backend.zhifei_autoplan.local_adapter_shim import (
        normalize_input as _local_adapter_normalize_input,
    )
except Exception:
    logger.warning("local adapter normalizer unavailable; ingest will remain unnormalized", exc_info=True)
    _local_adapter_normalize_input = None

router = APIRouter(prefix="/ingest", tags=["文档解析"])

PARSE_CACHE_DIR = Path(
    os.environ.get("ZF_AUTOPLAN_INGEST_CACHE_DIR", "backend/data/autoplan/ingest_cache")
)
INGEST_SPOOL_DIR = Path(
    os.environ.get("ZF_AUTOPLAN_INGEST_SPOOL_DIR", "backend/data/autoplan/ingest_spool")
)
PARSER_VERSION = "2026.08.runtime-v5-source-aware-ocr-cache"
OCR_POLICY_ORDINARY = "ordinary_bounded"
OCR_POLICY_DRAWING = "drawing_full_page"
OCR_POLICY_STANDARD = "standard_full_page"
OCR_POLICIES_FULL_PAGE = frozenset({OCR_POLICY_DRAWING, OCR_POLICY_STANDARD})
PARSE_CACHE_TEXT_SIDECAR_BYTES = 256 * 1024
FILE_ID_SEARCH_ROOTS = (
    Path("backend/data/uploads"),
    Path("backend/data/workspaces"),
)
MAX_BATCH_FILES = 40
MAX_BATCH_BYTES = 512 * 1024 * 1024
PARSE_TIMEOUT_SECONDS = 180
INGEST_HEARTBEAT_SECONDS = 10
try:
    INGEST_WORKERS = max(1, min(2, int(os.getenv("ZF_AUTOPLAN_INGEST_WORKERS", "2"))))
except (TypeError, ValueError):
    INGEST_WORKERS = 2
_PARSE_EXECUTOR: concurrent.futures.ProcessPoolExecutor | None = None
_PARSE_EXECUTOR_LOCK = threading.Lock()


def _guarded_write(
    write_guard: Callable[..., Any] | None,
    callback: Callable[..., Any],
    *args: Any,
    **kwargs: Any,
) -> Any:
    if callable(write_guard):
        return write_guard(callback, *args, **kwargs)
    return callback(*args, **kwargs)


def _parse_executor() -> concurrent.futures.ProcessPoolExecutor:
    global _PARSE_EXECUTOR
    with _PARSE_EXECUTOR_LOCK:
        if _PARSE_EXECUTOR is None:
            _PARSE_EXECUTOR = concurrent.futures.ProcessPoolExecutor(
                max_workers=INGEST_WORKERS,
                mp_context=multiprocessing.get_context("spawn"),
            )
    return _PARSE_EXECUTOR


def _discard_parse_executor(
    executor: concurrent.futures.ProcessPoolExecutor,
    *,
    terminate: bool = False,
) -> None:
    """Remove an unhealthy native-parser pool without taking down FastAPI."""

    global _PARSE_EXECUTOR
    with _PARSE_EXECUTOR_LOCK:
        if _PARSE_EXECUTOR is executor:
            _PARSE_EXECUTOR = None
    if terminate:
        processes = getattr(executor, "_processes", None)
        if isinstance(processes, dict):
            for process in list(processes.values()):
                try:
                    process.terminate()
                except Exception:
                    logger.warning("native parser worker could not be terminated", exc_info=True)
    try:
        executor.shutdown(wait=False, cancel_futures=True)
    except Exception:
        logger.warning("native parser executor shutdown failed", exc_info=True)


async def _run_isolated_process(
    func: Any,
    *args: Any,
    timeout: float = PARSE_TIMEOUT_SECONDS,
) -> Any:
    """Run native/CPU parsing outside the API process and recycle bad pools."""

    executor = _parse_executor()
    loop = asyncio.get_running_loop()
    future = loop.run_in_executor(executor, func, *args)
    try:
        return await asyncio.wait_for(future, timeout=timeout)
    except (asyncio.TimeoutError, BrokenProcessPool):
        _discard_parse_executor(executor, terminate=True)
        raise


async def _extract_text_path_bounded(ext: str, path: Path, total_bytes: int) -> dict[str, Any]:
    if ext in {"pdf", "doc", "docx"} and int(total_bytes) >= 1024 * 1024:
        return await _run_isolated_process(
            _extract_text_path,
            ext,
            path,
            timeout=PARSE_TIMEOUT_SECONDS,
        )
    else:
        future = asyncio.to_thread(_extract_text_path, ext, path)
    return await asyncio.wait_for(future, timeout=PARSE_TIMEOUT_SECONDS)
try:
    MAX_UPLOAD_BYTES = max(1, int(os.getenv("ZHIFEI_MAX_UPLOAD_BYTES", str(100 * 1024 * 1024))))
except (TypeError, ValueError):
    MAX_UPLOAD_BYTES = 100 * 1024 * 1024


def _sha256(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def _file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _publish_temp_exclusive(
    temp_path: Path,
    out_path: Path,
    *,
    write_guard: Callable[..., Any] | None = None,
) -> bool:
    """Publish a same-directory temp file without replacing an existing path."""

    try:
        _guarded_write(write_guard, os.link, temp_path, out_path)
    except FileExistsError:
        return False
    return True


def _ext(name: str) -> str:
    return (name.rsplit(".", 1)[-1].lower() if "." in name else "")


def _resolve_workspace_context(
    session_id: str | None = None,
    workspace_dir: str | None = None,
    *,
    write_guard: Callable[..., Any] | None = None,
) -> dict[str, str]:
    if workspace_dir:
        root = Path(workspace_dir)
    elif session_id:
        safe_session = "".join(ch if ch.isalnum() or ch in {"-", "_"} else "_" for ch in str(session_id))
        root = Path("backend/data/workspaces") / (safe_session or "default")
    else:
        root = Path("backend/data")
    _guarded_write(write_guard, root.mkdir, parents=True, exist_ok=True)
    return {"session_id": str(session_id or ""), "workspace_dir": str(root)}


def _attach_local_adapter_ingest_result(result: dict[str, Any], project_id: str | None = None) -> dict[str, Any]:
    if _local_adapter_normalize_input is None or not isinstance(result, dict):
        return result
    try:
        saved = result.get("saved") if isinstance(result.get("saved"), list) else []
        source_files = [
            {
                "filename": rec.get("filename"),
                "sha256": rec.get("sha256"),
                "doc_type": rec.get("doc_type"),
            }
            for rec in saved
            if isinstance(rec, dict)
        ]
        envelope = _local_adapter_normalize_input({"project_id": project_id, "source_files": source_files})
        result["local_adapter"] = {
            "status": "normalized",
            "source_file_count": len(envelope.get("source_files") or []),
            "missing_params": envelope.get("missing_params") or [],
        }
    except Exception as exc:  # noqa: BLE001 - adapter boundary emits a stable fail-closed issue
        logger.warning(
            "local adapter input normalization failed: %s",
            type(exc).__name__,
        )
        result["local_adapter"] = {
            "status": "normalizer_error",
            "issues": [
                {
                    "code": "LOCAL_ADAPTER_INPUT_NORMALIZE_FAILED",
                    "message": "local adapter input normalization failed",
                    "error_type": type(exc).__name__,
                }
            ],
        }
    return result


def workspace_paths(
    workspace_dir: str | Path,
    *,
    write_guard: Callable[..., Any] | None = None,
) -> dict[str, Path]:
    root = Path(workspace_dir)
    paths = {
        "uploads": root / "uploads",
        "extracts": root / "extracts",
        "previews": root / "previews",
        "ingest_audit": root / "audit" / "ingest.jsonl",
    }
    for key, path in paths.items():
        if key == "ingest_audit":
            _guarded_write(
                write_guard,
                path.parent.mkdir,
                parents=True,
                exist_ok=True,
            )
        else:
            _guarded_write(write_guard, path.mkdir, parents=True, exist_ok=True)
    return paths


def _extract_text_path(ext: str, path: Path) -> dict[str, Any]:
    if ext in {"txt", "md"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return {
            "doc_type": ext,
            "pages": 1,
            "text_bytes": len(text.encode("utf-8")),
            "extract_text": text,
        }
    if ext == "pdf":
        reader = PdfReader(str(path))
        texts = []
        for page in reader.pages:
            texts.append(page.extract_text() or "")
        text = "\n\n\f\n\n".join(texts)
        return {
            "doc_type": "pdf",
            "pages": len(reader.pages),
            "text_bytes": len(text.encode("utf-8")),
            "extract_text": text,
        }
    if ext == "docx":
        from docx import Document

        document = Document(str(path))
        lines = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if values:
                    lines.append(" | ".join(values))
        text = "\n".join(lines)
        return {
            "doc_type": "docx",
            "pages": None,
            "text_bytes": len(text.encode("utf-8")),
            "extract_text": text,
        }
    if ext == "doc":
        office = shutil.which("soffice") or shutil.which("libreoffice")
        if not office:
            raise RuntimeError("LEGACY_DOC_CONVERTER_UNAVAILABLE")
        with tempfile.TemporaryDirectory(prefix="zhifei-doc-convert-") as temp_dir:
            completed = subprocess.run(
                [
                    office,
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    temp_dir,
                    str(path),
                ],
                capture_output=True,
                timeout=60,
                check=False,
            )
            converted = Path(temp_dir) / f"{path.stem}.docx"
            if completed.returncode != 0 or not converted.exists():
                raise RuntimeError("LEGACY_DOC_CONVERSION_FAILED")
            parsed = _extract_text_path("docx", converted)
            parsed["doc_type"] = "doc"
            parsed["converted_via"] = "libreoffice"
            return parsed
    return {"doc_type": ext or "unknown", "pages": None, "text_bytes": None}


def _ocr_cache_policy(source_hint: str | None = None) -> str:
    normalized_hint = _normalize_source_hint(source_hint)
    if normalized_hint == "drawing_standard":
        return OCR_POLICY_DRAWING
    if normalized_hint == "standard":
        return OCR_POLICY_STANDARD
    return OCR_POLICY_ORDINARY


def _parse_cache_path(digest: str, source_hint: str | None = None) -> Path:
    cache_identity = f"{PARSER_VERSION}:{_ocr_cache_policy(source_hint)}"
    version_digest = hashlib.sha256(cache_identity.encode("utf-8")).hexdigest()[:12]
    return PARSE_CACHE_DIR / f"{digest!s}.{version_digest}.json"


def _parse_cache_text_path(digest: str, source_hint: str | None = None) -> Path:
    """Return the versioned sidecar used for potentially large extracted text."""

    return _parse_cache_path(digest, source_hint).with_suffix(".txt")


def _resolve_ingested_files(file_ids: list[str] | None) -> list[tuple[str, Path]]:
    """Resolve full-SHA file IDs to content-verified local upload paths."""

    resolved: list[tuple[str, Path]] = []
    for raw in file_ids or []:
        digest = str(raw or "").strip().lower()
        if len(digest) != 64 or any(ch not in "0123456789abcdef" for ch in digest):
            raise HTTPException(status_code=400, detail={"code": "INVALID_FILE_ID"})
        match: Path | None = None
        for root in FILE_ID_SEARCH_ROOTS:
            if not root.exists():
                continue
            full_named = list(root.rglob(f"{digest}_*"))
            for candidate in full_named:
                if candidate.is_symlink() or not candidate.is_file():
                    continue
                if _file_sha256(candidate) != digest:
                    raise HTTPException(
                        status_code=409,
                        detail={
                            "code": "CONTENT_ADDRESS_COLLISION",
                            "file_id": digest,
                        },
                    )
                match = candidate
                break
            if match is None:
                # Read-only compatibility for historical eight-character
                # names.  Every candidate is byte-verified; new writes never
                # use this namespace.
                for candidate in root.rglob(f"{digest[:8]}_*"):
                    if candidate.is_symlink() or not candidate.is_file():
                        continue
                    if _file_sha256(candidate) == digest:
                        match = candidate
                        break
            if match is not None:
                break
        if match is None:
            raise HTTPException(
                status_code=404,
                detail={"code": "FILE_ID_NOT_FOUND", "file_id": digest},
            )
        resolved.append((digest, match))
    return resolved


def resolve_ingested_file_ids(file_ids: list[str] | None) -> list[str]:
    """Resolve full-SHA file IDs to verified local upload paths."""

    return [str(path) for _digest, path in _resolve_ingested_files(file_ids)]


def resolve_ingested_tender_sources(
    file_ids: list[str] | None,
) -> list[dict[str, str | None]]:
    """Resolve Tender sources and reuse only the current validated text cache."""

    sources: list[dict[str, str | None]] = []
    for digest, path in _resolve_ingested_files(file_ids):
        cached = _load_parse_cache(digest, source_hint="tender_qa")
        base = cached.get("base") if isinstance(cached, dict) else None
        extract_text = base.get("extract_text") if isinstance(base, dict) else None
        sources.append(
            {
                "path": str(path),
                "cached_text": extract_text if isinstance(extract_text, str) else None,
            }
        )
    return sources


def _valid_positive_page_count(value: Any) -> int | None:
    if isinstance(value, bool) or not isinstance(value, int) or value <= 0:
        return None
    return value


def _page_text_sha256(
    text: Any,
    declared_pages: int | None,
) -> list[str] | None:
    page_count = _valid_positive_page_count(declared_pages)
    if page_count is None or not isinstance(text, str):
        return None
    pages = text.split("\f")
    if len(pages) != page_count:
        return None
    return [hashlib.sha256(page.encode("utf-8")).hexdigest() for page in pages]


def _full_page_ocr_result_proof(
    result: Any,
    declared_pages: int | None,
) -> dict[str, Any] | None:
    page_count = _valid_positive_page_count(declared_pages)
    if page_count is None or getattr(result, "error", None) not in {None, ""}:
        return None
    page_texts = getattr(result, "page_texts", None)
    statuses = getattr(result, "page_statuses", None)
    image_sha256 = getattr(result, "page_image_sha256", None)
    source_pages = _valid_positive_page_count(
        getattr(result, "source_pages", None)
    )
    if (
        not isinstance(page_texts, tuple)
        or not isinstance(statuses, tuple)
        or not isinstance(image_sha256, tuple)
        or len(page_texts) != page_count
        or len(statuses) != page_count
        or len(image_sha256) != page_count
        or source_pages != page_count
    ):
        return None
    normalized_texts = [str(text or "").strip() for text in page_texts]
    text_sha256 = [
        hashlib.sha256(text.encode("utf-8")).hexdigest()
        for text in normalized_texts
    ]
    for status, page_text, image_digest in zip(
        statuses,
        normalized_texts,
        image_sha256,
        strict=True,
    ):
        if status not in {"text", "blank"}:
            return None
        if status == "text" and not page_text:
            return None
        if status == "blank" and page_text:
            return None
        if not isinstance(image_digest, str) or not re.fullmatch(
            r"[0-9a-f]{64}", image_digest
        ):
            return None
    return {
        "ocr_page_statuses": list(statuses),
        "ocr_page_image_sha256": list(image_sha256),
        "ocr_page_text_sha256": text_sha256,
        "ocr_page_proof_version": "ocr-page-proof-v1",
        "ocr_source_pages": source_pages,
        "ocr_error": None,
        "ocr_blank_pages": [
            index
            for index, status in enumerate(statuses, start=1)
            if status == "blank"
        ],
    }


def _full_page_pdf_parse_proof_valid(
    parsed: dict[str, Any],
    *,
    expected_policy: str,
) -> bool:
    declared_pages = _valid_positive_page_count(parsed.get("pages"))
    return _parse_cache_policy_valid(
        {
            "ocr_policy": expected_policy,
            "declared_pages": declared_pages,
            "page_text_count": parsed.get("ocr_page_text_count"),
        },
        {"base": parsed},
        expected_policy=expected_policy,
    )


def _parse_cache_policy_valid(
    data: dict[str, Any],
    parsed: dict[str, Any],
    *,
    expected_policy: str,
) -> bool:
    if data.get("ocr_policy") != expected_policy:
        return False
    if expected_policy not in OCR_POLICIES_FULL_PAGE:
        return True
    base = parsed.get("base")
    if not isinstance(base, dict):
        return False
    if base.get("doc_type") != "pdf":
        # A one-page image or native text document is already parsed in full;
        # source-page coverage fields apply only to multi-page PDF OCR.
        return True
    declared_pages = _valid_positive_page_count(base.get("pages"))
    cached_declared_pages = _valid_positive_page_count(data.get("declared_pages"))
    ocr_pages = _valid_positive_page_count(base.get("ocr_pages"))
    page_text_count = _valid_positive_page_count(base.get("ocr_page_text_count"))
    cached_page_text_count = _valid_positive_page_count(data.get("page_text_count"))
    common_valid = (
        declared_pages is not None
        and cached_declared_pages == declared_pages
        and ocr_pages == declared_pages
        and page_text_count == declared_pages
        and cached_page_text_count == declared_pages
        and base.get("ocr_page_mapping") == "source_page_all"
    )
    if not common_valid:
        return common_valid
    statuses = base.get("ocr_page_statuses")
    image_sha256 = base.get("ocr_page_image_sha256")
    text_sha256 = base.get("ocr_page_text_sha256")
    extract_page_sha256 = base.get("ocr_extract_page_sha256")
    blank_pages = base.get("ocr_blank_pages")
    if (
        base.get("ocr_page_proof_version") != "ocr-page-proof-v1"
        or _valid_positive_page_count(base.get("ocr_source_pages"))
        != declared_pages
        or base.get("ocr_error") not in {None, ""}
        or not isinstance(statuses, list)
        or not isinstance(image_sha256, list)
        or not isinstance(text_sha256, list)
        or not isinstance(extract_page_sha256, list)
        or not isinstance(blank_pages, list)
        or len(statuses) != declared_pages
        or len(image_sha256) != declared_pages
        or len(text_sha256) != declared_pages
        or len(extract_page_sha256) != declared_pages
    ):
        return False
    empty_text_sha256 = hashlib.sha256(b"").hexdigest()
    for status, image_digest, text_digest in zip(
        statuses,
        image_sha256,
        text_sha256,
        strict=True,
    ):
        if status not in {"text", "blank"}:
            return False
        if not isinstance(image_digest, str) or not re.fullmatch(
            r"[0-9a-f]{64}", image_digest
        ):
            return False
        if not isinstance(text_digest, str) or not re.fullmatch(
            r"[0-9a-f]{64}", text_digest
        ):
            return False
        if status == "blank" and text_digest != empty_text_sha256:
            return False
        if status == "text" and text_digest == empty_text_sha256:
            return False
    if any(
        not isinstance(digest, str)
        or not re.fullmatch(r"[0-9a-f]{64}", digest)
        for digest in extract_page_sha256
    ):
        return False
    if blank_pages != [
        index
        for index, status in enumerate(statuses, start=1)
        if status == "blank"
    ]:
        return False
    extract_text = base.get("extract_text")
    return not isinstance(extract_text, str) or _page_text_sha256(
        extract_text,
        declared_pages,
    ) == extract_page_sha256


def _load_parse_cache(
    digest: str,
    *,
    source_hint: str | None = None,
) -> dict[str, Any] | None:
    expected_policy = _ocr_cache_policy(source_hint)
    path = _parse_cache_path(digest, source_hint)
    if not path.exists():
        return None
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError, UnicodeError):
        logger.warning("ignoring unreadable ingest parse cache: %s", path, exc_info=True)
        return None
    if (
        not isinstance(data, dict)
        or data.get("parser_version") != PARSER_VERSION
        or not isinstance(data.get("sha256"), str)
        or not hmac.compare_digest(data["sha256"].lower(), str(digest).lower())
    ):
        return None
    parsed = data.get("parsed")
    if not isinstance(parsed, dict) or not _parse_cache_policy_valid(
        data,
        parsed,
        expected_policy=expected_policy,
    ):
        return None
    result = dict(parsed)
    sidecar = data.get("extract_text_sidecar")
    if sidecar is None:
        # Compatibility with caches written before extracted text was split
        # from JSON metadata.  They remain valid until naturally replaced.
        base = result.get("base")
        if expected_policy in OCR_POLICIES_FULL_PAGE and (
            not isinstance(base, dict)
            or not isinstance(base.get("extract_text"), str)
        ):
            return None
        return result
    if not isinstance(sidecar, dict):
        return None
    text_path = _parse_cache_text_path(digest, source_hint)
    if str(sidecar.get("filename") or "") != text_path.name:
        return None
    try:
        expected_bytes = int(sidecar["bytes"])
        if expected_bytes < 0 or text_path.stat().st_size != expected_bytes:
            return None
        encoded_text = text_path.read_bytes()
        if len(encoded_text) != expected_bytes:
            return None
        expected_sha256 = sidecar["sha256"]
        if not isinstance(expected_sha256, str) or not hmac.compare_digest(
            hashlib.sha256(encoded_text).hexdigest(),
            expected_sha256.lower(),
        ):
            return None
        extracted_text = encoded_text.decode("utf-8")
    except (KeyError, OSError, UnicodeError, TypeError, ValueError):
        return None
    base = result.get("base")
    if not isinstance(base, dict):
        return None
    hydrated_base = dict(base)
    hydrated_base["extract_text"] = extracted_text
    result["base"] = hydrated_base
    if not _parse_cache_policy_valid(
        data,
        result,
        expected_policy=expected_policy,
    ):
        return None
    return result


def _save_parse_cache(
    digest: str,
    parsed: dict[str, Any],
    *,
    source_hint: str | None = None,
    write_guard: Callable[..., Any] | None = None,
) -> None:
    _guarded_write(
        write_guard,
        PARSE_CACHE_DIR.mkdir,
        parents=True,
        exist_ok=True,
        mode=0o700,
    )
    ocr_policy = _ocr_cache_policy(source_hint)
    path = _parse_cache_path(digest, source_hint)
    text_path = _parse_cache_text_path(digest, source_hint)
    temp_suffix = f".{os.getpid()}.{threading.get_ident()}.tmp"
    temp = path.with_suffix(path.suffix + temp_suffix)
    temp_text = text_path.with_suffix(text_path.suffix + temp_suffix)
    parsed_for_disk = dict(parsed)
    base = parsed_for_disk.get("base")
    extracted_text_bytes: bytes | None = None
    if isinstance(base, dict):
        persisted_base = dict(base)
        candidate = persisted_base.pop("extract_text", None)
        if isinstance(candidate, str):
            encoded_candidate = candidate.encode("utf-8")
            if len(encoded_candidate) >= PARSE_CACHE_TEXT_SIDECAR_BYTES:
                extracted_text_bytes = encoded_candidate
            else:
                persisted_base["extract_text"] = candidate
        parsed_for_disk["base"] = persisted_base
    payload = {
        "parser_version": PARSER_VERSION,
        "ocr_policy": ocr_policy,
        "sha256": str(digest),
        "saved_at": time.time(),
        "parsed": parsed_for_disk,
    }
    if ocr_policy in OCR_POLICIES_FULL_PAGE:
        persisted_base = parsed_for_disk.get("base")
        payload["declared_pages"] = (
            _valid_positive_page_count(persisted_base.get("pages"))
            if isinstance(persisted_base, dict)
            else None
        )
        payload["page_text_count"] = (
            _valid_positive_page_count(persisted_base.get("ocr_page_text_count"))
            if isinstance(persisted_base, dict)
            else None
        )
        # Never publish a drawing/standard cache unless it proves complete
        # source-page coverage.  A later upload must retry OCR instead of
        # reusing a partial result merely because the file SHA and policy match.
        if not _parse_cache_policy_valid(
            payload,
            parsed,
            expected_policy=ocr_policy,
        ):
            return
    try:
        if extracted_text_bytes is not None:
            temp_text.write_bytes(extracted_text_bytes)
            payload["extract_text_sidecar"] = {
                "filename": text_path.name,
                "bytes": len(extracted_text_bytes),
                "encoding": "utf-8",
                "sha256": hashlib.sha256(extracted_text_bytes).hexdigest(),
            }
        temp.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

        def _publish_cache() -> None:
            if extracted_text_bytes is not None:
                temp_text.replace(text_path)
                try:
                    text_path.chmod(0o600)
                except OSError:
                    pass
            temp.replace(path)
            try:
                path.chmod(0o600)
            except OSError:
                pass

        _guarded_write(write_guard, _publish_cache)
    finally:
        temp.unlink(missing_ok=True)
        temp_text.unlink(missing_ok=True)


def _extract_text_bytes(ext: str, content: bytes) -> dict[str, Any]:
    if ext in {"txt", "md"}:
        text = content.decode("utf-8", errors="ignore")
        return {
            "doc_type": ext,
            "pages": 1,
            "text_bytes": len(text.encode("utf-8")),
            "extract_text": text,
        }
    if ext == "pdf":
        reader = PdfReader(BytesIO(content))
        texts = []
        for page in reader.pages:
            # Use form-feed as a page boundary so evidence search can recover page numbers
            # without adding alnum/han noise (which would affect OCR heuristics).
            texts.append(page.extract_text() or "")
        text = "\n\n\f\n\n".join(texts)
        return {
            "doc_type": "pdf",
            "pages": len(reader.pages),
            "text_bytes": len(text.encode("utf-8")),
            "extract_text": text,
        }
    return {"doc_type": ext or "unknown", "pages": None, "text_bytes": None}


async def _persist_upload_file(
    uf: UploadFile,
    *,
    target_dir: Path,
    write_guard: Callable[..., Any] | None = None,
) -> tuple[Path | None, str | None, int]:
    filename = Path(str(uf.filename or "upload.bin")).name or "upload.bin"
    temp_name = f".upload_{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%S%f')}_{filename}"
    temp_path = target_dir / temp_name
    digest = hashlib.sha256()
    total_bytes = 0
    _guarded_write(write_guard, lambda: None)
    try:
        await uf.seek(0)
    except (OSError, ValueError):
        logger.warning("upload stream could not be rewound before persistence", exc_info=True)
    try:
        with temp_path.open("wb") as fh:
            while True:
                chunk = await uf.read(1024 * 1024)
                if not chunk:
                    break
                total_bytes += len(chunk)
                if total_bytes > MAX_UPLOAD_BYTES:
                    raise HTTPException(
                        status_code=413,
                        detail={
                            "code": "UPLOAD_TOO_LARGE",
                            "filename": filename,
                            "max_bytes": MAX_UPLOAD_BYTES,
                        },
                    )
                digest.update(chunk)
                fh.write(chunk)
        try:
            await uf.seek(0)
        except (OSError, ValueError):
            logger.warning("upload stream could not be rewound after persistence", exc_info=True)
        if total_bytes <= 0:
            return None, None, 0
        digest_hex = digest.hexdigest()
        out_path = target_dir / f"{digest_hex}_{filename}"
        published = _publish_temp_exclusive(
            temp_path,
            out_path,
            write_guard=write_guard,
        )
        if not published and (
            out_path.is_symlink()
            or not out_path.is_file()
            or _file_sha256(out_path) != digest_hex
        ):
            raise HTTPException(
                status_code=409,
                detail={
                    "code": "CONTENT_ADDRESS_COLLISION",
                    "sha256": digest_hex,
                },
            )
        return out_path, digest_hex, total_bytes
    finally:
        try:
            temp_path.unlink(missing_ok=True)
        except OSError:
            pass


def _meta_to_text(parsed_type: str | None, parsed_meta: Any) -> str:
    if parsed_type == "cad" and isinstance(parsed_meta, dict):
        layers = parsed_meta.get("layers_count")
        entities = parsed_meta.get("entities_count")
        inserts = parsed_meta.get("insert_blocks") or {}
        topo = parsed_meta.get("topology") if isinstance(parsed_meta.get("topology"), dict) else {}
        top_blocks = []
        if isinstance(inserts, dict):
            for k, v in sorted(inserts.items(), key=lambda x: x[1], reverse=True)[:8]:
                top_blocks.append(f"{k}:{v}")
        topo_lines = []
        if topo:
            topo_lines.append(f"拓扑节点: {topo.get('nodes_count')}")
            topo_lines.append(f"拓扑边: {topo.get('edges_count')}")
            topo_lines.append(f"连通分量: {topo.get('components_count')}")
            topo_lines.append(f"端点: {topo.get('endpoint_count')}")
            topo_lines.append(f"主干长度: {topo.get('trunk_length')}")
            topo_lines.append(f"建议流水段: {topo.get('suggested_flow_segments')}")
            if topo.get("topology_confidence"):
                topo_lines.append(f"拓扑置信度: {topo.get('topology_confidence')}")
        topo_text = ("\n" + "\n".join(topo_lines)) if topo_lines else ""
        return (
            f"图纸类型: CAD(DXF ASCII)\n"
            f"图层数量: {layers}\n"
            f"实体数量: {entities}\n"
            f"块引用: {'; '.join(top_blocks)}"
            f"{topo_text}"
        )
    if parsed_type == "image" and isinstance(parsed_meta, dict):
        return (
            f"图纸类型: 图片\n"
            f"尺寸: {parsed_meta.get('width')}x{parsed_meta.get('height')}\n"
            f"模式: {parsed_meta.get('mode')}"
        )
    if parsed_type == "dwg" and isinstance(parsed_meta, dict):
        return f"图纸类型: DWG\n说明: {parsed_meta.get('note') or '暂不支持解析'}"
    return ""


def _normalize_source_hint(source_hint: str | None) -> str:
    return normalize_source_hint(source_hint)


def _normalize_library_scope(library_scope: str | None, source_hint: str | None = None) -> str:
    raw = str(library_scope or "").strip().lower()
    if not raw:
        raw = _normalize_source_hint(source_hint)
    aliases = {
        "case": CASE_LIBRARY_SCOPE,
        "case_library": CASE_LIBRARY_SCOPE,
        "template_library": CASE_LIBRARY_SCOPE,
        "benchmark": CASE_LIBRARY_SCOPE,
        "样板库": CASE_LIBRARY_SCOPE,
        "案例库": CASE_LIBRARY_SCOPE,
        "image": IMAGE_LIBRARY_SCOPE,
        "image_library": IMAGE_LIBRARY_SCOPE,
        "图片库": IMAGE_LIBRARY_SCOPE,
        "图库": IMAGE_LIBRARY_SCOPE,
    }
    return aliases.get(raw, raw)


def _normalize_bool(value: Any, default: bool = True) -> bool:
    if value is None:
        return bool(default)
    if isinstance(value, bool):
        return value
    text = str(value).strip().lower()
    if text in {"1", "true", "yes", "y", "on", "启用", "可用"}:
        return True
    if text in {"0", "false", "no", "n", "off", "禁用", "不可用"}:
        return False
    return bool(default)


def _classify_tags(filename: str | None, ext: str, parsed_type: str | None, source_hint: str | None = None) -> list[str]:
    return classify_document_tags(
        filename,
        ext,
        parsed_type,
        source_hint=source_hint,
    )


def _make_preview_image(src_path: Path, dst_path: Path, max_width: int = 1400) -> str | None:
    try:
        from PIL import Image

        with Image.open(src_path) as im:
            im = im.convert("RGB")
            if im.width > max_width:
                h = int(im.height * (max_width / max(1, im.width)))
                im = im.resize((max_width, max(1, h)))
            im.save(dst_path, format="PNG")
        return str(dst_path)
    except Exception:
        logger.warning("image preview generation failed for %s", src_path, exc_info=True)
        return None


def _make_preview_pdf_first_page(src_path: Path, dst_path: Path, scale: float = 2.0, max_width: int = 1600) -> str | None:
    try:
        import pypdfium2 as pdfium

        pdf = pdfium.PdfDocument(str(src_path))
        if len(pdf) <= 0:
            return None
        page = pdf[0]
        bitmap = page.render(scale=float(scale))
        im = bitmap.to_pil()
        try:
            pdf.close()
        except Exception:
            logger.warning("PDF preview source could not be closed: %s", src_path, exc_info=True)
        if im.width > max_width:
            from PIL import Image

            h = int(im.height * (max_width / max(1, im.width)))
            im = im.resize((max_width, max(1, h)), resample=Image.BICUBIC)
        im.save(dst_path, format="PNG")
        return str(dst_path)
    except Exception:
        logger.warning("PDF preview generation failed for %s", src_path, exc_info=True)
        return None


async def _try_ocr(
    path: Path,
    ext: str,
    existing_text: str | None,
    *,
    source_hint: str | None = None,
    declared_pages: int | None = None,
) -> Any | None:
    """
    Best-effort OCR with a bounded, source-aware PDF policy.

    Drawing and standard PDFs are OCRed through the page count declared by the
    parser, even when embedded-font extraction produced long but unusable text.
    Other PDFs retain the fast ten-page scanned-document heuristic.
    """
    base = (existing_text or "").strip()
    normalized_hint = _normalize_source_hint(source_hint)
    full_page_ocr = ext == "pdf" and _ocr_cache_policy(
        normalized_hint
    ) in OCR_POLICIES_FULL_PAGE
    # Drawings and standards can contain plenty of embedded-font gibberish.
    # Do not treat length alone as proof that page text is machine-readable.
    if len(base) >= 200 and not full_page_ocr:
        return None
    try:
        from backend.zhifei_autoplan.ocr_runtime import (
            guess_ocr_lang,
            is_tesseract_available,
            is_text_probably_scanned,
            ocr_pdf_path,
        )
    except Exception:
        logger.warning("OCR runtime unavailable", exc_info=True)
        return None
    if not is_tesseract_available():
        return None

    if ext == "pdf":
        if full_page_ocr:
            # Accept only the page count produced by the PDF parser.  Never use
            # an arbitrary string/float/user option as an OCR expansion bound.
            if (
                isinstance(declared_pages, bool)
                or not isinstance(declared_pages, int)
                or declared_pages <= 0
            ):
                logger.warning(
                    "full-page OCR skipped because declared page count is unavailable"
                )
                return None
            max_pages = declared_pages
            stop_on_catalog = False
        else:
            max_pages = 10
            stop_on_catalog = True
        if not full_page_ocr and not is_text_probably_scanned(
            base,
            min_han=10,
            min_alnum=30,
        ):
            return None
        lang = guess_ocr_lang(prefer_chinese=True)
        res = await asyncio.to_thread(
            ocr_pdf_path,
            str(path),
            max_pages=max_pages,
            scale=2.2,
            lang=lang,
            stop_on_catalog=stop_on_catalog,
        )
        if full_page_ocr:
            result_pages = _valid_positive_page_count(getattr(res, "pages", None))
            page_texts = getattr(res, "page_texts", None)
            page_text_count = len(page_texts) if isinstance(page_texts, tuple) else None
            if result_pages != declared_pages or page_text_count != declared_pages:
                logger.warning(
                    "full-page OCR rejected because source-page coverage is incomplete "
                    "declared_pages=%s result_pages=%s page_text_count=%s",
                    declared_pages,
                    result_pages,
                    page_text_count,
                )
                return None
            if _full_page_ocr_result_proof(res, declared_pages) is None:
                logger.warning(
                    "full-page OCR rejected because per-page proof is incomplete"
                )
                return None
        return res if res and (res.text or res.page_texts) else None

    if ext in {"png", "jpg", "jpeg"}:
        try:
            import pytesseract
            from PIL import Image

            from backend.zhifei_autoplan.ocr_runtime import guess_ocr_lang

            lang = guess_ocr_lang(prefer_chinese=True)

            def _ocr():
                with Image.open(path) as im:
                    return pytesseract.image_to_string(im, lang=lang)

            txt = await asyncio.to_thread(_ocr)
            return txt.strip() if txt and txt.strip() else None
        except Exception:
            logger.warning("image OCR failed for %s", path, exc_info=True)
            return None
    return None


def _merge_pdf_ocr_pages(
    extracted_text: str | None,
    ocr_page_texts: Any,
    declared_pages: Any,
) -> tuple[str, bool]:
    """Overlay prefix-page OCR without changing source PDF page ordinals.

    PDF extraction already emits one form-feed-delimited entry per source page.
    OCR may cover only a bounded prefix, so it must be merged into those same
    page slots.  Appending a second OCR document would make page numbers and
    evidence offsets unverifiable.
    """

    try:
        page_count = int(declared_pages)
    except (TypeError, ValueError):
        return str(extracted_text or ""), False
    if page_count <= 0:
        return str(extracted_text or ""), False

    base_pages = str(extracted_text or "").split("\f")
    if len(base_pages) != page_count:
        return str(extracted_text or ""), False
    if not isinstance(ocr_page_texts, (list, tuple)):
        return str(extracted_text or ""), False

    merged_pages: list[str] = []
    for index, base_page in enumerate(base_pages):
        base = str(base_page or "").strip()
        ocr = (
            str(ocr_page_texts[index] or "").strip()
            if index < len(ocr_page_texts)
            else ""
        )
        if ocr and ocr not in base:
            merged_pages.append(f"{base}\n\n{ocr}".strip() if base else ocr)
        else:
            merged_pages.append(base)
    return "\n\n\f\n\n".join(merged_pages), True

async def _handle_upload(
    files: list[UploadFile],
    project_id: str | None = None,
    source_hint: str | None = None,
    session_id: str | None = None,
    workspace_dir: str | None = None,
    library_scope: str | None = None,
    project_type: str | None = None,
    title: str | None = None,
    tags: str | None = None,
    chapter_scope: str | None = None,
    process_scope: str | None = None,
    summary: str | None = None,
    style_profile: str | None = None,
    caption: str | None = None,
    description: str | None = None,
    usable: bool | str | None = True,
    _write_guard: Callable[..., Any] | None = None,
):
    if not files:
        raise HTTPException(status_code=400, detail="no files uploaded")

    workspace_kwargs: dict[str, Any] = {
        "session_id": session_id,
        "workspace_dir": workspace_dir,
    }
    if callable(_write_guard):
        workspace_kwargs["write_guard"] = _write_guard
    workspace = _resolve_workspace_context(**workspace_kwargs)
    paths_kwargs: dict[str, Any] = {}
    if callable(_write_guard):
        paths_kwargs["write_guard"] = _write_guard
    ws_paths = workspace_paths(workspace["workspace_dir"], **paths_kwargs)
    day = datetime.now(timezone.utc).strftime("%Y%m%d")
    target_dir = ws_paths["uploads"] / day
    _guarded_write(_write_guard, target_dir.mkdir, parents=True, exist_ok=True)
    extract_dir = ws_paths["extracts"]
    preview_dir = ws_paths["previews"]
    audit_file = ws_paths["ingest_audit"]
    _guarded_write(_write_guard, extract_dir.mkdir, parents=True, exist_ok=True)
    _guarded_write(_write_guard, preview_dir.mkdir, parents=True, exist_ok=True)
    _guarded_write(
        _write_guard,
        audit_file.parent.mkdir,
        parents=True,
        exist_ok=True,
    )

    records = []
    rejected: list[dict[str, Any]] = []
    warnings: list[dict[str, Any]] = []
    cache_hits = 0
    seen_digests: dict[str, Path] = {}
    pid = str(project_id).strip() if isinstance(project_id, str) and project_id.strip() else None
    normalized_hint = _normalize_source_hint(source_hint)
    normalized_library_scope = _normalize_library_scope(library_scope, normalized_hint)
    normalized_project_type = normalize_project_type(project_type)
    normalized_title = str(title or "").strip()[:240] or None
    normalized_tags = normalize_text_list(tags)
    normalized_chapter_scope = normalize_text_list(chapter_scope)
    normalized_process_scope = normalize_text_list(process_scope)
    normalized_summary = str(summary or "").strip()[:1000]
    normalized_style_profile = str(style_profile or "").strip()[:1000]
    normalized_caption = str(caption or "").strip()[:500]
    normalized_description = str(description or "").strip()[:1000]
    normalized_usable = _normalize_bool(usable, default=True)
    if normalized_library_scope in {CASE_LIBRARY_SCOPE, IMAGE_LIBRARY_SCOPE} and not normalized_project_type:
        if normalized_library_scope == IMAGE_LIBRARY_SCOPE:
            raise HTTPException(status_code=400, detail="image library upload requires valid project_type")
        raise HTTPException(status_code=400, detail="case library upload requires valid project_type")
    for uf in files:
        ext = _ext(uf.filename or "")
        if normalized_library_scope == IMAGE_LIBRARY_SCOPE and ext not in {"png", "jpg", "jpeg", "webp", "gif"}:
            raise HTTPException(status_code=400, detail="image library upload requires image files")

        out_path, digest, total_bytes = await _persist_upload_file(
            uf,
            target_dir=target_dir,
            write_guard=_write_guard,
        )
        if not out_path or not digest or total_bytes <= 0:
            rejected.append({"filename": uf.filename, "code": "EMPTY_FILE"})
            continue

        prior_path = seen_digests.get(digest)
        if prior_path is not None:
            if out_path != prior_path:
                _guarded_write(_write_guard, out_path.unlink, missing_ok=True)
            rejected.append(
                {
                    "filename": uf.filename,
                    "code": "DUPLICATE_FILE",
                    "sha256": digest,
                    "duplicate_of": str(prior_path),
                }
            )
            continue
        seen_digests[digest] = out_path

        cached = await asyncio.to_thread(
            _load_parse_cache,
            digest,
            source_hint=normalized_hint,
        )
        cache_hit = isinstance(cached, dict)
        if cache_hit:
            cache_hits += 1
            parsed = dict(cached.get("base") or {})
            parsed_type = cached.get("parsed_type")
            parsed_meta = cached.get("parsed_meta")
        else:
            try:
                parsed = await _extract_text_path_bounded(ext, out_path, total_bytes)
            except Exception as exc:  # noqa: BLE001 - parser boundary rejects and records the file
                logger.warning(
                    "ingest parser rejected %s: %s",
                    uf.filename,
                    type(exc).__name__,
                )
                rejected.append(
                    {
                        "filename": uf.filename,
                        "code": "FILE_PARSE_TIMEOUT" if isinstance(exc, asyncio.TimeoutError) else "FILE_PARSE_FAILED",
                        "sha256": digest,
                        "file_id": digest,
                        "saved_as": str(out_path),
                        "error_type": type(exc).__name__,
                    }
                )
                continue
            parsed_type = None
            parsed_meta = None
            if parsed.get("extract_text") is None and ext not in {"txt", "md", "pdf", "doc", "docx"}:
                try:
                    def _parse_unified(parse_path: Path) -> dict[str, Any]:
                        from modules.parser.parser_unify import UnifiedParser

                        return UnifiedParser(str(parse_path)).parse()

                    uret = await asyncio.wait_for(
                        asyncio.to_thread(_parse_unified, out_path),
                        timeout=PARSE_TIMEOUT_SECONDS,
                    )
                    parsed_type = uret.get("type")
                    parsed_meta = uret.get("meta")
                    utext = uret.get("text")
                    if isinstance(utext, str) and utext.strip():
                        parsed["extract_text"] = utext
                    else:
                        meta_text = _meta_to_text(parsed_type, parsed_meta)
                        if meta_text.strip():
                            parsed["extract_text"] = meta_text
                except Exception as exc:  # noqa: BLE001 - optional parser failure is recorded in metadata
                    logger.warning(
                        "unified parser failed for %s: %s",
                        uf.filename,
                        type(exc).__name__,
                    )
                    parsed_meta = {"error_type": type(exc).__name__}

            # OCR is cached with the parser result so repeated 268 MiB imports
            # do not repeat the most expensive extraction path.
            try:
                ocr_result = await _try_ocr(
                    out_path,
                    ext,
                    parsed.get("extract_text"),
                    source_hint=normalized_hint,
                    declared_pages=parsed.get("pages"),
                )
                if ocr_result:
                    if ext == "pdf" and hasattr(ocr_result, "page_texts"):
                        page_texts = tuple(getattr(ocr_result, "page_texts", ()))
                        merged, mapped = _merge_pdf_ocr_pages(
                            parsed.get("extract_text"),
                            page_texts,
                            parsed.get("pages"),
                        )
                        parsed["ocr_pages"] = int(
                            getattr(ocr_result, "pages", 0) or 0
                        )
                        parsed["ocr_page_text_count"] = len(page_texts)
                        if _ocr_cache_policy(
                            normalized_hint
                        ) in OCR_POLICIES_FULL_PAGE:
                            proof = _full_page_ocr_result_proof(
                                ocr_result,
                                _valid_positive_page_count(parsed.get("pages")),
                            )
                            if proof is not None:
                                parsed.update(proof)
                        if mapped and _ocr_cache_policy(
                            normalized_hint
                        ) in OCR_POLICIES_FULL_PAGE:
                            parsed["ocr_page_mapping"] = "source_page_all"
                        else:
                            parsed["ocr_page_mapping"] = (
                                "source_page_prefix" if mapped else "unavailable"
                            )
                        if mapped:
                            parsed["extract_text"] = merged
                            if _ocr_cache_policy(
                                normalized_hint
                            ) in OCR_POLICIES_FULL_PAGE:
                                extract_page_sha256 = _page_text_sha256(
                                    merged,
                                    _valid_positive_page_count(
                                        parsed.get("pages")
                                    ),
                                )
                                if extract_page_sha256 is not None:
                                    parsed["ocr_extract_page_sha256"] = (
                                        extract_page_sha256
                                    )
                    elif isinstance(ocr_result, str):
                        base = str(parsed.get("extract_text") or "").strip()
                        parsed["extract_text"] = (
                            f"{base}\n\n{ocr_result.strip()}".strip()
                            if base
                            else ocr_result.strip()
                        )
                    parsed["text_bytes"] = len(
                        str(parsed.get("extract_text") or "").encode("utf-8")
                    )
            except Exception as exc:  # noqa: BLE001 - OCR boundary preserves source-page text on failure
                logger.warning(
                    "OCR enrichment failed for %s: %s",
                    uf.filename,
                    type(exc).__name__,
                )
            await asyncio.to_thread(
                _save_parse_cache,
                digest,
                {
                    "base": dict(parsed),
                    "parsed_type": parsed_type,
                    "parsed_meta": parsed_meta,
                },
                source_hint=normalized_hint,
                write_guard=_write_guard,
            )

        full_page_policy = _ocr_cache_policy(normalized_hint)
        if (
            ext == "pdf"
            and full_page_policy in OCR_POLICIES_FULL_PAGE
            and not _full_page_pdf_parse_proof_valid(
                parsed,
                expected_policy=full_page_policy,
            )
        ):
            rejected.append(
                {
                    "filename": uf.filename,
                    "code": (
                        "STANDARD_FULL_PAGE_OCR_REQUIRED"
                        if full_page_policy == OCR_POLICY_STANDARD
                        else "DRAWING_FULL_PAGE_OCR_REQUIRED"
                    ),
                    "sha256": digest,
                    "file_id": digest,
                    "saved_as": str(out_path),
                }
            )
            continue

        meaningful_text = str(parsed.get("extract_text") or "").strip()
        if normalized_hint in {"tender_qa", "boq"} and not meaningful_text:
            rejected.append(
                {
                    "filename": uf.filename,
                    "code": "UNPARSED_MANDATORY_FILE",
                    "sha256": digest,
                    "file_id": digest,
                    "saved_as": str(out_path),
                }
            )
            continue

        # Preview (best-effort): reuse a content-addressed thumbnail whenever
        # possible.  PDFium is native code and a malformed page can segfault;
        # keep it in the same disposable process boundary as document parsing.
        preview_path = None
        preview_cache_hit = False
        preview_warning: dict[str, Any] | None = None
        preview_temp: Path | None = None
        try:
            preview_name = f"{digest}_preview.png"
            preview_out = preview_dir / preview_name
            if (
                not preview_out.is_symlink()
                and preview_out.is_file()
                and preview_out.stat().st_size > 0
            ):
                preview_path = str(preview_out)
                preview_cache_hit = True
            elif ext in {"png", "jpg", "jpeg"}:
                preview_temp = preview_dir / (
                    f".{digest}.{os.getpid()}.{threading.get_ident()}.preview.tmp"
                )
                rendered = await asyncio.to_thread(
                    _make_preview_image,
                    out_path,
                    preview_temp,
                )
                if rendered:
                    published = _publish_temp_exclusive(
                        preview_temp,
                        preview_out,
                        write_guard=_write_guard,
                    )
                    if published or (
                        not preview_out.is_symlink()
                        and preview_out.is_file()
                        and preview_out.stat().st_size > 0
                    ):
                        preview_path = str(preview_out)
            elif ext == "pdf":
                preview_temp = preview_dir / (
                    f".{digest}.{os.getpid()}.{threading.get_ident()}.preview.tmp"
                )
                rendered = await _run_isolated_process(
                    _make_preview_pdf_first_page,
                    out_path,
                    preview_temp,
                    2.0,
                    timeout=PARSE_TIMEOUT_SECONDS,
                )
                if rendered:
                    published = _publish_temp_exclusive(
                        preview_temp,
                        preview_out,
                        write_guard=_write_guard,
                    )
                    if published or (
                        not preview_out.is_symlink()
                        and preview_out.is_file()
                        and preview_out.stat().st_size > 0
                    ):
                        preview_path = str(preview_out)
        except JobLeaseLostError:
            raise
        except BrokenProcessPool:
            preview_warning = {
                "code": "PDF_PREVIEW_WORKER_CRASHED",
                "message": "PDF 首页预览进程异常退出，正文解析结果已保留。",
                "action": "系统已隔离并回收异常进程；可继续使用正文，或重新生成预览。",
                "filename": str(uf.filename or ""),
            }
            preview_path = None
        except asyncio.TimeoutError:
            preview_warning = {
                "code": "PDF_PREVIEW_TIMEOUT",
                "message": "PDF 首页预览超时，正文解析结果已保留。",
                "action": "可继续使用正文，或稍后重新生成预览。",
                "filename": str(uf.filename or ""),
            }
            preview_path = None
        except Exception as exc:  # noqa: BLE001 - preview boundary records a stable warning
            logger.warning(
                "preview generation failed for %s: %s",
                uf.filename,
                type(exc).__name__,
            )
            preview_warning = {
                "code": "PDF_PREVIEW_UNAVAILABLE",
                "message": "PDF 首页预览生成失败，正文解析结果已保留。",
                "action": "可继续使用正文，或稍后重新生成预览。",
                "filename": str(uf.filename or ""),
            }
            preview_path = None
        finally:
            if preview_temp is not None:
                preview_temp.unlink(missing_ok=True)
        if ext == "pdf" and not preview_path and preview_warning is None:
            preview_warning = {
                "code": "PDF_PREVIEW_UNAVAILABLE",
                "message": "PDF 首页预览生成失败，正文解析结果已保留。",
                "action": "可继续使用正文，或稍后重新生成预览。",
                "filename": str(uf.filename or ""),
            }
        if preview_warning is not None:
            warnings.append(preview_warning)
        preview_sha256 = None
        if preview_path:
            try:
                preview_sha256 = _file_sha256(Path(preview_path))
            except OSError:
                preview_path = None
                preview_cache_hit = False
                warnings.append(
                    {
                        "code": "PDF_PREVIEW_INTEGRITY_UNAVAILABLE",
                        "message": "PDF 预览文件无法完成摘要校验。",
                        "filename": str(uf.filename or ""),
                    }
                )

        extract_path = None
        extract_text_sha256 = None
        if parsed.get("extract_text") is not None:
            extract_text = str(parsed["extract_text"])
            extract_bytes = extract_text.encode("utf-8")
            extract_text_sha256 = hashlib.sha256(extract_bytes).hexdigest()
            # Bind both the source identity and extracted-text identity into the
            # durable name.  The same source may legitimately yield different
            # policy-specific extracts (ordinary vs full-page OCR), and neither
            # variant may overwrite the other.
            extract_path = extract_dir / (
                f"{digest}_{extract_text_sha256}.txt"
            )
            extract_temp = extract_dir / (
                f".{digest}.{os.getpid()}.{threading.get_ident()}.extract.tmp"
            )
            try:
                await asyncio.to_thread(extract_temp.write_bytes, extract_bytes)
                published = _publish_temp_exclusive(
                    extract_temp,
                    extract_path,
                    write_guard=_write_guard,
                )
                if not published and (
                    extract_path.is_symlink()
                    or not extract_path.is_file()
                    or _file_sha256(extract_path) != extract_text_sha256
                ):
                    raise HTTPException(
                        status_code=409,
                        detail={
                            "code": "EXTRACT_CONTENT_COLLISION",
                            "sha256": digest,
                        },
                    )
            finally:
                extract_temp.unlink(missing_ok=True)
            parsed.pop("extract_text", None)

        rec = {
            "ts": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
            "module": "ingest",
            "project_id": pid,
            "workspace_dir": workspace["workspace_dir"],
            "filename": uf.filename,
            "saved_as": str(out_path),
            "bytes": int(total_bytes),
            "sha256": digest,
            "file_id": digest,
            "parser_version": PARSER_VERSION,
            "ocr_cache_policy": _ocr_cache_policy(normalized_hint),
            "cache_hit": cache_hit,
            "preview_cache_hit": preview_cache_hit,
            "extract_saved_as": str(extract_path) if extract_path else None,
            "extract_text_sha256": extract_text_sha256,
            "preview_saved_as": preview_path,
            "preview_sha256": preview_sha256,
            **parsed,
            "parsed_type": parsed_type,
            "parsed_meta": parsed_meta,
            "source_hint": normalized_hint or None,
            "project_type": normalized_project_type or None,
            "library_scope": normalized_library_scope or None,
            "library_title": normalized_title,
            "library_tags": normalized_tags,
            "chapter_scope": normalized_chapter_scope,
            "process_scope": normalized_process_scope,
            "library_summary": normalized_summary,
            "library_style_profile": normalized_style_profile,
            "library_caption": normalized_caption,
            "library_description": normalized_description,
            "enabled": normalized_usable,
            "usable": normalized_usable,
            "tags": _classify_tags(uf.filename, ext, parsed_type, normalized_hint),
        }
        records.append(rec)

        def _append_audit_record(record: dict[str, Any]) -> None:
            with audit_file.open("a", encoding="utf-8") as f:
                f.write(json.dumps(record, ensure_ascii=False) + "\n")

        _guarded_write(_write_guard, _append_audit_record, rec)

    if not records:
        if rejected and all(item.get("code") == "EMPTY_FILE" for item in rejected):
            raise HTTPException(status_code=400, detail="all files are empty")
        raise HTTPException(
            status_code=422,
            detail={"code": "ALL_FILES_REJECTED", "rejected": rejected},
        )
    if normalized_hint in {"tender_qa", "boq", "standard"} and rejected:
        raise HTTPException(
            status_code=422,
            detail={
                "code": "MANDATORY_SOURCE_REJECTED",
                "message": "强制资料存在未解析或未完成全页 OCR 的文件，已阻止进入生成。",
                "accepted": records,
                "rejected": rejected,
            },
        )
    if rejected:
        warnings.append(
            {
                "code": "OPTIONAL_SOURCE_DEGRADED",
                "message": f"{len(rejected)} 个可选资料未完成解析。",
            }
        )
    return {
        "saved": records,
        "accepted": records,
        "rejected": rejected,
        "warnings": warnings,
        "cache_hits": cache_hits,
        "parser_version": PARSER_VERSION,
    }


class _SpoolUpload:
    def __init__(self, path: Path, filename: str) -> None:
        self.filename = filename
        self._handle = path.open("rb")

    async def read(self, size: int = -1) -> bytes:
        return self._handle.read(size)

    async def seek(self, offset: int, whence: int = 0) -> int:
        return self._handle.seek(offset, whence)

    async def close(self) -> None:
        self._handle.close()


def _process_spooled_entry(
    job_id: str,
    index: int,
    total: int,
    entry: dict[str, Any],
    options: dict[str, Any],
    lease_attempt_id: str,
    lease_owner_instance_id: str,
) -> dict[str, Any]:
    def _running_lease_active() -> bool:
        return job_lease_active(
            job_id,
            attempt_id=lease_attempt_id,
            owner_instance_id=lease_owner_instance_id,
            allowed_statuses={"running"},
        )

    def _running_write_guard(
        callback: Callable[..., Any],
        *args: Any,
        **kwargs: Any,
    ) -> Any:
        return run_with_job_lease(
            job_id,
            attempt_id=lease_attempt_id,
            owner_instance_id=lease_owner_instance_id,
            allowed_statuses={"running"},
            callback=callback,
            callback_args=tuple(args),
            callback_kwargs=dict(kwargs),
        )

    def _append_active_event(event: str, **fields: Any) -> bool:
        try:
            run_with_job_lease(
                job_id,
                attempt_id=lease_attempt_id,
                owner_instance_id=lease_owner_instance_id,
                callback=append_runtime_event,
                callback_args=(job_id, event),
                callback_kwargs=fields,
                allowed_statuses={"running"},
            )
            return True
        except JobLeaseLostError:
            return False

    current = get_job(job_id) or {}
    if (
        str(current.get("status") or "").strip().lower()
        in {"cancel_requested", "cancelled"}
        or not _running_lease_active()
    ):
        return {"cancelled": True, "index": index}
    filename = str(entry.get("filename") or "upload.bin")
    started = time.monotonic()
    started_record = transition_job(
        job_id,
        allowed_from={"running"},
        status="running",
        expected_attempt_id=lease_attempt_id,
        expected_owner_instance_id=lease_owner_instance_id,
        progress={
            "phase": "ingest",
            "stage": "ingest",
            "work_state": "processing_file",
            "current_file": filename,
        },
    )
    if started_record is None:
        return {"cancelled": True, "index": index}
    if not _append_active_event(
        "ingest_file_started",
        filename=filename,
        file_index=index,
        files_total=total,
    ):
        return {"cancelled": True, "index": index}
    upload = _SpoolUpload(Path(str(entry["path"])), filename)
    try:
        worker_options = dict(options)
        worker_options["_write_guard"] = _running_write_guard
        result = asyncio.run(_handle_upload([upload], **worker_options))
        if not _running_lease_active():
            return {"cancelled": True, "index": index}
        rows = result.get("accepted") if isinstance(result.get("accepted"), list) else result.get("saved") or []
        elapsed_seconds = round(time.monotonic() - started, 3)
        outcome = {
            "index": index,
            "filename": filename,
            "elapsed_seconds": elapsed_seconds,
            "accepted": [
                {**item, "elapsed_seconds": elapsed_seconds}
                for item in rows
                if isinstance(item, dict)
            ],
            "rejected": [
                {**item, "elapsed_seconds": elapsed_seconds}
                for item in (result.get("rejected") or [])
                if isinstance(item, dict)
            ],
            "warnings": [item for item in (result.get("warnings") or []) if isinstance(item, dict)],
            "cache_hits": int(result.get("cache_hits") or 0),
        }
        if not _append_active_event(
            "ingest_file_finished",
            filename=filename,
            ok=True,
            cache_hit=bool(result.get("cache_hits")),
            elapsed_seconds=elapsed_seconds,
        ):
            return {"cancelled": True, "index": index}
        for warning in outcome["warnings"]:
            if not _append_active_event(
                "ingest_warning",
                filename=filename,
                code=str(warning.get("code") or "INGEST_WARNING"),
                message=str(warning.get("message") or "")[:500],
            ):
                return {"cancelled": True, "index": index}
        return outcome
    except HTTPException as exc:
        if not _running_lease_active():
            return {"cancelled": True, "index": index}
        detail = exc.detail if isinstance(exc.detail, dict) else {
            "code": "FILE_PARSE_FAILED",
            "message": str(exc.detail),
        }
        rows = detail.get("rejected") if isinstance(detail.get("rejected"), list) else []
        rejected = [item for item in rows if isinstance(item, dict)]
        if not rejected:
            rejected = [{"filename": filename, "code": str(detail.get("code") or "FILE_PARSE_FAILED")}]
        elapsed_seconds = round(time.monotonic() - started, 3)
        rejected = [{**item, "elapsed_seconds": elapsed_seconds} for item in rejected]
        if not _append_active_event(
            "ingest_file_finished",
            filename=filename,
            ok=False,
            code=detail.get("code"),
            elapsed_seconds=elapsed_seconds,
        ):
            return {"cancelled": True, "index": index}
        return {
            "index": index,
            "filename": filename,
            "elapsed_seconds": elapsed_seconds,
            "accepted": [],
            "rejected": rejected,
            "warnings": [],
            "cache_hits": 0,
        }
    except Exception as exc:  # noqa: BLE001 - per-file worker returns a stable rejection envelope
        logger.warning(
            "ingest file worker failed for %s: %s",
            filename,
            type(exc).__name__,
        )
        if not _running_lease_active():
            return {"cancelled": True, "index": index}
        elapsed_seconds = round(time.monotonic() - started, 3)
        if not _append_active_event(
            "ingest_file_finished",
            filename=filename,
            ok=False,
            error_type=type(exc).__name__,
            elapsed_seconds=elapsed_seconds,
        ):
            return {"cancelled": True, "index": index}
        return {
            "index": index,
            "filename": filename,
            "elapsed_seconds": elapsed_seconds,
            "accepted": [],
            "rejected": [
                {
                    "filename": filename,
                    "code": "FILE_PARSE_FAILED",
                    "error_type": type(exc).__name__,
                    "elapsed_seconds": elapsed_seconds,
                }
            ],
            "warnings": [],
            "cache_hits": 0,
        }
    finally:
        asyncio.run(upload.close())


def _run_ingest_job(
    job_id: str,
    entries: list[dict[str, Any]],
    options: dict[str, Any],
    initial_rejected: list[dict[str, Any]],
) -> None:
    lease_record = acquire_job_lease(job_id)
    if lease_record is None:
        return
    lease_attempt_id = str(lease_record.get("attempt_id") or "")
    lease_owner_instance_id = str(lease_record.get("owner_instance_id") or "")
    if not lease_attempt_id or not lease_owner_instance_id:
        raise RuntimeError("job_lease_acquisition_invalid")

    def _lease_active() -> bool:
        return job_lease_active(
            job_id,
            attempt_id=lease_attempt_id,
            owner_instance_id=lease_owner_instance_id,
        )

    def _status() -> str:
        return str((get_job(job_id) or {}).get("status") or "").strip().lower()

    def _cancel_requested() -> bool:
        return _status() in {"cancel_requested", "cancelled"} or not _lease_active()

    def _append_active_event(event: str, **fields: Any) -> bool:
        try:
            run_with_job_lease(
                job_id,
                attempt_id=lease_attempt_id,
                owner_instance_id=lease_owner_instance_id,
                callback=append_runtime_event,
                callback_args=(job_id, event),
                callback_kwargs=fields,
                allowed_statuses={"running"},
            )
            return True
        except JobLeaseLostError:
            return False

    def _merge_active(**fields: Any) -> dict[str, Any]:
        updated = transition_job(
            job_id,
            allowed_from={"running"},
            status="running",
            expected_attempt_id=lease_attempt_id,
            expected_owner_instance_id=lease_owner_instance_id,
            **fields,
        )
        if updated is None:
            raise JobLeaseLostError("job_lease_lost")
        return updated

    def _mark_cancelled(*, completed: int = 0) -> None:
        if not _lease_active():
            return
        prior_progress = (get_job(job_id) or {}).get("progress") or {}
        transitioned = transition_job(
            job_id,
            allowed_from={"running", "cancel_requested"},
            status="cancelled",
            expected_attempt_id=lease_attempt_id,
            expected_owner_instance_id=lease_owner_instance_id,
            revoke_lease=True,
            error={
                "code": "INGEST_CANCELLED",
                "message": "用户已取消资料导入。",
                "action": "已完成解析与缓存予以保留；如需继续请显式重新导入。",
            },
            progress={
                "percent": min(99, int(prior_progress.get("percent") or 0)),
                "phase": "ingest",
                "stage": "cancelled",
                "work_state": "idle",
                "current_file": None,
                "detail": "资料导入已取消。",
            },
        )
        if transitioned is not None:
            append_runtime_event(job_id, "ingest_cancelled", completed=completed)

    accepted: list[dict[str, Any]] = []
    rejected = list(initial_rejected)
    warnings: list[dict[str, Any]] = []
    cache_hits = 0
    total = len(entries)
    file_items: list[dict[str, Any]] = [
        {
            "index": index,
            "filename": str(entry.get("filename") or "upload.bin"),
            "file_id": entry.get("file_id"),
            "bytes": int(entry.get("bytes") or 0),
            "status": "queued",
            "cache_hit": False,
            "elapsed_seconds": None,
        }
        for index, entry in enumerate(entries, start=1)
    ]
    file_items.extend(
        {
            "index": total + index,
            "filename": str(item.get("filename") or "upload.bin"),
            "status": "rejected",
            "code": str(item.get("code") or "UPLOAD_REJECTED"),
            "cache_hit": False,
            "elapsed_seconds": 0.0,
        }
        for index, item in enumerate(initial_rejected, start=1)
        if isinstance(item, dict)
    )
    mandatory = _normalize_source_hint(options.get("source_hint")) in {
        "tender_qa",
        "boq",
        "standard",
    }
    try:
        if _cancel_requested():
            _mark_cancelled(completed=0)
            return
        _merge_active(
            progress={
                "phase": "ingest",
                "stage": "ingest",
                "work_state": "processing_file",
                "percent": 0,
                "files": {
                    "completed": 0,
                    "accepted": 0,
                    "rejected": len(rejected),
                    "total": total + len(initial_rejected),
                    "items": file_items,
                },
            },
        )
        if not _append_active_event("ingest_started", files_total=total):
            return
        completed = 0
        max_workers = min(INGEST_WORKERS, max(1, total))
        with concurrent.futures.ThreadPoolExecutor(
            max_workers=max_workers,
            thread_name_prefix=f"ingest-{job_id[:8]}",
        ) as executor:
            entry_iter = iter(enumerate(entries, start=1))
            inflight: dict[concurrent.futures.Future, int] = {}
            active_files: dict[int, str] = {}

            def _submit_next() -> bool:
                if _cancel_requested():
                    return False
                try:
                    index, entry = next(entry_iter)
                except StopIteration:
                    return False
                filename = str(entry.get("filename") or "upload.bin")
                file_items[index - 1]["status"] = "processing"
                file_items[index - 1]["started_at"] = time.time()
                active_files[index] = filename
                future = executor.submit(
                    _process_spooled_entry,
                    job_id,
                    index,
                    total,
                    entry,
                    options,
                    lease_attempt_id,
                    lease_owner_instance_id,
                )
                inflight[future] = index
                return True

            for _ in range(max_workers):
                _submit_next()

            stop_processing = False
            while inflight:
                if _cancel_requested():
                    _mark_cancelled(completed=completed)
                    for pending in inflight:
                        pending.cancel()
                    break
                finished, _pending = concurrent.futures.wait(
                    tuple(inflight),
                    timeout=INGEST_HEARTBEAT_SECONDS,
                    return_when=concurrent.futures.FIRST_COMPLETED,
                )
                if not finished:
                    if _cancel_requested():
                        _mark_cancelled(completed=completed)
                        for pending in inflight:
                            pending.cancel()
                        break
                    heartbeat = heartbeat_job(
                        job_id,
                        activity=f"{len(active_files)} 个文件正在解析",
                        progress_updates={
                            "phase": "ingest",
                            "stage": "ingest",
                            "work_state": "processing_file",
                            "current_file": list(active_files.values()),
                        },
                        expected_attempt_id=lease_attempt_id,
                        expected_owner_instance_id=lease_owner_instance_id,
                        allowed_statuses={"running"},
                    )
                    if heartbeat is None:
                        raise JobLeaseLostError("job_lease_lost")
                    continue
                for future in finished:
                    index = inflight.pop(future)
                    active_files.pop(index, None)
                    outcome = future.result()
                    item = file_items[index - 1]
                    item["elapsed_seconds"] = outcome.get("elapsed_seconds")
                    item["finished_at"] = time.time()
                    item["cache_hit"] = bool(outcome.get("cache_hits"))
                    if outcome.get("cancelled"):
                        item["status"] = "cancelled"
                        _mark_cancelled(completed=completed)
                        for pending in inflight:
                            pending.cancel()
                        stop_processing = True
                        break
                    else:
                        accepted_rows = outcome.get("accepted") or []
                        rejected_rows = outcome.get("rejected") or []
                        item["status"] = "accepted" if accepted_rows else "rejected"
                        if accepted_rows:
                            first_accepted = accepted_rows[0]
                            item["pages"] = first_accepted.get("pages")
                            item["file_id"] = first_accepted.get("file_id") or item.get("file_id")
                            item["preview_cache_hit"] = bool(
                                first_accepted.get("preview_cache_hit")
                            )
                        if rejected_rows:
                            item["code"] = str(rejected_rows[0].get("code") or "FILE_PARSE_FAILED")
                    accepted.extend(outcome.get("accepted") or [])
                    rejected.extend(outcome.get("rejected") or [])
                    warnings.extend(outcome.get("warnings") or [])
                    cache_hits += int(outcome.get("cache_hits") or 0)
                    completed += 1
                    _submit_next()
                    _merge_active(
                        progress={
                            "percent": int((completed / max(1, total)) * 95),
                            "current_file": list(active_files.values()),
                            "files": {
                                "completed": completed,
                                "accepted": len(accepted),
                                "rejected": len(rejected),
                                "total": total + len(initial_rejected),
                                "items": file_items,
                            },
                            "cache_hits": cache_hits,
                        },
                    )
                if stop_processing:
                    break
        if not _lease_active():
            return
        if _status() == "cancel_requested":
            _mark_cancelled(completed=completed)
            return

        result = {
            "accepted": accepted,
            "saved": accepted,
            "rejected": rejected,
            "warnings": warnings,
            "cache_hits": cache_hits,
            "parser_version": PARSER_VERSION,
        }
        if (mandatory and rejected) or not accepted:
            error = {
                "code": "MANDATORY_SOURCE_REJECTED" if mandatory else "ALL_FILES_REJECTED",
                "message": "强制资料存在未解析文件，已阻止进入生成。" if mandatory else "所有文件均未完成解析。",
                "action": "修复或转换失败文件后重新导入。",
            }
            failed = transition_job(
                job_id,
                allowed_from={"running"},
                status="failed",
                expected_attempt_id=lease_attempt_id,
                expected_owner_instance_id=lease_owner_instance_id,
                revoke_lease=True,
                error=error,
                result=result,
                progress={"work_state": "idle", "current_file": None, "detail": error["message"]},
            )
            if failed is not None:
                append_runtime_event(
                    job_id,
                    "ingest_failed",
                    code=error["code"],
                    rejected=len(rejected),
                )
            elif _status() == "cancel_requested":
                _mark_cancelled(completed=completed)
            return
        if rejected:
            warnings.append(
                {
                    "code": "OPTIONAL_SOURCE_DEGRADED",
                    "message": f"{len(rejected)} 个可选资料未完成解析。",
                }
            )
            result["warnings"] = warnings
        succeeded = transition_job(
            job_id,
            allowed_from={"running"},
            status="succeeded",
            expected_attempt_id=lease_attempt_id,
            expected_owner_instance_id=lease_owner_instance_id,
            revoke_lease=True,
            result=result,
            error=None,
            progress={
                "percent": 100,
                "phase": "ingest",
                "stage": "done",
                "work_state": "idle",
                "current_file": None,
                "detail": "资料导入完成。",
                "warnings": warnings,
            },
        )
        if succeeded is not None:
            append_runtime_event(
                job_id,
                "ingest_succeeded",
                accepted=len(accepted),
                rejected=len(rejected),
                cache_hits=cache_hits,
            )
        elif _status() == "cancel_requested":
            _mark_cancelled(completed=completed)
    except JobLeaseLostError:
        if _status() == "cancel_requested":
            _mark_cancelled(completed=locals().get("completed", 0))
        return
    except Exception as exc:  # noqa: BLE001 - job worker persists a stable terminal projection
        logger.warning("ingest job worker failed: %s", type(exc).__name__)
        if _cancel_requested():
            _mark_cancelled(completed=locals().get("completed", 0))
        else:
            failed = transition_job(
                job_id,
                allowed_from={"running"},
                status="failed",
                expected_attempt_id=lease_attempt_id,
                expected_owner_instance_id=lease_owner_instance_id,
                revoke_lease=True,
                error={
                    "code": "INGEST_WORKER_FAILED",
                    "message": "资料导入工作进程异常终止。",
                    "action": "检查失败文件与后台日志后显式重试导入。",
                    "error_type": type(exc).__name__,
                },
                progress={
                    "phase": "ingest",
                    "stage": "failed",
                    "work_state": "idle",
                    "current_file": None,
                },
            )
            if failed is not None:
                append_runtime_event(
                    job_id,
                    "ingest_failed",
                    code="INGEST_WORKER_FAILED",
                    error_type=type(exc).__name__,
                )
            elif _status() == "cancel_requested":
                _mark_cancelled(completed=locals().get("completed", 0))
    finally:
        spool_dir = INGEST_SPOOL_DIR / job_id
        if spool_dir.exists():
            shutil.rmtree(spool_dir, ignore_errors=True)


@router.post("/jobs")
async def create_ingest_job(
    files: Annotated[list[UploadFile], File()],
    project_id: str | None = None,
    source_hint: str | None = None,
    session_id: str | None = None,
    workspace_dir: str | None = None,
):
    if not files:
        raise HTTPException(status_code=400, detail={"code": "NO_FILES"})
    if len(files) > MAX_BATCH_FILES:
        raise HTTPException(
            status_code=413,
            detail={"code": "TOO_MANY_FILES", "max_files": MAX_BATCH_FILES},
        )
    job_id = create_job(
        {
            "action": "ingest",
            "project_id": project_id,
            "source_hint": source_hint,
            "file_count": len(files),
            "filenames": [Path(str(item.filename or "upload.bin")).name for item in files],
        }
    )
    spool_dir = INGEST_SPOOL_DIR / job_id
    spool_dir.mkdir(parents=True, exist_ok=False, mode=0o700)
    entries: list[dict[str, Any]] = []
    rejected: list[dict[str, Any]] = []
    total_bytes = 0
    seen: dict[str, str] = {}
    try:
        for index, upload in enumerate(files):
            filename = Path(str(upload.filename or "upload.bin")).name or "upload.bin"
            target = spool_dir / f"{index:03d}_{filename}"
            digest = hashlib.sha256()
            size = 0
            await upload.seek(0)
            with target.open("wb") as handle:
                while True:
                    chunk = await upload.read(1024 * 1024)
                    if not chunk:
                        break
                    size += len(chunk)
                    total_bytes += len(chunk)
                    if size > MAX_UPLOAD_BYTES:
                        raise HTTPException(
                            status_code=413,
                            detail={"code": "UPLOAD_TOO_LARGE", "filename": filename, "max_bytes": MAX_UPLOAD_BYTES},
                        )
                    if total_bytes > MAX_BATCH_BYTES:
                        raise HTTPException(
                            status_code=413,
                            detail={"code": "BATCH_TOO_LARGE", "max_bytes": MAX_BATCH_BYTES},
                        )
                    digest.update(chunk)
                    handle.write(chunk)
            if size <= 0:
                target.unlink(missing_ok=True)
                rejected.append({"filename": filename, "code": "EMPTY_FILE"})
                continue
            sha256 = digest.hexdigest()
            if sha256 in seen:
                target.unlink(missing_ok=True)
                rejected.append(
                    {
                        "filename": filename,
                        "code": "DUPLICATE_FILE",
                        "sha256": sha256,
                        "duplicate_of": seen[sha256],
                    }
                )
                continue
            seen[sha256] = filename
            entries.append(
                {
                    "filename": filename,
                    "path": str(target),
                    "bytes": size,
                    "sha256": sha256,
                    "file_id": sha256,
                }
            )
    except Exception:
        merge_job(job_id, status="failed", error={"code": "INGEST_UPLOAD_REJECTED"})
        shutil.rmtree(spool_dir, ignore_errors=True)
        raise

    if not entries:
        merge_job(
            job_id,
            status="failed",
            error={"code": "ALL_FILES_REJECTED", "message": "所有上传文件为空或重复。"},
            result={"accepted": [], "rejected": rejected},
        )
        shutil.rmtree(spool_dir, ignore_errors=True)
        raise HTTPException(
            status_code=422,
            detail={"code": "ALL_FILES_REJECTED", "job_id": job_id, "rejected": rejected},
        )

    options = {
        "project_id": project_id,
        "source_hint": source_hint,
        "session_id": session_id,
        "workspace_dir": workspace_dir,
    }
    queue_depth = submit_local_job(job_id, _run_ingest_job, job_id, entries, options, rejected)
    return {
        "ok": True,
        "job_id": job_id,
        "status": "queued",
        "queue_depth": queue_depth,
        "files_total": len(entries) + len(rejected),
        "bytes_total": total_bytes,
        "file_ids": [entry["file_id"] for entry in entries],
        "rejected": rejected,
    }


@router.get("/jobs/{job_id}")
async def get_ingest_job(job_id: str):
    job = get_job(job_id)
    if not job or str((job.get("payload") or {}).get("action") or "") != "ingest":
        raise HTTPException(status_code=404, detail={"code": "INGEST_JOB_NOT_FOUND"})
    return {
        "ok": True,
        "job": {
            "job_id": job.get("job_id"),
            "status": job.get("status"),
            "created_at": job.get("created_at"),
            "updated_at": job.get("updated_at"),
            "progress": job.get("progress") or {},
            "result": job.get("result") or {},
            "error": job.get("error"),
        },
    }


@router.get("/ping")
async def ping():
    return {"module": "ingest", "status": "ok"}


@router.post("/upload")
async def upload(
    files: Annotated[list[UploadFile], File()],
    project_id: str | None = None,
    source_hint: str | None = None,
    session_id: str | None = None,
    workspace_dir: str | None = None,
    library_scope: str | None = None,
    project_type: str | None = None,
    title: str | None = None,
    tags: str | None = None,
    chapter_scope: str | None = None,
    process_scope: str | None = None,
    summary: str | None = None,
    style_profile: str | None = None,
    caption: str | None = None,
    description: str | None = None,
    usable: bool | str | None = True,
):
    result = await _handle_upload(
        files,
        project_id=project_id,
        source_hint=source_hint,
        session_id=session_id,
        workspace_dir=workspace_dir,
        library_scope=library_scope,
        project_type=project_type,
        title=title,
        tags=tags,
        chapter_scope=chapter_scope,
        process_scope=process_scope,
        summary=summary,
        style_profile=style_profile,
        caption=caption,
        description=description,
        usable=usable,
    )
    return _attach_local_adapter_ingest_result(result, project_id=project_id)


@router.post("/ingest")
async def ingest(
    files: Annotated[list[UploadFile], File()],
    project_id: str | None = None,
    source_hint: str | None = None,
    session_id: str | None = None,
    workspace_dir: str | None = None,
    library_scope: str | None = None,
    project_type: str | None = None,
    title: str | None = None,
    tags: str | None = None,
    chapter_scope: str | None = None,
    process_scope: str | None = None,
    summary: str | None = None,
    style_profile: str | None = None,
    caption: str | None = None,
    description: str | None = None,
    usable: bool | str | None = True,
):
    result = await _handle_upload(
        files,
        project_id=project_id,
        source_hint=source_hint,
        session_id=session_id,
        workspace_dir=workspace_dir,
        library_scope=library_scope,
        project_type=project_type,
        title=title,
        tags=tags,
        chapter_scope=chapter_scope,
        process_scope=process_scope,
        summary=summary,
        style_profile=style_profile,
        caption=caption,
        description=description,
        usable=usable,
    )
    return _attach_local_adapter_ingest_result(result, project_id=project_id)
