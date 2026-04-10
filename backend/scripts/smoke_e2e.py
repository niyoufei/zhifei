#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from __future__ import annotations

import json
import math
import os
import subprocess
import sys
import tempfile
import time
import urllib.error
import urllib.request
from pathlib import Path
from typing import Any
from urllib.parse import urlencode

from docx import Document
from openpyxl import Workbook


# Current release smoke for the V2 page main chain:
# app.py -> backend.app.main:app -> /actions/* -> actions_bridge -> zhifei_autoplan/*
BASE = os.environ.get("ZF_SMOKE_BASE_URL", "http://127.0.0.1:8010").rstrip("/")
ACTIONS_KEY = os.environ.get("ZF_ACTIONS_KEY", "zf-webui-key").strip()
ROOT_DIR = Path(__file__).resolve().parents[2]
PIPELINE_SCRIPT = ROOT_DIR / "scripts" / "run_actions_pipeline.py"
PLAN_GLOBAL_INSTRUCTION = "release-smoke-plan"
DEFAULT_OUTLINE = [
    "工程概况",
    "施工部署",
    "主要施工方法",
    "确保工程质量的技术组织措施",
    "确保安全生产的技术组织措施",
]


def _env_float(name: str, default: float, *, minimum: float) -> float:
    try:
        value = float(os.environ.get(name, default))
    except Exception:
        value = float(default)
    return max(float(minimum), value)


POLL_INTERVAL_SEC = _env_float("ZF_SMOKE_POLL_INTERVAL_SEC", 0.5, minimum=0.2)
MAX_WAIT_SEC = _env_float("ZF_SMOKE_MAX_WAIT_SEC", 150.0, minimum=30.0)
STAGE_GRACE_SEC = _env_float("ZF_SMOKE_STAGE_GRACE_SEC", 90.0, minimum=0.0)


def http_request(
    method: str,
    path: str,
    payload: dict | None = None,
    timeout: int = 60,
    *,
    extra_headers: dict[str, str] | None = None,
) -> tuple[int, bytes, dict[str, Any]]:
    url = BASE + path
    headers = {"x-actions-key": ACTIONS_KEY}
    if isinstance(extra_headers, dict):
        headers.update({str(k): str(v) for k, v in extra_headers.items()})
    data = b""
    if payload is not None:
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        headers["Content-Type"] = "application/json"
    req = urllib.request.Request(url, data=data, headers=headers, method=method)
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            return resp.getcode(), resp.read(), dict(resp.headers)
    except urllib.error.HTTPError as e:
        return e.code, e.read() or b"", dict(e.headers or {})
    except Exception as e:
        print("[FATAL] cannot reach server:", repr(e))
        sys.exit(2)


def ensure(cond: bool, msg: str) -> None:
    if cond:
        print(f"[OK] {msg}")
        return
    print(f"[FAIL] {msg}")
    raise SystemExit(1)


def decode_json(raw: bytes) -> dict:
    try:
        return json.loads(raw.decode("utf-8", errors="replace"))
    except Exception:
        return {}


def _poll_budget(total_wait_sec: float, poll_interval_sec: float) -> int:
    total_wait = max(0.0, float(total_wait_sec))
    poll_interval = max(0.01, float(poll_interval_sec))
    return max(1, int(math.ceil(total_wait / poll_interval)))


def _job_status_tuple(job: dict[str, Any]) -> tuple[str, str, int]:
    status = str(job.get("status") or "").strip().lower()
    progress = job.get("progress") if isinstance(job.get("progress"), dict) else {}
    stage = str(progress.get("stage") or "").strip().lower()
    try:
        percent = int(progress.get("percent") or 0)
    except Exception:
        percent = 0
    return status, stage, percent


def _should_extend_poll_grace(job: dict[str, Any]) -> bool:
    status, stage, percent = _job_status_tuple(job if isinstance(job, dict) else {})
    if status not in {"queued", "running"}:
        return False
    if stage in {"variant_running", "exporting", "cross_variant_check"}:
        return True
    return percent >= 80


def _wait_server_ready() -> None:
    last_err: Exception | None = None
    for _ in range(20):
        try:
            st, body, _ = http_request("GET", "/health", timeout=3)
            if st == 200 and decode_json(body).get("ok") is True:
                return
        except Exception as exc:  # pragma: no cover - defensive only
            last_err = exc
        time.sleep(0.4)
    detail = repr(last_err) if last_err is not None else "status_not_200"
    ensure(False, f"server reachable via /health ({detail})")


def _write_smoke_fixtures(tmp_dir: Path) -> dict[str, Path]:
    fixtures_dir = tmp_dir / "fixtures"
    fixtures_dir.mkdir(parents=True, exist_ok=True)

    tender_path = fixtures_dir / "tender.docx"
    doc = Document()
    doc.add_paragraph("项目名称：E2E 冒烟综合楼工程")
    doc.add_paragraph("项目编号：E2E-SMOKE-001")
    doc.add_paragraph("技术文件详细评审标准")
    for line in DEFAULT_OUTLINE:
        doc.add_paragraph(line)
    doc.add_paragraph("施工组织设计排版要求：纸张 A4，正文宋体小四，1.5 倍行距。")
    doc.save(tender_path)

    boq_path = fixtures_dir / "boq.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "清单"
    ws.append(["序号", "项目编码", "项目名称", "项目特征描述", "计量单位", "工程量"])
    ws.append([1, "010101001001", "土方开挖", "机械开挖 基坑深度2m", "m3", 120.0])
    ws.append([2, "010201001001", "混凝土垫层", "C15 商品混凝土", "m3", 32.5])
    wb.save(boq_path)

    plan_path = fixtures_dir / "plan.json"
    plan_payload = {
        "outline": DEFAULT_OUTLINE,
        "global_instruction": PLAN_GLOBAL_INSTRUCTION,
        "style": {"body_font": "宋体", "paper": "A4"},
    }
    plan_path.write_text(json.dumps(plan_payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"tender": tender_path, "boq": boq_path, "plan": plan_path}


def _parse_pipeline_output(text: str) -> dict[str, str]:
    out: dict[str, str] = {}
    for raw in str(text or "").splitlines():
        line = raw.strip()
        if not line or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        if key in {"job_id", "saved_to"} and key not in out:
            out[key] = value.strip()
    return out


def _resolve_repo_path(raw_path: str) -> Path:
    path = Path(str(raw_path or "").strip())
    return path if path.is_absolute() else (ROOT_DIR / path)


def _download_artifact(job_id: str, kind: str, variant: int, out_path: Path, *, session_id: str | None = None) -> None:
    query = urlencode(
        {
            "job_id": job_id,
            "kind": kind,
            "variant": variant,
            "session_id": str(session_id or "").strip(),
        }
    )
    st, body, _ = http_request("GET", f"/actions/download?{query}", timeout=120)
    ensure(st == 200 and len(body) > 0, f"/actions/download kind={kind} variant={variant} status={st}")
    out_path.write_bytes(body)
    ensure(out_path.exists() and out_path.stat().st_size > 0, f"downloaded {kind} saved: {out_path}")


def main() -> None:
    print(f"[INFO] base={BASE}")
    print(f"[INFO] pipeline_script={PIPELINE_SCRIPT}")
    ensure(PIPELINE_SCRIPT.exists(), f"pipeline script exists: {PIPELINE_SCRIPT}")
    _wait_server_ready()

    with tempfile.TemporaryDirectory(prefix="docgen-release-smoke-") as tmp:
        tmp_dir = Path(tmp)
        fixtures = _write_smoke_fixtures(tmp_dir)
        project_id = f"e2e_smoke_{int(time.time())}"
        session_id = f"smoke-{project_id}"
        topic = "E2E自检-施工组织设计主链验收"

        print("[STEP] run actions main chain pipeline (parse/save/generate/download)")
        cmd = [
            sys.executable,
            str(PIPELINE_SCRIPT),
            "--base-url",
            BASE,
            "--actions-key",
            ACTIONS_KEY,
            "--topic",
            topic,
            "--project-id",
            project_id,
            "--session-id",
            session_id,
            "--tender",
            str(fixtures["tender"]),
            "--boq",
            str(fixtures["boq"]),
            "--plan-json",
            str(fixtures["plan"]),
            "--variants",
            "1",
            "--dry-run",
            "--no-gate",
            "--timeout-sec",
            str(int(MAX_WAIT_SEC)),
            "--poll-sec",
            str(POLL_INTERVAL_SEC),
        ]
        proc = subprocess.run(
            cmd,
            cwd=str(ROOT_DIR),
            capture_output=True,
            text=True,
            timeout=max(120, int(MAX_WAIT_SEC) + 60),
            env={**os.environ, "ZF_ACTIONS_KEY": ACTIONS_KEY},
        )
        if proc.returncode != 0:
            detail = (proc.stderr or proc.stdout or "").strip()[-1200:]
            ensure(False, f"run_actions_pipeline rc={proc.returncode} detail={detail}")

        parsed = _parse_pipeline_output(proc.stdout)
        job_id = str(parsed.get("job_id") or "").strip()
        saved_to = str(parsed.get("saved_to") or "").strip()
        ensure(bool(job_id), "run_actions_pipeline returned job_id")
        ensure(bool(saved_to), "run_actions_pipeline returned saved_to")
        print(f"[INFO] job_id={job_id}")

        output_dir = _resolve_repo_path(saved_to)
        ensure(output_dir.exists(), f"saved_to exists: {output_dir}")
        ensure((output_dir / f"autoplan_{job_id}.json").exists(), f"saved json exists in {saved_to}")
        ensure((output_dir / f"autoplan_{job_id}_v1.docx").exists(), f"saved docx exists in {saved_to}")
        ensure((output_dir / f"autoplan_{job_id}_compare_v1.docx").exists(), f"saved compare_docx exists in {saved_to}")

        print("[STEP] verify /actions/plan/get")
        plan_query = urlencode({"project_id": project_id, "session_id": session_id})
        st_plan, body_plan, _ = http_request("GET", f"/actions/plan/get?{plan_query}", timeout=30)
        ensure(st_plan == 200, f"/actions/plan/get status={st_plan}")
        plan_resp = decode_json(body_plan)
        ensure(plan_resp.get("ok") is True, "/actions/plan/get ok=true")
        plan_obj = plan_resp.get("plan") if isinstance(plan_resp.get("plan"), dict) else {}
        ensure(
            str(plan_obj.get("global_instruction") or "").strip() == PLAN_GLOBAL_INSTRUCTION,
            "plan/get returned expected global_instruction",
        )
        ensure(len(plan_obj.get("outline") or []) >= 3, "plan/get returned outline")

        print("[STEP] verify /actions/job_status and /actions/result")
        job_query = urlencode({"job_id": job_id, "session_id": session_id})
        st_job, body_job, _ = http_request("GET", f"/actions/job_status?{job_query}", timeout=30)
        ensure(st_job == 200, f"/actions/job_status status={st_job}")
        job_resp = decode_json(body_job)
        job = job_resp.get("job") if isinstance(job_resp.get("job"), dict) else {}
        status = str(job.get("status") or "").strip().lower()
        ensure(status == "done", f"/actions/job_status status=done (actual={status})")

        result_query = urlencode({"job_id": job_id, "variant": 1, "include_sections": "false", "session_id": session_id})
        st_result, body_result, _ = http_request("GET", f"/actions/result?{result_query}", timeout=30)
        ensure(st_result == 200, f"/actions/result status={st_result}")
        result_resp = decode_json(body_result)
        ensure(result_resp.get("ok") is True, "/actions/result ok=true")
        ensure(len(result_resp.get("outline") or []) >= 1, "/actions/result returned outline")
        files = result_resp.get("files") if isinstance(result_resp.get("files"), dict) else {}
        ensure(_resolve_repo_path(str(files.get("json") or "")).exists(), "/actions/result json artifact exists")
        ensure(_resolve_repo_path(str(files.get("docx") or "")).exists(), "/actions/result docx artifact exists")
        ensure(_resolve_repo_path(str(files.get("compare_docx") or "")).exists(), "/actions/result compare_docx artifact exists")

        print("[STEP] verify /actions/download")
        downloads_dir = tmp_dir / "downloads"
        downloads_dir.mkdir(parents=True, exist_ok=True)
        _download_artifact(job_id, "json", 1, downloads_dir / f"{job_id}.json", session_id=session_id)
        _download_artifact(job_id, "docx", 1, downloads_dir / f"{job_id}.docx", session_id=session_id)
        _download_artifact(job_id, "compare_docx", 1, downloads_dir / f"{job_id}_compare.docx", session_id=session_id)

    print("[SUCCESS] actions main-chain release smoke passed")


if __name__ == "__main__":
    main()
