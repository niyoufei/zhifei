from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import json

def format_export_to_word(json_path: str):
    """将导出的 JSON 文件转为专业排版 Word 文档"""
    p = Path(json_path)
    if not p.exists():
        print(f"[FAIL] 未找到导出文件：{json_path}")
        return
    data = json.loads(p.read_text(encoding="utf-8"))
    meta = data.get("meta", {})
    content = data.get("data", {})

    doc = Document()
    # 标题
    title = doc.add_heading("自动化评分导出报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息区
    doc.add_heading("Ⅰ. 元信息与追溯链", level=2)
    for k, v in meta.items():
        doc.add_paragraph(f"{k}：{v}")

    # 评分结果区
    doc.add_heading("Ⅱ. 评分结果", level=2)
    if isinstance(content, dict):
        for k, v in content.items():
            doc.add_paragraph(f"{k}：{v}")
    else:
        doc.add_paragraph(str(content))

    # 结束签章
    doc.add_paragraph()
    doc.add_paragraph("—— 本文档由『专业级可追溯文档自动化生成系统』导出 ——")\
       .alignment = WD_ALIGN_PARAGRAPH.CENTER

    out_path = p.with_suffix(".docx")
    doc.save(out_path)
    print(f"[OK] 已生成专业排版文档：{out_path}")

# 示例用法（可单独运行）
if __name__ == "__main__":
    latest = sorted(Path("exports").glob("export_*.json"))[-1]
    format_export_to_word(str(latest))
