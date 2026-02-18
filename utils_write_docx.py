from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
import re

def write_compose_to_docx(sections, style, output_path="build/compose_output.docx"):
    doc = Document()

    # === 纸张大小 ===
    if style.get("paper") == "A4":
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)

    # === 页边距（cm） ===
    margins = style.get("margins", [2.0,2.0,2.0,2.0])
    section.top_margin = Cm(margins[0])
    section.right_margin = Cm(margins[1])
    section.bottom_margin = Cm(margins[2])
    section.left_margin = Cm(margins[3])

    # === 字体与字号 ===
    font_name = style.get("font", "SimSun")
    font_size = style.get("font_size", 12)

    def apply_paragraph_style(p):
        # 字体
        for run in p.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)

        # 行距
        p.paragraph_format.line_spacing = style.get("line_spacing", 1.5)

    def extract_priority(text: str) -> float | None:
        if not isinstance(text, str):
            return None
        m = re.search(r"章节优先级权重[:：]\s*([0-9.]+)", text)
        if not m:
            return None
        try:
            return float(m.group(1))
        except Exception:
            return None

    # === 写入章节内容 ===
    high_priority_titles = []

    for sec in sections:
        # 标题（高权重章节加粗/放大，并标记为重点）
        priority = extract_priority(sec.get("content", ""))
        title_level = 1
        title_text = sec["title"]
        if isinstance(priority, float) and priority >= 0.7:
            title_text = f"【重点】{title_text}"
            high_priority_titles.append(sec["title"])
        title_p = doc.add_heading(title_text, level=title_level)
        apply_paragraph_style(title_p)
        if isinstance(priority, float) and priority >= 0.7:
            for run in title_p.runs:
                run.bold = True
                run.font.size = Pt(font_size + 2)

        # 正文
        p = doc.add_paragraph(sec["content"])
        apply_paragraph_style(p)

        if style.get("auto_page_break"):
            doc.add_page_break()

    # 末尾追加：高权重章节清单（评审速览）+ 缺口摘要
    if high_priority_titles:
        doc.add_page_break()
        p = doc.add_heading("高权重章节清单（评审速览）", level=1)
        apply_paragraph_style(p)
        for t in high_priority_titles:
            item = doc.add_paragraph(f"- {t}")
            apply_paragraph_style(item)

    # 缺口清单摘要（从章节中提取）
    gaps = []
    for sec in sections:
        content = sec.get("content") or ""
        if "缺口清单汇总" in content:
            for line in content.splitlines():
                if line.startswith("- "):
                    gaps.append(line.replace("- ", ""))
    if gaps:
        doc.add_page_break()
        p = doc.add_heading("缺口清单摘要（评审速览）", level=1)
        apply_paragraph_style(p)
        for g in gaps[:30]:
            item = doc.add_paragraph(f"- {g}")
            apply_paragraph_style(item)

    doc.save(output_path)
    return output_path
