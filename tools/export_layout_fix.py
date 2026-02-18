#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
export_layout_fix.py
- Single-shot, idempotent print layout optimizer for DOCX
- Features: page size/orientation, margins, heading page breaks, keep-with-next
- Writes sidecar JSON audit with all decisions.
"""
import argparse, json, os, sys
from datetime import datetime
from typing import Tuple

try:
    from docx import Document
    from docx.shared import Mm
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_BREAK
except Exception as e:
    print("ERROR: python-docx not installed or failed to import:", e, file=sys.stderr)
    sys.exit(2)

A4 = (210, 297)      # mm
LETTER = (216, 279)  # mm

def mm_tuple(text: str) -> Tuple[float,float,float,float]:
    parts = [p.strip() for p in text.split(',')]
    if len(parts) != 4:
        raise ValueError("Margins must have 4 comma-separated numbers (top,right,bottom,left) in mm.")
    return tuple(float(x) for x in parts)  # type: ignore

def apply_layout(doc: Document, paper: str, orientation: str, margins_mm: Tuple[float,float,float,float], auto_pagebreak: bool, audit: dict):
    sec = doc.sections[0]
    # Paper
    if paper.lower() == 'a4':
        w, h = A4
    elif paper.lower() == 'letter':
        w, h = LETTER
    else:
        raise ValueError("Unsupported paper size. Use A4 or Letter.")
    # Orientation
    if orientation == 'portrait':
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width, sec.page_height = Mm(w), Mm(h)
        orient_set = 'portrait'
    elif orientation == 'landscape':
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = Mm(h), Mm(w)
        orient_set = 'landscape'
    elif orientation == 'auto':
        many_cols = any(getattr(t, "columns", []) and len(t.columns) >= 7 for t in doc.tables)
        if many_cols:
            sec.orientation = WD_ORIENT.LANDSCAPE
            sec.page_width, sec.page_height = Mm(h), Mm(w)
            orient_set = 'landscape(auto_tables>=7cols)'
        else:
            sec.orientation = WD_ORIENT.PORTRAIT
            sec.page_width, sec.page_height = Mm(w), Mm(h)
            orient_set = 'portrait(auto_default)'
    else:
        raise ValueError("orientation must be one of: auto, portrait, landscape")
    # Margins
    t, r, b, l = margins_mm
    sec.top_margin = Mm(t); sec.right_margin = Mm(r); sec.bottom_margin = Mm(b); sec.left_margin = Mm(l)

    audit.update({
        "paper": paper,
        "orientation": orient_set,
        "margins_mm": {"top": t, "right": r, "bottom": b, "left": l},
    })

    # Heading pagination & keeps
    if auto_pagebreak:
        first_h1_seen = False
        for p in doc.paragraphs:
            name = (p.style.name or "").lower()
            is_h1 = ('heading 1' in name) or ('标题 1' in name)
            is_h2 = ('heading 2' in name) or ('标题 2' in name)
            if is_h1:
                p.paragraph_format.keep_with_next = True
                p.paragraph_format.keep_together = True
                if first_h1_seen:
                    run = p.insert_paragraph_before().add_run()
                    run.add_break(WD_BREAK.PAGE)
                else:
                    first_h1_seen = True
            elif is_h2:
                p.paragraph_format.keep_with_next = True
                p.paragraph_format.keep_together = True
    audit["heading_rules"] = "page_break_before(H1, except first); keep_with_next(H1,H2); keep_together(H1,H2)"

def save_with_audit(doc: Document, in_path: str, out_path: str, audit: dict):
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    meta = {
        "tool": "export_layout_fix.py",
        "version": "1.0.0",
        "timestamp": datetime.now().isoformat(),
        "input": os.path.abspath(in_path) if in_path else None,
        "output": os.path.abspath(out_path),
        "decisions": audit,
    }
    with open(out_path + ".layout.json", "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)

def process(in_path: str, out_path: str, paper: str, orientation: str, margins: Tuple[float,float,float,float], auto_pagebreak: bool):
    if not os.path.exists(in_path):
        print(f"ERROR: input file not found: {in_path}", file=sys.stderr)
        sys.exit(3)
    doc = Document(in_path)
    audit = {}
    apply_layout(doc, paper, orientation, margins, auto_pagebreak, audit)
    save_with_audit(doc, in_path, out_path, audit)
    print(f"✅ Export layout applied -> {out_path}")
    print(f"🧾 Audit -> {out_path}.layout.json")

def make_demo(demo_in: str):
    from docx import Document
    doc = Document()
    doc.add_paragraph("示例文档").style = doc.styles["Title"]
    h1 = doc.add_paragraph("第 1 章 引言"); h1.style = "Heading 1"
    doc.add_paragraph("这一段落用于演示分页和排版。")
    h2 = doc.add_paragraph("1.1 背景"); h2.style = "Heading 2"
    doc.add_paragraph("背景描述……")
    tbl = doc.add_table(rows=2, cols=7)
    for i, cell in enumerate(tbl.rows[0].cells): cell.text = f"列{i+1}"
    for i, cell in enumerate(tbl.rows[1].cells): cell.text = f"值{i+1}"
    h1b = doc.add_paragraph("第 2 章 方法"); h1b.style = "Heading 1"
    doc.add_paragraph("更多内容……")
    os.makedirs(os.path.dirname(demo_in) or ".", exist_ok=True)
    doc.save(demo_in)
    print(f"🧪 Demo DOCX generated -> {demo_in}")

def main():
    ap = argparse.ArgumentParser(description="Print/Export layout optimizer for DOCX with audit log.")
    ap.add_argument("--in", dest="in_path", default="", help="input .docx path")
    ap.add_argument("--out", dest="out_path", default="", help="output .docx path")
    ap.add_argument("--paper", default="A4", choices=["A4", "Letter"], help="paper size")
    ap.add_argument("--orientation", default="auto", choices=["auto","portrait","landscape"], help="page orientation policy")
    ap.add_argument("--margins", default="20,20,20,25", help="margins in mm: top,right,bottom,left")
    ap.add_argument("--no-pagebreak", action="store_true", help="disable auto page breaks before H1")
    ap.add_argument("--demo", action="store_true", help="generate a demo input and run on it")
    args = ap.parse_args()

    margins = mm_tuple(args.margins)

    if args.demo:
        demo_in = os.path.join("build", "_demo.docx")
        demo_out = os.path.join("build", "_demo.print.docx")
        make_demo(demo_in)
        process(demo_in, demo_out, args.paper, args.orientation, margins, auto_pagebreak=(not args.no_pagebreak))
        return

    if not args.in_path or not args.out_path:
        print("ERROR: must provide --in and --out, or use --demo", file=sys.stderr)
        sys.exit(1)
    process(args.in_path, args.out_path, args.paper, args.orientation, margins, auto_pagebreak=(not args.no_pagebreak))

if __name__ == "__main__":
    main()
